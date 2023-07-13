import win32com.client as win32
from datetime import datetime as dt
from docx2pdf import convert
from openpyxl import load_workbook
from difflib import SequenceMatcher
from Admissao.models import Colaborador, engine
from sqlalchemy.orm import sessionmaker
import docx
import os

enviar_email = int(str(input('Enviar e-mail? s/n ')).replace('s', '1').replace('n', '0'))

# Subistituir nome nos modellos de certificados e salvar como em uma pasta da área de trabalho
outlook = win32.Dispatch('outlook.application')
wb = load_workbook('Treinamento.xlsx')
pasta = '.\\Certificados'
cert_terrestre = 'Treinamento Terrestre.docx'
cert_aquatico = 'Treinamento Aquático.docx'


def extenso(datacompleta):
    dia, mes, ano = datacompleta.split('/')
    mesext = {
        '01': 'janeiro',
        '02': 'fevereiro',
        '03': 'março',
        '04': 'abril',
        '05': 'maio',
        '06': 'junho',
        '07': 'julho',
        '08': 'agosto',
        '09': 'setembro',
        '10': 'outubro',
        '11': 'novembro',
        '12': 'dezembro'
    }
    return f'{dia} de {mesext[mes]} de {ano}.'


Sessions = sessionmaker(bind=engine)
session = Sessions()
pesq = session.query(Colaborador).all()
nomes = []
dicion = {}
for p in pesq:
    nomes.append(str(p.nome).upper())
# Certificado Terrestre
x = 2
sh = wb['Terrestre']
while x <= len(sh['B']):
    t1 = docx.Document(cert_terrestre)
    nomeplan = str(sh[f'B{x}'].value)
    for pessoa in nomes:
        dicion[pessoa] = SequenceMatcher(None, nomeplan, pessoa).ratio()
    nome = [i for i in dicion if dicion[i]==max(dicion.values())][0]
    endeletr = str(sh[f'C{x}'].value)
    dia = dt.strftime(dt.strptime(str(sh[f'D{x}'].value), '%Y-%m-%d %H:%M:%S'), '%d/%m/%Y')
    diaext = extenso(dia)
    doc = t1
    for p in doc.paragraphs:
        if '#nome' in p.text:
            inline = p.runs
        # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if '#nome' in inline[i].text:
                    text = inline[i].text.replace('#nome', nome.title()).replace('#data', dia).replace('#dataextens', diaext)
                    inline[i].text = text
    doc.save(pasta+f'\\{nome} PST1.docx')
    convert(pasta+f'\\{nome} PST1.docx', pasta+f'\\{nome} PST1.pdf')
    os.remove(pasta+f'\\{nome} PST1.docx')
    if enviar_email == 1:
        email = outlook.CreateItem(0)
        email.to = endeletr
        email.Subject = 'Certificado Curso Primeiros Socorros'
        email.HTMLBody = f'''
        <p>Boa tarde,</p>
        <p></p>
        <p>Segue certificado do curso de primeiros socorros.</p>
        <p></p>
        <p>Atenciosamente,</p>
        <p><img src="\\\Qnapcia\\rh\\01 - RH\\01 - Administração.Controles\\08 - Dados, Documentos e Declarações Diversas\\Logo Cia\\Assinatura.png"></p>
        '''
        anexo = pasta+f'\\{nome} PST1.pdf'
        email.Attachments.Add(anexo)
        email.Send()
    x += 1

# Certificado Aquático
x = 2
sh = wb['Aquatico']
while x <= len(sh['B']):
    a1 = docx.Document(cert_aquatico)
    nomeplan = str(sh[f'B{x}'].value)
    for pessoa in nomes:
        dicion[pessoa] = SequenceMatcher(None, nomeplan, pessoa).ratio()
    nome = [i for i in dicion if dicion[i]==max(dicion.values())][0]
    endeletr = str(sh[f'C{x}'].value)
    dia = dt.strftime(dt.strptime(str(sh[f'D{x}'].value), '%Y-%m-%d %H:%M:%S'), '%d/%m/%Y')
    diaext = extenso(dia)
    doc = a1
    for p in doc.paragraphs:
        if '#nome' in p.text:
            inline = p.runs
        # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if '#nome' in inline[i].text:
                    text = inline[i].text.replace('#nome', nome.title()).replace('#data', dia).replace('#dataextens', diaext)
                    inline[i].text = text
    doc.save(pasta+f'\\{nome} PSA1.docx')
    convert(pasta+f'\\{nome} PSA1.docx', pasta+f'\\{nome} PSA1.pdf')
    os.remove(pasta+f'\\{nome} PSA1.docx')
    if enviar_email == 1:
        email = outlook.CreateItem(0)
        email.to = endeletr
        email.Subject = 'Certificado Curso Primeiros Socorros - Aquático'
        email.HTMLBody = f'''
                <p>Boa tarde,</p>
                <p></p>
                <p>Segue certificado do curso de primeiros socorros.</p>
                <p></p>
                <p>Atenciosamente,</p>
                <p><img src="\\\Qnapcia\\rh\\01 - RH\\01 - Administração.Controles\\08 - Dados, Documentos e Declarações Diversas\\Logo Cia\\Assinatura.png"></p>
                '''
        anexo = pasta+f'\\{nome} PSA1.pdf'
        email.Attachments.Add(anexo)
        email.Send()
    x += 1
outlook.quit()
