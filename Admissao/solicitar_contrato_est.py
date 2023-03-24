import os
from openpyxl import load_workbook as l_w
import docx
import docx2pdf
from datetime import datetime
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders
from dados_servd import em_rem, em_ti, em_if, k1


hoje = datetime.today()

wb = l_w('Cadastro Funcionários.xlsm')
sh = wb['Admitidos']
lotacao = {
    'Unidade Park Sul - qualquer departamento': ['0013', 'Thais Feitosa', 'thais.morais@ciaathletica.com.br', 'Líder Park Sul'],
    'Kids': ['0010', 'Cindy Stefanie', 'cindy.neves@ciaathletica.com.br', 'Líder Kids'],
    'Musculação': ['0007', 'Aline Kanyó', 'aline.kanyo@soucia.com.br', 'Líder Musculação'],
    'Esportes e Lutas': ['0008', 'Morgana Rossini', 'morganalourenco@yahoo.com.br', 'Líder Natação'],
    'Crossfit': ['0012', 'Guilherme Salles', 'gmoreirasalles@gmail.com', 'Líder Crossfit'],
    'Ginástica': ['0006', 'Hugo Albuquerque', 'hugo.albuquerque@ciaathletica.com.br', 'Líder Ginástica'],
    'Gestantes': ['0006', 'Hugo Albuquerque', 'hugo.albuquerque@ciaathletica.com.br', 'Líder Ginástica'],
    'Recepção': ['0003', 'Paulo Renato', 'paulo.simoes@ciaathletica.com.br', 'Gerente Vendas'],
    'Administrativo': ['0001', 'Felipe Rodrigues', 'felipe.rodrigues@ciaathletica.com.br', 'Gerente RH'],
    'Manutenção': ['0004', 'José Aparecido', 'aparecido.grota@ciaathletica.com.br', 'Gerente Manutenção'],
}
cadastro = {'nome': str(sh["C3"].value).title().strip(), 'nasc_ed': sh["D3"].value,
            'genero': str(sh["E3"].value), 'est_civ': str(sh["F3"].value),
            'pai': str(sh["M3"].value), 'mae': str(sh["N3"].value), 'end': str(sh["O3"].value),
            'num': str(sh["P3"].value), 'bairro': str(sh["Q3"].value), 'cep': str(sh["R3"].value).replace('.','').replace('-',''),
            'cid_end': str(sh["S3"].value), 'uf_end': str(sh["T3"].value), 'tel': str(sh["U3"].value).replace('(','').replace(')','').replace('-','').replace(' ',''),
            'mun_end': str(sh["AP3"].value),
            'cpf': str(sh["V3"].value).strip().replace('.', '').replace('-', '').replace(' ','').zfill(11),
            'rg': str(sh["W3"].value).strip().replace('.', '').replace('-', '').replace(' ',''), 'emissor': str(sh["X3"].value),
            'lotacao': str(lotacao[f'{sh["AG3"].value}'][0]).zfill(4),
            'cargo': str(sh["AH3"].value), 'horario': str(sh["AI3"].value), 'email': str(sh["B3"].value).strip(),
            'admissao_ed': str(sh["AL3"].value),
            'faculdade': str(sh["AV3"].value), 'semestre': str(sh["AS3"].value),
            'turno':str(sh["AT3"].value), 'conclusao': str(sh["AU3"].value),'salario': str(sh["AM3"].value),
            'hrsemanais': str(sh["AQ3"].value),'hrmensais': str(sh["AR3"].value)}
email_remetente = em_rem
email_destinatario = 'felipe.rodrigs09@gmail.com'
senha = k1
lot = lotacao[f'{sh["AG3"].value}']
caminho = r'\192.168.0.250'
modelo = f'\\{caminho}\\rh\\01 - RH\\01 - Administração.Controles\\02 - Funcionários, Departamentos e Férias\\000 - Pastas Funcionais\\00 - ATIVOS\\0 - Estagiários\\Modelo'
os.makedirs(f'\\{caminho}\\rh\\01 - RH\\01 - Administração.Controles\\02 - Funcionários, Departamentos e Férias\\000 - Pastas Funcionais\\00 - ATIVOS\\0 - Estagiários\\0 - Ainda nao iniciaram\\{str(cadastro["nome"]).upper()}')
os.makedirs(f'\\{caminho}\\rh\\01 - RH\\01 - Administração.Controles\\02 - Funcionários, Departamentos e Férias\\000 - Pastas Funcionais\\00 - ATIVOS\\0 - Estagiários\\0 - Ainda nao iniciaram\\{str(cadastro["nome"]).upper()}\\Atestados')
os.makedirs(f'\\{caminho}\\rh\\01 - RH\\01 - Administração.Controles\\02 - Funcionários, Departamentos e Férias\\000 - Pastas Funcionais\\00 - ATIVOS\\0 - Estagiários\\0 - Ainda nao iniciaram\\{str(cadastro["nome"]).upper()}\\Diversos')
os.makedirs(f'\\{caminho}\\rh\\01 - RH\\01 - Administração.Controles\\02 - Funcionários, Departamentos e Férias\\000 - Pastas Funcionais\\00 - ATIVOS\\0 - Estagiários\\0 - Ainda nao iniciaram\\{str(cadastro["nome"]).upper()}\\Contratuais')
os.makedirs(f'\\{caminho}\\rh\\01 - RH\\01 - Administração.Controles\\02 - Funcionários, Departamentos e Férias\\000 - Pastas Funcionais\\00 - ATIVOS\\0 - Estagiários\\0 - Ainda nao iniciaram\\{str(cadastro["nome"]).upper()}\\Ferias')
os.makedirs(f'\\{caminho}\\rh\\01 - RH\\01 - Administração.Controles\\02 - Funcionários, Departamentos e Férias\\000 - Pastas Funcionais\\00 - ATIVOS\\0 - Estagiários\\0 - Ainda nao iniciaram\\{str(cadastro["nome"]).upper()}\\Ponto')
os.makedirs(f'\\{caminho}\\rh\\01 - RH\\01 - Administração.Controles\\02 - Funcionários, Departamentos e Férias\\000 - Pastas Funcionais\\00 - ATIVOS\\0 - Estagiários\\0 - Ainda nao iniciaram\\{str(cadastro["nome"]).upper()}\\Recibo')
os.makedirs(f'\\{caminho}\\rh\\01 - RH\\01 - Administração.Controles\\02 - Funcionários, Departamentos e Férias\\000 - Pastas Funcionais\\00 - ATIVOS\\0 - Estagiários\\0 - Ainda nao iniciaram\\{str(cadastro["nome"]).upper()}\\Rescisao')
pasta_contratuais = f'\\{caminho}\\rh\\01 - RH\\01 - Administração.Controles\\02 - Funcionários, Departamentos e Férias\\000 - Pastas Funcionais\\00 - ATIVOS\\0 - Estagiários\\0 - Ainda nao iniciaram\\{str(cadastro["nome"]).upper()}\\Contratuais'

solicitacao = docx.Document(modelo + r'\Solicitacao MODELO - Copia.docx')
solicitacao.tables[0].rows[4].cells[0].paragraphs[0].text = str(solicitacao.tables[0].rows[4].cells[0].paragraphs[0].text).replace('#supervisor_estagio', f'{lot[1]}')
solicitacao.tables[0].rows[5].cells[1].paragraphs[0].text = str(solicitacao.tables[0].rows[5].cells[1].paragraphs[0].text).replace('#cargo', f'{lot[3]}')
solicitacao.tables[0].rows[6].cells[0].paragraphs[0].text = str(solicitacao.tables[0].rows[6].cells[0].paragraphs[0].text).replace('#email_supervisor', f'{lot[2]}')
solicitacao.tables[0].rows[9].cells[0].paragraphs[0].text = str(solicitacao.tables[0].rows[9].cells[0].paragraphs[0].text).replace('#horario', cadastro['horario'])
solicitacao.tables[0].rows[14].cells[0].paragraphs[0].text = str(solicitacao.tables[0].rows[14].cells[0].paragraphs[0].text).replace('#nome_completo', cadastro['nome'])
solicitacao.tables[0].rows[15].cells[0].paragraphs[0].text = str(solicitacao.tables[0].rows[15].cells[0].paragraphs[0].text).replace('#nasc', datetime.strftime(cadastro['nasc_ed'], '%d/%m/%Y')).replace('#rg', cadastro['rg']).replace('#cpf', cadastro['cpf'])
solicitacao.tables[0].rows[16].cells[0].paragraphs[0].text = str(solicitacao.tables[0].rows[16].cells[0].paragraphs[0].text).replace('#sexo', cadastro['genero'])
solicitacao.tables[0].rows[17].cells[0].paragraphs[0].text = str(solicitacao.tables[0].rows[17].cells[0].paragraphs[0].text).replace('#endereco', cadastro['end'])
solicitacao.tables[0].rows[18].cells[0].paragraphs[0].text = str(solicitacao.tables[0].rows[18].cells[0].paragraphs[0].text).replace('#cep', cadastro['cep']).replace('#bairro', cadastro['bairro'])
solicitacao.tables[0].rows[19].cells[0].paragraphs[0].text = str(solicitacao.tables[0].rows[19].cells[0].paragraphs[0].text).replace('#telefone', cadastro['tel'])
solicitacao.tables[0].rows[20].cells[0].paragraphs[0].text = str(solicitacao.tables[0].rows[20].cells[0].paragraphs[0].text).replace('#end_eletr', cadastro['email'])
solicitacao.tables[0].rows[22].cells[0].paragraphs[0].text = str(solicitacao.tables[0].rows[22].cells[0].paragraphs[0].text).replace('#semestre', cadastro['semestre'])
solicitacao.tables[0].rows[23].cells[0].paragraphs[0].text = str(solicitacao.tables[0].rows[23].cells[0].paragraphs[0].text).replace('#turno', cadastro['turno']).replace('#ano_concl', cadastro['conclusao'])
solicitacao.tables[0].rows[24].cells[0].paragraphs[0].text = str(solicitacao.tables[0].rows[24].cells[0].paragraphs[0].text).replace('#faculdade', cadastro['faculdade'])
solicitacao.save(pasta_contratuais + f'\\Pedido TCE {str(cadastro["nome"]).split(" ")[0]}.docx')
docx2pdf.convert(pasta_contratuais + f'\\Pedido TCE {str(cadastro["nome"]).split(" ")[0]}.docx', pasta_contratuais + f'\\Pedido TCE {str(cadastro["nome"]).split(" ")[0]}.pdf')
os.remove(pasta_contratuais + f'\\Pedido TCE {str(cadastro["nome"]).split(" ")[0]}.docx')

ficha_cadastral = docx.Document(modelo + r'\Ficha Cadastral MODELO - Copia.docx')
ficha_cadastral.tables[1].rows[0].cells[0].paragraphs[0].text = str(ficha_cadastral.tables[1].rows[0].cells[0].paragraphs[0].text).replace('#nome_completo', cadastro['nome'])
ficha_cadastral.tables[1].rows[1].cells[0].paragraphs[0].text = str(ficha_cadastral.tables[1].rows[1].cells[0].paragraphs[0].text).replace('#nasc', datetime.strftime(cadastro['nasc_ed'], '%d/%m/%Y'))
ficha_cadastral.tables[1].rows[1].cells[2].paragraphs[0].text = str(ficha_cadastral.tables[1].rows[1].cells[2].paragraphs[0].text).replace('#gen', cadastro['genero'])
ficha_cadastral.tables[1].rows[1].cells[4].paragraphs[0].text = str(ficha_cadastral.tables[1].rows[1].cells[4].paragraphs[0].text).replace('#est_civ', cadastro['est_civ'])
ficha_cadastral.tables[1].rows[2].cells[0].paragraphs[0].text = str(ficha_cadastral.tables[1].rows[2].cells[0].paragraphs[0].text).replace('#local', cadastro['end'])
ficha_cadastral.tables[1].rows[2].cells[4].paragraphs[0].text = str(ficha_cadastral.tables[1].rows[2].cells[4].paragraphs[0].text).replace('#qd', cadastro['bairro'])
ficha_cadastral.tables[1].rows[2].cells[7].paragraphs[0].text = str(ficha_cadastral.tables[1].rows[2].cells[7].paragraphs[0].text).replace('#codigo', cadastro['cep'])
ficha_cadastral.tables[1].rows[4].cells[1].paragraphs[0].text = str(ficha_cadastral.tables[1].rows[4].cells[1].paragraphs[0].text).replace('#telefone', cadastro['tel'])
ficha_cadastral.tables[1].rows[4].cells[5].paragraphs[0].text = str(ficha_cadastral.tables[1].rows[4].cells[5].paragraphs[0].text).replace('#ident', cadastro['rg'])
ficha_cadastral.tables[1].rows[5].cells[1].paragraphs[0].text = str(ficha_cadastral.tables[1].rows[5].cells[1].paragraphs[0].text).replace('#cpf#', cadastro['cpf'])
ficha_cadastral.tables[1].rows[6].cells[3].paragraphs[0].text = str(ficha_cadastral.tables[1].rows[6].cells[3].paragraphs[0].text).replace('#pai#', cadastro['pai'])
ficha_cadastral.tables[1].rows[7].cells[1].paragraphs[0].text = str(ficha_cadastral.tables[1].rows[7].cells[1].paragraphs[0].text).replace('#mae#', cadastro['mae'])
ficha_cadastral.tables[1].rows[8].cells[0].paragraphs[0].text = str(ficha_cadastral.tables[1].rows[8].cells[0].paragraphs[0].text).replace('#end_eletr', cadastro['email'])
ficha_cadastral.tables[1].rows[8].cells[1].paragraphs[0].text = str(ficha_cadastral.tables[1].rows[8].cells[1].paragraphs[0].text).replace('#depart', str(sh["AG3"].value))
ficha_cadastral.save(pasta_contratuais + f'\\Ficha Cadastral {str(cadastro["nome"]).split(" ")[0]}.docx')
docx2pdf.convert(pasta_contratuais + f'\\Ficha Cadastral {str(cadastro["nome"]).split(" ")[0]}.docx', pasta_contratuais + f'\\Ficha Cadastral {str(cadastro["nome"]).split(" ")[0]}.pdf')
os.remove(pasta_contratuais + f'\\Ficha Cadastral {str(cadastro["nome"]).split(" ")[0]}.docx')

carta_banco = docx.Document(modelo + r'\Abertura Conta MODELO.docx') # alterar parágrafo 14:  #bairro #desde(alterar modelo) alterar parágrafo 24: data de hoje
carta_banco.paragraphs[14].text = str(carta_banco.paragraphs[14].text).replace('#nome_completo', cadastro['nome']).replace('#rg', cadastro['rg']).replace('#cpf', cadastro['cpf']).replace('#endereço', cadastro['end']).replace('#cep', cadastro['cep']).replace('#bairro', cadastro['bairro']).replace('#desde#', datetime.strftime(hoje,'%d/%m/%Y'))
carta_banco.save(pasta_contratuais + f'\\Carta Banco {str(cadastro["nome"]).split(" ")[0]}.docx')
docx2pdf.convert(pasta_contratuais + f'\\Carta Banco {str(cadastro["nome"]).split(" ")[0]}.docx', pasta_contratuais + f'\\Carta Banco {str(cadastro["nome"]).split(" ")[0]}.pdf')
os.remove(pasta_contratuais + f'\\Carta Banco {str(cadastro["nome"]).split(" ")[0]}.docx')


s = smtplib.SMTP(host='smtp.office365.com', port=587)
s.starttls()
s.login(email_remetente, senha)

# Enviar carta do banco para estag
msg = MIMEMultipart()
arquivo = pasta_contratuais + f'\\Carta Banco {str(cadastro["nome"]).split(" ")[0]}.pdf'
message = f'''
            Olá, {str(cadastro["nome"]).split(" ")[0]}!\n
            \n
            Segue sua carta para abertura de conta bancária no Itaú.\n
            Você deve abrir a conta antes de iniciar os trabalhos no estágio. Ok?\n
            Você já pode buscar seu contrato no IF. Será necessário levar uma declaração de matrícula do seu curso.\n
            \n
            Atenciosamente,\n
            Felipe Rodrigues
            '''
html = r'''
    <html>
        <body>
            <img src="C:\Users\Felipe\PycharmProjects\AutomacaoCia\Admissao\static\assinatura.png">
        </body>
    </html>
    '''

# setup the parameters of the message
msg['From'] = email_remetente
msg['To'] = cadastro['email']
msg['Subject'] = "Carta para Abertura de conta"
msg.attach(MIMEText(message, 'plain', _charset='utf-8'))
msg.attach(MIMEText(html, "html"))
# Anexo pdf
part = MIMEBase('application', "octet-stream")
part.set_payload(open(arquivo, "rb").read())
encoders.encode_base64(part)
part.add_header('Content-Disposition', 'attachment', filename=f'Carta Banco {str(cadastro["nome"]).split(" ")[0]}.pdf')
msg.attach(part)
s.send_message(msg)
del msg


# Enviar ficha cadastral para Wallace
msg = MIMEMultipart()
arquivo = pasta_contratuais + f'\\Ficha Cadastral {str(cadastro["nome"]).split(" ")[0]}.pdf'
message = f'''Oi, Wallace!\n\nSegue a ficha cadastral do(a) estagiário(a) {cadastro["nome"]}.\n\nAbs.,\nFelipe'''
html = r'''
    <html>
        <body>
            <img src="C:\Users\Felipe\PycharmProjects\AutomacaoCia\Admissao\static\assinatura.png">
        </body>
    </html>
    '''

# setup the parameters of the message
msg['From'] = email_remetente
msg['To'] = em_ti
msg['Subject'] = f"Ficha Cadastral {str(cadastro['nome']).split(' ')[0]}"
msg.attach(MIMEText(message, 'plain', _charset='utf-8'))
msg.attach(MIMEText(html, "html"))
# Anexo pdf
part = MIMEBase('application', "octet-stream")
part.set_payload(open(arquivo, "rb").read())
encoders.encode_base64(part)
part.add_header('Content-Disposition', 'attachment', filename=f'Ficha Cadastral {str(cadastro["nome"]).split(" ")[0]}.pdf')
msg.attach(part)
s.send_message(msg)
del msg

# Enviar pedido de TCE para IF
msg = MIMEMultipart()
arquivo = pasta_contratuais + f'\\Pedido TCE {str(cadastro["nome"]).split(" ")[0]}.pdf'
message = f'''
            Olá!\n
            \n
            Segue pedido de TCE do(a) estagiário(a) {cadastro["nome"]}.\n
            \n
            Atenciosamente,\n
            Felipe Rodrigues
            '''
html = r'''
    <html>
        <body>
            <img src="C:\Users\Felipe\PycharmProjects\AutomacaoCia\Admissao\static\assinatura.png">
        </body>
    </html>
    '''
# setup the parameters of the message
msg['From'] = email_remetente
msg['To'] = em_if
msg['Subject'] = f"Pedido TCE {str(cadastro['nome']).split(' ')[0]}"
msg.attach(MIMEText(message, 'plain', _charset='utf-8'))
msg.attach(MIMEText(html, "html"))
# Anexo pdf
part = MIMEBase('application', "octet-stream")
part.set_payload(open(arquivo, "rb").read())
encoders.encode_base64(part)
part.add_header('Content-Disposition', 'attachment', filename=f'Pedido TCE {str(cadastro["nome"]).split(" ")[0]}.pdf')
msg.attach(part)
s.send_message(msg)
del msg
s.quit()
