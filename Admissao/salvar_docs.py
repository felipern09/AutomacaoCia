import docx
import docx2pdf
import os
import pyautogui as pa
import pyperclip as pp
import time as t
from sqlalchemy.orm import sessionmaker
from models import Colaborador, engine

matricula = int(input('Digite a matrícula do colaborador: '))
Sessions = sessionmaker(bind=engine)
session = Sessions()
pessoa = session.query(Colaborador).filter_by(matricula=matricula).first()

p_pessoa = r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e Férias' \
          r'\000 - Pastas Funcionais\00 - ATIVOS\{}'.format(pessoa.nome)
p_atestado = r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e Férias' \
          r'\000 - Pastas Funcionais\00 - ATIVOS\{}\Atestado'.format(pessoa.nome)
p_contr = r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e Férias' \
          r'\000 - Pastas Funcionais\00 - ATIVOS\{}\Contratuais'.format(pessoa.nome)
p_diversos = r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e Férias' \
          r'\000 - Pastas Funcionais\00 - ATIVOS\{}\Diversos'.format(pessoa.nome)
p_ferias = r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e Férias' \
          r'\000 - Pastas Funcionais\00 - ATIVOS\{}\Férias'.format(pessoa.nome)
p_ponto = r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e Férias' \
          r'\000 - Pastas Funcionais\00 - ATIVOS\{}\Pontos'.format(pessoa.nome)
p_rec = r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e Férias' \
          r'\000 - Pastas Funcionais\00 - ATIVOS\{}\Recibos'.format(pessoa.nome)
p_rescisao = r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e Férias' \
          r'\000 - Pastas Funcionais\00 - ATIVOS\{}\Rescisão'.format(pessoa.nome)
p_ac = r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e Férias' \
       r'\000 - Pastas Funcionais\00 - ATIVOS\Modelo\AC Modelo.docx'
p_abconta = r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e Férias' \
            r'\000 - Pastas Funcionais\00 - ATIVOS\Modelo\Abertura Conta MODELO.docx'
p_fcadas = r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e Férias' \
           r'\000 - Pastas Funcionais\00 - ATIVOS\Modelo\Ficha Cadastral MODELO.docx'
p_recibos = r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e Férias' \
            r'\000 - Pastas Funcionais\00 - ATIVOS\Modelo\Recibo Crachá e Uniformes MODELO.docx'
p_recibovt = r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e Férias' \
             r'\000 - Pastas Funcionais\00 - ATIVOS\Modelo\Recibo VT MODELO.docx'
p_codetic = r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e Férias' \
            r'\000 - Pastas Funcionais\00 - ATIVOS\Modelo\Cod Etica MODELO.docx'
ps_contr = r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e Férias' \
           r'\000 - Pastas Funcionais\00 - ATIVOS\{}\Contratuais\Contrato.pdf'.format(pessoa.nome)
ps_acordo = r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e Férias' \
            r'\000 - Pastas Funcionais\00 - ATIVOS\{}\Contratuais\Acordo Banco de Horas.pdf'.format(pessoa.nome)
ps_recctps = r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e Férias' \
            r'\000 - Pastas Funcionais\00 - ATIVOS\{}\Contratuais\Recibo de Entrega e Dev CTPS.pdf'.format(pessoa.nome)
ps_anotctps = r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e Férias' \
            r'\000 - Pastas Funcionais\00 - ATIVOS\{}\Contratuais\Anotacoes CTPS.pdf'.format(pessoa.nome)
ps_termovt = r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e Férias' \
            r'\000 - Pastas Funcionais\00 - ATIVOS\{}\Contratuais\Termo Opcao VT.pdf'.format(pessoa.nome)
ps_contrato = r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e Férias' \
            r'\000 - Pastas Funcionais\00 - ATIVOS\{}\Contratuais\Contrato de Trabalho.pdf'.format(pessoa.nome)
ps_ficha = r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e Férias' \
            r'\000 - Pastas Funcionais\00 - ATIVOS\{}\Contratuais\Folha de Registro.pdf'.format(pessoa.nome)
try:
    os.mkdir(p_pessoa)
    os.mkdir(p_atestado)
    os.mkdir(p_contr)
    os.mkdir(p_diversos)
    os.mkdir(p_ferias)
    os.mkdir(p_ponto)
    os.mkdir(p_rec)
    os.mkdir(p_rescisao)
except:
    pass

lotacao = {
    'Unidade Park Sul - Qualquer Departamento': ['0013', 'Thais Feitosa', 'thais.morais@ciaathletica.com.br', 'Líder Park Sul'],
    'Kids': ['0010', 'Cindy Stefanie', 'cindy.neves@ciaathletica.com.br', 'Líder Kids'],
    'Musculação': ['0007', 'Aline Kanyó', 'aline.kanyo@soucia.com.br', 'Líder Musculação'],
    'Esportes e Lutas': ['0008', 'Morgana Rossini', 'morganalourenco@yahoo.com.br', 'Líder Natação'],
    'Crossfit': ['0012', 'Guilherme Salles', 'gmoreirasalles@gmail.com', 'Líder Crossfit'],
    'Ginástica': ['0006', 'Hugo Albuquerque', 'hugo.albuquerque@ciaathletica.com.br', 'Líder Ginástica'],
    'Gestantes': ['0006', 'Hugo Albuquerque', 'hugo.albuquerque@ciaathletica.com.br', 'Líder Ginástica'],
    'Recepção': ['0003', 'Paulo Renato', 'paulo.simoes@ciaathletica.com.br', 'Gerente Vendas'],
    'Administrativo': ['0001', 'Felipe Rodrigues', 'felipe.rodrigues@ciaathletica.com.br', 'Gerente RH'],
    'Manutenção': ['0004', 'José Aparecido', 'aparecido.grota@ciaathletica.com.br', 'Gerente Manutenção'],
}

abert_c = docx.Document(p_abconta)
ac = docx.Document(p_ac)
fch_c = docx.Document(p_fcadas)
recibos = docx.Document(p_recibos)
recibovt = docx.Document(p_recibovt)
codetic = docx.Document(p_codetic)

# click dexion
pa.click(pa.center(pa.locateOnScreen('./static/Dexion.png')))

# # imprimir recibo entrega e devolução de ctps
pa.press('alt'), pa.press('r'), pa.press('a'), pa.press('r'), pa.press('e'), pa.press('tab'), pa.write(str(
    pessoa.matricula))
pa.press('tab', 3), pa.write(str(pessoa.admiss).replace('/','')), pa.press('tab'), t.sleep(0.5), pa.press('space')
t.sleep(0.5), pa.press('tab'), pa.write(str(pessoa.admiss).replace('/','')), pa.press('tab', 2)
t.sleep(1), pa.press('enter'), t.sleep(2)

# # clique no endereço de salvamento do recibo
pa.click(pa.center(pa.locateOnScreen('./static/salvar.png'))), t.sleep(1)
pp.copy(ps_recctps), pa.hotkey('ctrl', 'v'), t.sleep(0.5)
pa.press('tab', 2), t.sleep(0.5), pa.press('enter')
t.sleep(5)
# # clique para fechar recibo ctps
pa.click(pa.center(pa.locateOnScreen('./static/fechar_janela.png'))), t.sleep(0.5)
pa.click(pa.center(pa.locateOnScreen('./static/fechar_janela.png')))

# # Imprimir Acordo de Banco de horas
pa.press('alt'), pa.press('r'), pa.press('z'), pa.press('d'), pa.press('d')
pa.write("(matricula = '00{}')".format(str(pessoa.matricula))), t.sleep(1), pa.press('tab'), pa.write('2')
pa.press('tab'), pa.press('enter'), t.sleep(10)
pa.click(pa.center(pa.locateOnScreen('./static/salvar.png')))
t.sleep(1), pp.copy(ps_acordo)
pa.hotkey('ctrl', 'v'), t.sleep(1), pa.press('enter'), t.sleep(15)
# # clique para fechar acordo
pa.click(pa.center(pa.locateOnScreen('./static/fechar_janela.png'))), t.sleep(0.5)
pa.click(pa.center(pa.locateOnScreen('./static/fechar_janela.png')))

# # imprimir Anotações em CTPS
pa.press('alt'), pa.press('r'), pa.press('a'), pa.press('c'), pa.press('e'), pa.press('tab')
pa.write(str(pessoa.matricula))
pa.press('tab', 2), pa.write(str(pessoa.admiss).replace('/','')), pa.press('tab')
pa.write(str(pessoa.admiss).replace('/','')), pa.press('tab', 4), pa.press('space')
pa.press('tab'), pa.press('enter'), t.sleep(1.5)
pa.click(pa.center(pa.locateOnScreen('./static/salvar.png'))), t.sleep(1)
pp.copy(ps_anotctps), pa.hotkey('ctrl', 'v'), t.sleep(0.5), pa.press('enter')
t.sleep(2), pa.click(pa.center(pa.locateOnScreen('./static/fechar_janela.png'))), t.sleep(0.5)
pa.click(pa.center(pa.locateOnScreen('./static/fechar_janela.png')))

# # imprimir Termo VT
pa.press('alt'), pa.press('r'), pa.press('a'), pa.press('v'), pa.press('e'), pa.press('tab')
pa.write(str(pessoa.matricula)), pa.press('tab', 2), pa.write(str(pessoa.admiss).replace('/',''))
pa.press('tab'), pa.write('d'), pa.press('tab', 4), pa.press('space')
pa.press('tab', 6), pa.press('enter'), t.sleep(1.5)
pa.click(pa.center(pa.locateOnScreen('./static/salvar.png')))
pp.copy(ps_termovt), pa.hotkey('ctrl', 'v'), t.sleep(0.5), pa.press('enter')
t.sleep(2), pa.click(pa.center(pa.locateOnScreen('./static/fechar_janela.png'))), t.sleep(0.5)
pa.click(pa.center(pa.locateOnScreen('./static/fechar_janela.png')))

# # Imprimir Contrato
pa.press('alt'), pa.press('r'), pa.press('z'), pa.press('d')
if pessoa.tipo_contr == 'Horista':
    pa.press('c')
else:
    pa.press('o')

pa.write("(matricula = '00{}')".format(str(pessoa.matricula))), t.sleep(1), pa.press('tab'), pa.write('2')
pa.press('tab'), pa.press('enter'), t.sleep(5)
pa.click(pa.center(pa.locateOnScreen('./static/salvar.png')))
pp.copy(ps_contrato), pa.hotkey('ctrl', 'v'), t.sleep(0.5), pa.press('enter')
t.sleep(10), pa.click(pa.center(pa.locateOnScreen('./static/fechar_janela.png'))), t.sleep(0.5)
pa.click(pa.center(pa.locateOnScreen('./static/fechar_janela.png')))

# # Imprimir Folha de rosto de Cadastro
pa.press('alt'), pa.press('r'), pa.press('i'), pa.press('o'), pa.press('r'), pa.press('e'), pa.press('tab')
pa.write(str(pessoa.matricula)), pa.press('tab', 2), pa.write(str(pessoa.admiss).replace('/','')), pa.press('tab', 2)
pa.press('enter'), t.sleep(3)
pa.click(pa.center(pa.locateOnScreen('./static/salvar.png')))
pp.copy(ps_ficha), pa.hotkey('ctrl', 'v'), t.sleep(0.5), pa.press('enter')
t.sleep(3), pa.click(pa.center(pa.locateOnScreen('./static/fechar_janela.png'))), t.sleep(0.5)
pa.click(pa.center(pa.locateOnScreen('./static/fechar_janela.png')))

# # Alterar AC e Salvar na pasta
ac.paragraphs[1].text = str(ac.paragraphs[1].text).replace('#gerente', lotacao[str(pessoa.depto).title()][1])
ac.paragraphs[2].text = str(ac.paragraphs[2].text).replace('#nome_completo', pessoa.nome)
ac.paragraphs[3].text = str(ac.paragraphs[3].text).replace('#cargo', pessoa.cargo)
ac.paragraphs[11].text = str(ac.paragraphs[11].text).replace('#salario', pessoa.salario)
ac.save(p_contr + '\\AC.docx')
docx2pdf.convert(p_contr + '\\AC.docx', p_contr + '\\AC.pdf')
os.remove(p_contr + '\\AC.docx')

# # Alterar Abertura de Conta e salvar na pasta
abert_c.paragraphs[14].text = str(abert_c.paragraphs[14].text).replace('#nome_completo', pessoa.nome).replace(
    '#rg', pessoa.rg).replace(
    '#cpf', pessoa.cpf).replace('#endereco', pessoa.endereco).replace('#cep', pessoa.cep).replace('#bairro',pessoa.bairro).replace(
    '#cargo', pessoa.cargo).replace('#data', pessoa.nascimento)
abert_c.save(p_contr + '\\Abertura Conta.docx')
docx2pdf.convert(p_contr + '\\Abertura Conta.docx', p_contr + '\\Abertura Conta.pdf')
os.remove(p_contr + '\\Abertura Conta.docx')


# # Alterar Ficha cadastral e salvar na pasta
fch_c.paragraphs[34].text = str(fch_c.paragraphs[34].text).replace('#gerente#', lotacao[str(pessoa.depto).title()][1])
fch_c.paragraphs[9].text = str(fch_c.paragraphs[9].text).replace('#nome_completo', pessoa.nome)
fch_c.paragraphs[21].text = str(fch_c.paragraphs[21].text).replace('#cargo', pessoa.cargo).replace('#depart',
                                                                                                        str(pessoa.depto).title())
fch_c.paragraphs[19].text = str(fch_c.paragraphs[19].text).replace('#end_eletr', pessoa.email)
fch_c.paragraphs[17].text = str(fch_c.paragraphs[17].text).replace('#mae#', pessoa.mae)
fch_c.paragraphs[16].text = str(fch_c.paragraphs[16].text).replace('#pai#', pessoa.pai)
fch_c.paragraphs[15].text = str(fch_c.paragraphs[15].text).replace('#ident', pessoa.rg).replace('#cpf#',
                                                                                                     pessoa.cpf)
fch_c.paragraphs[13].text = str(fch_c.paragraphs[13].text).replace('#telefone', pessoa.tel)
fch_c.paragraphs[12].text = str(fch_c.paragraphs[12].text).replace('#codigo', pessoa.cep).replace('#cid', pessoa.cidade).replace('#uf',
                        pessoa.uf)
fch_c.paragraphs[11].text = str(fch_c.paragraphs[11].text).replace('#local', pessoa.endereco).replace('#qd', pessoa.bairro)
fch_c.paragraphs[10].text = str(fch_c.paragraphs[10].text).replace('#nasc', pessoa.nascimento).replace('#gen', pessoa.genero).replace('#est_civ', pessoa.estado_civil)
fch_c.save(p_contr + '\\Ficha Cadastral.docx')
docx2pdf.convert(p_contr + '\\Ficha Cadastral.docx', p_contr + '\\Ficha Cadastral.pdf')
os.remove(p_contr + '\\Ficha Cadastral.docx')

# # Alterar Recibos e salvar na pasta
recibos.paragraphs[4].text = str(recibos.paragraphs[4].text).replace('#nome_completo', pessoa.nome)
recibos.paragraphs[12].text = str(recibos.paragraphs[12].text).replace('#nome_completo', pessoa.nome)
recibos.paragraphs[19].text = str(recibos.paragraphs[19].text).replace('#nome_completo', pessoa.nome)
recibos.paragraphs[27].text = str(recibos.paragraphs[27].text).replace('#nome_completo', pessoa.nome)
recibos.paragraphs[40].text = str(recibos.paragraphs[40].text).replace('#nome_completo', pessoa.nome)
recibos.paragraphs[48].text = str(recibos.paragraphs[48].text).replace('#nome_completo', pessoa.nome)
recibos.save(p_contr + '\\Recibos.docx')
docx2pdf.convert(p_contr + '\\Recibos.docx', p_contr + '\\Recibos.pdf')
os.remove(p_contr + '\\Recibos.docx')

# # Alterar Código de Ética e Salvar na pasta
codetic.paragraphs[534].text = str(codetic.paragraphs[534].text).replace('#nome_completo', pessoa.nome)
codetic.paragraphs[535].text = str(codetic.paragraphs[535].text).replace('#func', pessoa.cargo)
codetic.paragraphs[537].text = str(codetic.paragraphs[537].text).replace('#nome_completo', pessoa.nome)
codetic.paragraphs[541].text = str(codetic.paragraphs[541].text).replace('#admiss', pessoa.admiss)
codetic.save(p_contr + '\\Cod Etica.docx')
docx2pdf.convert(p_contr + '\\Cod Etica.docx', p_contr + '\\Cod Etica.pdf')
os.remove(p_contr + '\\Cod Etica.docx')
