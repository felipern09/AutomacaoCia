import tkinter as tk
from tkcalendar import DateEntry
from datetime import datetime
import tkinter.filedialog
import pyautogui as pa
import pyperclip as pp
import time as t
from models import Colaborador, engine
from sqlalchemy.orm import sessionmaker
from openpyxl import load_workbook as l_w
from listas import horarios, cargos, departamentos, tipodecontrato, municipios
import os
import tkinter.filedialog
from tkinter import ttk, messagebox
from tkinter import *
import docx
import docx2pdf
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.image import MIMEImage
from email.mime.base import MIMEBase
from email import encoders
from dados_servd import em_rem, em_ti, em_if, k1
from difflib import SequenceMatcher


class MainApplication(tk.Tk):
    def __init__(self):
        super().__init__()

        self.title("Atividades DP - Cia BSB")
        self.geometry('661x550')
        self.img = PhotoImage(file='./static/Icone.png')
        self.iconphoto(False, self.img)
        self.columnconfigure(0, weight=5)
        self.rowconfigure(0, weight=5)
        for child in self.winfo_children():
            child.grid_configure(padx=1, pady=3)
        self.notebook = ttk.Notebook(self)

        self.Frame1 = Frame1(self.notebook)
        self.Frame2 = Frame2(self.notebook)
        self.Frame3 = Frame3(self.notebook)
        self.Frame4 = Frame4(self.notebook)

        self.notebook.add(self.Frame1, text='Cadastrar Funcionário')
        self.notebook.add(self.Frame2, text='Cadastrar Estagiário')
        self.notebook.add(self.Frame3, text='Cadastrar Autônomo')
        self.notebook.add(self.Frame4, text='Desligamento')

        self.notebook.pack()


class Frame1(ttk.Frame):
    def __init__(self, container):
        super().__init__()

        self.hoje = datetime.today()
        self.geral = StringVar()
        self.caminho = StringVar()
        self.labelescolh = ttk.Label(self, width=40, text="Escolher planilha de novos funcionários")
        self.labelescolh.grid(column=1, row=1, padx=25, pady=1, sticky=W)
        self.botaoescolha = ttk.Button(self, text="Escolha a planilha", command=self.selecionarfunc)
        self.botaoescolha.grid(column=1, row=1, padx=350, pady=1, sticky=W)
        self.nome = StringVar()
        self.horario = StringVar()
        self.cargo = StringVar()
        self.departamento = StringVar()
        self.tipocontr = StringVar()
        self.nomesplan = []
        # aparecer dropdown com nomes da plan
        self.labelnome = ttk.Label(self, width=20, text="Nome:")
        self.labelnome.grid(column=1, row=10, padx=25, pady=1, sticky=W)
        self.combonome = ttk.Combobox(self, values=self.nomesplan, textvariable=self.nome, width=50)
        self.combonome.grid(column=1, row=10, padx=125, pady=1, sticky=W)
        # aparecer entry para preencher matricula
        self.labelmatr = ttk.Label(self, width=20, text="Matrícula:")
        self.labelmatr.grid(column=1, row=11, padx=25, pady=1, sticky=W)
        self.entrymatr = ttk.Entry(self, width=20)
        self.entrymatr.grid(column=1, row=11, padx=125, pady=1, sticky=W)
        # aparecer entry para preencher admissao
        self.labeladmiss = ttk.Label(self, width=20, text="Admissão:")
        self.labeladmiss.grid(column=1, row=12, padx=25, pady=1, sticky=W)
        self.entryadmiss = DateEntry(self, selectmode='day', year=self.hoje.year, month=self.hoje.month,
                                     day=self.hoje.day, locale='pt_BR')
        self.entryadmiss.grid(column=1, row=12, padx=125, pady=1, sticky=W)
        # aparecer horario preenchido e dropdown para escolher horario
        self.labelhor = ttk.Label(self, width=55, text="Horário preenchido: ")
        self.labelhor.grid(column=1, row=14, padx=25, pady=1, sticky=W)
        self.combohor = ttk.Combobox(self, values=horarios, textvariable=self.horario, width=50)
        self.combohor.grid(column=1, row=15, padx=125, pady=1, sticky=W)
        # aparecer entry para preencher salario
        self.labelsal = ttk.Label(self, width=20, text="Salário:")
        self.labelsal.grid(column=1, row=16, padx=25, pady=1, sticky=W)
        self.entrysal = ttk.Entry(self, width=20)
        self.entrysal.grid(column=1, row=16, padx=125, pady=1, sticky=W)
        # aparecer dropdown para escolher cargo
        self.labelcargo = ttk.Label(self, width=20, text="Cargo")
        self.labelcargo.grid(column=1, row=18, padx=25, pady=1, sticky=W)
        self.combocargo = ttk.Combobox(self, values=cargos, textvariable=self.cargo, width=50)
        self.combocargo.grid(column=1, row=18, padx=125, pady=1, sticky=W)
        # aparecer dropdown para escolher depto
        self.labeldepto = ttk.Label(self, width=20, text="Departamento:")
        self.labeldepto.grid(column=1, row=19, padx=25, pady=1, sticky=W)
        self.combodepto = ttk.Combobox(self, values=departamentos, textvariable=self.departamento, width=50)
        self.combodepto.grid(column=1, row=19, padx=125, pady=1, sticky=W)
        # aparecer dropdown para escolher tipo_contr
        self.labelcontr = ttk.Label(self, width=20, text="Tipo de contrato:")
        self.labelcontr.grid(column=1, row=21, padx=25, pady=1, sticky=W)
        self.combocontr = ttk.Combobox(self, values=tipodecontrato, textvariable=self.tipocontr, width=50)
        self.combocontr.grid(column=1, row=21, padx=125, pady=1, sticky=W)
        self.hrs = StringVar()
        self.hrm = StringVar()
        # aparecer entry para preencher hrsem
        self.labelhrsem = ttk.Label(self, width=20, text="Hrs Sem.:")
        self.labelhrsem.grid(column=1, row=24, padx=25, pady=1, sticky=W)
        self.entryhrsem = ttk.Entry(self, width=20, textvariable=self.hrs)
        self.entryhrsem.grid(column=1, row=24, padx=125, pady=1, sticky=W)
        # aparecer entry para preencher hrmens
        self.labelhrmen = ttk.Label(self, width=20, text="Hrs Mens.:")
        self.labelhrmen.grid(column=1, row=25, padx=25, pady=1, sticky=W)
        self.entryhrmen = ttk.Entry(self, width=20, textvariable=self.hrm)
        self.entryhrmen.grid(column=1, row=25, padx=125, pady=1, sticky=W)
        self.agencia = StringVar()
        self.conta = StringVar()
        self.digito = StringVar()
        # aparecer entry para agencia
        self.labelag = ttk.Label(self, width=20, text="Agência:")
        self.labelag.grid(column=1, row=24, padx=260, pady=1, sticky=W)
        self.entryag = ttk.Entry(self, width=20, textvariable=self.agencia)
        self.entryag.grid(column=1, row=24, padx=320, pady=1, sticky=W)
        # aparecer entry para conta
        self.labelcc = ttk.Label(self, width=20, text="Conta:")
        self.labelcc.grid(column=1, row=25, padx=260, pady=1, sticky=W)
        self.entrycc = ttk.Entry(self, width=20, textvariable=self.conta)
        self.entrycc.grid(column=1, row=25, padx=320, pady=1, sticky=W)
        # aparecer entry para ditigo
        self.labeldig = ttk.Label(self, width=20, text="Dígito:")
        self.labeldig.grid(column=1, row=26, padx=260, pady=1, sticky=W)
        self.entrydig = ttk.Entry(self, width=20, textvariable=self.digito)
        self.entrydig.grid(column=1, row=26, padx=320, pady=1, sticky=W)
        self.edicao = IntVar()
        self.editar = ttk.Checkbutton(self, text='Editar cadastro feito manualmente.', variable=self.edicao)
        self.editar.grid(column=1, row=26, padx=26, pady=1, sticky=W)
        self.feitonde = IntVar()
        self.onde = ttk.Checkbutton(self, text='Cadastro realizado fora da Cia.', variable=self.feitonde)
        self.onde.grid(column=1, row=27, padx=26, pady=1, sticky=W)

        def mostrarhorario(event):
            nome = event.widget.get()
            num, name = nome.split(' - ')
            linha = int(num)
            planwb = l_w(self.caminho.get())
            plansh = planwb['Respostas ao formulário 1']
            self.labelhor.config(text='Horário preenchido: ' + str(plansh[f'AI{linha}'].value))

        self.combonome.bind("<<ComboboxSelected>>", mostrarhorario)

        def carregarfunc(local):
            planwb = l_w(local)
            plansh = planwb['Respostas ao formulário 1']
            lista = []
            for x, pessoa in enumerate(plansh):
                lista.append(f'{x + 1} - {pessoa[2].value}')
            self.combonome.config(values=lista)

        self.botaocarregar = ttk.Button(self, text="Carregar planilha",
                                        command=lambda: [carregarfunc(self.caminho.get())])
        self.botaocarregar.grid(column=1, row=9, padx=350, pady=25, sticky=W)
        self.botaocadastrar = ttk.Button(self, width=20, text="Cadastrar no Dexion",
                                         command=lambda: [cadastro_funcionario(self.caminho.get(), self.edicao.get(),
                                                                               self.feitonde.get(),
                                                                               self.combonome.get(),
                                                                               self.entrymatr.get(),
                                                                               self.entryadmiss.get(),
                                                                               self.combohor.get(),
                                                                               self.entrysal.get(),
                                                                               self.combocargo.get(),
                                                                               self.combodepto.get(),
                                                                               self.combocontr.get(),
                                                                               self.hrs.get(),
                                                                               self.hrm.get(),
                                                                               self.agencia.get(),
                                                                               self.conta.get(),
                                                                               self.digito.get())])
        self.botaocadastrar.grid(column=1, row=28, padx=520, pady=1, sticky=W)
        self.botaosalvar = ttk.Button(self, width=20, text="Salvar Docs",
                                      command=lambda: [salvadocsfunc(self.entrymatr.get())])
        self.botaosalvar.grid(column=1, row=29, padx=520, pady=1, sticky=W)
        self.botaoenviaemail = ttk.Button(self, width=20, text="Enviar e-mails",
                                          command=lambda: [enviaemailsfunc(self.entrymatr.get())])
        self.botaoenviaemail.grid(column=1, row=30, padx=520, pady=1, sticky=W)

    def selecionarfunc(self):
        try:
            caminhoplan = tkinter.filedialog.askopenfilename(title='Planilha Funcionários')
            self.caminho.set(str(caminhoplan))
        except ValueError:
            pass


def cadastro_funcionario(caminho='', editar=0, ondestou=0, nome='', matricula='', admissao='',
                         horario='', salario='', cargo='', depto='', tipo_contr='',
                         hrsem='', hrmens='', agencia='', conta='', digito=''):
    sessions = sessionmaker(bind=engine)
    session = sessions()
    if caminho == '' or nome == '' or matricula == '' or admissao == '' or horario == '' or salario == '' or \
            cargo == '' or depto == '' or tipo_contr == '' or hrsem == '' or hrmens == '':
        tkinter.messagebox.showinfo(
            title='Erro de preenchimento',
            message='Preencha todos os campos antes de cadastrar o funcionário!'
        )
    else:
        wb = l_w(caminho, read_only=False)
        sh = wb['Respostas ao formulário 1']
        num, name = nome.strip().split(' - ')
        linha = int(num)

        # search for the highest compatibility between the city filled in the form and the cities in the lists to
        # define codmunnas value
        est = str(sh[f'AJ{linha}'].value)
        cidade = str(sh[f'L{linha}'].value).title()
        lista = []
        dicion = {}
        for cid in municipios[est]:
            dicion[SequenceMatcher(None, cidade, cid).ratio()] = cid
            lista.append(SequenceMatcher(None, cidade, cid).ratio())
        codmunnas = municipios[str(sh[f'AJ{linha}'].value).upper().strip()][dicion[max(lista)]]

        # search for the highest compatibility between the city filled in the form and the cities in the lists to
        # define codmunend value
        est = str(sh[f'T{linha}'].value)
        cidade = str(sh[f'S{linha}'].value).title()
        listaend = []
        dicionend = {}
        for cid in municipios[est]:
            dicionend[SequenceMatcher(None, cidade, cid).ratio()] = cid
            listaend.append(SequenceMatcher(None, cidade, cid).ratio())
        codmunend = municipios[str(sh[f'T{linha}'].value).upper().strip()][dicionend[max(listaend)]]

        lotacao = {'UNIDADE PARK SUL - QUALQUER DEPARTAMENTO': '0013', 'KIDS': '0010', 'MUSCULAÇÃO': '0007',
                   'ESPORTES E LUTAS': '0008', 'CROSSFIT': '0012', 'GINÁSTICA': '0006', 'GESTANTES': '0008',
                   'RECEPÇÃO': '0003',
                   'FINANCEIRO': '0001', 'TI': '0001', 'MARKETING': '0001', 'MANUTENÇÃO': '0004'}
        if editar == 0:
            if ondestou == 0:
                # Cadastro iniciado na Cia
                if linha:
                    pess = Colaborador(matricula=matricula, nome=name.upper(), admiss=admissao,
                                       nascimento=str(sh[f'D{linha}'].value),
                                       pis=str(int(sh[f'Y{linha}'].value)).zfill(11),
                                       cpf=str(int(sh[f'V{linha}'].value)).zfill(11),
                                       rg=str(int(sh[f'W{linha}'].value)),
                                       emissor=str(sh[f'X{linha}'].value), email=str(sh[f'B{linha}'].value),
                                       genero=str(sh[f'E{linha}'].value),
                                       estado_civil=str(sh[f'F{linha}'].value), cor=str(sh[f'G{linha}'].value),
                                       instru=str(sh[f'J{linha}'].value),
                                       nacional=str(sh[f'K{linha}'].value),
                                       cod_municipionas=codmunnas,
                                       cid_nas=str(sh[f'L{linha}'].value), uf_nas=str(sh[f'AJ{linha}'].value),
                                       pai=str(sh[f'M{linha}'].value).upper(),
                                       mae=str(sh[f'N{linha}'].value).upper(), endereco=str(sh[f'O{linha}'].value),
                                       num=str(int(sh[f'P{linha}'].value)),
                                       bairro=str(sh[f'Q{linha}'].value), cep=str(int(sh[f'R{linha}'].value)),
                                       cidade=str(sh[f'S{linha}'].value),
                                       uf=str(sh[f'T{linha}'].value),
                                       cod_municipioend=codmunend,
                                       tel=str(int(sh[f'U{linha}'].value)),
                                       tit_eleit=str(sh[f'Z{linha}'].value), zona_eleit=str(sh[f'AA{linha}'].value),
                                       sec_eleit=str(sh[f'AB{linha}'].value),
                                       ctps=str(int(sh[f'AC{linha}'].value)), serie_ctps=str(sh[f'AD{linha}'].value),
                                       uf_ctps=str(sh[f'AE{linha}'].value),
                                       emiss_ctps=str(sh[f'AF{linha}'].value), depto=depto,
                                       cargo=cargo,
                                       horario=horario, salario=salario, tipo_contr=tipo_contr, hr_sem=hrsem,
                                       hr_mens=hrmens,
                                       ag=agencia, conta=conta, cdigito=digito
                                       )
                    session.add(pess)
                    session.commit()
                    pessoa = session.query(Colaborador).filter_by(matricula=matricula).first()
                    pa.click(pa.center(pa.locateOnScreen('./static/Dexion.png')))
                    pa.press('alt'), pa.press('a'), pa.press('t'), t.sleep(2), pa.press(
                        'i'), t.sleep(5), pa.write(str(pessoa.matricula)), pa.press('enter'), t.sleep(20)
                    pp.copy(pessoa.nome), pa.hotkey('ctrl', 'v'), pa.press('tab'), pa.write(pessoa.cpf)
                    t.sleep(1), pa.press('tab', 3), pa.write(pessoa.genero), pa.press('tab'), pa.write(pessoa.cor)
                    t.sleep(1), pa.press('tab', 2), pa.write(pessoa.instru)
                    t.sleep(1), pa.press('tab'), pa.write(pessoa.estado_civil)
                    if str(pessoa.estado_civil) == '2 - Casado(a)':
                        pa.press('tab', 6)
                    else:
                        pa.press('tab', 5)
                    pa.write(datetime.strftime(datetime.strptime(pessoa.nascimento, '%Y-%m-%d %H:%M:%S'), '%d%m%Y'))
                    t.sleep(1), pa.press('tab'), pp.copy(pessoa.cid_nas), pa.hotkey('ctrl', 'v'), pa.press('tab')
                    t.sleep(1), pa.write(pessoa.uf_nas), pa.press('tab'), pa.write(pessoa.cod_municipionas), pa.press(
                        'tab')
                    t.sleep(1), pp.copy(pessoa.pai), pa.hotkey('ctrl', 'v'), pa.press('tab'), pa.write('105'), pa.press(
                        'tab')
                    t.sleep(1), pp.copy(pessoa.mae), pa.hotkey('ctrl', 'v'), pa.press('tab'), pa.write('105')
                    # #clique em documentos
                    pa.click(pa.center(pa.locateOnScreen('./static/Documentos.png')))
                    pa.press('tab'), pa.write(str(pessoa.rg)), pa.press('tab'), pa.write(
                        pessoa.emissor), pa.press('tab', 3), pa.write(pessoa.cod_municipioend),
                    pa.press('tab'), pa.write(pessoa.pis), pa.press('enter')
                    pa.press('tab'), pa.write(pessoa.tit_eleit), pa.press('tab'), pa.write(pessoa.zona_eleit), pa.press(
                        'tab'), pa.write(
                        pessoa.sec_eleit), pa.press('tab'), pa.write(pessoa.ctps)
                    pa.press('tab'), pa.write(pessoa.serie_ctps), pa.press('tab'), pa.write(pessoa.uf_ctps), pa.press(
                        'tab'), pa.write(datetime.strftime(datetime.strptime(pessoa.emiss_ctps, '%Y-%m-%d %H:%M:%S'),
                                                           '%d%m%Y'))
                    pa.press('tab', 5), pa.write('3'), pa.press('tab'), pa.write('341'), pa.press('tab')
                    pa.write(pessoa.ag), pa.press('tab'), pa.write(f'{pessoa.conta}-{pessoa.cdigito}'), pa.press('tab')

                    # #clique em endereço
                    pa.click(pa.center(pa.locateOnScreen('./static/Endereco.png')))
                    pa.press('tab', 2), pp.copy(pessoa.endereco), pa.hotkey('ctrl', 'v'), pa.press(
                        'tab'), pa.write(pessoa.num), pa.press('tab', 2)
                    pp.copy(pessoa.bairro), pa.hotkey('ctrl', 'v'), pa.press('tab'), pp.copy(pessoa.cidade), pa.hotkey(
                        'ctrl', 'v'), pa.press('tab'), pa.write(pessoa.uf)
                    pa.press('tab'), pa.write(pessoa.cep), pa.press('tab'), pa.write(pessoa.cod_municipioend),pa.press(
                        'tab'), pa.write(str(pessoa.tel)), pa.press('tab', 2), pa.write(pessoa.email)
                    # #clique em dados contratuais
                    pa.click(pa.center(pa.locateOnScreen('./static/Contratuais.png')))
                    pa.press('tab'), pa.write(str(pessoa.admiss).replace('/', '')), t.sleep(1)
                    pa.press('tab', 10), pa.write('2')
                    # #clique em Contrato de Experiência
                    try:
                        pa.click(pa.center(pa.locateOnScreen('./static/Experiencia.png')))
                    except:
                        t.sleep(5)
                        pa.click(pa.center(pa.locateOnScreen('./static/Experiencia.png')))
                    pa.press('tab'), pa.write('45'), pa.press('tab'), pa.write('45'), pa.press(
                        'tab'), pa.press(
                        'space'), pa.press('tab', 2), pa.write('003')
                    pa.press('tab'), pa.write(str(pessoa.matricula))
                    # #clique em Outros
                    try:
                        pa.click(pa.center(pa.locateOnScreen('./static/Outros.png')))
                    except:
                        t.sleep(5)
                        pa.click(pa.center(pa.locateOnScreen('./static/Outros.png')))
                    t.sleep(2), pa.write('CARGO GERAL')
                    # #clique em lupa de descrição de cargos
                    pa.click(pa.center(pa.locateOnScreen('./static/Lupa.png')))
                    t.sleep(2), pp.copy(pessoa.cargo), pa.hotkey('ctrl', 'v'), t.sleep(1.5), pa.press('enter', 2)
                    t.sleep(1), pa.press('tab'), pa.write(pessoa.salario), pa.press('tab')
                    if str(pessoa.tipo_contr) == 'Horista':
                        pa.press('1')
                    else:
                        pa.press('5')
                    pa.press('tab', 2), t.sleep(1), pa.write(str(pessoa.hr_sem)), pa.press('tab'), t.sleep(1)
                    pa.write(str(pessoa.hr_mens))
                    pa.press('tab', 5), pa.write('00395419000190'), pa.press('tab', 2), pa.write('5')
                    # #clique em eventos trabalhistas
                    pa.click(pa.center(pa.locateOnScreen('./static/EVTrab.png')))
                    t.sleep(1)
                    # #clique em lotação
                    pa.click(pa.center(pa.locateOnScreen('./static/Lotacoes.png')))
                    pa.press('tab'), pa.press('tab'), pa.write('i'), pa.write(str(pessoa.admiss).replace('/', ''))
                    t.sleep(1), pa.press('enter'), t.sleep(1)
                    pp.copy(lotacao[f'{pessoa.depto}']), pa.hotkey('ctrl', 'v'), pa.press('enter'), pa.write('f')
                    pa.press('tab'), pa.write('4')
                    pa.press('tab', 6), pa.write('i'), t.sleep(2), pa.press('tab'), pa.write(pessoa.horario)
                    t.sleep(3), pa.press('tab', 3), pa.press('enter'), t.sleep(3)
                    # #clique em cancelar novo registro de horario
                    pa.click(pa.center(pa.locateOnScreen('./static/Cancelarhor.png'))), t.sleep(2.5)
                    # #clique em salvar lotação
                    pa.click(pa.center(pa.locateOnScreen('./static/Salvarlot.png'))), t.sleep(1)
                    # #clique em fechar lotação
                    try:
                        pa.click(pa.center(pa.locateOnScreen('./static/Fecharlot.png'))), t.sleep(1)
                    except:
                        t.sleep(4)
                        pa.click(pa.center(pa.locateOnScreen('./static/Fecharlot.png')))
                    # #clique em Compatibilidade
                    pa.click(pa.center(pa.locateOnScreen('./static/Compatibilidade.png'))), t.sleep(1)
                    # #clique em Compatibilidade de novo
                    pa.click(pa.center(pa.locateOnScreen('./static/Compatibilidade2.png'))), t.sleep(1)
                    # #clique em CAGED
                    pa.click(pa.center(pa.locateOnScreen('./static/CAGED.png')))
                    pa.press('tab', 2), pa.write('20'), pa.press('tab'), pa.write('1'), t.sleep(0.5)
                    # #clique em RAIS
                    pa.click(pa.center(pa.locateOnScreen('./static/RAIS.png')))
                    pp.copy(pessoa.nome), pa.hotkey('ctrl', 'v'), pa.press('tab', 2), pa.write('2')
                    pa.press('tab'), pa.write('10')
                    # #clique em Salvar
                    pa.click(pa.center(pa.locateOnScreen('./static/Salvarcadastro.png'))), t.sleep(10)
                    # #clique em fechar novo cadastro
                    pa.click(pa.center(pa.locateOnScreen('./static/Fecharnovo1.png'))), t.sleep(2)
                    # #clique em fechar trabalhadores
                    pa.click(pa.center(pa.locateOnScreen('./static/Fechartrab1.png'))), t.sleep(0.5)
                    os.makedirs(
                        r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e '
                        r'Férias\000 - Pastas Funcionais\00 - ATIVOS\{}'.format(pessoa.nome))
                    os.makedirs(
                        r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e '
                        r'Férias\000 - Pastas Funcionais\00 - ATIVOS\{}\Atestados'.format(pessoa.nome))
                    os.makedirs(
                        r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e '
                        r'Férias\000 - Pastas Funcionais\00 - ATIVOS\{}\Diversos'.format(pessoa.nome))
                    os.makedirs(
                        r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e '
                        r'Férias\000 - Pastas Funcionais\00 - ATIVOS\{}\Contratuais'.format(pessoa.nome))
                    os.makedirs(
                        r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e '
                        r'Férias\000 - Pastas Funcionais\00 - ATIVOS\{}\Férias'.format(pessoa.nome))
                    os.makedirs(
                        r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e '
                        r'Férias\000 - Pastas Funcionais\00 - ATIVOS\{}\Pontos'.format(pessoa.nome))
                    os.makedirs(
                        r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e '
                        r'Férias\000 - Pastas Funcionais\00 - ATIVOS\{}\Recibos'.format(pessoa.nome))
                    os.makedirs(
                        r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e '
                        r'Férias\000 - Pastas Funcionais\00 - ATIVOS\{}\Rescisão'.format(pessoa.nome))
                    tkinter.messagebox.showinfo(
                        title='Cadastro ok!',
                        message='Cadastro realizado com sucesso!'
                    )
            else:
                # Cadastro iniciado em casa
                wb = l_w(caminho, read_only=False)
                sh = wb['Respostas ao formulário 1']
                num, name = nome.strip().split(' - ')
                linha = int(num)
                if linha:
                    pess = Colaborador(matricula=matricula, nome=name.upper(), admiss=admissao,
                                       nascimento=str(sh[f'D{linha}'].value),
                                       pis=str(int(sh[f'Y{linha}'].value)).zfill(11),
                                       cpf=str(int(sh[f'V{linha}'].value)).zfill(11),
                                       rg=str(int(sh[f'W{linha}'].value)),
                                       emissor=str(sh[f'X{linha}'].value), email=str(sh[f'B{linha}'].value),
                                       genero=str(sh[f'E{linha}'].value),
                                       estado_civil=str(sh[f'F{linha}'].value), cor=str(sh[f'G{linha}'].value),
                                       instru=str(sh[f'J{linha}'].value),
                                       nacional=str(sh[f'K{linha}'].value),
                                       cod_municipionas=codmunnas,
                                       cid_nas=str(sh[f'L{linha}'].value), uf_nas=str(sh[f'AJ{linha}'].value),
                                       pai=str(sh[f'M{linha}'].value).upper(),
                                       mae=str(sh[f'N{linha}'].value).upper(), endereco=str(sh[f'O{linha}'].value),
                                       num=str(int(sh[f'P{linha}'].value)),
                                       bairro=str(sh[f'Q{linha}'].value), cep=str(int(sh[f'R{linha}'].value)),
                                       cidade=str(sh[f'S{linha}'].value),
                                       uf=str(sh[f'T{linha}'].value),
                                       cod_municipioend=codmunend,
                                       tel=str(int(sh[f'U{linha}'].value)),
                                       tit_eleit=str(sh[f'Z{linha}'].value), zona_eleit=str(sh[f'AA{linha}'].value),
                                       sec_eleit=str(sh[f'AB{linha}'].value),
                                       ctps=str(int(sh[f'AC{linha}'].value)), serie_ctps=str(sh[f'AD{linha}'].value),
                                       uf_ctps=str(sh[f'AE{linha}'].value),
                                       emiss_ctps=str(sh[f'AF{linha}'].value), depto=depto,
                                       cargo=cargo,
                                       horario=horario, salario=salario, tipo_contr=tipo_contr, hr_sem=hrsem,
                                       hr_mens=hrmens,
                                       ag=agencia, conta=conta, cdigito=digito
                                       )
                    session.add(pess)
                    session.commit()
                    pessoa = session.query(Colaborador).filter_by(matricula=matricula).first()
                    pa.click(pa.center(pa.locateOnScreen('./static/Dexion.png')))
                    pa.press('alt'), pa.press('a'), pa.press('t'), t.sleep(2), pa.press(
                        'i'), t.sleep(5), pa.write(str(pessoa.matricula)), pa.press('enter'), t.sleep(60)
                    pp.copy(pessoa.nome), pa.hotkey('ctrl', 'v'), pa.press('tab'), pa.write(pessoa.cpf)
                    t.sleep(1), pa.press('tab', 3), pa.write(pessoa.genero), pa.press('tab'), pa.write(pessoa.cor)
                    t.sleep(1), pa.press('tab', 2), pa.write(pessoa.instru)
                    t.sleep(1), pa.press('tab'), pa.write(pessoa.estado_civil)
                    if str(pessoa.estado_civil) == '2 - Casado(a)':
                        pa.press('tab', 6)
                    else:
                        pa.press('tab', 5)
                    pa.write(datetime.strftime(datetime.strptime(pessoa.nascimento, '%Y-%m-%d %H:%M:%S'), '%d%m%Y'))
                    t.sleep(1), pa.press('tab'), pp.copy(pessoa.cid_nas), pa.hotkey('ctrl', 'v'), pa.press('tab')
                    t.sleep(1), pa.write(pessoa.uf_nas), pa.press('tab'), pa.write(pessoa.cod_municipionas), pa.press(
                        'tab')
                    t.sleep(1), pp.copy(pessoa.pai), pa.hotkey('ctrl', 'v'), pa.press('tab'), pa.write('105'), pa.press(
                        'tab')
                    t.sleep(1), pp.copy(pessoa.mae), pa.hotkey('ctrl', 'v'), pa.press('tab'), pa.write('105')
                    # #clique em documentos
                    pa.click(pa.center(pa.locateOnScreen('./static/Documentos.png')))
                    pa.press('tab'), pa.write(str(pessoa.rg)), pa.press('tab'), pa.write(
                        pessoa.emissor), pa.press('tab', 3), pa.write(pessoa.cod_municipioend),
                    pa.press('tab'), pa.write(pessoa.pis), pa.press('enter')
                    pa.press('tab'), pa.write(pessoa.tit_eleit), pa.press('tab'), pa.write(pessoa.zona_eleit), pa.press(
                        'tab'), pa.write(
                        pessoa.sec_eleit), pa.press('tab'), pa.write(pessoa.ctps)
                    pa.press('tab'), pa.write(pessoa.serie_ctps), pa.press('tab'), pa.write(pessoa.uf_ctps), pa.press(
                        'tab'), pa.write(datetime.strftime(datetime.strptime(pessoa.emiss_ctps, '%Y-%m-%d %H:%M:%S'),
                                                           '%d%m%Y'))
                    pa.press('tab', 5), pa.write('3'), pa.press('tab'), pa.write('341'), pa.press('tab')
                    pa.write(pessoa.ag), pa.press('tab'), pa.write(f'{pessoa.conta}-{pessoa.cdigito}'), pa.press('tab')
                    # #clique em endereço
                    pa.click(pa.center(pa.locateOnScreen('./static/Endereco.png')))
                    pa.press('tab', 2), pp.copy(pessoa.endereco), pa.hotkey('ctrl', 'v'), pa.press(
                        'tab'), pa.write(pessoa.num), pa.press('tab', 2)
                    pp.copy(pessoa.bairro), pa.hotkey('ctrl', 'v'), pa.press('tab'), pp.copy(pessoa.cidade), pa.hotkey(
                        'ctrl', 'v'), pa.press('tab'), pa.write(pessoa.uf)
                    pa.press('tab'), pa.write(pessoa.cep), pa.press('tab'), pa.write(pessoa.cod_municipioend), pa.press(
                        'tab'), pa.write(str(pessoa.tel)), pa.press('tab', 2), pa.write(pessoa.email)
                    # #clique em dados contratuais
                    pa.click(pa.center(pa.locateOnScreen('./static/Contratuais.png')))
                    pa.press('tab'), pa.write(str(pessoa.admiss).replace('/', '')), t.sleep(1)
                    pa.press('tab', 10), pa.write('2')
                    # #clique em Contrato de Experiência
                    try:
                        pa.click(pa.center(pa.locateOnScreen('./static/Experiencia.png')))
                    except:
                        t.sleep(5)
                        pa.click(pa.center(pa.locateOnScreen('./static/Experiencia.png')))
                    pa.press('tab'), pa.write('45'), pa.press('tab'), pa.write('45'), pa.press(
                        'tab'), pa.press(
                        'space'), pa.press('tab', 2), pa.write('003')
                    pa.press('tab'), pa.write(str(pessoa.matricula))
                    # #clique em Outros
                    try:
                        pa.click(pa.center(pa.locateOnScreen('./static/Outros.png')))
                    except:
                        t.sleep(5)
                        pa.click(pa.center(pa.locateOnScreen('./static/Outros.png')))
                    t.sleep(2), pa.write('CARGO GERAL')
                    # #clique em lupa de descrição de cargos
                    pa.click(pa.center(pa.locateOnScreen('./static/Lupa.png')))
                    t.sleep(2), pp.copy(pessoa.cargo), pa.hotkey('ctrl', 'v'), t.sleep(1.5), pa.press('enter', 2)
                    t.sleep(1), pa.press('tab'), pa.write(pessoa.salario), pa.press('tab')
                    if str(pessoa.tipo_contr) == 'Horista':
                        pa.press('1')
                    else:
                        pa.press('5')
                    pa.press('tab', 2), t.sleep(1), pa.write(str(pessoa.hr_sem)), pa.press('tab'), t.sleep(1)
                    pa.write(str(pessoa.hr_mens))
                    pa.press('tab', 5), pa.write('00395419000190'), pa.press('tab', 2), pa.write('5')
                    # #clique em eventos trabalhistas
                    pa.click(pa.center(pa.locateOnScreen('./static/EVTrab.png')))
                    t.sleep(1)
                    # #clique em lotação
                    pa.click(pa.center(pa.locateOnScreen('./static/Lotacoes.png')))
                    pa.press('tab'), pa.press('tab'), pa.write('i'), pa.write(str(pessoa.admiss).replace('/', ''))
                    t.sleep(1), pa.press('enter'), t.sleep(1)
                    pp.copy(lotacao[f'{pessoa.depto}']), pa.hotkey('ctrl', 'v'), pa.press('enter'), pa.write('f')
                    pa.press('tab'), pa.write('4')
                    pa.press('tab', 6), pa.write('i'), t.sleep(2), pa.press('tab'), pa.write(pessoa.horario)
                    t.sleep(3), pa.press('tab', 3), pa.press('enter'), t.sleep(3)
                    # #clique em cancelar novo registro de horario
                    pa.click(pa.center(pa.locateOnScreen('./static/Cancelarhor.png'))), t.sleep(2.5)
                    # #clique em salvar lotação
                    pa.click(pa.center(pa.locateOnScreen('./static/Salvarlot.png'))), t.sleep(1)
                    # #clique em fechar lotação
                    try:
                        pa.click(pa.center(pa.locateOnScreen('./static/Fecharlot.png'))), t.sleep(1)
                    except:
                        t.sleep(4)
                        pa.click(pa.center(pa.locateOnScreen('./static/Fecharlot.png')))
                    # #clique em Compatibilidade
                    pa.click(pa.center(pa.locateOnScreen('./static/Compatibilidade.png'))), t.sleep(1)
                    # #clique em Compatibilidade de novo
                    pa.click(pa.center(pa.locateOnScreen('./static/Compatibilidade2.png'))), t.sleep(1)
                    # #clique em CAGED
                    pa.click(pa.center(pa.locateOnScreen('./static/CAGED.png')))
                    pa.press('tab', 2), pa.write('20'), pa.press('tab'), pa.write('1'), t.sleep(0.5)
                    # #clique em RAIS
                    pa.click(pa.center(pa.locateOnScreen('./static/RAIS.png')))
                    pp.copy(pessoa.nome), pa.hotkey('ctrl', 'v'), pa.press('tab', 2), pa.write('2')
                    pa.press('tab'), pa.write('10')
                    # #clique em Salvar
                    pa.click(pa.center(pa.locateOnScreen('./static/Salvarcadastro.png'))), t.sleep(10)
                    # #clique em fechar novo cadastro
                    pa.click(pa.center(pa.locateOnScreen('./static/Fecharnovo1.png'))), t.sleep(2)
                    # #clique em fechar trabalhadores
                    pa.click(pa.center(pa.locateOnScreen('./static/Fechartrab1.png'))), t.sleep(0.5)

                    os.makedirs(
                        r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e '
                        r'Férias\000 - Pastas Funcionais\00 - ATIVOS\{}'.format(pessoa.nome))
                    os.makedirs(
                        r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e '
                        r'Férias\000 - Pastas Funcionais\00 - ATIVOS\{}\Atestados'.format(pessoa.nome))
                    os.makedirs(
                        r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e '
                        r'Férias\000 - Pastas Funcionais\00 - ATIVOS\{}\Diversos'.format(pessoa.nome))
                    os.makedirs(
                        r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e '
                        r'Férias\000 - Pastas Funcionais\00 - ATIVOS\{}\Contratuais'.format(pessoa.nome))
                    os.makedirs(
                        r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e '
                        r'Férias\000 - Pastas Funcionais\00 - ATIVOS\{}\Férias'.format(pessoa.nome))
                    os.makedirs(
                        r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e '
                        r'Férias\000 - Pastas Funcionais\00 - ATIVOS\{}\Pontos'.format(pessoa.nome))
                    os.makedirs(
                        r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e '
                        r'Férias\000 - Pastas Funcionais\00 - ATIVOS\{}\Recibos'.format(pessoa.nome))
                    os.makedirs(
                        r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e '
                        r'Férias\000 - Pastas Funcionais\00 - ATIVOS\{}\Rescisão'.format(pessoa.nome))
                    tkinter.messagebox.showinfo(
                        title='Cadastro ok!',
                        message='Cadastro realizado com sucesso!'
                    )

        else:
            if ondestou == 0:
                # Cadastro EDITADO na Cia
                num, name = nome.strip().split(' - ')
                linha = int(num)
                if linha:
                    pessoa = session.query(Colaborador).filter_by(matricula=matricula).first()
                    pa.click(pa.center(pa.locateOnScreen('./static/Dexion.png')))
                    pa.press('alt'), pa.press('a'), pa.press('t'), t.sleep(2), pa.press(
                        'a'), t.sleep(5), pa.write(str(pessoa.matricula)), pa.press('enter'), t.sleep(15)
                    pp.copy(pessoa.nome), pa.hotkey('ctrl', 'v'), pa.press('tab'), pa.write(pessoa.cpf)
                    pa.press('tab', 3), pa.write(pessoa.genero), pa.press('tab'), pa.write(pessoa.cor)
                    pa.press('tab', 2), pa.write(pessoa.instru)
                    pa.press('tab'), pa.write(pessoa.estado_civil)
                    if str(pessoa.estado_civil) == '2 - Casado(a)':
                        pa.press('tab', 6)
                    else:
                        pa.press('tab', 5)
                    pa.write(datetime.strftime(datetime.strptime(pessoa.nascimento, '%Y-%m-%d %H:%M:%S'), '%d%m%Y'))
                    pa.press('tab'), pp.copy(pessoa.cid_nas), pa.hotkey('ctrl', 'v'), pa.press('tab')
                    pa.write(pessoa.uf_nas), pa.press('tab'), pa.write(pessoa.cod_municipionas), pa.press('tab')
                    pp.copy(pessoa.pai), pa.hotkey('ctrl', 'v'), pa.press('tab'), pa.write('105'), pa.press('tab')
                    pp.copy(pessoa.mae), pa.hotkey('ctrl', 'v'), pa.press('tab'), pa.write('105')
                    # #clique em documentos
                    pa.click(pa.center(pa.locateOnScreen('./static/Documentos.png')))
                    pa.press('tab'), pa.write(str(pessoa.rg)), pa.press('tab'), pa.write(
                        pessoa.emissor), pa.press('tab', 3), pa.write(pessoa.cod_municipioend),
                    pa.press('tab'), pa.write(pessoa.pis), pa.press('enter')
                    pa.press('tab'), pa.write(pessoa.tit_eleit), pa.press('tab'), pa.write(pessoa.zona_eleit), pa.press(
                        'tab'), pa.write(
                        pessoa.sec_eleit), pa.press('tab'), pa.write(pessoa.ctps)
                    pa.press('tab'), pa.write(pessoa.serie_ctps), pa.press('tab'), pa.write(pessoa.uf_ctps), pa.press(
                        'tab'), pa.write(datetime.strftime(datetime.strptime(pessoa.emiss_ctps, '%Y-%m-%d %H:%M:%S'),
                                                           '%d%m%Y'))
                    pa.press('tab', 5), pa.write('3'), pa.press('tab'), pa.write('341'), pa.press('tab')
                    pa.write(pessoa.ag), pa.press('tab'), pa.write(f'{pessoa.conta}-{pessoa.cdigito}'), pa.press('tab')

                    # #clique em endereço
                    pa.click(pa.center(pa.locateOnScreen('./static/Endereco.png')))
                    pa.press('tab', 2), pp.copy(pessoa.endereco), pa.hotkey('ctrl', 'v'), pa.press(
                        'tab'), pa.write(pessoa.num), pa.press('tab', 2)
                    pp.copy(pessoa.bairro), pa.hotkey('ctrl', 'v'), pa.press('tab'), pp.copy(pessoa.cidade), pa.hotkey(
                        'ctrl', 'v'), pa.press('tab'), pa.write(pessoa.uf)
                    pa.press('tab'), pa.write(pessoa.cep), pa.press('tab'), pa.write(pessoa.cod_municipioend), pa.press(
                        'tab'), pa.write(str(pessoa.tel)), pa.press('tab', 2), pa.write(pessoa.email)
                    # #clique em dados contratuais
                    pa.click(pa.center(pa.locateOnScreen('./static/Contratuais.png')))
                    pa.press('tab'), pa.write(str(pessoa.admiss).replace('/', '')), t.sleep(1)
                    pa.press('tab', 10), pa.write('2')
                    # #clique em Contrato de Experiência
                    try:
                        pa.click(pa.center(pa.locateOnScreen('./static/Experiencia.png')))
                    except:
                        t.sleep(5)
                        pa.click(pa.center(pa.locateOnScreen('./static/Experiencia.png')))
                    pa.press('tab'), pa.write('45'), pa.press('tab'), pa.write('45'), pa.press(
                        'tab'), pa.press(
                        'space'), pa.press('tab', 2), pa.write('003')
                    pa.press('tab'), pa.write(str(pessoa.matricula))
                    # #clique em Outros
                    try:
                        pa.click(pa.center(pa.locateOnScreen('./static/Outros.png')))
                    except:
                        t.sleep(5)
                        pa.click(pa.center(pa.locateOnScreen('./static/Outros.png')))
                    t.sleep(2), pa.write('CARGO GERAL')
                    # #clique em lupa de descrição de cargos
                    pa.click(pa.center(pa.locateOnScreen('./static/Lupa.png')))
                    t.sleep(2), pp.copy(pessoa.cargo), pa.hotkey('ctrl', 'v'), t.sleep(1.5), pa.press('enter', 2)
                    t.sleep(1), pa.press('tab'), pa.write(pessoa.salario), pa.press('tab')
                    if str(pessoa.tipo_contr) == 'Horista':
                        pa.press('1')
                    else:
                        pa.press('5')
                    pa.press('tab', 2), t.sleep(1), pa.write(str(pessoa.hr_sem)), pa.press('tab'), t.sleep(1)
                    pa.write(str(pessoa.hr_mens))
                    pa.press('tab', 5), pa.write('00395419000190'), pa.press('tab', 2), pa.write('5')
                    # #clique em Compatibilidade
                    pa.click(pa.center(pa.locateOnScreen('./static/Compatibilidade1.png'))), t.sleep(1)
                    # #clique em Compatibilidade de novo
                    pa.click(pa.center(pa.locateOnScreen('./static/Compatibilidade2.png'))), t.sleep(1)
                    # #clique em CAGED
                    pa.click(pa.center(pa.locateOnScreen('./static/CAGED.png')))
                    pa.press('tab', 2), pa.write('20'), pa.press('tab'), pa.write('1'), t.sleep(0.5)
                    # #clique em RAIS
                    pa.click(pa.center(pa.locateOnScreen('./static/RAIS.png')))
                    pp.copy(pessoa.nome), pa.hotkey('ctrl', 'v'), pa.press('tab', 2), pa.write('2')
                    pa.press('tab'), pa.write('10')
                    # #clique em Salvar
                    pa.click(pa.center(pa.locateOnScreen('./static/Salvarcadastro.png'))), t.sleep(10)
                    # #clique em fechar novo cadastro
                    pa.click(pa.center(pa.locateOnScreen('./static/Fecharnovo1.png'))), t.sleep(2)
                    # #clique em fechar trabalhadores
                    pa.click(pa.center(pa.locateOnScreen('./static/Fechartrab1.png'))), t.sleep(0.5)

                    os.makedirs(
                        r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e '
                        r'Férias\000 - Pastas Funcionais\00 - ATIVOS\{}'.format(pessoa.nome))
                    os.makedirs(
                        r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e '
                        r'Férias\000 - Pastas Funcionais\00 - ATIVOS\{}\Atestados'.format(pessoa.nome))
                    os.makedirs(
                        r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e '
                        r'Férias\000 - Pastas Funcionais\00 - ATIVOS\{}\Diversos'.format(pessoa.nome))
                    os.makedirs(
                        r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e '
                        r'Férias\000 - Pastas Funcionais\00 - ATIVOS\{}\Contratuais'.format(pessoa.nome))
                    os.makedirs(
                        r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e '
                        r'Férias\000 - Pastas Funcionais\00 - ATIVOS\{}\Férias'.format(pessoa.nome))
                    os.makedirs(
                        r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e '
                        r'Férias\000 - Pastas Funcionais\00 - ATIVOS\{}\Pontos'.format(pessoa.nome))
                    os.makedirs(
                        r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e '
                        r'Férias\000 - Pastas Funcionais\00 - ATIVOS\{}\Recibos'.format(pessoa.nome))
                    os.makedirs(
                        r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e '
                        r'Férias\000 - Pastas Funcionais\00 - ATIVOS\{}\Rescisão'.format(pessoa.nome))
                    tkinter.messagebox.showinfo(title='Cadastro ok!',
                                                message='Cadastro editado com sucesso!')
                else:
                    # Cadastro EDITADO em casa
                    num, name = nome.strip().split(' - ')
                    linha = int(num)
                    if linha:
                        pessoa = session.query(Colaborador).filter_by(matricula=matricula).first()
                        pa.click(pa.center(pa.locateOnScreen('./static/Dexion.png')))
                        pa.press('alt'), pa.press('a'), pa.press('t'), t.sleep(2), pa.press(
                            'a'), t.sleep(5), pa.write(str(pessoa.matricula)), pa.press('enter'), t.sleep(20)
                        pp.copy(pessoa.nome), pa.hotkey('ctrl', 'v'), pa.press('tab'), pa.write(pessoa.cpf)
                        t.sleep(1), pa.press('tab', 3), pa.write(pessoa.genero), pa.press('tab'), pa.write(pessoa.cor)
                        t.sleep(1), pa.press('tab', 2), pa.write(pessoa.instru)
                        t.sleep(1), pa.press('tab'), pa.write(pessoa.estado_civil)
                        if str(pessoa.estado_civil) == '2 - Casado(a)':
                            pa.press('tab', 6)
                        else:
                            pa.press('tab', 5)
                        pa.write(datetime.strftime(datetime.strptime(pessoa.nascimento, '%Y-%m-%d %H:%M:%S'), '%d%m%Y'))
                        t.sleep(1), pa.press('tab'), pp.copy(pessoa.cid_nas), pa.hotkey('ctrl', 'v'), pa.press('tab')
                        t.sleep(1), pa.write(pessoa.uf_nas), pa.press('tab'), pa.write(
                            pessoa.cod_municipionas), pa.press('tab')
                        t.sleep(1), pp.copy(pessoa.pai), pa.hotkey('ctrl', 'v'), pa.press('tab'), pa.write(
                            '105'), pa.press('tab')
                        t.sleep(1), pp.copy(pessoa.mae), pa.hotkey('ctrl', 'v'), pa.press('tab'), pa.write('105')
                        # #clique em documentos
                        pa.click(pa.center(pa.locateOnScreen('./static/Documentos.png')))
                        pa.press('tab'), pa.write(str(pessoa.rg)), pa.press('tab'), pa.write(
                            pessoa.emissor), pa.press('tab', 3), pa.write(pessoa.cod_municipioend),
                        pa.press('tab'), pa.write(pessoa.pis), pa.press('enter')
                        pa.press('tab'), pa.write(pessoa.tit_eleit), pa.press('tab'), pa.write(
                            pessoa.zona_eleit), pa.press('tab'), pa.write(
                            pessoa.sec_eleit), pa.press('tab'), pa.write(pessoa.ctps)
                        pa.press('tab'), pa.write(pessoa.serie_ctps), pa.press('tab'), pa.write(
                            pessoa.uf_ctps), pa.press(
                            'tab'), pa.write(
                            datetime.strftime(datetime.strptime(pessoa.emiss_ctps, '%Y-%m-%d %H:%M:%S'),
                                              '%d%m%Y'))
                        pa.press('tab', 5), pa.write('3'), pa.press('tab'), pa.write('341'), pa.press('tab')
                        pa.write(pessoa.ag), pa.press('tab'), pa.write(f'{pessoa.conta}-{pessoa.cdigito}'), pa.press(
                            'tab')
                        # #clique em endereço
                        pa.click(pa.center(pa.locateOnScreen('./static/Endereco.png')))
                        pa.press('tab', 2), pp.copy(pessoa.endereco), pa.hotkey('ctrl', 'v'), pa.press(
                            'tab'), pa.write(pessoa.num), pa.press('tab', 2)
                        pp.copy(pessoa.bairro), pa.hotkey('ctrl', 'v'), pa.press('tab'), pp.copy(
                            pessoa.cidade), pa.hotkey(
                            'ctrl', 'v'), pa.press('tab'), pa.write(pessoa.uf)
                        pa.press('tab'), pa.write(pessoa.cep), pa.press('tab'), pa.write(
                            pessoa.cod_municipioend), pa.press(
                            'tab'), pa.write(str(pessoa.tel)), pa.press('tab', 2), pa.write(pessoa.email)
                        # #clique em dados contratuais
                        pa.click(pa.center(pa.locateOnScreen('./static/Contratuais.png')))
                        pa.press('tab'), pa.write(str(pessoa.admiss).replace('/', '')), t.sleep(1)
                        pa.press('tab', 10), pa.write('2')
                        # #clique em Contrato de Experiência
                        try:
                            pa.click(pa.center(pa.locateOnScreen('./static/Experiencia.png')))
                        except:
                            t.sleep(5)
                            pa.click(pa.center(pa.locateOnScreen('./static/Experiencia.png')))
                        pa.press('tab'), pa.write('45'), pa.press('tab'), pa.write('45'), pa.press(
                            'tab'), pa.press(
                            'space'), pa.press('tab', 2), pa.write('003')
                        pa.press('tab'), pa.write(str(pessoa.matricula))
                        # #clique em Outros
                        try:
                            pa.click(pa.center(pa.locateOnScreen('./static/Outros.png')))
                        except:
                            t.sleep(5)
                            pa.click(pa.center(pa.locateOnScreen('./static/Outros.png')))
                        t.sleep(2), pa.write('CARGO GERAL')
                        # #clique em lupa de descrição de cargos
                        pa.click(pa.center(pa.locateOnScreen('./static/Lupa.png')))
                        t.sleep(2), pp.copy(pessoa.cargo), pa.hotkey('ctrl', 'v'), t.sleep(1.5), pa.press('enter', 2)
                        t.sleep(1), pa.press('tab'), pa.write(pessoa.salario), pa.press('tab')
                        if str(pessoa.tipo_contr) == 'Horista':
                            pa.press('1')
                        else:
                            pa.press('5')
                        pa.press('tab', 2), t.sleep(1), pa.write(str(pessoa.hr_sem)), pa.press('tab'), t.sleep(1)
                        pa.write(str(pessoa.hr_mens))
                        pa.press('tab', 5), pa.write('00395419000190'), pa.press('tab', 2), pa.write('5')
                        # #clique em Compatibilidade
                        pa.click(pa.center(pa.locateOnScreen('./static/Compatibilidade.png'))), t.sleep(1)
                        # #clique em Compatibilidade de novo
                        pa.click(pa.center(pa.locateOnScreen('./static/Compatibilidade2.png'))), t.sleep(1)
                        # #clique em CAGED
                        pa.click(pa.center(pa.locateOnScreen('./static/CAGED.png')))
                        pa.press('tab', 2), pa.write('20'), pa.press('tab'), pa.write('1'), t.sleep(0.5)
                        # #clique em RAIS
                        pa.click(pa.center(pa.locateOnScreen('./static/RAIS.png')))
                        pp.copy(pessoa.nome), pa.hotkey('ctrl', 'v'), pa.press('tab', 2), pa.write('2')
                        pa.press('tab'), pa.write('10')
                        # #clique em Salvar
                        pa.click(pa.center(pa.locateOnScreen('./static/Salvarcadastro.png'))), t.sleep(10)
                        # #clique em fechar novo cadastro
                        pa.click(pa.center(pa.locateOnScreen('./static/Fecharnovo1.png'))), t.sleep(2)
                        # #clique em fechar trabalhadores
                        pa.click(pa.center(pa.locateOnScreen('./static/Fechartrab1.png'))), t.sleep(0.5)

                        os.makedirs(
                            r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\
                            02 - Funcionários, Departamentos e '
                            r'Férias\000 - Pastas Funcionais\00 - ATIVOS\{}'.format(pessoa.nome))
                        os.makedirs(
                            r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\
                            02 - Funcionários, Departamentos e '
                            r'Férias\000 - Pastas Funcionais\00 - ATIVOS\{}\Atestados'.format(pessoa.nome))
                        os.makedirs(
                            r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\
                            02 - Funcionários, Departamentos e '
                            r'Férias\000 - Pastas Funcionais\00 - ATIVOS\{}\Diversos'.format(pessoa.nome))
                        os.makedirs(
                            r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\
                            02 - Funcionários, Departamentos e '
                            r'Férias\000 - Pastas Funcionais\00 - ATIVOS\{}\Contratuais'.format(pessoa.nome))
                        os.makedirs(
                            r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\
                            02 - Funcionários, Departamentos e '
                            r'Férias\000 - Pastas Funcionais\00 - ATIVOS\{}\Férias'.format(pessoa.nome))
                        os.makedirs(
                            r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\
                            02 - Funcionários, Departamentos e '
                            r'Férias\000 - Pastas Funcionais\00 - ATIVOS\{}\Pontos'.format(pessoa.nome))
                        os.makedirs(
                            r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\
                            02 - Funcionários, Departamentos e '
                            r'Férias\000 - Pastas Funcionais\00 - ATIVOS\{}\Recibos'.format(pessoa.nome))
                        os.makedirs(
                            r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\
                            02 - Funcionários, Departamentos e '
                            r'Férias\000 - Pastas Funcionais\00 - ATIVOS\{}\Rescisão'.format(pessoa.nome))
                        tkinter.messagebox.showinfo(
                            title='Cadastro ok!',
                            message='Cadastro editado com sucesso!'
                        )


def salvadocsfunc(matricula):
    sessions = sessionmaker(bind=engine)
    session = sessions()
    pessoa = session.query(Colaborador).filter_by(matricula=matricula).first()
    pa.click(pa.center(pa.locateOnScreen('./static/Dexion.png')))
    p_pessoa = r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e Férias' \
               r'\000 - Pastas Funcionais\00 - ATIVOS\{}'.format(pessoa.nome)
    p_atestado = r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e Férias' \
                 r'\000 - Pastas Funcionais\00 - ATIVOS\{}\Atestado'.format(pessoa.nome)
    p_contr = r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e Férias' \
              r'\000 - Pastas Funcionais\00 - ATIVOS\{}\Contratuais'.format(pessoa.nome)
    p_diversos = r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e Férias' \
                 r'\000 - Pastas Funcionais\00 - ATIVOS\{}\Diversos'.format(pessoa.nome)
    p_ferias = r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e Férias' \
               r'\000 - Pastas Funcionais\00 - ATIVOS\{}\Férias'.format(pessoa.nome)
    p_ponto = r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e Férias' \
              r'\000 - Pastas Funcionais\00 - ATIVOS\{}\Pontos'.format(pessoa.nome)
    p_rec = r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e Férias' \
            r'\000 - Pastas Funcionais\00 - ATIVOS\{}\Recibos'.format(pessoa.nome)
    p_rescisao = r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e Férias' \
                 r'\000 - Pastas Funcionais\00 - ATIVOS\{}\Rescisão'.format(pessoa.nome)
    p_ac = r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e Férias' \
           r'\000 - Pastas Funcionais\00 - ATIVOS\Modelo\AC Modelo.docx'
    p_abconta = r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e Férias' \
                r'\000 - Pastas Funcionais\00 - ATIVOS\Modelo\Abertura Conta MODELO.docx'
    p_fcadas = r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e Férias' \
               r'\000 - Pastas Funcionais\00 - ATIVOS\Modelo\Ficha Cadastral MODELO.docx'
    p_recibos = r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e Férias' \
                r'\000 - Pastas Funcionais\00 - ATIVOS\Modelo\Recibo Crachá e Uniformes MODELO.docx'
    p_recibovt = r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e Férias' \
                 r'\000 - Pastas Funcionais\00 - ATIVOS\Modelo\Recibo VT MODELO.docx'
    p_codetic = r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e Férias' \
                r'\000 - Pastas Funcionais\00 - ATIVOS\Modelo\Cod Etica MODELO.docx'
    ps_acordo = r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e Férias' \
                r'\000 - Pastas Funcionais\00 - ATIVOS\{}\Contratuais\Acordo Banco de Horas.pdf'.format(pessoa.nome)
    ps_recctps = r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e Férias' \
                 r'\000 - Pastas Funcionais\00 - ATIVOS\{}\Contratuais\Recibo de Entrega e Dev CTPS.pdf'.format(
        pessoa.nome)
    ps_anotctps = r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e Férias' \
                  r'\000 - Pastas Funcionais\00 - ATIVOS\{}\Contratuais\Anotacoes CTPS.pdf'.format(pessoa.nome)
    ps_termovt = r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e Férias' \
                 r'\000 - Pastas Funcionais\00 - ATIVOS\{}\Contratuais\Termo Opcao VT.pdf'.format(pessoa.nome)
    ps_contrato = r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e Férias' \
                  r'\000 - Pastas Funcionais\00 - ATIVOS\{}\Contratuais\Contrato de Trabalho.pdf'.format(pessoa.nome)
    ps_ficha = r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e Férias' \
               r'\000 - Pastas Funcionais\00 - ATIVOS\{}\Contratuais\Folha de Registro.pdf'.format(pessoa.nome)
    try:
        os.mkdir(p_pessoa)
        os.mkdir(p_atestado)
        os.mkdir(p_contr)
        os.mkdir(p_diversos)
        os.mkdir(p_ferias)
        os.mkdir(p_ponto)
        os.mkdir(p_rec)
        os.mkdir(p_rescisao)
    except:
        pass

    lotacao = {
        'Unidade Park Sul - Qualquer Departamento': ['0013', 'Thais Feitosa', 'thais.morais@ciaathletica.com.br',
                                                     'Líder Park Sul'],
        'Kids': ['0010', 'Cindy Stefanie', 'cindy.neves@ciaathletica.com.br', 'Líder Kids'],
        'Musculação': ['0007', 'Thaís Feitosa', 'thais.morais@ciaathletica.com.br', 'Líder Musculação'],
        'Esportes e Lutas': ['0008', 'Morgana Rossini', 'morganalourenco@yahoo.com.br', 'Líder Natação'],
        'Crossfit': ['0012', 'Guilherme Salles', 'gmoreirasalles@gmail.com', 'Líder Crossfit'],
        'Ginástica': ['0006', 'Morgana Rossini', 'morganalourenco@yahoo.com.br', 'Líder Ginástica'],
        'Gestantes': ['0006', 'Filipe Feijó', 'filipe.feijo@ciaathletica.com.br', 'Líder Ginástica'],
        'Recepção': ['0003', 'Paulo Renato', 'paulo.simoes@ciaathletica.com.br', 'Gerente Vendas'],
        'Administrativo': ['0001', 'Felipe Rodrigues', 'felipe.rodrigues@ciaathletica.com.br', 'Gerente RH'],
        'Manutenção': ['0004', 'José Aparecido', 'aparecido.grota@ciaathletica.com.br', 'Gerente Manutenção'],
    }

    abert_c = docx.Document(p_abconta)
    ac = docx.Document(p_ac)
    fch_c = docx.Document(p_fcadas)
    recibos = docx.Document(p_recibos)
    codetic = docx.Document(p_codetic)

    # # imprimir recibo entrega e devolução de ctps
    pa.press('alt'), pa.press('r'), pa.press('a'), pa.press('r'), pa.press('e'), pa.press('tab'), pa.write(str(
        pessoa.matricula))
    pa.press('tab', 3), pa.write(str(pessoa.admiss).replace('/', '')), pa.press('tab'), t.sleep(0.5), pa.press('space')
    t.sleep(0.5), pa.press('tab'), pa.write(str(pessoa.admiss).replace('/', '')), pa.press('tab', 2)
    t.sleep(1), pa.press('enter'), t.sleep(2)

    # # clique no endereço de salvamento do recibo
    pa.click(pa.center(pa.locateOnScreen('./static/salvar.png'))), t.sleep(1)
    pp.copy(ps_recctps), pa.hotkey('ctrl', 'v'), t.sleep(0.5)
    pa.press('tab', 2), t.sleep(0.5), pa.press('enter')
    t.sleep(5)
    # # clique para fechar recibo ctps
    pa.click(pa.center(pa.locateOnScreen('./static/fechar_janela.png'))), t.sleep(0.5)
    pa.click(pa.center(pa.locateOnScreen('./static/fechar_janela.png')))

    # # Imprimir Acordo de Banco de horas
    pa.press('alt'), pa.press('r'), pa.press('z'), pa.press('d'), pa.press('d')
    pa.write("(matricula = '00{}')".format(str(pessoa.matricula))), t.sleep(1), pa.press('tab'), pa.write('2')
    pa.press('tab'), pa.press('enter'), t.sleep(10)
    pa.click(pa.center(pa.locateOnScreen('./static/salvar.png')))
    t.sleep(1), pp.copy(ps_acordo)
    pa.hotkey('ctrl', 'v'), t.sleep(1), pa.press('enter'), t.sleep(15)
    # # clique para fechar acordo
    pa.click(pa.center(pa.locateOnScreen('./static/fechar_janela.png'))), t.sleep(0.5)
    pa.click(pa.center(pa.locateOnScreen('./static/fechar_janela.png')))

    # # imprimir Anotações em CTPS
    pa.press('alt'), pa.press('r'), pa.press('a'), pa.press('c'), pa.press('e'), pa.press('tab')
    pa.write(str(pessoa.matricula))
    pa.press('tab', 2), pa.write(str(pessoa.admiss).replace('/', '')), pa.press('tab')
    pa.write(str(pessoa.admiss).replace('/', '')), pa.press('tab', 4), pa.press('space')
    pa.press('tab'), pa.press('enter'), t.sleep(1.5)
    pa.click(pa.center(pa.locateOnScreen('./static/salvar.png'))), t.sleep(1)
    pp.copy(ps_anotctps), pa.hotkey('ctrl', 'v'), t.sleep(0.5), pa.press('enter')
    t.sleep(2), pa.click(pa.center(pa.locateOnScreen('./static/fechar_janela.png'))), t.sleep(0.5)
    pa.click(pa.center(pa.locateOnScreen('./static/fechar_janela.png')))

    # # imprimir Termo VT
    pa.press('alt'), pa.press('r'), pa.press('a'), pa.press('v'), pa.press('e'), pa.press('tab')
    pa.write(str(pessoa.matricula)), pa.press('tab', 2), pa.write(str(pessoa.admiss).replace('/', ''))
    pa.press('tab'), pa.write('d'), pa.press('tab', 4), pa.press('space')
    pa.press('tab', 6), pa.press('enter'), t.sleep(1.5)
    pa.click(pa.center(pa.locateOnScreen('./static/salvar.png')))
    pp.copy(ps_termovt), pa.hotkey('ctrl', 'v'), t.sleep(0.5), pa.press('enter')
    t.sleep(2), pa.click(pa.center(pa.locateOnScreen('./static/fechar_janela.png'))), t.sleep(0.5)
    pa.click(pa.center(pa.locateOnScreen('./static/fechar_janela.png')))

    # # Imprimir Contrato
    pa.press('alt'), pa.press('r'), pa.press('z'), pa.press('d')
    if pessoa.tipo_contr == 'Horista':
        pa.press('c')
    else:
        pa.press('o')

    pa.write("(matricula = '00{}')".format(str(pessoa.matricula))), t.sleep(1), pa.press('tab'), pa.write('2')
    pa.press('tab'), pa.press('enter'), t.sleep(5)
    pa.click(pa.center(pa.locateOnScreen('./static/salvar.png')))
    pp.copy(ps_contrato), pa.hotkey('ctrl', 'v'), t.sleep(0.5), pa.press('enter')
    t.sleep(10), pa.click(pa.center(pa.locateOnScreen('./static/fechar_janela.png'))), t.sleep(0.5)
    pa.click(pa.center(pa.locateOnScreen('./static/fechar_janela.png')))

    # # Imprimir Folha de rosto de Cadastro
    pa.press('alt'), pa.press('r'), pa.press('i'), pa.press('o'), pa.press('r'), pa.press('e'), pa.press('tab')
    pa.write(str(pessoa.matricula)), pa.press('tab', 2), pa.write(str(pessoa.admiss).replace('/', '')), pa.press('tab',
                                                                                                                 2)
    pa.press('enter'), t.sleep(3)
    pa.click(pa.center(pa.locateOnScreen('./static/salvar.png')))
    pp.copy(ps_ficha), pa.hotkey('ctrl', 'v'), t.sleep(0.5), pa.press('enter')
    t.sleep(3), pa.click(pa.center(pa.locateOnScreen('./static/fechar_janela.png'))), t.sleep(0.5)
    pa.click(pa.center(pa.locateOnScreen('./static/fechar_janela.png')))

    # # Alterar AC e Salvar na pasta
    ac.paragraphs[1].text = str(ac.paragraphs[1].text).replace('#gerente', lotacao[str(pessoa.depto).title()][1])
    ac.paragraphs[2].text = str(ac.paragraphs[2].text).replace('#nome_completo', pessoa.nome)
    ac.paragraphs[3].text = str(ac.paragraphs[3].text).replace('#cargo', pessoa.cargo)
    ac.paragraphs[11].text = str(ac.paragraphs[11].text).replace('#salario', pessoa.salario)
    ac.save(p_contr + '\\AC.docx')
    docx2pdf.convert(p_contr + '\\AC.docx', p_contr + '\\AC.pdf')
    os.remove(p_contr + '\\AC.docx')

    # # Alterar Abertura de Conta e salvar na pasta
    abert_c.paragraphs[14].text = str(abert_c.paragraphs[14].text).replace('#nome_completo', pessoa.nome)\
        .replace('#rg', pessoa.rg).replace('#cpf', pessoa.cpf).replace('#endereco', pessoa.endereco)\
        .replace('#cep', pessoa.cep).replace('#bairro', pessoa.bairro).replace('#cargo', pessoa.cargo)\
        .replace('#data', pessoa.admiss)
    abert_c.save(p_contr + '\\Abertura Conta.docx')
    docx2pdf.convert(p_contr + '\\Abertura Conta.docx', p_contr + '\\Abertura Conta.pdf')
    os.remove(p_contr + '\\Abertura Conta.docx')

    # # Alterar Ficha cadastral e salvar na pasta
    fch_c.paragraphs[34].text = str(fch_c.paragraphs[34].text).replace('#gerente#',
                                                                       lotacao[str(pessoa.depto).title()][1])
    fch_c.paragraphs[9].text = str(fch_c.paragraphs[9].text).replace('#nome_completo', pessoa.nome)
    fch_c.paragraphs[21].text = str(fch_c.paragraphs[21].text).replace('#cargo', pessoa.cargo)\
        .replace('#depart', str(pessoa.depto).title())
    fch_c.paragraphs[19].text = str(fch_c.paragraphs[19].text).replace('#end_eletr', pessoa.email)
    fch_c.paragraphs[17].text = str(fch_c.paragraphs[17].text).replace('#mae#', pessoa.mae)
    fch_c.paragraphs[16].text = str(fch_c.paragraphs[16].text).replace('#pai#', pessoa.pai)
    fch_c.paragraphs[15].text = str(fch_c.paragraphs[15].text).replace('#ident', pessoa.rg)\
        .replace('#cpf#', pessoa.cpf)
    fch_c.paragraphs[13].text = str(fch_c.paragraphs[13].text).replace('#telefone', pessoa.tel)
    fch_c.paragraphs[12].text = str(fch_c.paragraphs[12].text).replace('#codigo', pessoa.cep)\
        .replace('#cid', pessoa.cidade).replace('#uf', pessoa.uf)
    fch_c.paragraphs[11].text = str(fch_c.paragraphs[11].text).replace('#local', pessoa.endereco)\
        .replace('#qd', pessoa.bairro)
    fch_c.paragraphs[10].text = str(fch_c.paragraphs[10].text)\
        .replace('#nasc', datetime.strftime(datetime.strptime(pessoa.nascimento, '%Y-%m-%d %H:%M:%S'), '%d/%m/%Y'))\
        .replace('#gen', pessoa.genero)\
        .replace('#est_civ', str(pessoa.estado_civil).replace('1 - ', '').replace('2 - ', '').replace('3 - ', '')
                 .replace('4 - ', ''))
    fch_c.save(p_contr + '\\Ficha Cadastral.docx')
    docx2pdf.convert(p_contr + '\\Ficha Cadastral.docx', p_contr + '\\Ficha Cadastral.pdf')
    os.remove(p_contr + '\\Ficha Cadastral.docx')

    # # Alterar Recibos e salvar na pasta
    recibos.paragraphs[4].text = str(recibos.paragraphs[4].text).replace('#nome_completo', pessoa.nome)
    recibos.paragraphs[12].text = str(recibos.paragraphs[12].text).replace('#nome_completo', pessoa.nome)
    recibos.paragraphs[19].text = str(recibos.paragraphs[19].text).replace('#nome_completo', pessoa.nome)
    recibos.paragraphs[27].text = str(recibos.paragraphs[27].text).replace('#nome_completo', pessoa.nome)
    recibos.paragraphs[40].text = str(recibos.paragraphs[40].text).replace('#nome_completo', pessoa.nome)
    recibos.paragraphs[48].text = str(recibos.paragraphs[48].text).replace('#nome_completo', pessoa.nome)
    recibos.save(p_contr + '\\Recibos.docx')
    docx2pdf.convert(p_contr + '\\Recibos.docx', p_contr + '\\Recibos.pdf')
    os.remove(p_contr + '\\Recibos.docx')

    # # Alterar Código de Ética e Salvar na pasta
    codetic.paragraphs[534].text = str(codetic.paragraphs[534].text).replace('#nome_completo', pessoa.nome)
    codetic.paragraphs[535].text = str(codetic.paragraphs[535].text).replace('#func', pessoa.cargo)
    codetic.paragraphs[537].text = str(codetic.paragraphs[537].text).replace('#nome_completo', pessoa.nome)
    codetic.paragraphs[541].text = str(codetic.paragraphs[541].text).replace('#admiss', pessoa.admiss)
    codetic.save(p_contr + '\\Cod Etica.docx')
    docx2pdf.convert(p_contr + '\\Cod Etica.docx', p_contr + '\\Cod Etica.pdf')
    os.remove(p_contr + '\\Cod Etica.docx')
    tkinter.messagebox.showinfo(
        title='Documentos ok!',
        message='Documentos salvos com sucesso!'
    )


def enviaemailsfunc(matricula):
    sessions = sessionmaker(bind=engine)
    session = sessions()
    pessoa = session.query(Colaborador).filter_by(matricula=matricula).first()
    p_contr = r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e Férias' \
              r'\000 - Pastas Funcionais\00 - ATIVOS\{}\Contratuais'.format(pessoa.nome)
    email_remetente = em_rem
    senha = k1
    # set up smtp connection
    s = smtplib.SMTP(host='smtp.office365.com', port=587)
    s.starttls()
    s.login(email_remetente, senha)

    # send e-mail to employee with a pdf file so he/she can go to bank to open an account
    msg = MIMEMultipart('alternative')
    msg['From'] = email_remetente
    msg['To'] = pessoa.email
    msg['Subject'] = "Carta para Abertura de conta"
    arquivo = p_contr + '\\Abertura Conta.pdf'
    text = MIMEText(f'''Olá, {str(pessoa.nome).title().split(" ")[0]}!<br><br>
    Segue sua carta para abertura de conta bancária no Itaú.<br>
    Você deve abrir a conta antes de iniciar seu contrato de trabalho. Ok?<br><br>
    Atenciosamente,<br>
    <img src="cid:image1">''', 'html')

    # set up the parameters of the message
    msg.attach(text)
    image = MIMEImage(
        open(r'C:\Users\Felipe Rodrigues\PycharmProjects\AutomacaoCia\Admissao\static\assinatura.png', 'rb').read())
    image.add_header('Content-ID', '<image1>')
    msg.attach(image)
    # attach pdf file
    part = MIMEBase('application', "octet-stream")
    part.set_payload(open(arquivo, "rb").read())
    encoders.encode_base64(part)
    part.add_header('Content-Disposition', 'attachment',
                    filename=f'Carta Banco {str(pessoa.nome).title().split(" ")[0]}.pdf')
    msg.attach(part)
    s.sendmail(email_remetente, pessoa.email, msg.as_string())
    del msg

    # send e-mail to coworker asking to register the ner employee
    msg = MIMEMultipart('alternative')
    arquivo = p_contr + '\\Ficha Cadastral.pdf'
    text = MIMEText(
        f'''Oi, Wallace!<br><br>Segue a ficha cadastral do(a) {pessoa.nome}.<br><br>Abs.,<br><img src="cid:image1">''',
        'html')
    msg.attach(text)
    image = MIMEImage(
        open(r'C:\Users\Felipe Rodrigues\PycharmProjects\AutomacaoCia\Admissao\static\assinatura.png', 'rb').read())
    image.add_header('Content-ID', '<image1>')
    msg.attach(image)
    # set up the parameters of the message
    msg['From'] = email_remetente
    msg['To'] = em_ti
    msg['Subject'] = f"Ficha Cadastral {str(pessoa.nome).title().split(' ')[0]}"
    # attach pdf file
    part = MIMEBase('application', "octet-stream")
    part.set_payload(open(arquivo, "rb").read())
    encoders.encode_base64(part)
    part.add_header('Content-Disposition', 'attachment',
                    filename=f'Ficha Cadastral {str(pessoa.nome).title().split(" ")[0]}.pdf')
    msg.attach(part)
    s.sendmail(email_remetente, em_ti, msg.as_string())
    del msg
    s.quit()
    tkinter.messagebox.showinfo(
        title='E-mails ok!',
        message='E-mails enviados com sucesso'
    )


class Frame2(ttk.Frame):
    def __init__(self, container):
        super().__init__()
        self.hoje = datetime.today()
        self.caminhoest = StringVar()
        self.nomeest = StringVar()
        self.horarioest = StringVar()
        self.cargoest = StringVar()
        self.departamentoest = StringVar()
        self.tipocontrest = StringVar()
        self.nomesplanest = []
        self.labelescolhest = ttk.Label(self, width=40, text="Escolher planilha de novos estagiários")
        self.labelescolhest.grid(column=1, row=2, padx=25, pady=1, sticky=W)
        self.botaoescolhest = ttk.Button(self, text="Escolha a planilha", command=self.selecionarest)
        self.botaoescolhest.grid(column=1, row=2, padx=350, pady=1, sticky=W)
        self.labelnomest = ttk.Label(self, width=20, text="Nome:")
        self.labelnomest.grid(column=1, row=10, padx=25, pady=1, sticky=W)
        self.combonomest = ttk.Combobox(self, values=self.nomesplanest, textvariable=self.nomeest, width=50)
        self.combonomest.grid(column=1, row=10, padx=125, pady=1, sticky=W)
        # aparecer entry para preencher matricula
        self.labelmatrest = ttk.Label(self, width=20, text="Matrícula:")
        self.labelmatrest.grid(column=1, row=11, padx=25, pady=1, sticky=W)
        self.entrymatrest = ttk.Entry(self, width=20)
        self.entrymatrest.grid(column=1, row=11, padx=125, pady=1, sticky=W)
        # aparecer entry para preencher admissao
        self.labeladmissest = ttk.Label(self, width=20, text="Admissão:")
        self.labeladmissest.grid(column=1, row=12, padx=25, pady=1, sticky=W)
        self.entryadmissest = DateEntry(self, selectmode='day', year=self.hoje.year, month=self.hoje.month,
                                        day=self.hoje.day, locale='pt_BR')
        self.entryadmissest.grid(column=1, row=12, padx=125, pady=1, sticky=W)
        #aparecer combo departamento
        self.labeldeptoest = ttk.Label(self, width=20, text="Departamento:")
        self.labeldeptoest.grid(column=1, row=19, padx=25, pady=1, sticky=W)
        self.combodeptoest = ttk.Combobox(self, values=departamentos, textvariable=self.departamentoest, width=50)
        self.combodeptoest.grid(column=1, row=19, padx=125, pady=1, sticky=W)
        self.agenciaest = StringVar()
        self.contaest = StringVar()
        self.digitoest = StringVar()
        # aparecer entry para agencia
        self.labelagest = ttk.Label(self, width=20, text="Agência:")
        self.labelagest.grid(column=1, row=24, padx=260, pady=1, sticky=W)
        self.entryagest = ttk.Entry(self, width=20, textvariable=self.agenciaest)
        self.entryagest.grid(column=1, row=24, padx=320, pady=1, sticky=W)
        # aparecer entry para conta
        self.labelccest = ttk.Label(self, width=20, text="Conta:")
        self.labelccest.grid(column=1, row=25, padx=260, pady=1, sticky=W)
        self.entryccest = ttk.Entry(self, width=20, textvariable=self.contaest)
        self.entryccest.grid(column=1, row=25, padx=320, pady=1, sticky=W)
        # aparecer entry para ditigo
        self.labeldigest = ttk.Label(self, width=20, text="Dígito:")
        self.labeldigest.grid(column=1, row=26, padx=260, pady=1, sticky=W)
        self.entrydigest = ttk.Entry(self, width=20, textvariable=self.digitoest)
        self.entrydigest.grid(column=1, row=26, padx=320, pady=1, sticky=W)
        self.solicitarest = IntVar()
        self.solictest = ttk.Checkbutton(self, text='Apenas solicitar contrato.', variable=self.solicitarest)
        self.solictest.grid(column=1, row=25, padx=26, pady=1, sticky=W)
        self.edicaoest = IntVar()
        self.editarest = ttk.Checkbutton(self, text='Editar cadastro feito manualmente.', variable=self.edicaoest)
        self.editarest.grid(column=1, row=26, padx=26, pady=1, sticky=W)
        self.feitondeest = IntVar()
        self.ondeest = ttk.Checkbutton(self, text='Cadastro realizado fora da Cia.', variable=self.feitondeest)
        self.ondeest.grid(column=1, row=27, padx=26, pady=1, sticky=W)
        self.cargoest = StringVar()
        self.botaocadastrarest = ttk.Button(self, width=20, text="Cadastrar Estagiário",
                                            command=lambda: [
                                                cadastro_estagiario(
                                                    self.solicitarest.get(), self.caminhoest.get(),
                                                    self.edicaoest.get(), self.feitondeest.get(),
                                                    self.combonomest.get(),
                                                    self.entrymatrest.get(), self.entryadmissest.get(),
                                                    '', self.combodeptoest.get(),
                                                    '', '', '',
                                                    self.agenciaest.get(),
                                                    self.contaest.get(),
                                                    self.digitoest.get()
                                                )
                                            ]
                                            )
        self.botaocadastrarest.grid(column=1, row=28, padx=520, pady=1, sticky=W)

        def carregarest(local):
            planwb = l_w(local)
            plansh = planwb['Respostas ao formulário 1']
            lista = []
            for x, pessoa in enumerate(plansh):
                lista.append(f'{x + 1} - {pessoa[2].value}')
            self.combonomest.config(values=lista)

        self.botaocarregest = ttk.Button(self, text="Carregar planilha",
                                         command=lambda: [carregarest(self.caminhoest.get())])
        self.botaocarregest.grid(column=1, row=4, padx=350, pady=25, sticky=W)

    def selecionarest(self):
        try:
            caminhoplanest = tkinter.filedialog.askopenfilename(title='Planilha Estagiários')
            self.caminhoest.set(str(caminhoplanest))
        except ValueError:
            pass


def cadastro_estagiario(solicitar_contr=0, caminho='', editar=0, ondestou=0, nome='', matricula='', admissao='',
                        cargo='', depto='', tipo_contr='Horista',
                        hrsem='25', hrmens='100', agencia='', conta='', digito=''):
    sessions = sessionmaker(bind=engine)
    session = sessions()
    pa.FAILSAFE = False
    salario = 5.10
    if solicitar_contr == 1:
        hoje = datetime.today()
        wb = l_w(caminho)
        sh = wb['Respostas ao formulário 1']
        num, name = nome.strip().split(' - ')
        linha = int(num)
        lotacao = {
            'Unidade Park Sul - qualquer departamento': ['0013', 'Thais Feitosa', 'thais.morais@ciaathletica.com.br',
                                                         'Líder Park Sul'],
            'Kids': ['0010', 'Cindy Stefanie', 'cindy.neves@ciaathletica.com.br', 'Líder Kids'],
            'Musculação': ['0007', 'Thais Feitosa', 'thais.morais@ciaathletica.com.br', 'Líder Musculação'],
            'Esportes e Lutas': ['0008', 'Morgana Rossini', 'morganalourenco@yahoo.com.br', 'Líder Natação'],
            'Crossfit': ['0012', 'Guilherme Salles', 'gmoreirasalles@gmail.com', 'Líder Crossfit'],
            'Ginástica': ['0006', 'Morgana Rossini', 'morganalourenco@yahoo.com.br', 'Líder Ginástica'],
            'Gestantes': ['0006', 'Filipe Feijó', 'filipe.feijo@ciaathletica.com.br', 'Gerente Técnico'],
            'Recepção': ['0003', 'Paulo Renato', 'paulo.simoes@ciaathletica.com.br', 'Gerente Vendas'],
            'Administrativo': ['0001', 'Felipe Rodrigues', 'felipe.rodrigues@ciaathletica.com.br', 'Gerente RH'],
            'Manutenção': ['0004', 'José Aparecido', 'aparecido.grota@ciaathletica.com.br', 'Gerente Manutenção'],
        }
        cadastro = {'nome': str(sh[f"C{linha}"].value).title().strip(), 'nasc_ed': sh[f"D{linha}"].value,
                    'genero': str(sh[f"E{linha}"].value), 'est_civ': str(sh[f"F{linha}"].value),
                    'pai': str(sh[f"M{linha}"].value), 'mae': str(sh[f"N{linha}"].value),
                    'end': str(sh[f"O{linha}"].value),
                    'num': str(sh[f"P{linha}"].value), 'bairro': str(sh[f"Q{linha}"].value),
                    'cep': str(sh[f"R{linha}"].value).replace('.', '').replace('-', ''),
                    'cid_end': str(sh[f"S{linha}"].value), 'uf_end': str(sh[f"T{linha}"].value),
                    'tel': str(sh[f"U{linha}"].value).replace('(', '').replace(')', '').replace('-', '').replace(' ',
                                                                                                                 ''),
                    'mun_end': str(sh[f"AP{linha}"].value),
                    'cpf': str(sh[f"V{linha}"].value).strip().replace('.', '').replace('-', '').replace(' ', '').zfill(
                        11),
                    'rg': str(sh[f"W{linha}"].value).strip().replace('.', '').replace('-', '').replace(' ', ''),
                    'emissor': str(sh[f"X{linha}"].value),
                    'lotacao': str(lotacao[f'{sh[f"AG{linha}"].value}'][0]).zfill(4),
                    'cargo': str(sh[f"AH{linha}"].value), 'horario': str(sh[f"AI{linha}"].value),
                    'email': str(sh[f"B{linha}"].value).strip(),
                    'admissao_ed': str(sh[f"AL{linha}"].value),
                    'faculdade': str(sh[f"AV{linha}"].value), 'semestre': str(sh[f"AS{linha}"].value),
                    'turno': str(sh[f"AT{linha}"].value), 'conclusao': str(sh[f"AU{linha}"].value),
                    'salario': str(sh[f"AM{linha}"].value),
                    'hrsemanais': str(sh[f"AQ{linha}"].value), 'hrmensais': str(sh[f"AR{linha}"].value)}
        email_remetente = em_rem
        senha = k1
        lot = lotacao[f'{sh[f"AG{linha}"].value}']
        pasta = r'\192.168.0.250'
        modelo = f'\\{pasta}\\rh\\01 - RH\\01 - Administração.Controles\\02 - Funcionários, Departamentos e Férias\\' \
                 f'000 - Pastas Funcionais\\00 - ATIVOS\\0 - Estagiários\\Modelo'
        os.makedirs(
            f'\\{pasta}\\rh\\01 - RH\\01 - Administração.Controles\\02 - Funcionários, Departamentos e Férias\\'
            f'000 - Pastas Funcionais\\00 - ATIVOS\\0 - Estagiários\\'
            f'0 - Ainda nao iniciaram\\{str(cadastro["nome"]).upper()}')
        os.makedirs(
            f'\\{pasta}\\rh\\01 - RH\\01 - Administração.Controles\\02 - Funcionários, Departamentos e Férias\\'
            f'000 - Pastas Funcionais\\00 - ATIVOS\\0 - Estagiários\\'
            f'0 - Ainda nao iniciaram\\{str(cadastro["nome"]).upper()}\\Atestados')
        os.makedirs(
            f'\\{pasta}\\rh\\01 - RH\\01 - Administração.Controles\\02 - Funcionários, Departamentos e Férias\\'
            f'000 - Pastas Funcionais\\00 - ATIVOS\\0 - Estagiários\\'
            f'0 - Ainda nao iniciaram\\{str(cadastro["nome"]).upper()}\\Diversos')
        os.makedirs(
            f'\\{pasta}\\rh\\01 - RH\\01 - Administração.Controles\\02 - Funcionários, Departamentos e Férias\\'
            f'000 - Pastas Funcionais\\00 - ATIVOS\\0 - Estagiários\\'
            f'0 - Ainda nao iniciaram\\{str(cadastro["nome"]).upper()}\\Contratuais')
        os.makedirs(
            f'\\{pasta}\\rh\\01 - RH\\01 - Administração.Controles\\02 - Funcionários, Departamentos e Férias\\'
            f'000 - Pastas Funcionais\\00 - ATIVOS\\0 - Estagiários\\'
            f'0 - Ainda nao iniciaram\\{str(cadastro["nome"]).upper()}\\Ferias')
        os.makedirs(
            f'\\{pasta}\\rh\\01 - RH\\01 - Administração.Controles\\02 - Funcionários, Departamentos e Férias\\'
            f'000 - Pastas Funcionais\\00 - ATIVOS\\0 - Estagiários\\'
            f'0 - Ainda nao iniciaram\\{str(cadastro["nome"]).upper()}\\Ponto')
        os.makedirs(
            f'\\{pasta}\\rh\\01 - RH\\01 - Administração.Controles\\02 - Funcionários, Departamentos e Férias\\'
            f'000 - Pastas Funcionais\\00 - ATIVOS\\0 - Estagiários\\'
            f'0 - Ainda nao iniciaram\\{str(cadastro["nome"]).upper()}\\Recibo')
        os.makedirs(
            f'\\{pasta}\\rh\\01 - RH\\01 - Administração.Controles\\02 - Funcionários, Departamentos e Férias\\'
            f'000 - Pastas Funcionais\\00 - ATIVOS\\0 - Estagiários\\'
            f'0 - Ainda nao iniciaram\\{str(cadastro["nome"]).upper()}\\Rescisao')
        pasta_contratuais = f'\\{pasta}\\rh\\01 - RH\\01 - Administração.Controles\\' \
                            f'02 - Funcionários, Departamentos e Férias\\000 - Pastas Funcionais\\' \
                            f'00 - ATIVOS\\0 - Estagiários\\' \
                            f'0 - Ainda nao iniciaram\\{str(cadastro["nome"]).upper()}\\Contratuais'

        # change tree docx model files with intern data and save pdfs files
        solicitacao = docx.Document(modelo + r'\Solicitacao MODELO - Copia.docx')
        solicitacao.tables[0].rows[4].cells[0].paragraphs[0].text = str(
            solicitacao.tables[0].rows[4].cells[0].paragraphs[0].text).replace('#supervisor_estagio', f'{lot[1]}')
        solicitacao.tables[0].rows[5].cells[1].paragraphs[0].text = str(
            solicitacao.tables[0].rows[5].cells[1].paragraphs[0].text).replace('#cargo', f'{lot[3]}')
        solicitacao.tables[0].rows[6].cells[0].paragraphs[0].text = str(
            solicitacao.tables[0].rows[6].cells[0].paragraphs[0].text).replace('#email_supervisor', f'{lot[2]}')
        solicitacao.tables[0].rows[9].cells[0].paragraphs[0].text = str(
            solicitacao.tables[0].rows[9].cells[0].paragraphs[0].text).replace('#horario', cadastro['horario'])
        solicitacao.tables[0].rows[14].cells[0].paragraphs[0].text = str(
            solicitacao.tables[0].rows[14].cells[0].paragraphs[0].text).replace('#nome_completo', cadastro['nome'])
        solicitacao.tables[0].rows[15].cells[0].paragraphs[0].text = str(
            solicitacao.tables[0].rows[15].cells[0].paragraphs[0].text).replace('#nasc',
                                                                                datetime.strftime(cadastro['nasc_ed'],
                                                                                                  '%d/%m/%Y')
                                                                                ).replace('#rg',
                                                                                          cadastro['rg']).replace(
            '#cpf', cadastro['cpf'])
        solicitacao.tables[0].rows[16].cells[0].paragraphs[0].text = str(
            solicitacao.tables[0].rows[16].cells[0].paragraphs[0].text).replace('#sexo', cadastro['genero'])
        solicitacao.tables[0].rows[17].cells[0].paragraphs[0].text = str(
            solicitacao.tables[0].rows[17].cells[0].paragraphs[0].text).replace('#endereco', cadastro['end'])
        solicitacao.tables[0].rows[18].cells[0].paragraphs[0].text = str(
            solicitacao.tables[0].rows[18].cells[0].paragraphs[0].text).replace('#cep', cadastro['cep']).replace(
            '#bairro', cadastro['bairro'])
        solicitacao.tables[0].rows[19].cells[0].paragraphs[0].text = str(
            solicitacao.tables[0].rows[19].cells[0].paragraphs[0].text).replace('#telefone', cadastro['tel'])
        solicitacao.tables[0].rows[20].cells[0].paragraphs[0].text = str(
            solicitacao.tables[0].rows[20].cells[0].paragraphs[0].text).replace('#end_eletr', cadastro['email'])
        solicitacao.tables[0].rows[22].cells[0].paragraphs[0].text = str(
            solicitacao.tables[0].rows[22].cells[0].paragraphs[0].text).replace('#semestre', cadastro['semestre'])
        solicitacao.tables[0].rows[23].cells[0].paragraphs[0].text = str(
            solicitacao.tables[0].rows[23].cells[0].paragraphs[0].text).replace('#turno', cadastro['turno']).replace(
            '#ano_concl', cadastro['conclusao'])
        solicitacao.tables[0].rows[24].cells[0].paragraphs[0].text = str(
            solicitacao.tables[0].rows[24].cells[0].paragraphs[0].text).replace('#faculdade', cadastro['faculdade'])
        solicitacao.save(pasta_contratuais + f'\\Pedido TCE {str(cadastro["nome"]).split(" ")[0]}.docx')
        docx2pdf.convert(pasta_contratuais + f'\\Pedido TCE {str(cadastro["nome"]).split(" ")[0]}.docx',
                         pasta_contratuais + f'\\Pedido TCE {str(cadastro["nome"]).split(" ")[0]}.pdf')
        os.remove(pasta_contratuais + f'\\Pedido TCE {str(cadastro["nome"]).split(" ")[0]}.docx')

        ficha_cadastral = docx.Document(modelo + r'\Ficha Cadastral MODELO - Copia.docx')
        ficha_cadastral.tables[1].rows[0].cells[0].paragraphs[0].text = str(
            ficha_cadastral.tables[1].rows[0].cells[0].paragraphs[0].text).replace('#nome_completo', cadastro['nome'])
        ficha_cadastral.tables[1].rows[1].cells[0].paragraphs[0].text = str(
            ficha_cadastral.tables[1].rows[1].cells[0].paragraphs[0].text).replace('#nasc', datetime.strftime(
            cadastro['nasc_ed'], '%d/%m/%Y'))
        ficha_cadastral.tables[1].rows[1].cells[2].paragraphs[0].text = str(
            ficha_cadastral.tables[1].rows[1].cells[2].paragraphs[0].text).replace('#gen', cadastro['genero'])
        ficha_cadastral.tables[1].rows[1].cells[4].paragraphs[0].text = str(
            ficha_cadastral.tables[1].rows[1].cells[4].paragraphs[0].text).replace('#est_civ', cadastro['est_civ'])
        ficha_cadastral.tables[1].rows[2].cells[0].paragraphs[0].text = str(
            ficha_cadastral.tables[1].rows[2].cells[0].paragraphs[0].text).replace('#local', cadastro['end'])
        ficha_cadastral.tables[1].rows[2].cells[4].paragraphs[0].text = str(
            ficha_cadastral.tables[1].rows[2].cells[4].paragraphs[0].text).replace('#qd', cadastro['bairro'])
        ficha_cadastral.tables[1].rows[2].cells[7].paragraphs[0].text = str(
            ficha_cadastral.tables[1].rows[2].cells[7].paragraphs[0].text).replace('#codigo', cadastro['cep'])
        ficha_cadastral.tables[1].rows[4].cells[1].paragraphs[0].text = str(
            ficha_cadastral.tables[1].rows[4].cells[1].paragraphs[0].text).replace('#telefone', cadastro['tel'])
        ficha_cadastral.tables[1].rows[4].cells[5].paragraphs[0].text = str(
            ficha_cadastral.tables[1].rows[4].cells[5].paragraphs[0].text).replace('#ident', cadastro['rg'])
        ficha_cadastral.tables[1].rows[5].cells[1].paragraphs[0].text = str(
            ficha_cadastral.tables[1].rows[5].cells[1].paragraphs[0].text).replace('#cpf#', cadastro['cpf'])
        ficha_cadastral.tables[1].rows[6].cells[3].paragraphs[0].text = str(
            ficha_cadastral.tables[1].rows[6].cells[3].paragraphs[0].text).replace('#pai#', cadastro['pai'])
        ficha_cadastral.tables[1].rows[7].cells[1].paragraphs[0].text = str(
            ficha_cadastral.tables[1].rows[7].cells[1].paragraphs[0].text).replace('#mae#', cadastro['mae'])
        ficha_cadastral.tables[1].rows[8].cells[0].paragraphs[0].text = str(
            ficha_cadastral.tables[1].rows[8].cells[0].paragraphs[0].text).replace('#end_eletr', cadastro['email'])
        ficha_cadastral.tables[1].rows[8].cells[1].paragraphs[0].text = str(
            ficha_cadastral.tables[1].rows[8].cells[1].paragraphs[0].text).replace('#depart', str(sh["AG3"].value))
        ficha_cadastral.save(pasta_contratuais + f'\\Ficha Cadastral {str(cadastro["nome"]).split(" ")[0]}.docx')
        docx2pdf.convert(pasta_contratuais + f'\\Ficha Cadastral {str(cadastro["nome"]).split(" ")[0]}.docx',
                         pasta_contratuais + f'\\Ficha Cadastral {str(cadastro["nome"]).split(" ")[0]}.pdf')
        os.remove(pasta_contratuais + f'\\Ficha Cadastral {str(cadastro["nome"]).split(" ")[0]}.docx')

        carta_banco = docx.Document(
            modelo + r'\Abertura Conta MODELO.docx')
        carta_banco.paragraphs[14].text = str(carta_banco.paragraphs[14].text).replace('#nome_completo',
                                                                                       cadastro['nome']
                                                                                       ).replace('#rg', cadastro['rg']
                                                                                                 ).replace(
            '#cpf', cadastro['cpf']).replace('#endereço', cadastro['end']).replace('#cep', cadastro['cep']).replace(
            '#bairro', cadastro['bairro']).replace('#desde#', datetime.strftime(hoje, '%d/%m/%Y'))
        carta_banco.save(pasta_contratuais + f'\\Carta Banco {str(cadastro["nome"]).split(" ")[0]}.docx')
        docx2pdf.convert(pasta_contratuais + f'\\Carta Banco {str(cadastro["nome"]).split(" ")[0]}.docx',
                         pasta_contratuais + f'\\Carta Banco {str(cadastro["nome"]).split(" ")[0]}.pdf')
        os.remove(pasta_contratuais + f'\\Carta Banco {str(cadastro["nome"]).split(" ")[0]}.docx')

        # set up smtpp connection
        s = smtplib.SMTP(host='smtp.office365.com', port=587)
        s.starttls()
        s.login(email_remetente, senha)

        # send e-mail to intern with a pdf file so he/she can go to bank to open an account
        msg = MIMEMultipart('alternative')
        arquivo = pasta_contratuais + f'\\Carta Banco {str(cadastro["nome"]).split(" ")[0]}.pdf'
        text = MIMEText(f'''Olá, {str(cadastro["nome"]).split(" ")[0]}!<br><br>
        Segue sua carta para abertura de conta bancária no Itaú.<br>
        Você deve abrir a conta antes de iniciar os trabalhos no estágio. Ok?<br>
        Você já pode buscar seu contrato no IF. Será necessário levar uma declaração de matrícula do seu curso.<br><br>
        Atenciosamente,<br>
        <img src="cid:image1">''', 'html')
        msg.attach(text)
        image = MIMEImage(
            open(r'C:\Users\Felipe Rodrigues\PycharmProjects\AutomacaoCia\Admissao\static\assinatura.png', 'rb').read())
        image.add_header('Content-ID', '<image1>')
        msg.attach(image)
        # set up the parameters of the message
        msg['From'] = email_remetente
        msg['To'] = cadastro['email']
        msg['Subject'] = "Carta para Abertura de conta"
        # attach pdf file
        part = MIMEBase('application', "octet-stream")
        part.set_payload(open(arquivo, "rb").read())
        encoders.encode_base64(part)
        part.add_header('Content-Disposition', 'attachment',
                        filename=f'Carta Banco {str(cadastro["nome"]).split(" ")[0]}.pdf')
        msg.attach(part)
        s.sendmail(email_remetente, cadastro['email'], msg.as_string())
        del msg

        # send e-mail to coworker asking to register the intern
        msg = MIMEMultipart('alternative')
        arquivo = pasta_contratuais + f'\\Ficha Cadastral {str(cadastro["nome"]).split(" ")[0]}.pdf'
        text = MIMEText(f'''Oi, Wallace!<br><br>
        Segue a ficha cadastral do(a) estagiário(a) {cadastro["nome"]}.<br><br>
        Abs.,<br>
        <img src="cid:image1">''', 'html')
        msg.attach(text)
        image = MIMEImage(
            open(r'C:\Users\Felipe Rodrigues\PycharmProjects\AutomacaoCia\Admissao\static\assinatura.png', 'rb').read())
        # Define the image's ID as referenced in the HTML body above
        image.add_header('Content-ID', '<image1>')
        msg.attach(image)
        # set up the parameters of the message
        msg['From'] = email_remetente
        msg['To'] = em_ti
        msg['Subject'] = f"Ficha Cadastral {str(cadastro['nome']).split(' ')[0]}"
        # attach pdf file
        part = MIMEBase('application', "octet-stream")
        part.set_payload(open(arquivo, "rb").read())
        encoders.encode_base64(part)
        part.add_header('Content-Disposition', 'attachment',
                        filename=f'Ficha Cadastral {str(cadastro["nome"]).split(" ")[0]}.pdf')
        msg.attach(part)
        s.sendmail(email_remetente, em_ti, msg.as_string())
        del msg

        # send document asking for the intern contract
        msg = MIMEMultipart('alternative')
        arquivo = pasta_contratuais + f'\\Pedido TCE {str(cadastro["nome"]).split(" ")[0]}.pdf'
        text = MIMEText(
            f'''Olá!\n\nSegue pedido de TCE do(a) estagiário(a) {cadastro["nome"]}.\n\n
            Atenciosamente,<br><img src="cid:image1">''',
            'html')
        msg.attach(text)
        image = MIMEImage(
            open(r'C:\Users\Felipe Rodrigues\PycharmProjects\AutomacaoCia\Admissao\static\assinatura.png', 'rb').read())
        # Define the image's ID as referenced in the HTML body above
        image.add_header('Content-ID', '<image1>')
        msg.attach(image)
        # set up the parameters of the message
        msg['From'] = email_remetente
        msg['To'] = em_if
        msg['Subject'] = f"Pedido TCE {str(cadastro['nome']).split(' ')[0]}"
        # attach pdf file
        part = MIMEBase('application', "octet-stream")
        part.set_payload(open(arquivo, "rb").read())
        encoders.encode_base64(part)
        part.add_header('Content-Disposition', 'attachment',
                        filename=f'Pedido TCE {str(cadastro["nome"]).split(" ")[0]}.pdf')
        msg.attach(part)
        s.sendmail(email_remetente, em_if, msg.as_string())
        del msg
        s.quit()
        tkinter.messagebox.showinfo(
            title='E-mails ok!',
            message='E-mails enviados com sucesso'
        )

    else:
        if editar == 0:
            if ondestou == 0:
                # Cadastro iniciado na Cia
                wb = l_w(caminho)
                sh = wb['Respostas ao formulário 1']
                num, name = nome.strip().split(' - ')
                linha = int(num)
                lotacao = {'UNIDADE PARK SUL - QUALQUER DEPARTAMENTO': '0013', 'KIDS': '0010', 'MUSCULAÇÃO': '0007',
                           'ESPORTES E LUTAS': '0008', 'CROSSFIT': '0012', 'GINÁSTICA': '0006', 'GESTANTES': '0008',
                           'RECEPÇÃO': '0003',
                           'FINANCEIRO': '0001', 'TI': '0001', 'MARKETING': '0001', 'MANUTENÇÃO': '0004'}
                if str(sh[f'E{linha}'].value) == 'Masculino':
                    cargo = 'ESTAGIARIO'
                else:
                    cargo = 'ESTAGIARIA'

                estag = session.query(Colaborador).filter_by(matricula=matricula).first()
                if estag:
                    pass
                else:
                    estag_cadastrado = Colaborador(
                        matricula=matricula, nome=name.upper(), admiss=admissao,
                        nascimento=str(sh[f'D{linha}'].value),
                        cpf=str(sh[f'V{linha}'].value).replace('.', '').replace('-', '').zfill(11),
                        rg=str(int(sh[f'W{linha}'].value)),
                        emissor='SSP/DF', email=str(sh[f'B{linha}'].value),
                        genero=str(sh[f'E{linha}'].value),
                        estado_civil=str(sh[f'F{linha}'].value), cor='9',
                        instru='08 - Educação Superior Incompleta',
                        nacional='Brasileiro(a)',
                        pai=str(sh[f'M{linha}'].value).upper(),
                        mae=str(sh[f'N{linha}'].value).upper(), endereco=str(sh[f'O{linha}'].value),
                        num='1',
                        bairro=str(sh[f'Q{linha}'].value),
                        cep=str(sh[f'R{linha}'].value).replace('.', '').replace('-', ''),
                        cidade='Brasília', cid_nas='Brasília - DF',
                        uf='DF',
                        cod_municipioend=municipios['DF']['Brasília'],
                        tel=str(sh[f'U{linha}'].value).replace('(', '').replace(')', '').replace('.', '').replace('-',
                                                                                                                  ''),
                        depto=depto, cargo=cargo,
                        horario=str(sh[f'AI{linha}'].value), salario=salario, tipo_contr=tipo_contr,
                        hr_sem='25', hr_mens='100',
                        est_semestre=str(sh[f'AS{linha}'].value),
                        est_turno=str(sh[f'AT{linha}'].value),
                        est_prev_conclu=str(sh[f'AU{linha}'].value),
                        est_faculdade=str(sh[f'AV{linha}'].value),
                        est_endfacul='End',
                        est_numendfacul='1',
                        est_bairroendfacul='Bairro',
                        ag=agencia, conta=conta, cdigito=digito

                    )
                    session.add(estag_cadastrado)
                    session.commit()
                pasta = r'\192.168.0.250'
                try:
                    os.makedirs(
                        f'\\{pasta}\\rh\\01 - RH\\01 - Administração.Controles\\'
                        f'02 - Funcionários, Departamentos e Férias\\000 - Pastas Funcionais\\'
                        f'00 - ATIVOS\\0 - Estagiários\\0 - Ainda nao iniciaram\\{estag.nome}')
                    os.makedirs(
                        f'\\{pasta}\\rh\\01 - RH\\01 - Administração.Controles\\'
                        f'02 - Funcionários, Departamentos e Férias\\000 - Pastas Funcionais\\'
                        f'00 - ATIVOS\\0 - Estagiários\\0 - Ainda nao iniciaram\\{estag.nome}\\Atestados')
                    os.makedirs(
                        f'\\{pasta}\\rh\\01 - RH\\01 - Administração.Controles\\'
                        f'02 - Funcionários, Departamentos e Férias\\000 - Pastas Funcionais\\'
                        f'00 - ATIVOS\\0 - Estagiários\\0 - Ainda nao iniciaram\\{estag.nome}\\Diversos')
                    os.makedirs(
                        f'\\{pasta}\\rh\\01 - RH\\01 - Administração.Controles\\'
                        f'02 - Funcionários, Departamentos e Férias\\000 - Pastas Funcionais\\'
                        f'00 - ATIVOS\\0 - Estagiários\\0 - Ainda nao iniciaram\\{estag.nome}\\Contratuais')
                    os.makedirs(
                        f'\\{pasta}\\rh\\01 - RH\\01 - Administração.Controles\\'
                        f'02 - Funcionários, Departamentos e Férias\\000 - Pastas Funcionais\\'
                        f'00 - ATIVOS\\0 - Estagiários\\0 - Ainda nao iniciaram\\{estag.nome}\\Ferias')
                    os.makedirs(
                        f'\\{pasta}\\rh\\01 - RH\\01 - Administração.Controles\\'
                        f'02 - Funcionários, Departamentos e Férias\\000 - Pastas Funcionais\\'
                        f'00 - ATIVOS\\0 - Estagiários\\0 - Ainda nao iniciaram\\{estag.nome}\\Ponto')
                    os.makedirs(
                        f'\\{pasta}\\rh\\01 - RH\\01 - Administração.Controles\\'
                        f'02 - Funcionários, Departamentos e Férias\\000 - Pastas Funcionais\\'
                        f'00 - ATIVOS\\0 - Estagiários\\0 - Ainda nao iniciaram\\{estag.nome}\\Recibo')
                    os.makedirs(
                        f'\\{pasta}\\rh\\01 - RH\\01 - Administração.Controles\\'
                        f'02 - Funcionários, Departamentos e Férias\\000 - Pastas Funcionais\\'
                        f'00 - ATIVOS\\0 - Estagiários\\0 - Ainda nao iniciaram\\{estag.nome}\\Rescisao')
                except:
                    pass
                # abrir cadastro no dexion e atualizar informações campo a campo
                pessoa = session.query(Colaborador).filter_by(matricula=matricula).first()
                pastapessoa = f'\\{pasta}\\rh\\01 - RH\\01 - Administração.Controles\\' \
                              f'02 - Funcionários, Departamentos e Férias\\000 - Pastas Funcionais\\' \
                              f'00 - ATIVOS\\0 - Estagiários\\0 - Ainda nao iniciaram\\{pessoa.nome}'
                pa.click(pa.center(pa.locateOnScreen('./static/Dexion.png')))
                pa.press('alt'), pa.press('a'), pa.press('t'), t.sleep(2), pa.press(
                    'i'), t.sleep(5), pa.write(str(pessoa.matricula)), pa.press('enter'), t.sleep(20)
                pp.copy(pessoa.nome), pa.hotkey('ctrl', 'v'), pa.press('tab'), pa.write(pessoa.cpf)
                t.sleep(1), pa.press('tab', 3), pa.write(pessoa.genero), pa.press('tab'), pa.write(pessoa.cor)
                t.sleep(1), pa.press('tab', 2), pa.write(pessoa.instru)
                t.sleep(1), pa.press('tab')
                if str(pessoa.estado_civil) == 'Casado(a)':
                    pa.write('2')
                    pa.press('tab', 6)
                else:
                    pa.write('1')
                    pa.press('tab', 5)
                pa.write(datetime.strftime(datetime.strptime(pessoa.nascimento, '%Y-%m-%d %H:%M:%S'), '%d%m%Y'))
                t.sleep(1), pa.press('tab'), pp.copy(pessoa.cid_nas), pa.hotkey('ctrl', 'v'), pa.press('tab')
                t.sleep(1), pa.write(pessoa.uf), pa.press('tab'), pa.write(pessoa.cod_municipioend), pa.press('tab')
                t.sleep(1), pp.copy(pessoa.pai), pa.hotkey('ctrl', 'v'), pa.press('tab'), pa.write('105'), pa.press(
                    'tab')
                t.sleep(1), pp.copy(pessoa.mae), pa.hotkey('ctrl', 'v'), pa.press('tab'), pa.write('105')
                # # clique em documentos
                pa.click(pa.center(pa.locateOnScreen('./static/Documentos.png')))
                pa.press('tab'), pa.write(str(pessoa.rg)), pa.press('tab'), pa.write(
                    pessoa.emissor), pa.press('tab', 3), pa.write(pessoa.cod_municipioend)
                pa.press('tab', 9), pa.write('3'), pa.press('tab'), pa.write('341'), pa.press('tab')
                pa.write(pessoa.ag), pa.press('tab'), pa.write(f'{pessoa.conta}-{pessoa.cdigito}'), pa.press('tab')
                # #clique em endereço
                pa.click(pa.center(pa.locateOnScreen('./static/Endereco.png')))
                pa.press('tab', 2), pp.copy(pessoa.endereco), pa.hotkey('ctrl', 'v'), pa.press(
                    'tab'), pa.write(pessoa.num), pa.press('tab', 2)
                pp.copy(pessoa.bairro), pa.hotkey('ctrl', 'v'), pa.press('tab'), pp.copy(pessoa.cidade), pa.hotkey(
                    'ctrl', 'v'), pa.press('tab'), pa.write(pessoa.uf)
                pa.press('tab'), pa.write(pessoa.cep), pa.press('tab'), pa.write(pessoa.cod_municipioend), pa.press(
                    'tab'), pa.write(str(pessoa.tel)), pa.press('tab', 2), pa.write(pessoa.email)
                # #clique em dados contratuais
                pa.click(pa.center(pa.locateOnScreen('./static/Contratuais.png')))
                pa.press('tab'), pa.write(str(pessoa.admiss).replace('/', '')), t.sleep(1)
                pa.press('tab'), pa.write('9')
                pa.press('tab', 7), pa.write('n'), pa.press('tab'), pa.write('4')
                pa.press('tab'), pa.write('Ed. Fisica')
                pa.press('tab', 2), pa.write(str(int(str(pessoa.admiss).replace('/', '')) + 2).zfill(8))
                # #clique em instituição de ensino
                try:
                    pa.click(pa.center(pa.locateOnScreen('./static/faculdade.png')))
                except:
                    t.sleep(3)
                    pa.click(pa.center(pa.locateOnScreen('./static/faculdade.png')))
                pa.press('tab'), pp.copy(pessoa.est_faculdade), pa.hotkey('ctrl', 'v')
                pa.press('tab'), pp.copy(pessoa.est_endfacul), pa.hotkey('ctrl', 'v')
                pa.press('tab'), pp.copy(pessoa.est_numendfacul), pa.hotkey('ctrl', 'v')
                pa.press('tab'), pp.copy(pessoa.est_bairroendfacul), pa.hotkey('ctrl', 'v')
                # #clique em Outros
                try:
                    pa.click(pa.center(pa.locateOnScreen('./static/Outros.png')))
                except:
                    t.sleep(5)
                    pa.click(pa.center(pa.locateOnScreen('./static/Outros.png')))
                t.sleep(2), pa.write('CARGO GERAL')
                pa.press('tab'), pa.write(pessoa.cargo)
                t.sleep(1), pa.press('tab'), pa.write(pessoa.salario), pa.press('tab')
                if str(pessoa.tipo_contr) == 'Horista':
                    pa.press('1')
                else:
                    pa.press('5')
                pa.press('tab', 2), t.sleep(1), pa.write(str(pessoa.hr_sem)), pa.press('tab'), t.sleep(1)
                pa.write(str(pessoa.hr_mens))
                # #clique em eventos trabalhistas
                pa.click(pa.center(pa.locateOnScreen('./static/EVTrab.png')))
                t.sleep(1)
                # #clique em lotação
                pa.click(pa.center(pa.locateOnScreen('./static/Lotacoes.png')))
                pa.press('tab'), pa.press('tab'), pa.write('i'), pa.write(str(pessoa.admiss).replace('/', ''))
                t.sleep(1), pa.press('enter'), t.sleep(1)
                pp.copy(lotacao[f'{pessoa.depto}']), pa.hotkey('ctrl', 'v'), pa.press('enter'), pa.write('f')
                pa.press('tab'), pa.write('4')
                # #clique em salvar lotação
                pa.click(pa.center(pa.locateOnScreen('./static/Salvarbtn.png'))), t.sleep(1)
                # #clique em fechar lotação
                try:
                    pa.click(pa.center(pa.locateOnScreen('./static/Fecharlot.png'))), t.sleep(1)
                except:
                    t.sleep(4)
                    pa.click(pa.center(pa.locateOnScreen('./static/Fecharlot.png')))
                # #clique em Compatibilidade
                pa.click(pa.center(pa.locateOnScreen('./static/Compatibilidade3.png'))), t.sleep(1)
                # #clique em Compatibilidade de novo
                pa.click(pa.center(pa.locateOnScreen('./static/Compatibilidade2.png'))), t.sleep(1)
                pa.press('tab', 2), pa.write('9')
                # #clique em Salvar
                pa.click(pa.center(pa.locateOnScreen('./static/Salvarcadastro.png'))), t.sleep(10)
                # #clique em fechar novo cadastro
                pa.click(pa.center(pa.locateOnScreen('./static/Fecharnovo1.png'))), t.sleep(2)
                # #clique em fechar trabalhadores
                pa.click(pa.center(pa.locateOnScreen('./static/Fechartrab1.png'))), t.sleep(0.5)
                os.rename(pastapessoa,
                          f'\\{pasta}\\rh\\01 - RH\\01 - Administração.Controles\\'
                          f'02 - Funcionários, Departamentos e Férias\\000 - Pastas Funcionais\\'
                          f'00 - ATIVOS\\0 - Estagiários\\{pessoa.nome}')
                tkinter.messagebox.showinfo(
                    title='Cadastro ok!',
                    message='Cadastro realizado com sucesso!'
                )
            else:
                # Cadastro iniciado em casa
                wb = l_w(caminho)
                sh = wb['Respostas ao formulário 1']
                num, name = nome.strip().split(' - ')
                linha = int(num)
                lotacao = {'UNIDADE PARK SUL - QUALQUER DEPARTAMENTO': '0013', 'KIDS': '0010', 'MUSCULAÇÃO': '0007',
                           'ESPORTES E LUTAS': '0008', 'CROSSFIT': '0012', 'GINÁSTICA': '0006', 'GESTANTES': '0008',
                           'RECEPÇÃO': '0003',
                           'FINANCEIRO': '0001', 'TI': '0001', 'MARKETING': '0001', 'MANUTENÇÃO': '0004'}
                if str(sh[f'E{linha}'].value) == 'Masculino':
                    cargo = 'ESTAGIARIO'
                else:
                    cargo = 'ESTAGIARIA'

                estag = session.query(Colaborador).filter_by(matricula=matricula).first()
                if estag:
                    pass
                else:
                    estag_cadastrado = Colaborador(
                        matricula=matricula, nome=name.upper(), admiss=admissao,
                        nascimento=str(sh[f'D{linha}'].value),
                        cpf=str(int(sh[f'V{linha}'].value)).zfill(11),
                        rg=str(int(sh[f'W{linha}'].value)),
                        emissor='SSP/DF', email=str(sh[f'B{linha}'].value),
                        genero=str(sh[f'E{linha}'].value),
                        estado_civil=str(sh[f'F{linha}'].value), cor='9',
                        instru='08 - Educação Superior Incompleta',
                        nacional='Brasileiro(a)',
                        pai=str(sh[f'M{linha}'].value).upper(),
                        mae=str(sh[f'N{linha}'].value).upper(), endereco=str(sh[f'O{linha}'].value),
                        num='1',
                        bairro=str(sh[f'Q{linha}'].value),
                        cep=str(sh[f'R{linha}'].value).replace('.', '').replace('-', ''),
                        cidade='Brasília', cid_nas='Brasília - DF',
                        uf='DF',
                        cod_municipioend=municipios['DF']['Brasília'],
                        tel=str(sh[f'U{linha}'].value).replace('(', '').replace(')', '').replace('.', '').replace('-',
                                                                                                                  ''),
                        depto=depto, cargo=cargo,
                        horario=str(sh[f'AI{linha}'].value), salario=salario, tipo_contr=tipo_contr,
                        hr_sem='25', hr_mens='100',
                        est_semestre=str(sh[f'AS{linha}'].value),
                        est_turno=str(sh[f'AT{linha}'].value),
                        est_prev_conclu=str(sh[f'AU{linha}'].value),
                        est_faculdade=str(sh[f'AV{linha}'].value),
                        est_endfacul='End',
                        est_numendfacul='1',
                        est_bairroendfacul='Bairro',
                        ag=agencia, conta=conta, cdigito=digito
                    )
                    session.add(estag_cadastrado)
                    session.commit()
                pasta = r'\192.168.0.250'
                try:
                    os.makedirs(
                        f'\\{pasta}\\rh\\01 - RH\\01 - Administração.Controles\\'
                        f'02 - Funcionários, Departamentos e Férias\\000 - Pastas Funcionais\\'
                        f'00 - ATIVOS\\0 - Estagiários\\0 - Ainda nao iniciaram\\{estag.nome}')
                    os.makedirs(
                        f'\\{pasta}\\rh\\01 - RH\\01 - Administração.Controles\\'
                        f'02 - Funcionários, Departamentos e Férias\\000 - Pastas Funcionais\\'
                        f'00 - ATIVOS\\0 - Estagiários\\0 - Ainda nao iniciaram\\{estag.nome}\\Atestados')
                    os.makedirs(
                        f'\\{pasta}\\rh\\01 - RH\\01 - Administração.Controles\\'
                        f'02 - Funcionários, Departamentos e Férias\\000 - Pastas Funcionais\\'
                        f'00 - ATIVOS\\0 - Estagiários\\0 - Ainda nao iniciaram\\{estag.nome}\\Diversos')
                    os.makedirs(
                        f'\\{pasta}\\rh\\01 - RH\\01 - Administração.Controles\\'
                        f'02 - Funcionários, Departamentos e Férias\\000 - Pastas Funcionais\\'
                        f'00 - ATIVOS\\0 - Estagiários\\0 - Ainda nao iniciaram\\{estag.nome}\\Contratuais')
                    os.makedirs(
                        f'\\{pasta}\\rh\\01 - RH\\01 - Administração.Controles\\'
                        f'02 - Funcionários, Departamentos e Férias\\000 - Pastas Funcionais\\'
                        f'00 - ATIVOS\\0 - Estagiários\\0 - Ainda nao iniciaram\\{estag.nome}\\Ferias')
                    os.makedirs(
                        f'\\{pasta}\\rh\\01 - RH\\01 - Administração.Controles\\'
                        f'02 - Funcionários, Departamentos e Férias\\000 - Pastas Funcionais\\'
                        f'00 - ATIVOS\\0 - Estagiários\\0 - Ainda nao iniciaram\\{estag.nome}\\Ponto')
                    os.makedirs(
                        f'\\{pasta}\\rh\\01 - RH\\01 - Administração.Controles\\'
                        f'02 - Funcionários, Departamentos e Férias\\000 - Pastas Funcionais\\'
                        f'00 - ATIVOS\\0 - Estagiários\\0 - Ainda nao iniciaram\\{estag.nome}\\Recibo')
                    os.makedirs(
                        f'\\{pasta}\\rh\\01 - RH\\01 - Administração.Controles\\'
                        f'02 - Funcionários, Departamentos e Férias\\000 - Pastas Funcionais\\'
                        f'00 - ATIVOS\\0 - Estagiários\\0 - Ainda nao iniciaram\\{estag.nome}\\Rescisao')
                except:
                    pass
                # abrir cadastro no dexion e atualizar informações campo a campo
                pessoa = session.query(Colaborador).filter_by(matricula=matricula).first()
                pastapessoa = f'\\{pasta}\\rh\\01 - RH\\01 - Administração.Controles\\' \
                              f'02 - Funcionários, Departamentos e Férias\\000 - Pastas Funcionais\\' \
                              f'00 - ATIVOS\\0 - Estagiários\\0 - Ainda nao iniciaram\\{pessoa.nome}'
                pa.click(pa.center(pa.locateOnScreen('./static/Dexion.png')))
                pa.press('alt'), pa.press('a'), pa.press('t'), t.sleep(2), pa.press(
                    'i'), t.sleep(5), pa.write(str(pessoa.matricula)), pa.press('enter'), t.sleep(20)
                pp.copy(pessoa.nome), pa.hotkey('ctrl', 'v'), pa.press('tab'), pa.write(pessoa.cpf)
                t.sleep(1), pa.press('tab', 3), pa.write(pessoa.genero), pa.press('tab'), pa.write(pessoa.cor)
                t.sleep(1), pa.press('tab', 2), pa.write(pessoa.instru)
                t.sleep(1), pa.press('tab')
                if str(pessoa.estado_civil) == 'Casado(a)':
                    pa.write('2')
                    pa.press('tab', 6)
                else:
                    pa.write('1')
                    pa.press('tab', 5)
                pa.write(datetime.strftime(datetime.strptime(pessoa.nascimento, '%Y-%m-%d %H:%M:%S'), '%d%m%Y'))
                t.sleep(1), pa.press('tab'), pp.copy(pessoa.cid_nas), pa.hotkey('ctrl', 'v'), pa.press('tab')
                t.sleep(1), pa.write(pessoa.uf), pa.press('tab'), pa.write(pessoa.cod_municipioend), pa.press('tab')
                t.sleep(1), pp.copy(pessoa.pai), pa.hotkey('ctrl', 'v'), pa.press('tab'), pa.write('105'), pa.press(
                    'tab')
                t.sleep(1), pp.copy(pessoa.mae), pa.hotkey('ctrl', 'v'), pa.press('tab'), pa.write('105')
                # # clique em documentos
                pa.click(pa.center(pa.locateOnScreen('./static/Documentos.png')))
                pa.press('tab'), pa.write(str(pessoa.rg)), pa.press('tab'), pa.write(
                    pessoa.emissor), pa.press('tab', 3), pa.write(pessoa.cod_municipioend)
                pa.press('tab', 9), pa.write('3'), pa.press('tab'), pa.write('341'), pa.press('tab')
                pa.write(pessoa.ag), pa.press('tab'), pa.write(f'{pessoa.conta}-{pessoa.cdigito}'), pa.press('tab')
                # #clique em endereço
                pa.click(pa.center(pa.locateOnScreen('./static/Endereco.png')))
                pa.press('tab', 2), pp.copy(pessoa.endereco), pa.hotkey('ctrl', 'v'), pa.press(
                    'tab'), pa.write(pessoa.num), pa.press('tab', 2)
                pp.copy(pessoa.bairro), pa.hotkey('ctrl', 'v'), pa.press('tab'), pp.copy(pessoa.cidade), pa.hotkey(
                    'ctrl', 'v'), pa.press('tab'), pa.write(pessoa.uf)
                pa.press('tab'), pa.write(pessoa.cep), pa.press('tab'), pa.write(pessoa.cod_municipioend), pa.press(
                    'tab'), pa.write(str(pessoa.tel)), pa.press('tab', 2), pa.write(pessoa.email)
                # #clique em dados contratuais
                pa.click(pa.center(pa.locateOnScreen('./static/Contratuais.png')))
                pa.press('tab'), pa.write(str(pessoa.admiss).replace('/', '')), t.sleep(1)
                pa.press('tab'), pa.write('9')
                pa.press('tab', 7), pa.write('n'), pa.press('tab'), pa.write('4')
                pa.press('tab'), pa.write('Ed. Fisica')
                pa.press('tab', 2), pa.write(str(int(str(pessoa.admiss).replace('/', '')) + 2).zfill(8))
                # #clique em instituição de ensino
                try:
                    pa.click(pa.center(pa.locateOnScreen('./static/faculdade.png')))
                except:
                    t.sleep(3)
                    pa.click(pa.center(pa.locateOnScreen('./static/faculdade.png')))
                pa.press('tab'), pp.copy(pessoa.est_faculdade), pa.hotkey('ctrl', 'v')
                pa.press('tab'), pp.copy(pessoa.est_endfacul), pa.hotkey('ctrl', 'v')
                pa.press('tab'), pp.copy(pessoa.est_numendfacul), pa.hotkey('ctrl', 'v')
                pa.press('tab'), pp.copy(pessoa.est_bairroendfacul), pa.hotkey('ctrl', 'v')
                # #clique em Outros
                try:
                    pa.click(pa.center(pa.locateOnScreen('./static/Outros.png')))
                except:
                    t.sleep(5)
                    pa.click(pa.center(pa.locateOnScreen('./static/Outros.png')))
                t.sleep(2), pa.write('CARGO GERAL')
                pa.press('tab'), pa.write(pessoa.cargo)
                t.sleep(1), pa.press('tab'), pa.write(pessoa.salario), pa.press('tab')
                if str(pessoa.tipo_contr) == 'Horista':
                    pa.press('1')
                else:
                    pa.press('5')
                pa.press('tab', 2), t.sleep(1), pa.write(str(pessoa.hr_sem)), pa.press('tab'), t.sleep(1)
                pa.write(str(pessoa.hr_mens))
                # #clique em eventos trabalhistas
                pa.click(pa.center(pa.locateOnScreen('./static/EVTrab.png')))
                t.sleep(1)
                # #clique em lotação
                pa.click(pa.center(pa.locateOnScreen('./static/Lotacoes.png')))
                pa.press('tab'), pa.press('tab'), pa.write('i'), pa.write(str(pessoa.admiss).replace('/', ''))
                t.sleep(1), pa.press('enter'), t.sleep(1)
                pp.copy(lotacao[f'{pessoa.depto}']), pa.hotkey('ctrl', 'v'), pa.press('enter'), pa.write('f')
                pa.press('tab'), pa.write('4')
                # #clique em salvar lotação
                pa.click(pa.center(pa.locateOnScreen('./static/Salvarbtn.png'))), t.sleep(1)
                # #clique em fechar lotação
                try:
                    pa.click(pa.center(pa.locateOnScreen('./static/Fecharlot.png'))), t.sleep(1)
                except:
                    t.sleep(4)
                    pa.click(pa.center(pa.locateOnScreen('./static/Fecharlot.png')))
                # #clique em Compatibilidade
                pa.click(pa.center(pa.locateOnScreen('./static/Compatibilidade3.png'))), t.sleep(1)
                # #clique em Compatibilidade de novo
                pa.click(pa.center(pa.locateOnScreen('./static/Compatibilidade2.png'))), t.sleep(1)
                pa.press('tab', 2), pa.write('9')
                # #clique em Salvar
                pa.click(pa.center(pa.locateOnScreen('./static/Salvarcadastro.png'))), t.sleep(10)
                # #clique em fechar novo cadastro
                pa.click(pa.center(pa.locateOnScreen('./static/Fecharnovo1.png'))), t.sleep(2)
                # #clique em fechar trabalhadores
                pa.click(pa.center(pa.locateOnScreen('./static/Fechartrab1.png'))), t.sleep(0.5)
                os.rename(pastapessoa,
                          f'\\{pasta}\\rh\\01 - RH\\01 - Administração.Controles\\'
                          f'02 - Funcionários, Departamentos e Férias\\000 - Pastas Funcionais\\'
                          f'00 - ATIVOS\\0 - Estagiários\\{pessoa.nome}')
                tkinter.messagebox.showinfo(
                    title='Cadastro ok!',
                    message='Cadastro realizado com sucesso!'
                )
        else:
            if ondestou == 0:
                # Editando o cadastro na Cia
                hoje = datetime.today()
                wb = l_w(caminho)
                sh = wb['Respostas ao formulário 1']
                num, name = nome.strip().split(' - ')
                linha = int(num)
                lotacao = {
                    'Unidade Park Sul - qualquer departamento': ['0013', 'Thais Feitosa',
                                                                 'thais.morais@ciaathletica.com.br',
                                                                 'Líder Park Sul'],
                    'Kids': ['0010', 'Cindy Stefanie', 'cindy.neves@ciaathletica.com.br', 'Líder Kids'],
                    'Musculação': ['0007', 'Aline Kanyó', 'aline.kanyo@soucia.com.br', 'Líder Musculação'],
                    'Esportes e Lutas': ['0008', 'Morgana Rossini', 'morganalourenco@yahoo.com.br', 'Líder Natação'],
                    'Crossfit': ['0012', 'Guilherme Salles', 'gmoreirasalles@gmail.com', 'Líder Crossfit'],
                    'Ginástica': ['0006', 'Hugo Albuquerque', 'hugo.albuquerque@ciaathletica.com.br',
                                  'Líder Ginástica'],
                    'Gestantes': ['0006', 'Hugo Albuquerque', 'hugo.albuquerque@ciaathletica.com.br',
                                  'Líder Ginástica'],
                    'Recepção': ['0003', 'Paulo Renato', 'paulo.simoes@ciaathletica.com.br', 'Gerente Vendas'],
                    'Administrativo': ['0001', 'Felipe Rodrigues', 'felipe.rodrigues@ciaathletica.com.br',
                                       'Gerente RH'],
                    'Manutenção': ['0004', 'José Aparecido', 'aparecido.grota@ciaathletica.com.br',
                                   'Gerente Manutenção'],
                }
                cadastro = {'nome': str(sh[f"C{linha}"].value).title().strip(), 'nasc_ed': sh[f"D{linha}"].value,
                            'genero': str(sh[f"E{linha}"].value), 'est_civ': str(sh[f"F{linha}"].value),
                            'pai': str(sh[f"M{linha}"].value), 'mae': str(sh[f"N{linha}"].value),
                            'end': str(sh[f"O{linha}"].value),
                            'num': str(sh[f"P{linha}"].value), 'bairro': str(sh[f"Q{linha}"].value),
                            'cep': str(sh[f"R{linha}"].value).replace('.', '').replace('-', ''),
                            'cid_end': str(sh[f"S{linha}"].value), 'uf_end': str(sh[f"T{linha}"].value),
                            'tel': str(sh[f"U{linha}"].value).replace('(', '')
                            .replace(')', '').replace('-','').replace(' ',''),
                            'mun_end': str(sh[f"AP{linha}"].value),
                            'cpf': str(sh[f"V{linha}"].value).strip().replace('.', '')
                            .replace('-', '').replace(' ','').zfill(11),
                            'rg': str(sh[f"W{linha}"].value).strip().replace('.', '').replace('-', '').replace(' ', ''),
                            'emissor': str(sh[f"X{linha}"].value),
                            'lotacao': str(lotacao[f'{sh[f"AG{linha}"].value}'][0]).zfill(4),
                            'cargo': str(sh[f"AH{linha}"].value), 'horario': str(sh[f"AI{linha}"].value),
                            'email': str(sh[f"B{linha}"].value).strip(),
                            'admissao_ed': str(sh[f"AL{linha}"].value),
                            'faculdade': str(sh[f"AV{linha}"].value), 'semestre': str(sh[f"AS{linha}"].value),
                            'turno': str(sh[f"AT{linha}"].value), 'conclusao': str(sh[f"AU{linha}"].value),
                            'salario': str(sh[f"AM{linha}"].value),
                            'hrsemanais': str(sh[f"AQ{linha}"].value), 'hrmensais': str(sh[f"AR{linha}"].value)}
            else:
                # Editando o cadastro em Casa
                wb = l_w(caminho)
                sh = wb['Respostas ao formulário 1']
                num, name = nome.strip().split(' - ')
                linha = int(num)
                lotacao = {'UNIDADE PARK SUL - QUALQUER DEPARTAMENTO': '0013', 'KIDS': '0010', 'MUSCULAÇÃO': '0007',
                           'ESPORTES E LUTAS': '0008', 'CROSSFIT': '0012', 'GINÁSTICA': '0006', 'GESTANTES': '0008',
                           'RECEPÇÃO': '0003',
                           'FINANCEIRO': '0001', 'TI': '0001', 'MARKETING': '0001', 'MANUTENÇÃO': '0004'}
                if str(sh[f'E{linha}'].value) == 'Masculino':
                    cargo = 'ESTAGIARIO'
                else:
                    cargo = 'ESTAGIARIA'

                estag = session.query(Colaborador).filter_by(matricula=matricula).first()
                if estag:
                    pass
                else:
                    estag_cadastrado = Colaborador(
                        matricula=matricula, nome=name.upper(), admiss=admissao,
                        nascimento=str(sh[f'D{linha}'].value),
                        cpf=str(int(sh[f'V{linha}'].value)).zfill(11),
                        rg=str(int(sh[f'W{linha}'].value)),
                        emissor='SSP/DF', email=str(sh[f'B{linha}'].value),
                        genero=str(sh[f'E{linha}'].value),
                        estado_civil=str(sh[f'F{linha}'].value), cor='9',
                        instru='08 - Educação Superior Incompleta',
                        nacional='Brasileiro(a)',
                        pai=str(sh[f'M{linha}'].value).upper(),
                        mae=str(sh[f'N{linha}'].value).upper(), endereco=str(sh[f'O{linha}'].value),
                        num='1',
                        bairro=str(sh[f'Q{linha}'].value),
                        cep=str(sh[f'R{linha}'].value).replace('.', '').replace('-', ''),
                        cidade='Brasília', cid_nas='Brasília - DF',
                        uf='DF',
                        cod_municipioend=municipios['DF']['Brasília'],
                        tel=str(sh[f'U{linha}'].value).replace('(', '').replace(')', '').replace('.', '').replace('-',
                                                                                                                  ''),
                        depto=depto, cargo=cargo,
                        horario=str(sh[f'AI{linha}'].value), salario=salario, tipo_contr=tipo_contr,
                        hr_sem='25', hr_mens='100',
                        est_semestre=str(sh[f'AS{linha}'].value),
                        est_turno=str(sh[f'AT{linha}'].value),
                        est_prev_conclu=str(sh[f'AU{linha}'].value),
                        est_faculdade=str(sh[f'AV{linha}'].value),
                        est_endfacul='End',
                        est_numendfacul='1',
                        est_bairroendfacul='Bairro',
                        ag=agencia, conta=conta, cdigito=digito
                    )
                    session.add(estag_cadastrado)
                    session.commit()
                pasta = r'\192.168.0.250'
                try:
                    os.makedirs(
                        f'\\{pasta}\\rh\\01 - RH\\01 - Administração.Controles\\'
                        f'02 - Funcionários, Departamentos e Férias\\000 - Pastas Funcionais\\'
                        f'00 - ATIVOS\\0 - Estagiários\\0 - Ainda nao iniciaram\\{estag.nome}')
                    os.makedirs(
                        f'\\{pasta}\\rh\\01 - RH\\01 - Administração.Controles\\'
                        f'02 - Funcionários, Departamentos e Férias\\000 - Pastas Funcionais\\'
                        f'00 - ATIVOS\\0 - Estagiários\\0 - Ainda nao iniciaram\\{estag.nome}\\Atestados')
                    os.makedirs(
                        f'\\{pasta}\\rh\\01 - RH\\01 - Administração.Controles\\'
                        f'02 - Funcionários, Departamentos e Férias\\000 - Pastas Funcionais\\'
                        f'00 - ATIVOS\\0 - Estagiários\\0 - Ainda nao iniciaram\\{estag.nome}\\Diversos')
                    os.makedirs(
                        f'\\{pasta}\\rh\\01 - RH\\01 - Administração.Controles\\'
                        f'2 - Funcionários, Departamentos e Férias\\000 - Pastas Funcionais\\'
                        f'00 - ATIVOS\\0 - Estagiários\\0 - Ainda nao iniciaram\\{estag.nome}\\Contratuais')
                    os.makedirs(
                        f'\\{pasta}\\rh\\01 - RH\\01 - Administração.Controles\\'
                        f'02 - Funcionários, Departamentos e Férias\\000 - Pastas Funcionais\\'
                        f'00 - ATIVOS\\0 - Estagiários\\0 - Ainda nao iniciaram\\{estag.nome}\\Ferias')
                    os.makedirs(
                        f'\\{pasta}\\rh\\01 - RH\\01 - Administração.Controles\\'
                        f'02 - Funcionários, Departamentos e Férias\\000 - Pastas Funcionais\\'
                        f'00 - ATIVOS\\0 - Estagiários\\0 - Ainda nao iniciaram\\{estag.nome}\\Ponto')
                    os.makedirs(
                        f'\\{pasta}\\rh\\01 - RH\\01 - Administração.Controles\\'
                        f'02 - Funcionários, Departamentos e Férias\\000 - Pastas Funcionais\\'
                        f'00 - ATIVOS\\0 - Estagiários\\0 - Ainda nao iniciaram\\{estag.nome}\\Recibo')
                    os.makedirs(
                        f'\\{pasta}\\rh\\01 - RH\\01 - Administração.Controles\\'
                        f'02 - Funcionários, Departamentos e Férias\\000 - Pastas Funcionais\\'
                        f'00 - ATIVOS\\0 - Estagiários\\0 - Ainda nao iniciaram\\{estag.nome}\\Rescisao')
                except:
                    pass
                # abrir cadastro no dexion e atualizar informações campo a campo
                pessoa = session.query(Colaborador).filter_by(matricula=matricula).first()
                pa.click(pa.center(pa.locateOnScreen('./static/Dexion.png')))
                pa.press('alt'), pa.press('a'), pa.press('t'), t.sleep(2), pa.press(
                    'a'), t.sleep(5), pa.write(str(pessoa.matricula)), pa.press('enter'), t.sleep(20)
                pp.copy(pessoa.nome), pa.hotkey('ctrl', 'v'), pa.press('tab'), pa.write(pessoa.cpf)
                t.sleep(1), pa.press('tab', 3), pa.write(pessoa.genero), pa.press('tab'), pa.write(pessoa.cor)
                t.sleep(1), pa.press('tab', 2), pa.write(pessoa.instru)
                t.sleep(1), pa.press('tab')
                if str(pessoa.estado_civil) == 'Casado(a)':
                    pa.write('2')
                    pa.press('tab', 6)
                else:
                    pa.write('1')
                    pa.press('tab', 5)
                pa.write(datetime.strftime(datetime.strptime(pessoa.nascimento, '%Y-%m-%d %H:%M:%S'), '%d%m%Y'))
                t.sleep(1), pa.press('tab'), pp.copy(pessoa.cid_nas), pa.hotkey('ctrl', 'v'), pa.press('tab')
                t.sleep(1), pa.write(pessoa.uf), pa.press('tab'), pa.write(pessoa.cod_municipioend), pa.press('tab')
                t.sleep(1), pp.copy(pessoa.pai), pa.hotkey('ctrl', 'v'), pa.press('tab'), pa.write('105'), pa.press(
                    'tab')
                t.sleep(1), pp.copy(pessoa.mae), pa.hotkey('ctrl', 'v'), pa.press('tab'), pa.write('105')
                # # clique em documentos
                pa.click(pa.center(pa.locateOnScreen('./static/Documentos.png')))
                pa.press('tab'), pa.write(str(pessoa.rg)), pa.press('tab'), pa.write(
                    pessoa.emissor), pa.press('tab', 3), pa.write(pessoa.cod_municipioend)
                pa.press('tab', 9), pa.write('3'), pa.press('tab'), pa.write('341'), pa.press('tab')
                pa.write(pessoa.ag), pa.press('tab'), pa.write(f'{pessoa.conta}-{pessoa.cdigito}'), pa.press('tab')
                # #clique em endereço
                pa.click(pa.center(pa.locateOnScreen('./static/Endereco.png')))
                pa.press('tab', 2), pp.copy(pessoa.endereco), pa.hotkey('ctrl', 'v'), pa.press(
                    'tab'), pa.write(pessoa.num), pa.press('tab', 2)
                pp.copy(pessoa.bairro), pa.hotkey('ctrl', 'v'), pa.press('tab'), pp.copy(pessoa.cidade), pa.hotkey(
                    'ctrl', 'v'), pa.press('tab'), pa.write(pessoa.uf)
                pa.press('tab'), pa.write(pessoa.cep), pa.press('tab'), pa.write(pessoa.cod_municipioend), pa.press(
                    'tab'), pa.write(str(pessoa.tel)), pa.press('tab', 2), pa.write(pessoa.email)
                # #clique em dados contratuais
                pa.click(pa.center(pa.locateOnScreen('./static/Contratuais.png')))
                pa.press('tab'), pa.write(str(pessoa.admiss).replace('/', '')), t.sleep(1)
                pa.press('tab'), pa.write('9')
                pa.press('tab', 7), pa.write('n'), pa.press('tab'), pa.write('4')
                pa.press('tab'), pa.write('Ed. Fisica')
                pa.press('tab', 2), pa.write(str(int(str(pessoa.admiss).replace('/', '')) + 2).zfill(8))
                # #clique em instituição de ensino
                pa.click(pa.center(pa.locateOnScreen('./static/faculdade.png')))
                pa.press('tab'), pp.copy(pessoa.est_faculdade), pa.hotkey('ctrl', 'v')
                pa.press('tab'), pp.copy(pessoa.est_endfacul), pa.hotkey('ctrl', 'v')
                pa.press('tab'), pp.copy(pessoa.est_numendfacul), pa.hotkey('ctrl', 'v')
                pa.press('tab'), pp.copy(pessoa.est_bairroendfacul), pa.hotkey('ctrl', 'v')
                # #clique em Outros
                try:
                    pa.click(pa.center(pa.locateOnScreen('./static/Outros.png')))
                except:
                    t.sleep(5)
                    pa.click(pa.center(pa.locateOnScreen('./static/Outros.png')))
                t.sleep(2), pa.write('CARGO GERAL')
                pa.press('tab'), pa.write(pessoa.cargo)
                t.sleep(1), pa.press('tab'), pa.write(pessoa.salario), pa.press('tab')
                if str(pessoa.tipo_contr) == 'Horista':
                    pa.press('1')
                else:
                    pa.press('5')
                pa.press('tab', 2), t.sleep(1), pa.write(str(pessoa.hr_sem)), pa.press('tab'), t.sleep(1)
                pa.write(str(pessoa.hr_mens))
                # #clique em Compatibilidade
                pa.click(pa.center(pa.locateOnScreen('./static/Compatibilidade3.png'))), t.sleep(1)
                # #clique em Compatibilidade de novo
                pa.click(pa.center(pa.locateOnScreen('./static/Compatibilidade2.png'))), t.sleep(1)
                pa.press('tab', 2), pa.write('9')
                # #clique em Salvar
                pa.click(pa.center(pa.locateOnScreen('./static/Salvarcadastro.png'))), t.sleep(10)
                # #clique em fechar novo cadastro
                pa.click(pa.center(pa.locateOnScreen('./static/Fecharnovo1.png'))), t.sleep(2)
                # #clique em fechar trabalhadores
                pa.click(pa.center(pa.locateOnScreen('./static/Fechartrab1.png'))), t.sleep(0.5)
                tkinter.messagebox.showinfo(
                    title='Cadastro ok!',
                    message='Cadastro editado com sucesso!'
                )


class Frame3(ttk.Frame):
    def __init__(self, container):
        super().__init__()
        self.hoje = datetime.today()
        self.caminhoaut = StringVar()
        self.nomeaut = StringVar()
        self.horarioaut = StringVar()
        self.cargoaut = StringVar()
        self.departamentoaut = StringVar()
        self.tipocontraut = StringVar()
        self.nomesplanaut = []
        self.labelescolh = ttk.Label(self, width=40, text="Escolher planilha de autônomos")
        self.labelescolh.grid(column=1, row=1, padx=25, pady=1, sticky=W)
        self.botaoescolh = ttk.Button(self, text="Escolha a planilha", command=self.selecionaraut)
        self.botaoescolh.grid(column=1, row=1, padx=350, pady=1, sticky=W)
        self.nomeaut = StringVar()
        self.cargo = StringVar()
        self.departamento = StringVar()
        self.nomesplanaut = []
        # aparecer dropdown com nomes da plan
        self.labelnomeaut = ttk.Label(self, width=20, text="Nome:")
        self.labelnomeaut.grid(column=1, row=10, padx=25, pady=1, sticky=W)
        self.combonomeaut = ttk.Combobox(self, values=self.nomesplanaut, textvariable=self.nomeaut, width=50)
        self.combonomeaut.grid(column=1, row=10, padx=125, pady=1, sticky=W)
        # aparecer entry para preencher matricula
        self.labelmatraut = ttk.Label(self, width=20, text="Matrícula:")
        self.labelmatraut.grid(column=1, row=11, padx=25, pady=1, sticky=W)
        self.entrymatraut = ttk.Entry(self, width=20)
        self.entrymatraut.grid(column=1, row=11, padx=125, pady=1, sticky=W)
        # aparecer entry para preencher admissao
        self.labeladmissaut = ttk.Label(self, width=20, text="Admissão:")
        self.labeladmissaut.grid(column=1, row=12, padx=25, pady=1, sticky=W)
        self.entryadmissaut = DateEntry(self, selectmode='day', year=self.hoje.year, month=self.hoje.month,
                                        day=self.hoje.day, locale='pt_BR')
        self.entryadmissaut.grid(column=1, row=12, padx=125, pady=1, sticky=W)
        # aparecer dropdown para escolher cargo
        self.labelcargoaut = ttk.Label(self, width=20, text="Cargo")
        self.labelcargoaut.grid(column=1, row=18, padx=25, pady=1, sticky=W)
        self.combocargoaut = ttk.Combobox(self, values=cargos, textvariable=self.cargo, width=50)
        self.combocargoaut.grid(column=1, row=18, padx=125, pady=1, sticky=W)
        # aparecer dropdown para escolher depto
        self.labeldeptoaut = ttk.Label(self, width=20, text="Departamento:")
        self.labeldeptoaut.grid(column=1, row=19, padx=25, pady=1, sticky=W)
        self.combodeptoaut = ttk.Combobox(self, values=departamentos, textvariable=self.departamento, width=50)
        self.combodeptoaut.grid(column=1, row=19, padx=125, pady=1, sticky=W)
        self.feitondeaut = IntVar()
        self.ondeaut = ttk.Checkbutton(self, text='Cadastro realizado fora da Cia.', variable=self.feitondeaut)
        self.ondeaut.grid(column=1, row=27, padx=26, pady=1, sticky=W)

        def carregaraut(local):
            planwb = l_w(local)
            plansh = planwb['Respostas ao formulário 1']
            lista = []
            for x, pessoa in enumerate(plansh):
                lista.append(f'{x + 1} - {pessoa[2].value}')
            self.combonomeaut.config(values=lista)

        self.botaocarregar = ttk.Button(self, text="Carregar planilha",
                                        command=lambda: [carregaraut(self.caminhoaut.get())])
        self.botaocarregar.grid(column=1, row=9, padx=350, pady=25, sticky=W)
        self.botaovalidarpis = ttk.Button(self, width=20, text="Validar PIS",
                                          command=lambda: [validarpis(self.caminhoaut.get(), self.combonomeaut.get())])
        self.botaovalidarpis.grid(column=1, row=10, padx=520, pady=1, sticky=W)

        self.botaocarregar = ttk.Button(self, width=20, text="Cadastrar autônomo",
                                        command=lambda: [
                                            cadastrar_autonomo(self.caminhoaut.get(), self.combonomeaut.get(),
                                                               self.entrymatraut.get(), self.entryadmissaut.get(),
                                                               self.combocargoaut.get(),
                                                               self.combodeptoaut.get(), self.feitondeaut.get())])
        self.botaocarregar.grid(column=1, row=28, padx=520, pady=1, sticky=W)

    def selecionaraut(self):
        try:
            caminhoplanaut = tkinter.filedialog.askopenfilename(title='Planilha Autônomos')
            self.caminhoaut.set(str(caminhoplanaut))
        except ValueError:
            pass


def cadastrar_autonomo(caminhoaut, nomeaut, matriculaaut, admissaoaut, cargoaut, deptoaut, ondeaut):
    sessions = sessionmaker(bind=engine)
    session = sessions()
    # Cadastro iniciado em casa
    wb = l_w(caminhoaut)
    sh = wb['Respostas ao formulário 1']
    num, name = nomeaut.strip().split(' - ')
    linha = int(num)
    lotacao = {'UNIDADE PARK SUL - QUALQUER DEPARTAMENTO': '0013', 'KIDS': '0010', 'MUSCULAÇÃO': '0007',
               'ESPORTES E LUTAS': '0008', 'CROSSFIT': '0012', 'GINÁSTICA': '0006', 'GESTANTES': '0008',
               'RECEPÇÃO': '0003',
               'FINANCEIRO': '0001', 'TI': '0001', 'MARKETING': '0001', 'MANUTENÇÃO': '0004'}
    aut = session.query(Colaborador).filter_by(matricula=matriculaaut).first()
    if aut:
        pass
    else:
        aut_cadastrado = Colaborador(
            matricula=matriculaaut, nome=name.upper(), admiss=admissaoaut,
            nascimento=str(sh[f'D{linha}'].value),
            pis=str(sh[f'S{linha}'].value).replace('.', '').replace('-', '').zfill(11),
            cpf=str(int(sh[f'P{linha}'].value)).zfill(11),
            rg=str(int(sh[f'Q{linha}'].value)),
            emissor='SSP/DF', email=str(sh[f'B{linha}'].value),
            genero=str(sh[f'E{linha}'].value), cor='9',
            instru=str(sh[f'F{linha}'].value),
            nacional='Brasileiro(a)', estado_civil='Solteiro(a)',
            endereco=str(sh[f'I{linha}'].value),
            num=str(sh[f'J{linha}'].value),
            bairro=str(sh[f'K{linha}'].value), cep=str(sh[f'L{linha}'].value).replace('.', '').replace('-', ''),
            cidade=str(sh[f'M{linha}'].value), cid_nas='Brasília - DF', uf='DF',
            cod_municipioend=municipios['DF']['Brasília'],
            tel=str(sh[f'O{linha}'].value).replace('(', '').replace(')', '').replace('.', '').replace('-', ''),
            depto=deptoaut, cargo=cargoaut,
        )
        session.add(aut_cadastrado)
        session.commit()
    # abrir cadastro no dexion e atualizar informações campo a campo
    pessoa = session.query(Colaborador).filter_by(matricula=matriculaaut).first()
    pa.click(pa.center(pa.locateOnScreen('./static/Dexion.png')))
    pa.press('alt'), pa.press('a'), pa.press('t'), t.sleep(2), pa.press(
        'i'), t.sleep(5), pa.write(str(pessoa.matricula)), pa.press('enter'), t.sleep(20)
    pp.copy(pessoa.nome), pa.hotkey('ctrl', 'v'), pa.press('tab'), pa.write(pessoa.cpf)
    t.sleep(1), pa.press('tab', 3), pa.write(pessoa.genero), pa.press('tab'), pa.write(pessoa.cor)
    t.sleep(1), pa.press('tab', 2), pa.write(pessoa.instru)
    t.sleep(1), pa.press('tab')
    if str(pessoa.estado_civil) == 'Casado(a)':
        pa.write('2')
        pa.press('tab', 6)
    else:
        pa.write('1')
        pa.press('tab', 5)
    pa.write(datetime.strftime(datetime.strptime(pessoa.nascimento, '%Y-%m-%d %H:%M:%S'), '%d%m%Y'))
    t.sleep(1), pa.press('tab'), pp.copy(pessoa.cid_nas), pa.hotkey('ctrl', 'v')
    # # clique em documentos
    pa.click(pa.center(pa.locateOnScreen('./static/Documentos.png')))
    pa.press('tab'), pa.write(str(pessoa.rg)), pa.press('tab'), pa.write(
        pessoa.emissor), pa.press('tab', 3), pa.write(pessoa.cod_municipioend), pa.press('tab')
    pa.write(pessoa.pis)
    # #clique em endereço
    pa.click(pa.center(pa.locateOnScreen('./static/Endereco.png')))
    pa.press('tab', 2), pp.copy(pessoa.endereco), pa.hotkey('ctrl', 'v'), pa.press(
        'tab'), pa.write(pessoa.num), pa.press('tab', 2)
    pp.copy(pessoa.bairro), pa.hotkey('ctrl', 'v'), pa.press('tab'), pp.copy(pessoa.cidade), pa.hotkey(
        'ctrl', 'v'), pa.press('tab'), pa.write(pessoa.uf)
    pa.press('tab'), pa.write(pessoa.cep), pa.press('tab'), pa.write(pessoa.cod_municipioend), pa.press(
        'tab'), pa.write(str(pessoa.tel)), pa.press('tab', 2), pa.write(pessoa.email)
    # #clique em dados contratuais
    pa.click(pa.center(pa.locateOnScreen('./static/Contratuais.png')))
    pa.press('tab'), pa.write(str(pessoa.admiss).replace('/', '')), t.sleep(1)
    pa.press('tab'), pa.write('7')
    # #clique em Outros
    try:
        pa.click(pa.center(pa.locateOnScreen('./static/Outros.png')))
    except:
        t.sleep(5)
        pa.click(pa.center(pa.locateOnScreen('./static/Outros.png')))
    t.sleep(2), pa.write('CARGO GERAL')
    pa.press('tab'), pa.write(pessoa.cargo)
    # #clique em eventos trabalhistas
    pa.click(pa.center(pa.locateOnScreen('./static/EVTrab.png')))
    t.sleep(1)
    # #clique em lotação
    pa.click(pa.center(pa.locateOnScreen('./static/Lotacoes.png')))
    pa.press('tab'), pa.press('tab'), pa.write('i'), pa.write(str(pessoa.admiss).replace('/', ''))
    t.sleep(1), pa.press('enter'), t.sleep(1)
    pp.copy(lotacao[f'{pessoa.depto}']), pa.hotkey('ctrl', 'v'), pa.press('enter'), pa.write('3')
    # #clique em salvar lotação
    pa.click(pa.center(pa.locateOnScreen('./static/Salvarbtn.png'))), t.sleep(1)
    # #clique em fechar lotação
    try:
        pa.click(pa.center(pa.locateOnScreen('./static/Fecharlot.png'))), t.sleep(1)
    except:
        t.sleep(4)
        pa.click(pa.center(pa.locateOnScreen('./static/Fecharlot.png')))
    # #clique em Compatibilidade
    pa.click(pa.center(pa.locateOnScreen('./static/Compatibilidade3.png'))), t.sleep(1)
    # #clique em Compatibilidade de novo
    pa.click(pa.center(pa.locateOnScreen('./static/Compatibilidade2.png'))), t.sleep(1)
    pa.press('tab', 2), pa.write('13')
    # #clique em Salvar
    pa.click(pa.center(pa.locateOnScreen('./static/Salvarcadastro.png'))), t.sleep(10)
    # #clique em fechar novo cadastro
    pa.click(pa.center(pa.locateOnScreen('./static/Fecharnovo1.png'))), t.sleep(2)
    # #clique em fechar trabalhadores
    pa.click(pa.center(pa.locateOnScreen('./static/Fechartrab1.png'))), t.sleep(0.5)
    tkinter.messagebox.showinfo(
        title='Cadastro ok!',
        message='Cadastro realizado com sucesso!'
    )


def validarpis(local, nome):
    wb = l_w(local, read_only=False)
    sh = wb['Respostas ao formulário 1']
    num, name = nome.strip().split(' - ')
    x = int(num)
    pis = str(sh[f'S{x}'].value).replace('-', '').replace('.', '').zfill(11)
    v1 = int(pis[0]) * 3
    v2 = int(pis[1]) * 2
    v3 = int(pis[2]) * 9
    v4 = int(pis[3]) * 8
    v5 = int(pis[4]) * 7
    v6 = int(pis[5]) * 6
    v7 = int(pis[6]) * 5
    v8 = int(pis[7]) * 4
    v9 = int(pis[8]) * 3
    v10 = int(pis[9]) * 2
    d = int(pis[10])
    soma = v1 + v2 + v3 + v4 + v5 + v6 + v7 + v8 + v9 + v10
    divisao = soma % 11
    resultado = 11 - divisao
    if resultado != d:
        if resultado == 10 & d == 0 | resultado == 11 & d == 0:
            pass
        else:
            tkinter.messagebox.showinfo(
                title='Erro!',
                message='PIS inválido!'
            )
    else:
        tkinter.messagebox.showinfo(
            title='Ok!',
            message='PIS ok!'
        )


class Frame4(ttk.Frame):
    def __init__(self, container):
        super().__init__()


if __name__ == '__main__':
    app = MainApplication()
    app.mainloop()
