from datetime import datetime as dt, timedelta as td
import docx
from docx.shared import Pt, Cm
import docx2pdf
import locale
import numpy as np
from openpyxl import load_workbook as l_w
from openpyxl.styles import NamedStyle
import os
import pandas as pd
import pyautogui as pa
import pyperclip as pp
from PIL import ImageGrab
from src.models.modelsponto import engineponto, BasePonto
from src.models.models import Colaborador, engine
from sqlalchemy.orm import sessionmaker
from src.models.dados_servd import rede
import tkinter.filedialog
from tkinter import messagebox
import tkinter.filedialog
import time
import win32com.client as client

locale.setlocale(locale.LC_MONETARY, 'pt_BR.UTF-8')
locale.setlocale(locale.LC_ALL, 'pt_BR.UTF-8')
locale.setlocale(locale.LC_MONETARY, 'pt_BR.UTF-8')


def gerar_excel_from_ponto_secullum() -> None:
    """
    This function generates Excel files from data extracted from external aplication.
    Its generetaes one Excel file per employee with schedules information.
    :return: one .xlsx file per employee
    """
    locale.setlocale(locale.LC_ALL, 'pt_pt.UTF-8')
    data_inicial = dt.strptime('01/10/2022', '%d/%m/%Y')
    data_final = dt.strptime('13/11/2022', '%d/%m/%Y')

    # Ler planilha geral
    geral = pd.read_excel('../Ponto/xls/zzPonto Geral.xls')
    geral = geral.rename(
        columns={'CARTÃO PONTO': 'Dia', 'Unnamed: 1': 0, 'Unnamed: 2': 1, 'Unnamed: 3': 2, 'Unnamed: 4': 3,
                 'Unnamed: 5': 4,
                 'Unnamed: 6': 5})
    geral = geral.drop(['Unnamed: 7', 'Unnamed: 8'], axis=1)
    geral = geral[geral.Dia.notnull()]

    # Pegar index onde aparece 'Nome'
    linhasnomes = geral.index[geral['Dia'].str.contains('Nome')]

    # salvar plan com nome do funcionário pasta ponto (na pasta automação)
    for linha in linhasnomes:
        geral = geral.rename(
            columns={'CARTÃO PONTO': 'Dia', 'Unnamed: 1': 0, 'Unnamed: 2': 1, 'Unnamed: 3': 2, 'Unnamed: 4': 3,
                     'Unnamed: 5': 4, 'Unnamed: 6': 5})
        geral = geral[geral.Dia.notnull()]
        geral = geral[geral['Dia'].str.contains(' - ') | geral['Dia'].str.contains('Nome')]
        geral2 = geral.loc[linha:(linha + (linhasnomes[1] - linhasnomes[0] - 1))] \
            .to_excel(f'../Ponto/xls/{geral[0][linha]}.xlsx')

        # Verifica a hora exata na planilha zzBase.xlsx
        wb = l_w(f'../Ponto/xls/zzBase.xlsx')
        ws = wb.active
        for row in ws.rows:
            for cell in row:
                if cell.value == f'{geral[0][linha]}':
                    if ws.cell(row=cell.row, column=5).value is None:
                        ent1 = dt.strptime('00:00:00', '%H:%M:%S')
                    else:
                        ent1 = dt.strptime(str(ws.cell(row=cell.row, column=5).value), '%H:%M:%S')
                    if ws.cell(row=cell.row, column=6).value is None:
                        sai1 = dt.strptime('00:00:00', '%H:%M:%S')
                    else:
                        sai1 = dt.strptime(str(ws.cell(row=cell.row, column=6).value), '%H:%M:%S')
                    if ws.cell(row=cell.row, column=7).value is None:
                        ent2 = dt.strptime('00:00:00', '%H:%M:%S')
                    else:
                        ent2 = dt.strptime(str(ws.cell(row=cell.row, column=7).value), '%H:%M:%S')
                    if ws.cell(row=cell.row, column=8).value is None:
                        sai2 = dt.strptime('00:00:00', '%H:%M:%S')
                    else:
                        sai2 = dt.strptime(str(ws.cell(row=cell.row, column=8).value), '%H:%M:%S')
                    if ws.cell(row=cell.row, column=9).value is None:
                        ent3 = dt.strptime('00:00:00', '%H:%M:%S')
                    else:
                        ent3 = dt.strptime(str(ws.cell(row=cell.row, column=9).value), '%H:%M:%S')
                    if ws.cell(row=cell.row, column=10).value is None:
                        sai3 = dt.strptime('00:00:00', '%H:%M:%S')
                    else:
                        sai3 = dt.strptime(str(ws.cell(row=cell.row, column=10).value), '%H:%M:%S')
        # relacionar dia da semana
        entradacerta = td(hours=ent1.hour, minutes=ent1.minute, seconds=ent1.second)
        saidacerta = td(hours=sai1.hour, minutes=sai1.minute, seconds=sai1.second)
        entradacerta2 = td(hours=ent2.hour, minutes=ent2.minute, seconds=ent2.second)
        saidacerta2 = td(hours=sai2.hour, minutes=sai2.minute, seconds=sai2.second)
        entradacerta3 = td(hours=ent3.hour, minutes=ent3.minute, seconds=ent3.second)
        saidacerta3 = td(hours=sai3.hour, minutes=sai3.minute, seconds=sai3.second)
        regra = td(hours=0, minutes=10, seconds=0)

        # Salvar planilha adicionando colunas de diferenças
        plan = pd.read_excel(f'../Ponto/xls/{geral[0][linha]}.xlsx')
        plan = plan[plan['Dia'].str.contains(' - ')]
        plan = plan.drop(['Unnamed: 0'], axis=1)
        plan = plan.rename(columns={'Dia': 'Data'})
        plan[0] = pd.to_timedelta(plan[0].astype(str))
        plan[1] = pd.to_timedelta(plan[1].astype(str))
        plan[2] = pd.to_timedelta(plan[2].astype(str))
        plan[3] = pd.to_timedelta(plan[3].astype(str))
        plan[4] = pd.to_timedelta(plan[4].astype(str))
        plan[5] = pd.to_timedelta(plan[5].astype(str))
        plan['Entradacerta'] = entradacerta
        plan['Saídacerta'] = saidacerta
        plan['DifEntr'] = abs(plan[0] - plan['Entradacerta'])
        plan['DifSaida'] = abs(plan[1] - plan['Saídacerta'])
        dif = plan
        dif = dif[['Data', 'DifEntr', 'DifSaida']]
        dif = dif.loc[(dif['DifEntr'] >= regra) | (dif['DifEntr'] + dif['DifSaida'] >= regra)]
        dif = dif.astype(str)
        dif['DifEntr'] = dif['DifEntr'].map(
            lambda x: dt.strftime(dt.strptime(str(x).replace('0 days ', '').replace('NaT', '00:00:00'), '%H:%M:%S'),
                      '%H hora(s) e %M minutos')
        )
        dif['DifSaida'] = dif['DifSaida'].map(
            lambda x: dt.strftime(dt.strptime(str(x).replace('0 days ', '').replace('NaT', '00:00:00'),
                                                          '%H:%M:%S'), '%H hora(s) e %M minutos')
        )
        dif['DifEntr'] = dif['DifEntr'].map(
            lambda x: str(x).replace('00 hora(s) e 00 minutos', '-').replace('00 hora(s) e ', '')
        )
        dif['DifSaida'] = dif['DifSaida'].map(
            lambda x: str(x).replace('00 hora(s) e 00 minutos', '-').replace('00 hora(s) e ', '')
        )
        plan = plan[['Data', 0, 1, 2, 3]]
        plan = plan.merge(dif, on='Data', how='outer')
        plan = plan.astype(str)
        # plan = plan.fillna('-')
        # plan['Data'] = plan['Data'].map(lambda x: x.rstrip('- qua qui sex sab ter seg sá dom'))
        plan[0] = plan[0].map(lambda x: x.lstrip('0 days 00:'))
        plan[1] = plan[1].map(lambda x: x.lstrip('0 days 00:'))
        plan[2] = plan[2].map(lambda x: x.lstrip('0 days 00:'))
        plan[3] = plan[3].map(lambda x: x.lstrip('0 days 00:'))
        plan = plan.replace('NaT', '-').replace('nan', '-')
        plan = plan.rename(columns={'Data': 'Dia', 0: 'Entrada1', 1: 'Saída 1', 2: 'Entrada 2', 3: 'Saída 2', 'DifEntr':
            'Diferença Entrada', 'DifSaida': 'Diferença Saída 1'})
        plan = plan[plan.Entrada1 != '-']
        plan = plan.rename(columns={'Entrada1': 'Entrada 1'})
        plan = plan.astype(str)
        plan['Dia'] = plan['Dia'].map(lambda x: dt.strptime(x, '%d/%m/%y - %a'))
        plan = plan[plan.Dia >= data_inicial]
        plan = plan[plan.Dia <= data_final]
        plan['Dia'] = plan['Dia'].map(lambda x: dt.strftime(x, '%d/%m/%y - %a'))
        plan = plan.rename(columns={'Dia': 'Data'})
        plan = plan.set_index(['Data'])
        plan = plan.to_excel(f'../Ponto/xls/{geral[0][linha]}.xlsx')
        func = l_w(f'../Ponto/xls/{geral[0][linha]}.xlsx', read_only=False)
        sh = func['Sheet1']
        sh.column_dimensions['A'].width = 15
        sh.column_dimensions['F'].width = 22
        sh.column_dimensions['G'].width = 22
        sh.column_dimensions['H'].width = 22
        sh.column_dimensions['I'].width = 22
        sh.column_dimensions['J'].width = 22
        sh.column_dimensions['K'].width = 22
        func.save(f'../Ponto/xls/{geral[0][linha]}.xlsx')


def gerar_relatorios_de_atrasos_estagiarios():
    # Ler arquivo txt dos registror rejeitados
    geral = pd.read_csv('Rejeitados.txt', sep=' ', header=None, encoding='iso8859-1')
    geral = geral.rename(columns={0: 'matricula', 16: 'data', 17: 'dia', 18: 'hora'})
    geral = geral[geral.dia != 'Batida']
    geral = geral[geral.matricula < 9999]
    mat = []
    mat_unicas = []
    for matricula in geral['matricula']:
        mat.append(matricula)
        mat_unicas = list(set(mat))
    for matr in mat_unicas:
        geral2 = geral[geral.matricula == matr]
        geral2 = geral2.set_index('matricula')
        geral2 = geral2.dropna(axis='columns')
        geral2 = geral2.drop(geral2.iloc[:, [2, 3, 4, 5]], axis=1)
        geral2['dia'] = pd.to_datetime(geral2['dia'], format='%d/%m/%Y')
        geral2['dia'] = geral2['dia'].apply(lambda y: dt.strftime(y, '%d/%m/%Y'))
        geral2['hora'] = geral2['hora'].apply(lambda x: f'{x}:00')
        geral2['hora'] = pd.to_timedelta(geral2['hora'])
        # geral2 = geral2.set_index('dia')
        # geral2 = geral2.groupby('dia')
        geral2.to_excel(f'../Ponto/xls/Ponto Estágio - {matr}.xlsx')
        wb = l_w(f'../Ponto/xls/Ponto Estágio - {matr}.xlsx')
        sh = wb['Sheet1']
        sh['A1'].value = 'Matrícula'
        sh['B1'].value = 'Data'
        sh['C1'].value = 'Entrada 1'
        sh['D1'].value = 'Saída 1'
        sh['E1'].value = 'Entrada 2'
        sh['F1'].value = 'Saída 2'
        sh['G1'].value = 'Entrada 3'
        sh['H1'].value = 'Saída 3'
        x = 2
        for row in sh:
            if sh[f'B{x}'].value == sh[f'B{x - 1}'].value:
                if sh[f'D{x - 1}'].value is None:
                    sh[f'D{x - 1}'].value = sh[f'C{x}'].value
                    sh.delete_rows(x, 1)
                else:
                    if sh[f'E{x - 1}'].value is None:
                        sh[f'E{x - 1}'].value = sh[f'C{x}'].value
                        sh.delete_rows(x, 1)
                    else:
                        if sh[f'F{x - 1}'].value is None:
                            sh[f'F{x - 1}'].value = sh[f'C{x}'].value
                            sh.delete_rows(x, 1)
                        else:
                            if sh[f'G{x - 1}'].value is None:
                                sh[f'G{x - 1}'].value = sh[f'C{x}'].value
                                sh.delete_rows(x, 1)
                            else:
                                if sh[f'H{x - 1}'].value is None:
                                    sh[f'H{x - 1}'].value = sh[f'C{x}'].value
                                    sh.delete_rows(x, 1)
        for row in sh:
            if sh[f'B{x}'].value == sh[f'B{x - 1}'].value:
                if sh[f'D{x - 1}'].value is None:
                    sh[f'D{x - 1}'].value = sh[f'C{x}'].value
                    sh.delete_rows(x, 1)
                else:
                    if sh[f'E{x - 1}'].value is None:
                        sh[f'E{x - 1}'].value = sh[f'C{x}'].value
                        sh.delete_rows(x, 1)
                    else:
                        if sh[f'F{x - 1}'].value is None:
                            sh[f'F{x - 1}'].value = sh[f'C{x}'].value
                            sh.delete_rows(x, 1)
                        else:
                            if sh[f'G{x - 1}'].value is None:
                                sh[f'G{x - 1}'].value = sh[f'C{x}'].value
                                sh.delete_rows(x, 1)
                            else:
                                if sh[f'H{x - 1}'].value is None:
                                    sh[f'H{x - 1}'].value = sh[f'C{x}'].value
                                    sh.delete_rows(x, 1)
            for row in sh:
                if sh[f'B{x}'].value == sh[f'B{x - 1}'].value:
                    if sh[f'D{x - 1}'].value is None:
                        sh[f'D{x - 1}'].value = sh[f'C{x}'].value
                        sh.delete_rows(x, 1)
                    else:
                        if sh[f'E{x - 1}'].value is None:
                            sh[f'E{x - 1}'].value = sh[f'C{x}'].value
                            sh.delete_rows(x, 1)
                        else:
                            if sh[f'F{x - 1}'].value is None:
                                sh[f'F{x - 1}'].value = sh[f'C{x}'].value
                                sh.delete_rows(x, 1)
                            else:
                                if sh[f'G{x - 1}'].value is None:
                                    sh[f'G{x - 1}'].value = sh[f'C{x}'].value
                                    sh.delete_rows(x, 1)
                                else:
                                    if sh[f'H{x - 1}'].value is None:
                                        sh[f'H{x - 1}'].value = sh[f'C{x}'].value
                                        sh.delete_rows(x, 1)
            x += 1

        estilo_data = NamedStyle(name='data', number_format='DD/MM/YYYY')
        estilo_hora = NamedStyle(name='hora', number_format='HH:MM:SS')
        for cell in sh['B']:
            sh[f'B{int(cell.row)}'].style = estilo_data
        for item in sh['C']:
            sh[f'C{int(item.row)}'].style = estilo_hora
        for item in sh['D']:
            sh[f'D{int(item.row)}'].style = estilo_hora
        for item in sh['E']:
            sh[f'E{int(item.row)}'].style = estilo_hora
        for item in sh['F']:
            sh[f'F{int(item.row)}'].style = estilo_hora
        for item in sh['G']:
            sh[f'G{int(item.row)}'].style = estilo_hora
        for item in sh['H']:
            sh[f'H{int(item.row)}'].style = estilo_hora
        sh.column_dimensions['A'].width = 11
        sh.column_dimensions['B'].width = 11
        sh.column_dimensions['C'].width = 11
        sh.column_dimensions['D'].width = 11
        sh.column_dimensions['E'].width = 11
        sh.column_dimensions['F'].width = 11
        sh.column_dimensions['G'].width = 11
        sh.column_dimensions['H'].width = 11
        wb2 = l_w(f'../Ponto/xls/zzBase.xlsx')
        ws2 = wb2.active
        for row in ws2.rows:
            for cell in row:
                if cell.value == matr:
                    nome = ws2.cell(row=cell.row, column=3).value
                    wb.save(f'../Ponto/xls/Ponto Estágio - {nome}.xlsx')
                    os.remove(f'../Ponto/xls/Ponto Estágio - {matr}.xlsx')
        plan = l_w(f'../Ponto/xls/Ponto Estágio - {nome}.xlsx')
        splan = plan.active
        # Verifica a hora exata na planilha zzBase.xlsx
        wb = l_w(f'../Ponto/xls/zzBase.xlsx')
        ws = wb.active
        for row in ws.rows:
            for cell in row:
                if cell.value == f'{nome}':
                    for rowsplan in splan.rows:
                        for cellplan in rowsplan:
                            if str(splan.cell(row=cellplan.row, column=2).value) == 'Data':
                                pass
                            # se o dia no ponto for segunda
                            if dt.strptime(str(splan.cell(row=cellplan.row, column=2).value),
                                           '%d/%m/%Y').isoweekday() == 1:
                                pass
                            # se o dia no ponto for terça
                            if dt.strptime(str(splan.cell(row=cellplan.row, column=2).value),
                                           '%d/%m/%Y').isoweekday() == 2:
                                pass
                            # se o dia no ponto for quarta
                            if dt.strptime(str(splan.cell(row=cellplan.row, column=2).value),
                                           '%d/%m/%Y').isoweekday() == 3:
                                pass
                            # se o dia no ponto for quinta
                            if dt.strptime(str(splan.cell(row=cellplan.row, column=2).value),
                                           '%d/%m/%Y').isoweekday() == 4:
                                pass
                            # se o dia no ponto for sexta
                            if dt.strptime(str(splan.cell(row=cellplan.row, column=2).value),
                                           '%d/%m/%Y').isoweekday() == 5:
                                pass
                            # se o dia no ponto for sábado
                            if dt.strptime(str(splan.cell(row=cellplan.row, column=2).value),
                                           '%d/%m/%Y').isoweekday() == 6:
                                pass
                            # se o dia no ponto for domingo
                            if dt.strptime(str(splan.cell(row=cellplan.row, column=2).value),
                                           '%d/%m/%Y').isoweekday() == 7:
                                pass
                    if ws.cell(row=cell.row, column=5).value is None:
                        ent1 = dt.strptime('00:00:00', '%H:%M:%S')
                    else:
                        ent1 = dt.strptime(str(ws.cell(row=cell.row, column=5).value), '%H:%M:%S')
                    if ws.cell(row=cell.row, column=6).value is None:
                        sai1 = dt.strptime('00:00:00', '%H:%M:%S')
                    else:
                        sai1 = dt.strptime(str(ws.cell(row=cell.row, column=6).value), '%H:%M:%S')
                    if ws.cell(row=cell.row, column=7).value is None:
                        ent2 = dt.strptime('00:00:00', '%H:%M:%S')
                    else:
                        ent2 = dt.strptime(str(ws.cell(row=cell.row, column=7).value), '%H:%M:%S')
                    if ws.cell(row=cell.row, column=8).value is None:
                        sai2 = dt.strptime('00:00:00', '%H:%M:%S')
                    else:
                        sai2 = dt.strptime(str(ws.cell(row=cell.row, column=8).value), '%H:%M:%S')
                    if ws.cell(row=cell.row, column=9).value is None:
                        ent3 = dt.strptime('00:00:00', '%H:%M:%S')
                    else:
                        ent3 = dt.strptime(str(ws.cell(row=cell.row, column=9).value), '%H:%M:%S')
                    if ws.cell(row=cell.row, column=10).value is None:
                        sai3 = dt.strptime('00:00:00', '%H:%M:%S')
                    else:
                        sai3 = dt.strptime(str(ws.cell(row=cell.row, column=10).value), '%H:%M:%S')
                # relacionar dia da semana
                entradacerta = td(hours=ent1.hour, minutes=ent1.minute, seconds=ent1.second)
                saidacerta = td(hours=sai1.hour, minutes=sai1.minute, seconds=sai1.second)
                entradacerta2 = td(hours=ent2.hour, minutes=ent2.minute, seconds=ent2.second)
                saidacerta2 = td(hours=sai2.hour, minutes=sai2.minute, seconds=sai2.second)
                entradacerta3 = td(hours=ent3.hour, minutes=ent3.minute, seconds=ent3.second)
                saidacerta3 = td(hours=sai3.hour, minutes=sai3.minute, seconds=sai3.second)
                regra = td(hours=0, minutes=10, seconds=0)
    # # procurar na planilha base.xlsx nome e-mail matricula no ponto e horarios de entrada e saida


def gerar_excel_ponto_estagiarios():
    # Ler arquivo txt dos registror rejeitados
    geral = pd.read_csv(
        r'C:\Users\Felipe Rodrigues\Desktop\Relatorios Ponto\Rej - 16-03-2023.txt', sep=' ', header=None,
        encoding='iso8859-1'
    )
    geral = geral.rename(columns={0: 'matricula', 16: 'data', 17: 'dia', 18: 'hora'})
    geral = geral[geral.dia != 'Batida']
    geral = geral[geral.matricula < 9999]
    mat = []
    mat_unicas = []
    for matricula in geral['matricula']:
        mat.append(matricula)
        mat_unicas = list(set(mat))
    for matr in mat_unicas:
        geral2 = geral[geral.matricula == matr]
        geral2 = geral2.set_index('matricula')
        geral2 = geral2.dropna(axis='columns')
        # print(geral2)
        geral2 = geral2.drop(geral2.iloc[:, [2, 3, 4]], axis=1)
        geral2['dia'] = pd.to_datetime(geral2['dia'], format='%d/%m/%Y')
        geral2['dia'] = geral2['dia'].apply(lambda y: dt.strftime(y, '%d/%m/%Y'))
        geral2['hora'] = geral2['hora'].apply(lambda x: f'{x}:00')
        geral2['hora'] = pd.to_timedelta(geral2['hora'])
        # geral2 = geral2.set_index('dia')
        # geral2 = geral2.groupby('dia')
        geral2.to_excel(f'../Ponto/xls/Ponto Estágio - {matr}.xlsx')
        wb = l_w(f'../Ponto/xls/Ponto Estágio - {matr}.xlsx')
        sh = wb['Sheet1']
        sh['A1'].value = 'Matrícula'
        sh['B1'].value = 'Data'
        sh['C1'].value = 'Entrada 1'
        sh['D1'].value = 'Saída 1'
        sh['E1'].value = 'Entrada 2'
        sh['F1'].value = 'Saída 2'
        sh['G1'].value = 'Entrada 3'
        sh['H1'].value = 'Saída 3'
        x = 2
        for row in sh:
            if sh[f'B{x}'].value == sh[f'B{x - 1}'].value:
                if sh[f'D{x - 1}'].value is None:
                    sh[f'D{x - 1}'].value = sh[f'C{x}'].value
                    sh.delete_rows(x, 1)
                else:
                    if sh[f'E{x - 1}'].value is None:
                        sh[f'E{x - 1}'].value = sh[f'C{x}'].value
                        sh.delete_rows(x, 1)
                    else:
                        if sh[f'F{x - 1}'].value is None:
                            sh[f'F{x - 1}'].value = sh[f'C{x}'].value
                            sh.delete_rows(x, 1)
                        else:
                            if sh[f'G{x - 1}'].value is None:
                                sh[f'G{x - 1}'].value = sh[f'C{x}'].value
                                sh.delete_rows(x, 1)
                            else:
                                if sh[f'H{x - 1}'].value is None:
                                    sh[f'H{x - 1}'].value = sh[f'C{x}'].value
                                    sh.delete_rows(x, 1)
        for row in sh:
            if sh[f'B{x}'].value == sh[f'B{x - 1}'].value:
                if sh[f'D{x - 1}'].value is None:
                    sh[f'D{x - 1}'].value = sh[f'C{x}'].value
                    sh.delete_rows(x, 1)
                else:
                    if sh[f'E{x - 1}'].value is None:
                        sh[f'E{x - 1}'].value = sh[f'C{x}'].value
                        sh.delete_rows(x, 1)
                    else:
                        if sh[f'F{x - 1}'].value is None:
                            sh[f'F{x - 1}'].value = sh[f'C{x}'].value
                            sh.delete_rows(x, 1)
                        else:
                            if sh[f'G{x - 1}'].value is None:
                                sh[f'G{x - 1}'].value = sh[f'C{x}'].value
                                sh.delete_rows(x, 1)
                            else:
                                if sh[f'H{x - 1}'].value is None:
                                    sh[f'H{x - 1}'].value = sh[f'C{x}'].value
                                    sh.delete_rows(x, 1)
            for row in sh:
                if sh[f'B{x}'].value == sh[f'B{x - 1}'].value:
                    if sh[f'D{x - 1}'].value is None:
                        sh[f'D{x - 1}'].value = sh[f'C{x}'].value
                        sh.delete_rows(x, 1)
                    else:
                        if sh[f'E{x - 1}'].value is None:
                            sh[f'E{x - 1}'].value = sh[f'C{x}'].value
                            sh.delete_rows(x, 1)
                        else:
                            if sh[f'F{x - 1}'].value is None:
                                sh[f'F{x - 1}'].value = sh[f'C{x}'].value
                                sh.delete_rows(x, 1)
                            else:
                                if sh[f'G{x - 1}'].value is None:
                                    sh[f'G{x - 1}'].value = sh[f'C{x}'].value
                                    sh.delete_rows(x, 1)
                                else:
                                    if sh[f'H{x - 1}'].value is None:
                                        sh[f'H{x - 1}'].value = sh[f'C{x}'].value
                                        sh.delete_rows(x, 1)
            x += 1

        estilo_data = NamedStyle(name='data', number_format='DD/MM/YYYY')
        estilo_hora = NamedStyle(name='hora', number_format='HH:MM:SS')
        for cell in sh['B']:
            sh[f'B{int(cell.row)}'].style = estilo_data
        for item in sh['C']:
            sh[f'C{int(item.row)}'].style = estilo_hora
        for item in sh['D']:
            sh[f'D{int(item.row)}'].style = estilo_hora
        for item in sh['E']:
            sh[f'E{int(item.row)}'].style = estilo_hora
        for item in sh['F']:
            sh[f'F{int(item.row)}'].style = estilo_hora
        for item in sh['G']:
            sh[f'G{int(item.row)}'].style = estilo_hora
        for item in sh['H']:
            sh[f'H{int(item.row)}'].style = estilo_hora
        sh.column_dimensions['A'].width = 11
        sh.column_dimensions['B'].width = 11
        sh.column_dimensions['C'].width = 11
        sh.column_dimensions['D'].width = 11
        sh.column_dimensions['E'].width = 11
        sh.column_dimensions['F'].width = 11
        sh.column_dimensions['G'].width = 11
        sh.column_dimensions['H'].width = 11
        wb2 = l_w(f'../Ponto/xls/zzBase.xlsx')
        ws2 = wb2.active
        for row in ws2.rows:
            for cell in row:
                if cell.value == matr:
                    nome = ws2.cell(row=cell.row, column=3).value
                    linha = int((len(sh['A']) / 2) + 1)
                    wb.save(f'../Ponto/xls/Ponto Estágio - {nome}.xlsx')
                    os.remove(f'../Ponto/xls/Ponto Estágio - {matr}.xlsx')
                    # t.sleep(2)
                    # inserir código de envio de email aqui
                    excel = client.Dispatch('Excel.Application')
                    plan = excel.Workbooks.Open(
                        r'C:\Users\Felipe Rodrigues\PycharmProjects\AutomacaoCia\Ponto\xls\Ponto Estágio - {}.xlsx'
                        .format(nome)
                    )
                    folha = plan.Sheets['Sheet1']
                    excel.visible = 0
                    copyrange = folha.Range(f'B1:F{linha}')
                    copyrange.CopyPicture(Format=2)
                    ImageGrab.grabclipboard().save(f'Ponto {str(nome).title().split(" ")[0]}.png')
                    excel.Quit()
                    html_body = f'''
                    <p>Olá {str(nome).title().split(" ")[0]},</p>
                    <p> Segue em anexo seu relatório de ponto do mês Fevereiro/2023.</p><br>
                    Atenciosamente,<br>
                    Felipe Rodrigues,<br>
                    '''
                    outlook = client.Dispatch('Outlook.Application')
                    message = outlook.CreateItem(0)
                    message.To = 'felipe.rodrigs09@gmail.com'
                    message.Subject = 'Relatório de Ponto'
                    message.HTMLBody = html_body
                    message.Attachments.Add(
                        f'C:\\Users\\{os.getlogin()}\\PycharmProjects\\AutomacaoCia\\Ponto\\'
                        f'Ponto {str(nome).title().split(" ")[0]}.png'
                    )
                    message.Send()
    # # procurar na planilha base.xlsx nome e-mail matricula no ponto e horarios de entrada e saida


def cadastrar_funcionario_no_secullum():
    wb = l_w('PlanPonto.xlsx')
    sh = wb['Planilha1']
    x = 2
    while x <= len(sh['A']):
        matricula = str(sh[f'A{x}'].value)
        nome = str(sh[f'B{x}'].value)
        pis = str(sh[f'C{x}'].value)
        horario = '1'
        funcao = str(sh[f'D{x}'].value)
        depto = str(sh[f'E{x}'].value)
        admiss = str(sh[f'F{x}'].value)
        # # clicar incluir
        pa.click(-1507, 147), time.sleep(3)
        pa.write(matricula), pa.press('tab')
        pp.copy(nome), pa.hotkey('ctrl', 'v'), pa.press('tab')
        pa.write(pis), pa.press('tab', 4)
        pa.write(horario), pa.press('tab')
        pp.copy(funcao), pa.hotkey('ctrl', 'v'), pa.press('tab')
        pp.copy(depto), pa.hotkey('ctrl', 'v'), pa.press('tab')
        pa.write(admiss)
        # click concluir
        pa.click(-1492, 617)
        time.sleep(2)
        x += 1


def gerar_relatorios_ponto_pdf(arq: str, datai: str, dataf: str, estag: int):
    """
    Generates working time reports in .pdf trought analysis of file .AFD.
    :param arq: path of .AFD file.
    :param datai: Beginning date.
    :param dataf: End date.
    :param estag: Intern
    :return: Worker reports.
    """
    # edit arq
    edic1 = open(arq, 'r')
    linhas = edic1.readlines()
    novalinha = [x for x in linhas if len(x) <= 39]
    out = open(arq, 'w')
    out.writelines(novalinha)
    out.close()
    arquivo = arq

    # define excel plans to work with
    base = os.path.relpath(rf'C:\Users\{os.getlogin()}\PycharmProjects\AutomacaoCia\src\models\static\files\zzBase.xlsx')
    wb = l_w(base)
    sh = wb['Funcionários e e-mail']
    hrsflh = os.path.relpath(rf'C:\Users\{os.getlogin()}\PycharmProjects\AutomacaoCia\src\models\static\files\Hrs Folha.xlsx')
    pl = l_w(hrsflh)
    fl = pl['Planilha1']
    dataipt = datai.replace('/', '.')
    datafpt = dataf.replace('/', '.')
    dia, mes, ano = dataf.split('/')
    datainicio = dt.strptime(datai, '%d/%m/%Y')
    datafim = dt.strptime(dataf, '%d/%m/%Y')
    mesext = {'01': 'JAN', '02': 'FEV', '03': 'MAR', '04': 'ABR', '05': 'MAI', '06': 'JUN',
              '07': 'JUL', '08': 'AGO', '09': 'SET', '10': 'OUT', '11': 'NOV', '12': 'DEZ'}

    def intervalo(inicio, fim):
        for n in range(int((fim - inicio).days) + 1):
            yield [inicio + td(n), np.nan, np.nan, np.nan, np.nan, np.nan, np.nan]

    # creating dicts to store employees datas
    planbase = {}
    planmat = {}
    planfunc = {}
    plandept = {}
    planemail = {}

    sessions = sessionmaker(engineponto)
    session = sessions()
    if estag == 1:
        nomes = session.query(BasePonto).filter_by(cargo='ESTAGIARIA').all()
        nomes2 = session.query(BasePonto).filter_by(cargo='ESTAGIARIO').all()
        nomes = nomes + nomes2
    else:
        nomes = session.query(BasePonto).all()

    for pessoa in nomes:
        planbase.update({pessoa.matrponto: pessoa.nome})
        planmat.update({pessoa.matrponto: pessoa.matricula})
        planfunc.update({pessoa.matrponto: pessoa.cargo})
        plandept.update({pessoa.matrponto: pessoa.departamento})
        planemail.update({pessoa.matrponto: pessoa.email})

    planbase.update({'21058461070': 'Teste Administrador'})
    planmat.update({'21058461070': '456'})
    planfunc.update({'21058461070': 'Gerente RH'})
    plandept.update({'21058461070': 'RH'})
    planemail.update({'21058461070': 'felipe.rodrigues@ciaathletica.com.br'})

    geral = pd.read_csv(arquivo, sep=' ', header=None, encoding='iso8859-1')
    geral.dropna(axis=1, inplace=True)
    geral = geral.rename(columns={0: 'Dados'})

    # dividir ultimos 11 caracteres em outra col
    geral['Matricula'] = geral['Dados'].str[-11:]
    geral = geral.drop('Dados', axis=1)
    mat = []
    matriculas_unicas = []
    for item in geral['Matricula']:
        mat.append(item)
    matnum = list(map(int, mat))
    matriculas_unicas = list(set(matnum))
    for matricula in matriculas_unicas:
        base = pd.DataFrame(list(intervalo(datainicio, datafim)),
                            columns=['Data', 'Entrada 1', 'Saída 1', 'Entrada 2', 'Saída 2', 'Entrada 3', 'Saída 3'])
        base = base.set_index('Data')
        geral = pd.read_csv(arquivo, sep=' ', header=None, encoding='iso8859-1')
        geral.dropna(axis=1, inplace=True)
        geral = geral.rename(columns={0: 'Dados'})
        # dividir ultimos 11 caracteres em outra col
        geral['Matricula'] = geral['Dados'].str[-11:]
        # dividir caracteres da data em outra col
        geral['Data'] = geral['Dados'].str[-24:-16]
        geral['Data'] = pd.to_datetime(geral['Data'], format='%d%m%Y')
        # dividir caracteres da hr em outra col
        geral['Hora'] = geral['Dados'].str[-16:-12]
        geral['Hora'] = pd.to_datetime(geral['Hora'], format='%H%M')

        geral1 = pd.DataFrame(geral)

        geral = geral.loc[(geral['Data'] < datafim) & (geral['Data'] > datainicio)]
        geral1 = geral1.loc[(geral1['Data'] < datafim) & (geral1['Data'] > datainicio)]

        geral = geral.loc[geral['Matricula'] == str(matricula).zfill(11)]
        geral1 = geral1.loc[geral1['Matricula'] == str(matricula).zfill(11)]

        geral['Hora'] = geral['Hora'].apply(lambda k: dt.strftime(k, '%H:%M'))
        geral1['Hora'] = geral1['Hora'].apply(lambda l: dt.strftime(l, '%H:%M:%S'))

        geral = geral.drop('Dados', axis=1)
        geral1 = geral1.drop('Dados', axis=1)

        geral = geral.drop('Matricula', axis=1)
        geral1 = geral1.drop('Matricula', axis=1)

        geral = geral.reset_index(drop=True)
        geral = geral.pivot_table(index='Data', columns=geral.groupby('Data').cumcount() + 1, values='Hora',
                                  aggfunc='first')
        geral = geral.reset_index(level=[0])
        geral = geral.rename(
            columns={0: 'Data', 1: 'Entrada 1', 2: 'Saída 1', 3: 'Entrada 2', 4: 'Saída 2', 5: 'Entrada 3',
                     6: 'Saída 3'})

        geral1 = geral1.reset_index(drop=True)
        geral1 = geral1.pivot_table(index='Data', columns=geral1.groupby('Data').cumcount() + 1, values='Hora',
                                    aggfunc='first')
        geral1 = geral1.reset_index(level=[0])
        geral1 = geral1.rename(
            columns={0: 'Data', 1: 'Entrada 1', 2: 'Saída 1', 3: 'Entrada 2', 4: 'Saída 2', 5: 'Entrada 3',
                     6: 'Saída 3'})
        try:
            geral1['Entrada 1a'] = geral1['Entrada 1'].apply(lambda z: pd.to_timedelta(str(z)))
            geral1['Saída 1a'] = geral1['Saída 1'].apply(lambda z: pd.to_timedelta(str(z)))
            geral1['Tot1'] = geral1['Saída 1a'] - geral1['Entrada 1a']
            total_horas = round(geral1['Tot1'].sum().total_seconds() / 3600, 2)

            if geral1.shape[1] > 6:
                if geral1.shape[1] < 10:
                    geral1['Entrada 2a'] = geral1['Entrada 2'].apply(lambda z: pd.to_timedelta(str(z)))
                    geral1['Saída 2a'] = geral1['Saída 2'].apply(lambda z: pd.to_timedelta(str(z)))
                    geral1['Tot2'] = geral1['Saída 2a'] - geral1['Entrada 2a']
                    geral1['Soma'] = geral1['Tot1'] + geral1['Tot2']
                    total_horas = round(geral1['Soma'].sum().total_seconds() / 3600, 2)
                    geral1 = geral1.drop('Entrada 2a', axis=1)
                    geral1 = geral1.drop('Saída 2a', axis=1)
                    geral1 = geral1.drop('Tot2', axis=1)
                    geral1 = geral1.drop('Soma', axis=1)

            if geral1.shape[1] > 8:
                geral1['Entrada 2a'] = geral1['Entrada 2'].apply(lambda z: pd.to_timedelta(str(z)))
                geral1['Saída 2a'] = geral1['Saída 2'].apply(lambda z: pd.to_timedelta(str(z)))
                geral1['Entrada 3a'] = geral1['Entrada 3'].apply(lambda z: pd.to_timedelta(str(z)))
                geral1['Saída 3a'] = geral1['Saída 3'].apply(lambda z: pd.to_timedelta(str(z)))
                geral1['Tot1'] = geral1['Saída 1a'] - geral1['Entrada 1a']
                geral1['Tot2'] = geral1['Saída 2a'] - geral1['Entrada 2a']
                geral1['Tot3'] = geral1['Saída 3a'] - geral1['Entrada 3a']
                geral1['Soma'] = geral1['Tot1'] + geral1['Tot2'] + geral1['Tot3']
                total_horas = round(geral1['Soma'].sum().total_seconds() / 3600, 2)
                geral1 = geral1.drop('Entrada 2a', axis=1)
                geral1 = geral1.drop('Saída 2a', axis=1)
                geral1 = geral1.drop('Entrada 3a', axis=1)
                geral1 = geral1.drop('Saída 3a', axis=1)
                geral1 = geral1.drop('Tot2', axis=1)
                geral1 = geral1.drop('Tot3', axis=1)
                geral1 = geral1.drop('Soma', axis=1)

            geral1 = geral1.drop('Entrada 1a', axis=1)
            geral1 = geral1.drop('Saída 1a', axis=1)
            geral1 = geral1.drop('Tot1', axis=1)
        except KeyError:
            total_horas = 0

        geral = geral.set_index('Data')
        base = base.combine_first(geral)
        base = base.reset_index(level=[0])
        base['Data'] = base['Data'].apply(lambda h: dt.strftime(h, '%d/%m/%Y - %a'))
        base = base[['Data', 'Entrada 1', 'Saída 1', 'Entrada 2', 'Saída 2', 'Entrada 3', 'Saída 3']]
        dias = base['Entrada 1'].count()
        base = base.fillna('-')
        ponto = docx.Document(rf'C:\Users\{os.getlogin()}\PycharmProjects\AutomacaoCia\src\models\static\files\modelo.docx')
        ponto.paragraphs[1].text = str(ponto.paragraphs[1].text).replace('#data1',
                                                                         dt.strftime(datainicio, '%d/%m/%Y')).replace(
            '#data2', dt.strftime(datafim, '%d/%m/%Y'))
        ponto.paragraphs[2].text = str(ponto.paragraphs[2].text).replace('#emissao',
                                                                         dt.strftime(dt.today(), '%d/%m/%Y'))
        try:
            ponto.tables[0].rows[3].cells[1].paragraphs[0].text = str(
                ponto.tables[0].rows[3].cells[1].paragraphs[0].text).replace('#nome', planbase[str(matricula)])
            ponto.tables[0].rows[4].cells[1].paragraphs[0].text = str(
                ponto.tables[0].rows[4].cells[1].paragraphs[0].text).replace('#cod', str(matricula))
            ponto.tables[0].rows[5].cells[1].paragraphs[0].text = str(
                ponto.tables[0].rows[5].cells[1].paragraphs[0].text).replace('#mat', str(planmat[str(matricula)]))
            ponto.tables[0].rows[6].cells[1].paragraphs[0].text = str(
                ponto.tables[0].rows[6].cells[1].paragraphs[0].text).replace('#func', planfunc[str(matricula)])
            ponto.tables[0].rows[7].cells[1].paragraphs[0].text = str(
                ponto.tables[0].rows[7].cells[1].paragraphs[0].text).replace('#depto', plandept[str(matricula)])
        except KeyError:
            pass
        # ponto.tables[0].rows[5].cells[4].paragraphs[0].text = 'HH:MM'
        style = ponto.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(10)
        style8 = ponto.styles['Default']
        font = style8.font
        font.name = 'Times New Roman'
        font.size = Pt(8)

        ponto.paragraphs[0].style = ponto.styles['Normal']
        ponto.paragraphs[1].style = ponto.styles['Normal']
        ponto.paragraphs[2].style = ponto.styles['Normal']
        ponto.paragraphs[3].style = ponto.styles['Normal']
        ponto.paragraphs[4].style = ponto.styles['Normal']
        ponto.tables[0].rows[5].cells[4].paragraphs[0].style = ponto.styles['Default']
        t = ponto.add_table(base.shape[0] + 1, base.shape[1])
        t.style = 'Estilo2'
        # add the header rows.
        for j in range(base.shape[-1]):
            t.cell(0, j).text = base.columns[j]
        # add the rest of the data frame
        for i in range(base.shape[0]):
            for j in range(base.shape[-1]):
                t.cell(i + 1, j).text = str(base.values[i, j])
                t.cell(i + 1, j).paragraphs[0].alignment = 1
                t.cell(i + 1, 0).paragraphs[0].alignment = 0

        ponto.tables[1].columns[0].cells[0].width = Cm(4.2)
        ponto.tables[1].columns[1].cells[0].width = Cm(2.6)
        ponto.tables[1].columns[2].cells[0].width = Cm(2.6)
        ponto.tables[1].columns[3].cells[0].width = Cm(2.6)
        ponto.tables[1].columns[4].cells[0].width = Cm(2.6)
        ponto.tables[1].columns[5].cells[0].width = Cm(2.6)
        ponto.tables[1].columns[6].cells[0].width = Cm(2.6)
        ponto.tables[1].rows[0].cells[0].paragraphs[0].alignment = 1
        ponto.tables[1].rows[0].cells[1].paragraphs[0].alignment = 1
        ponto.tables[1].rows[0].cells[2].paragraphs[0].alignment = 1
        ponto.tables[1].rows[0].cells[3].paragraphs[0].alignment = 1
        ponto.tables[1].rows[0].cells[4].paragraphs[0].alignment = 1
        ponto.tables[1].rows[0].cells[5].paragraphs[0].alignment = 1
        ponto.tables[1].rows[0].cells[6].paragraphs[0].alignment = 1

        y = len(fl['A']) + 1
        try:
            fl[f'A{y}'].value = planbase[str(matricula)]
            fl[f'B{y}'].value = total_horas
            pl.save(rf'C:\Users\{os.getlogin()}\PycharmProjects\AutomacaoCia\src\models\static\files\Hrs Folha.xlsx')
    
            ponto.add_paragraph(f'Total de dias trabalhados: {dias}', 'Default').alignment = 2
            ponto.add_paragraph('', 'Normal')
            ponto.add_paragraph('', 'Normal')
            ponto.add_paragraph('______________________________________________', 'Normal').alignment = 1
            ponto.add_paragraph(f'{planbase[str(matricula)]}', 'Normal').alignment = 1
            try:
                ponto.save(rede + rf'\04 - Folha de Pgto\{ano}\{mes} - {mesext[mes]}\Relatórios de Ponto\{dataipt} a {datafpt}\Ponto {planbase[str(matricula)]}.docx')
                docx2pdf.convert(
                    rede + rf'\04 - Folha de Pgto\{ano}\{mes} - {mesext[mes]}\Relatórios de Ponto\{dataipt} a {datafpt}\Ponto {planbase[str(matricula)]}.docx',
                    rede + rf'\04 - Folha de Pgto\{ano}\{mes} - {mesext[mes]}\Relatórios de Ponto\{dataipt} a {datafpt}\Ponto {planbase[str(matricula)]}.pdf')
                os.remove(
                    rede + rf'\04 - Folha de Pgto\{ano}\{mes} - {mesext[mes]}\Relatórios de Ponto\{dataipt} a {datafpt}\Ponto {planbase[str(matricula)]}.docx')
            except FileNotFoundError:
                os.makedirs(
                    rede + rf'\04 - Folha de Pgto\{ano}\{mes} - {mesext[mes]}\Relatórios de Ponto\{dataipt} a {datafpt}')
                ponto.save(rede + rf'\04 - Folha de Pgto\{ano}\{mes} - {mesext[mes]}\Relatórios de Ponto\{dataipt} a {datafpt}\Ponto {planbase[str(matricula)]}.docx')
                docx2pdf.convert(
                    rede + rf'\04 - Folha de Pgto\{ano}\{mes} - {mesext[mes]}\Relatórios de Ponto\{dataipt} a {datafpt}\Ponto {planbase[str(matricula)]}.docx',
                    rede + rf'\04 - Folha de Pgto\{ano}\{mes} - {mesext[mes]}\Relatórios de Ponto\{dataipt} a {datafpt}\Ponto {planbase[str(matricula)]}.pdf')
                os.remove(
                    rede + rf'\04 - Folha de Pgto\{ano}\{mes} - {mesext[mes]}\Relatórios de Ponto\{dataipt} a {datafpt}\Ponto {planbase[str(matricula)]}.docx')
        except KeyError:
            fl[f'A{y}'].value = str(matricula)
            fl[f'B{y}'].value = total_horas
            pl.save(rf'C:\Users\{os.getlogin()}\PycharmProjects\AutomacaoCia\src\models\static\files\Hrs Folha.xlsx')

            ponto.add_paragraph(f'Total de dias trabalhados: {dias}', 'Default').alignment = 2
            ponto.add_paragraph('', 'Normal')
            ponto.add_paragraph('', 'Normal')
            ponto.add_paragraph('______________________________________________', 'Normal').alignment = 1
            ponto.add_paragraph(f'{str(matricula)}', 'Normal').alignment = 1
            try:
                ponto.save(
                    rede + rf'\04 - Folha de Pgto\{ano}\{mes} - {mesext[mes]}\Relatórios de Ponto\{dataipt} a {datafpt}\Ponto {str(matricula)}.docx')
                docx2pdf.convert(
                    rede + rf'\04 - Folha de Pgto\{ano}\{mes} - {mesext[mes]}\Relatórios de Ponto\{dataipt} a {datafpt}\Ponto {str(matricula)}.docx',
                    rede + rf'\04 - Folha de Pgto\{ano}\{mes} - {mesext[mes]}\Relatórios de Ponto\{dataipt} a {datafpt}\Ponto {str(matricula)}.pdf')
                os.remove(
                    rede + rf'\04 - Folha de Pgto\{ano}\{mes} - {mesext[mes]}\Relatórios de Ponto\{dataipt} a {datafpt}\Ponto {str(matricula)}.docx')
            except FileNotFoundError:
                os.makedirs(
                    rede + rf'\04 - Folha de Pgto\{ano}\{mes} - {mesext[mes]}\Relatórios de Ponto\{dataipt} a {datafpt}')
                ponto.save(
                    rede + rf'\04 - Folha de Pgto\{ano}\{mes} - {mesext[mes]}\Relatórios de Ponto\{dataipt} a {datafpt}\Ponto {str(matricula)}.docx')
                docx2pdf.convert(
                    rede + rf'\04 - Folha de Pgto\{ano}\{mes} - {mesext[mes]}\Relatórios de Ponto\{dataipt} a {datafpt}\Ponto {str(matricula)}.docx',
                    rede + rf'\04 - Folha de Pgto\{ano}\{mes} - {mesext[mes]}\Relatórios de Ponto\{dataipt} a {datafpt}\Ponto {str(matricula)}.pdf')
                os.remove(
                    rede + rf'\04 - Folha de Pgto\{ano}\{mes} - {mesext[mes]}\Relatórios de Ponto\{dataipt} a {datafpt}\Ponto {str(matricula)}.docx')
    while 1:
        if pa.locateOnScreen('../models/static/imgs/pyt.png'):
            pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/pyt.png')))
            break
        else:
            time.sleep(5)
    tkinter.messagebox.showinfo(title='Relatórios ok!', message='Relatórios de ponto salvos com sucesso!')


def cadastrar_no_ponto(nome, altera, matrpt=''):
    sessions = sessionmaker(engine)
    session = sessions()
    sessionspt = sessionmaker(engineponto)
    sessionpt = sessionspt()
    pessoa = session.query(Colaborador).filter_by(nome=nome).order_by(Colaborador.matricula.desc()).first()
    if altera == 0:
        if 'ESTAG' in pessoa.cargo:
            pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/controlid.png'))), time.sleep(0.5)
            pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/novofunc.png'))), time.sleep(0.5)
            pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/adcusuario.png'))), time.sleep(0.5)
            pp.copy(pessoa.nome), pa.hotkey('ctrl', 'v'), pa.press('tab'), pa.write(str(pessoa.matricula)), time.sleep(0.5)
            pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/okponto.png'))), time.sleep(0.5)
            pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/okconfirmaponto.png')))
            estag = BasePonto(nome=pessoa.nome, matricula=pessoa.matricula, pis=pessoa.pis,
                              matrponto=pessoa.matricula, email=pessoa.email, cargo=pessoa.cargo,
                              departamento=pessoa.depto)
            sessionpt.add(estag)
            sessionpt.commit()
        else:
            pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/controlid.png'))), time.sleep(0.5)
            pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/novofunc.png'))), time.sleep(0.5)
            pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/adcusuario.png'))), time.sleep(0.5)
            pp.copy(pessoa.nome), pa.hotkey('ctrl', 'v'), pa.press('tab'), pa.write(str(pessoa.pis)), time.sleep(0.5)
            pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/okponto.png')))
            func = BasePonto(nome=pessoa.nome, matricula=pessoa.matricula, pis=pessoa.pis,
                              matrponto=pessoa.pis, email=pessoa.email, cargo=pessoa.cargo,
                              departamento=pessoa.depto)
            sessionpt.add(func)
            sessionpt.commit()

    else:
        pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/controlid.png'))), time.sleep(0.5)
        pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/novofunc.png'))), time.sleep(0.5)
        pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/adcusuario.png'))), time.sleep(0.5)
        pp.copy(pessoa.nome), pa.hotkey('ctrl', 'v'), pa.press('tab'), pa.write(str(matrpt)), time.sleep(0.5)
        pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/okponto.png'))), time.sleep(0.5)
        pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/okconfirmaponto.png')))
        estag = BasePonto(nome=pessoa.nome, matricula=pessoa.matricula, pis=pessoa.pis,
                          matrponto=matrpt, email=pessoa.email, cargo=pessoa.cargo,
                          departamento=pessoa.depto)
        sessionpt.add(estag)
        sessionpt.commit()
    tkinter.messagebox.showinfo(title='Cadastro no ponto ok!', message='Colaborador cadastrado com sucesso!')

