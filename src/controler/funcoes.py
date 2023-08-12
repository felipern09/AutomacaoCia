from datetime import datetime as dt, timedelta as td
import datetime
from dateutil.relativedelta import relativedelta
from difflib import SequenceMatcher
import docx
from docx.shared import Pt, Cm
import docx2pdf
from docx2pdf import convert
from email import encoders
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.image import MIMEImage
from email.mime.base import MIMEBase
import holidays
import locale
import numpy as np
import num2words as nw
from openpyxl import load_workbook as l_w
from openpyxl.styles import PatternFill, Font, NamedStyle
import openpyxl.utils.cell
import os
import pandas as pd
import pyautogui as pa
import pyperclip as pp
from PIL import ImageGrab
from src.models.modelsfolha import Aula, Folha, Aulas, Faltas, Ferias, Hrcomplement, Atestado, Desligados, \
    Escala, Substituicao, enginefolha
from src.models.models import Colaborador, engine
from sqlalchemy.orm import sessionmaker
from src.models.listas import municipios
import smtplib
from src.models.dados_servd import em_rem, em_ti, em_if, k1, host, port, rede
import tkinter.filedialog
from tkinter import messagebox
import tkinter.filedialog
import time as t
from typing import List, Dict, Tuple, Type
import urllib
from urllib import parse
import win32com.client as client

locale.setlocale(locale.LC_MONETARY, 'pt_BR.UTF-8')
locale.setlocale(locale.LC_ALL, 'pt_BR.UTF-8')
locale.setlocale(locale.LC_MONETARY, 'pt_BR.UTF-8')
# data = 7

atestado = PatternFill(start_color='A9D08E',
                       end_color='A9D08E',
                       fill_type='solid')
falta = PatternFill(start_color='FF0000',
                    end_color='FF0000',
                    fill_type='solid')
ferias = PatternFill(start_color='9BC2E6',
                     end_color='9BC2E6',
                     fill_type='solid')
feriado = PatternFill(start_color='F4B084',
                      end_color='F4B084',
                      fill_type='solid')
fds = PatternFill(start_color='BFBFBF',
                  end_color='BFBFBF',
                  fill_type='solid')
deslig = PatternFill(start_color='454545',
                     end_color='454545',
                     fill_type='solid')
subst = PatternFill(start_color='FFFF00',
                    end_color='FFFF00',
                    fill_type='solid')
comple = PatternFill(start_color='FFC000',
                     end_color='FFC000',
                     fill_type='solid')


def confirma_folha(comp: int):
    """
    Confirm that the user wnats to proceed with procedures to post values of payroll in external aplication.
    :param comp: Month reference to payroll calculate.
    :return: Call function lancar_folha_no_dexion().
    """
    resp = messagebox.askyesno(title='Tem certeza?',
                               message=f'Tem certeza que deseja lançar a folha do mês {comp} no Dexion?')
    if resp:
        lancar_folha_no_dexion(comp)


def confirma_grade(comp: int):
    """
    Confirm that the user wnats to proceed with procedures to register the payroll.
    :param comp: Month reference to payroll calculate.
    :return: Call function salvar_planilha_soma_final().
    """
    r = messagebox.askyesno(title='Tem certeza?',
                            message=f'Tem certeza que deseja gerar a folha do mês {comp}?\n'
                                    f'Essa ação irá sobrepor qualquer arquivo de folha já salvo na pasta dessa competência.')
    if r:
        salvar_planilha_soma_final(comp)


def lancar_folha_no_dexion(competencia):
    """
    Register values of payroll in external aplication through PyAutogui package.
    :param competencia: Month reference to payroll calculate.
    :return: External aplication with payroll values registred.
    """
    pa.hotkey('alt', 'tab'), pa.press('a'), t.sleep(2)
    folhagrd = os.path.relpath(rf'C:\Users\{os.getlogin()}\PycharmProjects\AutomacaoCia\src\view\Somafinal mes {competencia}.xlsx')
    wb = l_w(folhagrd, read_only=False)

    # lançamento de faltas
    sh = wb['Faltas']
    x = 2
    while x <= len(sh['A']):
        mat = str(sh[f'A{x}'].value)
        rub = str(sh[f'C{x}'].value)
        hr = str(sh[f'D{x}'].value)
        pa.write(mat), t.sleep(0.5), t.sleep(2), t.sleep(0.5), pa.press('enter', 2), t.sleep(0.5), t.sleep(
            1.5), t.sleep(0.5), pa.press('i'), t.sleep(0.5), t.sleep(0.5), t.sleep(0.5), pa.write(rub)
        t.sleep(0.5), t.sleep(0.5), pa.press('enter'), t.sleep(0.5), t.sleep(0.5), t.sleep(0.5), pa.write(hr), t.sleep(
            0.5), t.sleep(0.5), t.sleep(0.5), pa.press('enter', 65)
        x += 1

    # deletar férias antigas
    sh = wb['DeletarFerias']
    x = 2
    while x <= len(sh['A']):
        rub = ['1006', '1007', '1010', '1011', '1012', '1037']
        mat = str(sh[f'A{x}'].value)
        pa.write(mat), t.sleep(0.5), pa.press('enter', 2), t.sleep(1), pa.press('d')
        for r in rub:
            pa.write(r), t.sleep(0.8), pa.press('enter'), t.sleep(0.5), pa.press('left'), t.sleep(0.5), pa.press(
                'enter')
        pa.press('enter', 2)
        x += 1

    # lançamento de horistas
    sh = wb['Horistas']
    x = 2
    while x <= len(sh['A']):
        mat = str(sh[f'A{x}'].value)
        rub = str(sh[f'C{x}'].value)
        hr = str(sh[f'D{x}'].value)
        obshr = str(sh[f'E{x}'].value)
        pa.write(mat), t.sleep(0.5), pa.press('enter', 2), t.sleep(2.3)
        try:
            pa.center(pa.locateOnScreen('dsr.png'))
            dsrlancado = True
        except:
            try:
                pa.center(pa.locateOnScreen('dsr2.png'))
                dsrlancado = True
            except:
                dsrlancado = False
        pa.press('a'), t.sleep(0.5), pa.write(rub)
        pa.press('enter'), t.sleep(0.5), pa.write(hr), t.sleep(0.5), pa.press('enter', 2)
        if dsrlancado:
            pass
        else:
            if obshr != 'HORA AULA ESTÁGIO 5.10':
                pa.press('i'), t.sleep(0.5), pa.write('27'), t.sleep(0.5)
                pa.press('enter', 3), t.sleep(0.5)
        pa.press('enter')
        x += 1

    # lançamento de comissões
    sh = wb['Comissoes']
    x = 2
    while x <= len(sh['A']):
        mat = str(sh[f'A{x}'].value)
        rub = str(sh[f'C{x}'].value)
        hr = str(sh[f'D{x}'].value)
        pa.write(mat), t.sleep(0.5), pa.press('enter', 2), t.sleep(0.5), pa.press('i'), t.sleep(0.5), pa.write(rub)
        pa.press('enter'), t.sleep(0.5), pa.write(hr), t.sleep(0.5), pa.press('enter')
        pa.press('enter'), t.sleep(0.5), pa.press('enter')
        x += 1

    # # Lançamento de adiantamento
    sh = wb['Adiantamento']
    x = 2
    while x <= len(sh['A']):
        mat = str(sh[f'A{x}'].value)
        rub = '81'
        hr = str(sh[f'D{x}'].value)
        pa.write(mat), t.sleep(0.5), pa.press('enter', 2), t.sleep(0.5), pa.press('i'), t.sleep(0.5), pa.write(rub)
        pa.press('enter'), t.sleep(0.5), pa.write(hr), t.sleep(0.5), pa.press('enter')
        pa.press('enter'), t.sleep(0.5), pa.press('enter')
        x += 1

    # lançamento de desconto de VT
    sh = wb['DescontoVT']
    x = 2
    while x <= len(sh['A']):
        mat = str(sh[f'A{x}'].value)
        rub = '80'
        hr = str(sh[f'D{x}'].value)
        pa.write(mat), t.sleep(0.5), pa.press('enter', 2), t.sleep(0.5), pa.press('i'), t.sleep(0.5), pa.write(rub)
        pa.press('enter'), t.sleep(0.5), pa.press('enter')
        pa.press('enter'), t.sleep(0.5), pa.press('enter')
        x += 1

    # lançamento de plano de saúde
    sh = wb['Plano']
    x = 2
    while x <= len(sh['A']):
        mat = str(sh[f'A{x}'].value)
        rub = str(sh[f'C{x}'].value)
        hr = str(sh[f'D{x}'].value)
        sq = str(sh[f'E{x}'].value)
        pa.write(mat), t.sleep(0.5), pa.press('enter', 2), t.sleep(0.5), pa.press('i'), t.sleep(0.5), pa.write(rub)
        pa.press('enter'), t.sleep(0.5), pa.write(sq), t.sleep(0.5), pa.press('enter'), t.sleep(0.5), pa.write(hr)
        pa.press('enter', 3), t.sleep(0.5),
        x += 1
    tkinter.messagebox.showinfo(
        title='Folha ok!',
        message=f'Folha do mês {competencia} lançada no Dexion com sucesso!'
    )


def somar_aulas_da_grade(diasem: str, inic: datetime.datetime, fim: datetime.datetime, competencia: int, iniciograd: str, fimgrad: str) -> float:
    """
    Sum classes that have status 'Active' at the period selected for the payroll.
    :param diasem: Weekday.
    :param inic: Class start time.
    :param fim: End time of the class.
    :param competencia: Month of the payroll.
    :param iniciograd: Start day for the payroll.
    :param fimgrad: End day of the payroll.
    :return: Sum of all classes times for each day of the week.
    """
    horas = fim - inic
    hr, minut, seg = str(horas).split(':')
    igrad = dt.strptime(iniciograd, '%d/%m/%Y')
    if fimgrad is not None:
        fgrad = dt.strptime(fimgrad, '%d/%m/%Y')
    else:
        fgrad = ''
    somadia = int(hr) + int(minut) / 60 + int(seg) / 60 * 60
    somadia = round(somadia, 2)
    soma = 0
    comp = dt(day=1, month=competencia, year=dt.today().year)
    inicio = dt(day=21, month=(comp - relativedelta(months=1)).month, year=(comp - relativedelta(months=1)).year)
    fechamento = dt(day=20, month=comp.month, year=comp.year)

    def intervalo(inicio, fechamento):
        for n in range(int((fechamento - inicio).days) + 1):
            yield inicio + td(n)

    if fgrad != '':
        if diasem == 'Segunda':
            for dia in intervalo(inicio, fechamento):
                if dia.weekday() == 0 and igrad <= dia <= fgrad:
                    soma += somadia
        if diasem == 'Terça':
            for dia in intervalo(inicio, fechamento):
                if dia.weekday() == 1 and igrad <= dia <= fgrad:
                    soma += somadia
        if diasem == 'Quarta':
            for dia in intervalo(inicio, fechamento):
                if dia.weekday() == 2 and igrad <= dia <= fgrad:
                    soma += somadia
        if diasem == 'Quinta':
            for dia in intervalo(inicio, fechamento):
                if dia.weekday() == 3 and igrad <= dia <= fgrad:
                    soma += somadia
        if diasem == 'Sexta':
            for dia in intervalo(inicio, fechamento):
                if dia.weekday() == 4 and igrad <= dia <= fgrad:
                    soma += somadia
        if diasem == 'Sábado':
            for dia in intervalo(inicio, fechamento):
                if dia.weekday() == 5 and igrad <= dia <= fgrad:
                    soma += somadia
        if diasem == 'Domingo':
            for dia in intervalo(inicio, fechamento):
                if dia.weekday() == 6 and igrad <= dia <= fgrad:
                    soma += somadia
    else:
        if diasem == 'Segunda':
            for dia in intervalo(inicio, fechamento):
                if dia.weekday() == 0 and dia >= igrad:
                    soma += somadia
        if diasem == 'Terça':
            for dia in intervalo(inicio, fechamento):
                if dia.weekday() == 1 and dia >= igrad:
                    soma += somadia
        if diasem == 'Quarta':
            for dia in intervalo(inicio, fechamento):
                if dia.weekday() == 2 and dia >= igrad:
                    soma += somadia
        if diasem == 'Quinta':
            for dia in intervalo(inicio, fechamento):
                if dia.weekday() == 3 and dia >= igrad:
                    soma += somadia
        if diasem == 'Sexta':
            for dia in intervalo(inicio, fechamento):
                if dia.weekday() == 4 and dia >= igrad:
                    soma += somadia
        if diasem == 'Sábado':
            for dia in intervalo(inicio, fechamento):
                if dia.weekday() == 5 and dia >= igrad:
                    soma += somadia
        if diasem == 'Domingo':
            for dia in intervalo(inicio, fechamento):
                if dia.weekday() == 6 and dia >= igrad:
                    soma += somadia
    return round(soma, 2)


def listar_aulas_ativas() -> list:
    """
    List all classes with 'Active' status.
    :return: List of active classes.
    """
    sessions = sessionmaker(bind=enginefolha)
    session = sessions()
    aulasativasdb = session.query(Aulas).filter_by(status='Ativa').all()
    aula = []
    for i, a in enumerate(aulasativasdb):
        aula.append(i)
        aula[i] = Aula(a.nome, a.professor, a.departamento, a.diadasemana, a.inicio, a.fim, a.valor, a.iniciograde,
                       a.fimgrade)
        yield aula[i]
    return aula


def listar_departamentos_ativos() -> list:
    """
    List all departments with classes with "Active" status.
    :return: List of active departments.
    """
    sessions = sessionmaker(bind=enginefolha)
    session = sessions()
    aulasativasdb = session.query(Aulas).filter_by(status='Ativa').all()
    departamentos = []
    for i, a in enumerate(aulasativasdb):
        departamentos.append(a.departamento)
        departamentos = list(set(departamentos))
    return departamentos


def listar_professores_ativos() -> list:
    """
    List of all teachers who have active classes.
    :return: List of active teachers.
    """
    sessions = sessionmaker(bind=enginefolha)
    session = sessions()
    aulasativasdb = session.query(Aulas).filter_by(status='Ativa').all()
    professores = []
    for i, a in enumerate(aulasativasdb):
        professores.append(a.professor)
        professores = list(set(professores))
    return professores


def calcular_total_monetario_folha(compet: int) -> float:
    """
    Sum the total payroll payment.
    :param compet: Month of payroll
    :return: Payment amount.
    """
    somatorio = 0
    for al in listar_aulas_ativas():
        somatorio += somar_aulas_da_grade(al.dia, al.inicio, al.fim, compet, al.iniciograde, al.fimgrade) * float(
            str(al.valor).replace(',', '.')) * al.dsr
    return round(somatorio, 2)


def somar_horas_professor(folha: Folha, prof: str, depto: str, nome: str, compet: int) -> float:
    """
    Sum total of hours for each teacher.
    :param folha: Payroll reference.
    :param prof: Teacher.
    :param depto: Department
    :param nome: Name of the class.
    :param compet: Month of payroll
    :return: The amount of hours for each teacher on this payroll.
    """
    somahoras = 0
    for aula in folha.aulas:
        if aula.professor == prof and aula.departamento == depto and aula.nome == nome:
            somahoras += somar_aulas_da_grade(aula.dia, aula.inicio, aula.fim, compet, aula.iniciograde, aula.fimgrade)
    return somahoras


def consultar_faltas(comp) -> dict:
    sessions = sessionmaker(bind=enginefolha)
    session = sessions()
    falt = session.query(Faltas).all()
    inicio = dt(day=21, month=(dt(day=1, month=comp, year=dt.today().year) - relativedelta(months=1)).month,
                year=dt.today().year)
    fim = dt(day=20, month=comp, year=dt.today().year)
    dic = {}
    for f in falt:
        if inicio <= dt.strptime(f.data, '%d/%m/%Y') <= fim:
            if f.professor in dic:
                d2 = {f.professor: {f.data: {f.departamento: f.horas}}}
                dic[f.professor] = {**dic[f.professor], **d2[f.professor]}
            else:
                d2 = {f.professor: {f.data: {f.departamento: f.horas}}}
                dic = {**dic, **d2}
    return dic


def consultar_ferias(comp) -> dict:
    sessions = sessionmaker(bind=enginefolha)
    session = sessions()
    fer = session.query(Ferias).all()
    inicio = dt(day=21, month=(dt(day=1, month=comp, year=dt.today().year) - relativedelta(months=1)).month,
                year=dt.today().year)
    fim = dt(day=20, month=comp, year=dt.today().year)
    dic = {}
    for f in fer:
        if inicio <= dt.strptime(f.inicio, '%d/%m/%Y') <= fim:
            if f.professor in dic:
                d2 = {f.professor: {f.departamento: {f.inicio: f.fim}}}
                dic[f.professor] = {**dic[f.professor], **d2[f.professor]}
            else:
                d2 = {f.professor: {f.departamento: {f.inicio: f.fim}}}
                dic = {**dic, **d2}
    return dic


def consultar_atestados(comp) -> dict:
    sessions = sessionmaker(bind=enginefolha)
    session = sessions()
    ates = session.query(Atestado).all()
    inicio = dt(day=21, month=(dt(day=1, month=comp, year=dt.today().year) - relativedelta(months=1)).month,
                year=dt.today().year)
    fim = dt(day=20, month=comp, year=dt.today().year)
    dic = {}
    for a in ates:
        if inicio <= dt.strptime(a.data, '%d/%m/%Y') <= fim:
            if a.professor in dic:
                d2 = {a.professor: {a.data: a.departamento}}
                dic[a.professor] = {**dic[a.professor], **d2[a.professor]}
            else:
                d2 = {a.professor: {a.data: a.departamento}}
                dic = {**dic, **d2}
    return dic


def listar_feriados(comp: int) -> list:
    inicio = dt(day=21, month=(dt(day=1, month=comp, year=dt.today().year) - relativedelta(months=1)).month,
                year=dt.today().year)
    fim = dt(day=20, month=comp, year=dt.today().year)
    # Get the Bank Holidays for the given country
    feriados = holidays.country_holidays('BR')
    # Create a list of dates between the start and end date
    intervalo_datas = pd.date_range(inicio, fim)
    # Filter the dates to only include Bank Holidays
    feriados_nacionais = [date for date in intervalo_datas if date in feriados]
    return feriados_nacionais


def consultar_substituicoes(comp: int) -> dict:
    sessions = sessionmaker(bind=enginefolha)
    session = sessions()
    subst = session.query(Substituicao).all()
    inicio = dt(day=21, month=(dt(day=1, month=comp, year=dt.today().year) - relativedelta(months=1)).month,
                year=dt.today().year)
    fim = dt(day=20, month=comp, year=dt.today().year)
    dic = {}
    for s in subst:
        if inicio <= dt.strptime(s.data, '%d/%m/%Y') <= fim:
            if s.professorsubst in dic:
                dic2 = {s.professorsubst: {s.substituto: {s.departamento: {s.data: s.horas}}}}
                dic[s.professorsubst] = {**dic[s.professorsubst], **dic2[s.professorsubst]}
            else:
                d2 = {s.professorsubst: {s.substituto: {s.departamento: {s.data: s.horas}}}}
                dic = {**dic, **d2}
    return dic


def consultar_desligamentos(comp: int) -> dict:
    sessions = sessionmaker(bind=enginefolha)
    session = sessions()
    desl = session.query(Desligados).all()
    inicio = dt(day=21, month=(dt(day=1, month=comp, year=dt.today().year) - relativedelta(months=1)).month,
                year=dt.today().year)
    fim = dt(day=20, month=comp, year=dt.today().year)
    dic = {}
    for d in desl:
        if inicio <= dt.strptime(d.datadesligamento, '%d/%m/%Y') <= fim:
            if d.professor in dic:
                d2 = {d.professor: {d.departamento: d.datadesligamento}}
                dic[d.professor] = {**dic[d.professor], **d2[d.professor]}
            else:
                d2 = {d.professor: {d.departamento: d.datadesligamento}}
                dic = {**dic, **d2}
    return dic


def consultar_escalas(comp: int) -> dict:
    sessions = sessionmaker(bind=enginefolha)
    session = sessions()
    esc = session.query(Escala).all()
    inicio = dt(day=21, month=(dt(day=1, month=comp, year=dt.today().year) - relativedelta(months=1)).month,
                year=dt.today().year)
    fim = dt(day=20, month=comp, year=dt.today().year)
    dic = {}
    for e in esc:
        if inicio <= dt.strptime(e.data, '%d/%m/%Y') <= fim:
            if e.professor in dic:
                d2 = {e.professor: {e.data: {e.departamento: e.horas}}}
                dic[e.professor] = {**dic[e.professor], **d2[e.professor]}
            else:
                d2 = {e.professor: {e.data: {e.departamento: e.horas}}}
                dic = {**dic, **d2}
    return dic


def consultar_horas_complementares(comp: int) -> dict:
    sessions = sessionmaker(bind=enginefolha)
    session = sessions()
    hrsc = session.query(Hrcomplement).all()
    inicio = dt(day=21, month=(dt(day=1, month=comp, year=dt.today().year) - relativedelta(months=1)).month,
                year=dt.today().year)
    fim = dt(day=20, month=comp, year=dt.today().year)
    dic = {}
    for h in hrsc:
        if inicio <= dt.strptime(h.data, '%d/%m/%Y') <= fim:
            if h.professor in dic:
                d2 = {h.professor: {h.data: {h.departamento: h.horas}}}
                dic[h.professor] = {**dic[h.professor], **d2[h.professor]}
            else:
                d2 = {h.professor: {h.data: {h.departamento: h.horas}}}
                dic = {**dic, **d2}
    return dic


def salvar_planilha_grade_horaria(dic: dict, comp: int):
    grd = os.path.relpath(rf'C:\Users\{os.getlogin()}\PycharmProjects\AutomacaoCia\src\models\static\files\Grade.xlsx')
    grade = l_w(grd, read_only=False)
    plan1 = grade['Planilha1']
    flt = consultar_faltas(comp)
    subs = consultar_substituicoes(comp)
    dslg = consultar_desligamentos(comp)
    fer = consultar_ferias(comp)
    complem = consultar_horas_complementares(comp)
    atest = consultar_atestados(comp)
    feriad = listar_feriados(comp)
    escal = consultar_escalas(comp)
    competencia = dt(day=10, month=comp, year=dt.today().year)
    inicio = dt(day=21, month=(competencia - relativedelta(months=1)).month,
                year=(competencia - relativedelta(months=1)).year)
    fechamento = dt(day=20, month=competencia.month, year=competencia.year)
    # primeira linha deve aparecer 'Folha' na coluna A1 e 'Julho' de '2023' na B1
    plan1['A1'].value = 'Folha'
    plan1['B1'].value = f'{fechamento.month} de {fechamento.year}'

    def intervalo(inicio, fechamento):
        for n in range(int((fechamento - inicio).days) + 1):
            yield dt.strftime(inicio + td(n), '%d/%m/%Y')
    # inserir data no cabeçalho da grade
    col = 3
    for item in list(intervalo(inicio, fechamento)):
        plan1.cell(column=col, row=3, value=dt.strftime(dt.strptime(item, '%d/%m/%Y'), '%a'))
        plan1.cell(column=col, row=4, value=dt.strftime(dt.strptime(item, '%d/%m/%Y'), '%d/%m'))
        col += 1
    # indesir total na coluna ao final das datas
    plan1.cell(column=col, row=3, value='Total')
    # formatar coloração de fds
    for itens in plan1.iter_cols(min_row=3, min_col=3, max_row=3, max_col=35):
        for cell in itens:
            if cell.value != 'Total':
                if cell.value == 'sáb' or cell.value == 'dom':
                    letras = openpyxl.utils.cell.get_column_letter(cell.column)
                    for numero in range(3, 150):
                        plan1[f'{letras}{numero}'].fill = fds
    # separar cada prof de cada depto
    musculacao = []
    ginastica = []
    esportes = []
    kids = []
    cross = []
    for i in dic:
        for sub in dic[i]:
            if dic[i][sub] == {}:
                pass
            else:
                if sub == 'Musculação':
                    musculacao.append(i)
                    musculacao.sort()
                if sub == 'Ginástica':
                    ginastica.append(i)
                    ginastica.sort()
                if sub == 'Esportes':
                    esportes.append(i)
                    esportes.sort()
                if sub == 'Kids':
                    kids.append(i)
                    kids.sort()
                if sub == 'Cross Cia':
                    cross.append(i)
                    cross.sort()

    plan1['A4'].value = 'Musculação'
    novalinha = 5
    for i in musculacao:
        for row in plan1.iter_cols(min_row=3, min_col=3, max_row=3, max_col=35):
            for cell in row:
                if cell.value != 'Total':
                    if cell.value == 'seg':
                        plan1.cell(column=cell.column, row=novalinha, value=somar_aulas_de_segunda(i, 'Musculação'))
                    if cell.value == 'ter':
                        plan1.cell(column=cell.column, row=novalinha, value=somar_aulas_de_terca(i, 'Musculação'))
                    if cell.value == 'qua':
                        plan1.cell(column=cell.column, row=novalinha, value=somar_aulas_de_quarta(i, 'Musculação'))
                    if cell.value == 'qui':
                        plan1.cell(column=cell.column, row=novalinha, value=somar_aulas_de_quinta(i, 'Musculação'))
                    if cell.value == 'sex':
                        plan1.cell(column=cell.column, row=novalinha, value=somar_aulas_de_sexta(i, 'Musculação'))
                    if cell.value == 'sáb':
                        plan1.cell(column=cell.column, row=novalinha, value=somar_aulas_de_sabado(i, 'Musculação'))
                    if cell.value == 'dom':
                        plan1.cell(column=cell.column, row=novalinha, value=somar_aulas_de_domingo(i, 'Musculação'))
                else:
                    letra = openpyxl.utils.cell.get_column_letter(cell.column - 1)
                    plan1.cell(column=cell.column, row=novalinha, value=f'=SUM(C{novalinha}:{letra}{novalinha})')
                # aplica cor de falta
                for nome in flt:
                    for dia in flt[nome]:
                        for depart in flt[nome][dia]:
                            if depart == 'Musculação' and nome == i:
                                if plan1.cell(column=cell.column, row=cell.row + 1).value == dt.strftime(
                                        dt.strptime(dia, '%d/%m/%Y'), '%d/%m'):
                                    plan1.cell(column=cell.column, row=novalinha).fill = falta
                # aplica alterações de substituição
                # {s.professorsubst: {s.substituto: {s.departamento: {s.data: s.horas}}}}
                # {s.professorsubst: {s.data: {s.substituto: {s.departamento: s.horas}}}}
                for nome in subs:
                    for substituto in subs[nome]:
                        for depart in subs[nome][substituto]:
                            for dia in subs[nome][substituto][depart]:
                                if depart == 'Musculação' and nome == i:
                                    if plan1.cell(column=cell.column, row=cell.row + 1).value == dt.strftime(
                                            dt.strptime(dia, '%d/%m/%Y'), '%d/%m'):
                                        plan1.cell(column=cell.column, row=novalinha).fill = falta
                                if depart == 'Musculação' and substituto == i:
                                    if plan1.cell(column=cell.column, row=cell.row + 1).value == dt.strftime(
                                            dt.strptime(dia, '%d/%m/%Y'), '%d/%m'):
                                        plan1.cell(column=cell.column, row=novalinha).value = plan1.cell(
                                            column=cell.column, row=novalinha).value + float(
                                            str(subs[nome][substituto][depart][dia]).replace(',', '.'))
                                        plan1.cell(column=cell.column, row=novalinha).fill = subst

                # aplica alterações de desligamento
                # {d.professor: {d.departamento: d.datadesligamento}}
                # conferir se tem outras aulas ativas ou foi desligado de tudo
                # se desligado de tudo, alterar status das aulas para inativas
                for nome in dslg:
                    for depart in dslg[nome]:
                        for dia in dslg[nome][depart]:
                            if depart == 'Musculação' and nome == i and dt.strptime(dia, '%d/%m/%Y') <= fechamento:
                                if plan1.cell(column=cell.column, row=cell.row + 1).value is not None:
                                    if dt.strptime(dia, '%d/%m/%Y') <= dt(day=int(
                                            str(plan1.cell(column=cell.column, row=cell.row + 1).value).split('/')[0]),
                                                                          month=int(str(plan1.cell(column=cell.column,
                                                                                                   row=cell.row + 1).value).split(
                                                                                  '/')[1]),
                                                                          year=dt.today().year) <= fechamento:
                                        plan1.cell(column=cell.column, row=novalinha).value = 0
                                        plan1.cell(column=cell.column, row=novalinha).fill = deslig
                                        plan1.cell(column=cell.column, row=novalinha).font = Font(color='FFFFFF')

                # aplica talterações de férias
                # # {f.professor: {f.departamento: {f.inicio: f.fim}}}
                for nome in fer:
                    for depart in fer[nome]:
                        for inic in fer[nome][depart]:
                            if depart == 'Musculação' and nome == i and dt.strptime(inic, '%d/%m/%Y') <= fechamento:
                                if plan1.cell(column=cell.column, row=cell.row + 1).value is not None:
                                    if dt.strptime(inic, '%d/%m/%Y') <= dt(day=int(
                                            str(plan1.cell(column=cell.column, row=cell.row + 1).value).split('/')[0]),
                                                                           month=int(str(plan1.cell(column=cell.column,
                                                                                                    row=cell.row + 1).value).split(
                                                                                   '/')[1]),
                                                                           year=dt.today().year) <= dt.strptime(
                                            fer[nome][depart][inic], '%d/%m/%Y'):
                                        plan1.cell(column=cell.column, row=novalinha).fill = ferias
                                        plan1.cell(column=cell.column, row=novalinha).value = 0

                # aplica alterações de horas complementares
                # {h.professor: {h.data: {h.departamento: h.horas}}}
                for nome in complem:
                    for dia in complem[nome]:
                        for depart in complem[nome][dia]:
                            if depart == 'Musculação' and nome == i:
                                if plan1.cell(column=cell.column, row=cell.row + 1).value == dt.strftime(
                                        dt.strptime(dia, '%d/%m/%Y'), '%d/%m'):
                                    plan1.cell(column=cell.column, row=novalinha).value = plan1.cell(column=cell.column,
                                                                                                     row=novalinha).value + float(
                                        str(complem[nome][dia][depart]).replace(',', '.'))
                                    plan1.cell(column=cell.column, row=novalinha).fill = comple

                # aplica alterações de atestados
                # {a.professor: {a.data: a.departamento}}
                for nome in atest:
                    for d in atest[nome]:
                        if atest[nome][d] == 'Musculação' and nome == i:
                            if plan1.cell(column=cell.column, row=cell.row + 1).value == dt.strftime(
                                    dt.strptime(d, '%d/%m/%Y'), '%d/%m'):
                                plan1.cell(column=cell.column, row=novalinha).fill = atestado

                # aplica alterações de feriado
                # [datas de feriado formato dt]
                for dia in feriad:
                    if plan1.cell(column=cell.column, row=cell.row + 1).value == dt.strftime(dia, '%d/%m'):
                        plan1.cell(column=cell.column, row=novalinha).fill = feriado
                        plan1.cell(column=cell.column, row=novalinha).value = 0

                # aplica alterações de escala
                # {e.professor: {e.data: {e.departamento: e.horas}}}
                for nome in escal:
                    for dia in escal[nome]:
                        for depart in escal[nome][dia]:
                            if depart == 'Musculação' and nome == i:
                                if plan1.cell(column=cell.column, row=cell.row + 1).value == dt.strftime(
                                        dt.strptime(dia, '%d/%m/%Y'), '%d/%m'):
                                    plan1.cell(column=cell.column, row=novalinha).value = plan1.cell(column=cell.column,
                                                                                                     row=novalinha).value + float(
                                        str(escal[nome][dia][depart]).replace(',', '.'))
        plan1.cell(column=2, row=novalinha, value=i)
        novalinha += 1

    plan1[f'A{novalinha}'].value = 'Ginástica'
    novalinha += 1
    for i in ginastica:
        for row in plan1.iter_cols(min_row=3, min_col=3, max_row=3, max_col=35):
            for cell in row:
                if cell.value != 'Total':
                    if cell.value == 'seg':
                        plan1.cell(column=cell.column, row=novalinha, value=somar_aulas_de_segunda(i, 'Ginástica'))
                    if cell.value == 'ter':
                        plan1.cell(column=cell.column, row=novalinha, value=somar_aulas_de_terca(i, 'Ginástica'))
                    if cell.value == 'qua':
                        plan1.cell(column=cell.column, row=novalinha, value=somar_aulas_de_quarta(i, 'Ginástica'))
                    if cell.value == 'qui':
                        plan1.cell(column=cell.column, row=novalinha, value=somar_aulas_de_quinta(i, 'Ginástica'))
                    if cell.value == 'sex':
                        plan1.cell(column=cell.column, row=novalinha, value=somar_aulas_de_sexta(i, 'Ginástica'))
                    if cell.value == 'sáb':
                        plan1.cell(column=cell.column, row=novalinha, value=somar_aulas_de_sabado(i, 'Ginástica'))
                    if cell.value == 'dom':
                        plan1.cell(column=cell.column, row=novalinha, value=somar_aulas_de_domingo(i, 'Ginástica'))
                else:
                    letra = openpyxl.utils.cell.get_column_letter(cell.column - 1)
                    plan1.cell(column=cell.column, row=novalinha, value=f'=SUM(C{novalinha}:{letra}{novalinha})')
                # aplica cor de falta
                for nome in flt:
                    for dia in flt[nome]:
                        for depart in flt[nome][dia]:
                            if depart == 'Ginástica' and nome == i:
                                if plan1.cell(column=cell.column, row=cell.row + 1).value == dt.strftime(
                                        dt.strptime(dia, '%d/%m/%Y'), '%d/%m'):
                                    plan1.cell(column=cell.column, row=novalinha).fill = falta
                # aplica alterações de substituição
                # {s.professorsubst: {s.substituto: {s.departamento: {s.data: s.horas}}}}
                for nome in subs:
                    for substituto in subs[nome]:
                        for depart in subs[nome][substituto]:
                            for dia in subs[nome][substituto][depart]:
                                if depart == 'Ginástica' and nome == i:
                                    if plan1.cell(column=cell.column, row=cell.row + 1).value == dt.strftime(
                                            dt.strptime(dia, '%d/%m/%Y'), '%d/%m'):
                                        plan1.cell(column=cell.column, row=novalinha).fill = falta
                                if depart == 'Ginástica' and substituto == i:
                                    if plan1.cell(column=cell.column, row=cell.row + 1).value == dt.strftime(
                                            dt.strptime(dia, '%d/%m/%Y'), '%d/%m'):
                                        plan1.cell(column=cell.column, row=novalinha).value = plan1.cell(
                                            column=cell.column, row=novalinha).value + float(
                                            str(subs[nome][substituto][depart][dia]).replace(',', '.'))
                                        plan1.cell(column=cell.column, row=novalinha).fill = subst

                # aplica alterações de desligamento
                # {d.professor: {d.departamento: d.datadesligamento}}
                # conferir se tem outras aulas ativas ou foi desligado de tudo
                # se desligado de tudo, alterar status das aulas para inativas
                for nome in dslg:
                    for depart in dslg[nome]:
                        for dia in dslg[nome][depart]:
                            if depart == 'Ginástica' and nome == i and dt.strptime(dia, '%d/%m/%Y') <= fechamento:
                                if plan1.cell(column=cell.column, row=cell.row + 1).value is not None:
                                    if dt.strptime(dia, '%d/%m/%Y') <= dt(day=int(
                                            str(plan1.cell(column=cell.column, row=cell.row + 1).value).split('/')[0]),
                                                                          month=int(str(plan1.cell(column=cell.column,
                                                                                                   row=cell.row + 1).value).split(
                                                                                  '/')[1]),
                                                                          year=dt.today().year) <= fechamento:
                                        plan1.cell(column=cell.column, row=novalinha).value = 0
                                        plan1.cell(column=cell.column, row=novalinha).fill = deslig
                                        plan1.cell(column=cell.column, row=novalinha).font = Font(color='FFFFFF')

                # aplica talterações de férias
                # # {f.professor: {f.departamento: {f.inicio: f.fim}}}
                for nome in fer:
                    for depart in fer[nome]:
                        for inic in fer[nome][depart]:
                            if depart == 'Ginástica' and nome == i and dt.strptime(inic, '%d/%m/%Y') <= fechamento:
                                if plan1.cell(column=cell.column, row=cell.row + 1).value is not None:
                                    if dt.strptime(inic, '%d/%m/%Y') <= dt(day=int(
                                            str(plan1.cell(column=cell.column, row=cell.row + 1).value).split('/')[0]),
                                                                           month=int(str(plan1.cell(column=cell.column,
                                                                                                    row=cell.row + 1).value).split(
                                                                                   '/')[1]),
                                                                           year=dt.today().year) <= dt.strptime(
                                            fer[nome][depart][inic], '%d/%m/%Y'):
                                        plan1.cell(column=cell.column, row=novalinha).fill = ferias
                                        plan1.cell(column=cell.column, row=novalinha).value = 0

                # aplica alterações de horas complementares
                # {h.professor: {h.data: {h.departamento: h.horas}}}
                for nome in complem:
                    for dia in complem[nome]:
                        for depart in complem[nome][dia]:
                            if depart == 'Ginástica' and nome == i:
                                if plan1.cell(column=cell.column, row=cell.row + 1).value == dt.strftime(
                                        dt.strptime(dia, '%d/%m/%Y'), '%d/%m'):
                                    plan1.cell(column=cell.column, row=novalinha).value = plan1.cell(column=cell.column,
                                                                                                     row=novalinha).value + float(
                                        str(complem[nome][dia][depart]).replace(',', '.'))
                                    plan1.cell(column=cell.column, row=novalinha).fill = comple

                # aplica alterações de atestados
                # {a.professor: {a.data: a.departamento}}
                for nome in atest:
                    for d in atest[nome]:
                        if atest[nome][d] == 'Ginástica' and nome == i:
                            if plan1.cell(column=cell.column, row=cell.row + 1).value == dt.strftime(
                                    dt.strptime(d, '%d/%m/%Y'), '%d/%m'):
                                plan1.cell(column=cell.column, row=novalinha).fill = atestado

                # aplica alterações de feriado
                # [datas de feriado formato dt]
                for dia in feriad:
                    if plan1.cell(column=cell.column, row=cell.row + 1).value == dt.strftime(dia, '%d/%m'):
                        plan1.cell(column=cell.column, row=novalinha).fill = feriado
                        plan1.cell(column=cell.column, row=novalinha).value = 0

                # aplica alterações de escala
                # {e.professor: {e.data: {e.departamento: e.horas}}}
                for nome in escal:
                    for dia in escal[nome]:
                        for depart in escal[nome][dia]:
                            if depart == 'Ginástica' and nome == i:
                                if plan1.cell(column=cell.column, row=cell.row + 1).value == dt.strftime(
                                        dt.strptime(dia, '%d/%m/%Y'), '%d/%m'):
                                    plan1.cell(column=cell.column, row=novalinha).value = plan1.cell(column=cell.column,
                                                                                                     row=novalinha).value + float(
                                        str(escal[nome][dia][depart]).replace(',', '.'))

        plan1.cell(column=2, row=novalinha, value=i)
        novalinha += 1

    plan1[f'A{novalinha}'].value = 'Kids'
    novalinha += 1
    for i in kids:
        for row in plan1.iter_cols(min_row=3, min_col=3, max_row=3, max_col=35):
            for cell in row:
                if cell.value != 'Total':
                    if cell.value == 'seg':
                        plan1.cell(column=cell.column, row=novalinha, value=somar_aulas_de_segunda(i, 'Kids'))
                    if cell.value == 'ter':
                        plan1.cell(column=cell.column, row=novalinha, value=somar_aulas_de_terca(i, 'Kids'))
                    if cell.value == 'qua':
                        plan1.cell(column=cell.column, row=novalinha, value=somar_aulas_de_quarta(i, 'Kids'))
                    if cell.value == 'qui':
                        plan1.cell(column=cell.column, row=novalinha, value=somar_aulas_de_quinta(i, 'Kids'))
                    if cell.value == 'sex':
                        plan1.cell(column=cell.column, row=novalinha, value=somar_aulas_de_sexta(i, 'Kids'))
                    if cell.value == 'sáb':
                        plan1.cell(column=cell.column, row=novalinha, value=somar_aulas_de_sabado(i, 'Kids'))
                    if cell.value == 'dom':
                        plan1.cell(column=cell.column, row=novalinha, value=somar_aulas_de_domingo(i, 'Kids'))
                else:
                    letra = openpyxl.utils.cell.get_column_letter(cell.column - 1)
                    plan1.cell(column=cell.column, row=novalinha, value=f'=SUM(C{novalinha}:{letra}{novalinha})')
                # aplica cor de falta
                for nome in flt:
                    for dia in flt[nome]:
                        for depart in flt[nome][dia]:
                            if depart == 'Kids' and nome == i:
                                if plan1.cell(column=cell.column, row=cell.row + 1).value == dt.strftime(
                                        dt.strptime(dia, '%d/%m/%Y'), '%d/%m'):
                                    plan1.cell(column=cell.column, row=novalinha).fill = falta
                # aplica alterações de substituição
                # {s.professorsubst: {s.substituto: {s.departamento: {s.data: s.horas}}}}
                for nome in subs:
                    for substituto in subs[nome]:
                        for depart in subs[nome][substituto]:
                            for dia in subs[nome][substituto][depart]:
                                if depart == 'Kids' and nome == i:
                                    if plan1.cell(column=cell.column, row=cell.row + 1).value == dt.strftime(
                                            dt.strptime(dia, '%d/%m/%Y'), '%d/%m'):
                                        plan1.cell(column=cell.column, row=novalinha).fill = falta
                                if depart == 'Kids' and substituto == i:
                                    if plan1.cell(column=cell.column, row=cell.row + 1).value == dt.strftime(
                                            dt.strptime(dia, '%d/%m/%Y'), '%d/%m'):
                                        plan1.cell(column=cell.column, row=novalinha).value = plan1.cell(
                                            column=cell.column, row=novalinha).value + float(
                                            str(subs[nome][substituto][depart][dia]).replace(',', '.'))
                                        plan1.cell(column=cell.column, row=novalinha).fill = subst

                # aplica alterações de desligamento
                # {d.professor: {d.departamento: d.datadesligamento}}
                # conferir se tem outras aulas ativas ou foi desligado de tudo
                # se desligado de tudo, alterar status das aulas para inativas
                for nome in dslg:
                    for depart in dslg[nome]:
                        if depart == 'Kids' and nome == i and dt.strptime(dslg[nome][depart], '%d/%m/%Y') <= fechamento:
                            if plan1.cell(column=cell.column, row=cell.row + 1).value is not None:
                                if dt.strptime(dslg[nome][depart], '%d/%m/%Y') <= dt(day=int(
                                        str(plan1.cell(column=cell.column, row=cell.row + 1).value).split('/')[0]),
                                                                                     month=int(str(plan1.cell(
                                                                                             column=cell.column,
                                                                                             row=cell.row + 1).value).split(
                                                                                             '/')[1]),
                                                                                     year=dt.today().year) <= fechamento:
                                    plan1.cell(column=cell.column, row=novalinha).value = 0
                                    plan1.cell(column=cell.column, row=novalinha).fill = deslig
                                    plan1.cell(column=cell.column, row=novalinha).font = Font(color='FFFFFF')

                # aplica talterações de férias
                # # {f.professor: {f.departamento: {f.inicio: f.fim}}}
                for nome in fer:
                    for depart in fer[nome]:
                        for inic in fer[nome][depart]:
                            if depart == 'Kids' and nome == i and dt.strptime(inic, '%d/%m/%Y') <= fechamento:
                                if plan1.cell(column=cell.column, row=cell.row + 1).value is not None:
                                    if dt.strptime(inic, '%d/%m/%Y') <= dt(day=int(
                                            str(plan1.cell(column=cell.column, row=cell.row + 1).value).split('/')[0]),
                                                                           month=int(str(plan1.cell(column=cell.column,
                                                                                                    row=cell.row + 1).value).split(
                                                                                   '/')[1]),
                                                                           year=dt.today().year) <= dt.strptime(
                                            fer[nome][depart][inic], '%d/%m/%Y'):
                                        plan1.cell(column=cell.column, row=novalinha).fill = ferias
                                        plan1.cell(column=cell.column, row=novalinha).value = 0

                # aplica alterações de horas complementares
                # {h.professor: {h.data: {h.departamento: h.horas}}}
                for nome in complem:
                    for dia in complem[nome]:
                        for depart in complem[nome][dia]:
                            if depart == 'Kids' and nome == i:
                                if plan1.cell(column=cell.column, row=cell.row + 1).value == dt.strftime(
                                        dt.strptime(dia, '%d/%m/%Y'), '%d/%m'):
                                    plan1.cell(column=cell.column, row=novalinha).value = plan1.cell(column=cell.column,
                                                                                                     row=novalinha).value + float(
                                        str(complem[nome][dia][depart]).replace(',', '.'))
                                    plan1.cell(column=cell.column, row=novalinha).fill = comple

                # aplica alterações de atestados
                # {a.professor: {a.data: a.departamento}}
                for nome in atest:
                    for d in atest[nome]:
                        if atest[nome][d] == 'Kids' and nome == i:
                            if plan1.cell(column=cell.column, row=cell.row + 1).value == dt.strftime(
                                    dt.strptime(d, '%d/%m/%Y'), '%d/%m'):
                                plan1.cell(column=cell.column, row=novalinha).fill = atestado

                # aplica alterações de feriado
                # [datas de feriado formato dt]
                for dia in feriad:
                    if plan1.cell(column=cell.column, row=cell.row + 1).value == dt.strftime(dia, '%d/%m'):
                        plan1.cell(column=cell.column, row=novalinha).fill = feriado
                        plan1.cell(column=cell.column, row=novalinha).value = 0

                # aplica alterações de escala
                # {e.professor: {e.data: {e.departamento: e.horas}}}
                for nome in escal:
                    for dia in escal[nome]:
                        for depart in escal[nome][dia]:
                            if depart == 'Kids' and nome == i:
                                if plan1.cell(column=cell.column, row=cell.row + 1).value == dt.strftime(
                                        dt.strptime(dia, '%d/%m/%Y'), '%d/%m'):
                                    plan1.cell(column=cell.column, row=novalinha).value = plan1.cell(column=cell.column,
                                                                                                     row=novalinha).value + float(
                                        str(escal[nome][dia][depart]).replace(',', '.'))
        plan1.cell(column=2, row=novalinha, value=i)
        novalinha += 1

    plan1[f'A{novalinha}'].value = 'Esportes'
    novalinha += 1
    for i in esportes:
        for row in plan1.iter_cols(min_row=3, min_col=3, max_row=3, max_col=35):
            for cell in row:
                if cell.value != 'Total':
                    if cell.value == 'seg':
                        plan1.cell(column=cell.column, row=novalinha, value=somar_aulas_de_segunda(i, 'Esportes'))
                    if cell.value == 'ter':
                        plan1.cell(column=cell.column, row=novalinha, value=somar_aulas_de_terca(i, 'Esportes'))
                    if cell.value == 'qua':
                        plan1.cell(column=cell.column, row=novalinha, value=somar_aulas_de_quarta(i, 'Esportes'))
                    if cell.value == 'qui':
                        plan1.cell(column=cell.column, row=novalinha, value=somar_aulas_de_quinta(i, 'Esportes'))
                    if cell.value == 'sex':
                        plan1.cell(column=cell.column, row=novalinha, value=somar_aulas_de_sexta(i, 'Esportes'))
                    if cell.value == 'sáb':
                        plan1.cell(column=cell.column, row=novalinha, value=somar_aulas_de_sabado(i, 'Esportes'))
                    if cell.value == 'dom':
                        plan1.cell(column=cell.column, row=novalinha, value=somar_aulas_de_domingo(i, 'Esportes'))
                else:
                    letra = openpyxl.utils.cell.get_column_letter(cell.column - 1)
                    plan1.cell(column=cell.column, row=novalinha, value=f'=SUM(C{novalinha}:{letra}{novalinha})')
                # aplica cor de falta
                for nome in flt:
                    for dia in flt[nome]:
                        for depart in flt[nome][dia]:
                            if depart == 'Esportes' and nome == i:
                                if plan1.cell(column=cell.column, row=cell.row + 1).value == dt.strftime(
                                        dt.strptime(dia, '%d/%m/%Y'), '%d/%m'):
                                    plan1.cell(column=cell.column, row=novalinha).fill = falta
                # aplica alterações de substituição
                # {s.professorsubst: {s.substituto: {s.departamento: {s.data: s.horas}}}}
                for nome in subs:
                    for substituto in subs[nome]:
                        for depart in subs[nome][substituto]:
                            for dia in subs[nome][substituto][depart]:
                                if depart == 'Esportes' and nome == i:
                                    if plan1.cell(column=cell.column, row=cell.row + 1).value == dt.strftime(
                                            dt.strptime(dia, '%d/%m/%Y'), '%d/%m'):
                                        plan1.cell(column=cell.column, row=novalinha).fill = falta
                                if depart == 'Esportes' and substituto == i:
                                    if plan1.cell(column=cell.column, row=cell.row + 1).value == dt.strftime(
                                            dt.strptime(dia, '%d/%m/%Y'), '%d/%m'):
                                        plan1.cell(column=cell.column, row=novalinha).value = plan1.cell(
                                            column=cell.column, row=novalinha).value + float(
                                            str(subs[nome][substituto][depart][dia]).replace(',', '.'))
                                        plan1.cell(column=cell.column, row=novalinha).fill = subst

                # aplica alterações de desligamento
                # {d.professor: {d.departamento: d.datadesligamento}}
                # conferir se tem outras aulas ativas ou foi desligado de tudo
                # se desligado de tudo, alterar status das aulas para inativas
                for nome in dslg:
                    for depart in dslg[nome]:
                        for dia in dslg[nome][depart]:
                            if depart == 'Esportes' and nome == i and dt.strptime(dia, '%d/%m/%Y') <= fechamento:
                                if plan1.cell(column=cell.column, row=cell.row + 1).value is not None:
                                    if dt.strptime(dia, '%d/%m/%Y') <= dt(day=int(
                                            str(plan1.cell(column=cell.column, row=cell.row + 1).value).split('/')[0]),
                                                                          month=int(str(plan1.cell(column=cell.column,
                                                                                                   row=cell.row + 1).value).split(
                                                                                  '/')[1]),
                                                                          year=dt.today().year) <= fechamento:
                                        plan1.cell(column=cell.column, row=novalinha).value = 0
                                        plan1.cell(column=cell.column, row=novalinha).fill = deslig
                                        plan1.cell(column=cell.column, row=novalinha).font = Font(color='FFFFFF')

                # aplica talterações de férias
                # {f.professor: {f.departamento: {f.inicio: f.fim}}}
                for nome in fer:
                    for depart in fer[nome]:
                        for inic in fer[nome][depart]:
                            if depart == 'Esportes' and nome == i and dt.strptime(inic, '%d/%m/%Y') <= fechamento:
                                if plan1.cell(column=cell.column, row=cell.row + 1).value is not None:
                                    if dt.strptime(inic, '%d/%m/%Y') <= dt(day=int(
                                            str(plan1.cell(column=cell.column, row=cell.row + 1).value).split('/')[0]),
                                                                           month=int(str(plan1.cell(column=cell.column,
                                                                                                    row=cell.row + 1).value).split(
                                                                                   '/')[1]),
                                                                           year=dt.today().year) <= dt.strptime(
                                            fer[nome][depart][inic], '%d/%m/%Y'):
                                        plan1.cell(column=cell.column, row=novalinha).fill = ferias
                                        plan1.cell(column=cell.column, row=novalinha).value = 0

                # aplica alterações de horas complementares
                # {h.professor: {h.data: {h.departamento: h.horas}}}
                for nome in complem:
                    for dia in complem[nome]:
                        for depart in complem[nome][dia]:
                            if depart == 'Esportes' and nome == i:
                                if plan1.cell(column=cell.column, row=cell.row + 1).value == dt.strftime(
                                        dt.strptime(dia, '%d/%m/%Y'), '%d/%m'):
                                    plan1.cell(column=cell.column, row=novalinha).value = plan1.cell(column=cell.column,
                                                                                                     row=novalinha).value + float(
                                        str(complem[nome][dia][depart]).replace(',', '.'))
                                    plan1.cell(column=cell.column, row=novalinha).fill = comple

                # aplica alterações de atestados
                # {a.professor: {a.data: a.departamento}}
                for nome in atest:
                    for d in atest[nome]:
                        if atest[nome][d] == 'Esportes' and nome == i:
                            if plan1.cell(column=cell.column, row=cell.row + 1).value == dt.strftime(
                                    dt.strptime(d, '%d/%m/%Y'), '%d/%m'):
                                plan1.cell(column=cell.column, row=novalinha).fill = atestado

                # aplica alterações de feriado
                # [datas de feriado formato dt]
                for dia in feriad:
                    if plan1.cell(column=cell.column, row=cell.row + 1).value == dt.strftime(dia, '%d/%m'):
                        plan1.cell(column=cell.column, row=novalinha).fill = feriado
                        plan1.cell(column=cell.column, row=novalinha).value = 0

                # aplica alterações de escala
                # {e.professor: {e.data: {e.departamento: e.horas}}}
                for nome in escal:
                    for dia in escal[nome]:
                        for depart in escal[nome][dia]:
                            if depart == 'Esportes' and nome == i:
                                if plan1.cell(column=cell.column, row=cell.row + 1).value == dt.strftime(
                                        dt.strptime(dia, '%d/%m/%Y'), '%d/%m'):
                                    plan1.cell(column=cell.column, row=novalinha).value = plan1.cell(column=cell.column,
                                                                                                     row=novalinha).value + float(
                                        str(escal[nome][dia][depart]).replace(',', '.'))

        plan1.cell(column=2, row=novalinha, value=i)
        novalinha += 1

    plan1[f'A{novalinha}'].value = 'Cross Cia'
    novalinha += 1
    for i in cross:
        for row in plan1.iter_cols(min_row=3, min_col=3, max_row=3, max_col=35):
            for cell in row:
                if cell.value != 'Total':
                    if cell.value == 'seg':
                        plan1.cell(column=cell.column, row=novalinha, value=somar_aulas_de_segunda(i, 'Cross Cia'))
                    if cell.value == 'ter':
                        plan1.cell(column=cell.column, row=novalinha, value=somar_aulas_de_terca(i, 'Cross Cia'))
                    if cell.value == 'qua':
                        plan1.cell(column=cell.column, row=novalinha, value=somar_aulas_de_quarta(i, 'Cross Cia'))
                    if cell.value == 'qui':
                        plan1.cell(column=cell.column, row=novalinha, value=somar_aulas_de_quinta(i, 'Cross Cia'))
                    if cell.value == 'sex':
                        plan1.cell(column=cell.column, row=novalinha, value=somar_aulas_de_sexta(i, 'Cross Cia'))
                    if cell.value == 'sáb':
                        plan1.cell(column=cell.column, row=novalinha, value=somar_aulas_de_sabado(i, 'Cross Cia'))
                    if cell.value == 'dom':
                        plan1.cell(column=cell.column, row=novalinha, value=somar_aulas_de_domingo(i, 'Cross Cia'))
                else:
                    letra = openpyxl.utils.cell.get_column_letter(cell.column - 1)
                    plan1.cell(column=cell.column, row=novalinha, value=f'=SUM(C{novalinha}:{letra}{novalinha})')
                # aplica cor de falta
                for nome in flt:
                    for dia in flt[nome]:
                        for depart in flt[nome][dia]:
                            if depart == 'Cross Cia' and nome == i:
                                if plan1.cell(column=cell.column, row=cell.row + 1).value == dt.strftime(
                                        dt.strptime(dia, '%d/%m/%Y'), '%d/%m'):
                                    plan1.cell(column=cell.column, row=novalinha).fill = falta
                # aplica alterações de substituição
                # {s.professorsubst: {s.substituto: {s.departamento: {s.data: s.horas}}}}
                for nome in subs:
                    for substituto in subs[nome]:
                        for depart in subs[nome][substituto]:
                            for dia in subs[nome][substituto][depart]:
                                if depart == 'Cross Cia' and nome == i:
                                    if plan1.cell(column=cell.column, row=cell.row + 1).value == dt.strftime(
                                            dt.strptime(dia, '%d/%m/%Y'), '%d/%m'):
                                        plan1.cell(column=cell.column, row=novalinha).fill = falta
                                if depart == 'Cross Cia' and substituto == i:
                                    if plan1.cell(column=cell.column, row=cell.row + 1).value == dt.strftime(
                                            dt.strptime(dia, '%d/%m/%Y'), '%d/%m'):
                                        plan1.cell(column=cell.column, row=novalinha).value = plan1.cell(
                                            column=cell.column, row=novalinha).value + float(
                                            str(subs[nome][substituto][depart][dia]).replace(',', '.'))
                                        plan1.cell(column=cell.column, row=novalinha).fill = subst

                # aplica alterações de desligamento
                # {d.professor: {d.departamento: d.datadesligamento}}
                # conferir se tem outras aulas ativas ou foi desligado de tudo
                # se desligado de tudo, alterar status das aulas para inativas
                for nome in dslg:
                    for depart in dslg[nome]:
                        for dia in dslg[nome][depart]:
                            if depart == 'Cross Cia' and nome == i and dt.strptime(dia, '%d/%m/%Y') <= fechamento:
                                if plan1.cell(column=cell.column, row=cell.row + 1).value is not None:
                                    if dt.strptime(dia, '%d/%m/%Y') <= dt(day=int(
                                            str(plan1.cell(column=cell.column, row=cell.row + 1).value).split('/')[0]),
                                                                          month=int(str(plan1.cell(column=cell.column,
                                                                                                   row=cell.row + 1).value).split(
                                                                                  '/')[1]),
                                                                          year=dt.today().year) <= fechamento:
                                        plan1.cell(column=cell.column, row=novalinha).value = 0
                                        plan1.cell(column=cell.column, row=novalinha).fill = deslig
                                        plan1.cell(column=cell.column, row=novalinha).font = Font(color='FFFFFF')

                # aplica talterações de férias
                # # {f.professor: {f.departamento: {f.inicio: f.fim}}}
                for nome in fer:
                    for depart in fer[nome]:
                        for inic in fer[nome][depart]:
                            if depart == 'Cross Cia' and nome == i and dt.strptime(inic, '%d/%m/%Y') <= fechamento:
                                if plan1.cell(column=cell.column, row=cell.row + 1).value is not None:
                                    if dt.strptime(inic, '%d/%m/%Y') <= dt(day=int(
                                            str(plan1.cell(column=cell.column, row=cell.row + 1).value).split('/')[0]),
                                                                           month=int(str(plan1.cell(column=cell.column,
                                                                                                    row=cell.row + 1).value).split(
                                                                                   '/')[1]),
                                                                           year=dt.today().year) <= dt.strptime(
                                            fer[nome][depart][inic], '%d/%m/%Y'):
                                        plan1.cell(column=cell.column, row=novalinha).fill = ferias
                                        plan1.cell(column=cell.column, row=novalinha).value = 0

                # aplica alterações de horas complementares
                # {h.professor: {h.data: {h.departamento: h.horas}}}
                for nome in complem:
                    for dia in complem[nome]:
                        for depart in complem[nome][dia]:
                            if depart == 'Cross Cia' and nome == i:
                                if plan1.cell(column=cell.column, row=cell.row + 1).value == dt.strftime(
                                        dt.strptime(dia, '%d/%m/%Y'), '%d/%m'):
                                    plan1.cell(column=cell.column, row=novalinha).value = plan1.cell(column=cell.column,
                                                                                                     row=novalinha).value + float(
                                        str(complem[nome][dia][depart]).replace(',', '.'))
                                    plan1.cell(column=cell.column, row=novalinha).fill = comple

                # aplica alterações de atestados
                # {a.professor: {a.data: a.departamento}}
                for nome in atest:
                    for d in atest[nome]:
                        if atest[nome][d] == 'Cross Cia' and nome == i:
                            if plan1.cell(column=cell.column, row=cell.row + 1).value == dt.strftime(
                                    dt.strptime(d, '%d/%m/%Y'), '%d/%m'):
                                plan1.cell(column=cell.column, row=novalinha).fill = atestado

                # aplica alterações de feriado
                # [datas de feriado formato dt]
                for dia in feriad:
                    if plan1.cell(column=cell.column, row=cell.row + 1).value == dt.strftime(dia, '%d/%m'):
                        plan1.cell(column=cell.column, row=novalinha).fill = feriado
                        plan1.cell(column=cell.column, row=novalinha).value = 0

                # aplica alterações de escala
                # {e.professor: {e.data: {e.departamento: e.horas}}}
                for nome in escal:
                    for dia in escal[nome]:
                        for depart in escal[nome][dia]:
                            if depart == 'Cross Cia' and nome == i:
                                if plan1.cell(column=cell.column, row=cell.row + 1).value == dt.strftime(
                                        dt.strptime(dia, '%d/%m/%Y'), '%d/%m'):
                                    plan1.cell(column=cell.column, row=novalinha).value = plan1.cell(column=cell.column,
                                                                                                     row=novalinha).value + float(
                                        str(escal[nome][dia][depart]).replace(',', '.'))

        plan1.cell(column=2, row=novalinha, value=i)
        novalinha += 1

    for i, coluna in enumerate(plan1.columns):
        max_length = 0
        column = coluna[0].column_letter
        for cell in coluna:
            try:
                if len(str(cell.value)) > max_length and len(str(cell.value)) != 18:
                    max_length = len(str(cell.value))
            except Exception:
                pass
        adjusted_width = max_length + 1
        plan1.column_dimensions[column].width = adjusted_width

    for row in plan1.iter_cols(min_row=3, min_col=3, max_row=3, max_col=35):
        for cell in row:
            if plan1.cell(column=cell.column, row=cell.row).value == 'Total':
                plan1.column_dimensions[openpyxl.utils.cell.get_column_letter(cell.column)].width = 8

    plan1['C1'].fill = atestado
    plan1['D1'].value = 'Atestado'
    plan1['C2'].fill = falta
    plan1['D2'].value = 'Falta'
    plan1['F1'].fill = ferias
    plan1['G1'].value = 'Férias'
    plan1['F2'].fill = feriado
    plan1['G2'].value = 'Feriado'
    plan1['I1'].fill = deslig
    plan1['J1'].value = 'Desligamento'
    plan1['I2'].fill = subst
    plan1['J2'].value = 'Substituiu'
    plan1['M1'].fill = comple
    plan1['N1'].value = 'Horas Complementares'
    grade.save(f'Grade {fechamento.month}-{fechamento.year}.xlsx')


def somar_aulas_de_segunda(nome: str, depto: str) -> float:
    sessions = sessionmaker(bind=enginefolha)
    session = sessions()
    aulasseg = session.query(Aulas).filter_by(professor=nome).filter_by(diadasemana='Segunda') \
        .filter_by(departamento=depto).all()
    somas = 0
    for aula in aulasseg:
        hr, minut, seg = str(dt.strptime(aula.fim, '%H:%M:%S') - dt.strptime(aula.inicio, '%H:%M:%S')).split(':')
        somadia = int(hr) + int(minut) / 60 + int(seg) / 60 * 60
        somadia = round(somadia, 2)
        somas += somadia
    return somas


def somar_aulas_de_terca(nome: str, depto: str) -> float:
    sessions = sessionmaker(bind=enginefolha)
    session = sessions()
    aulaster = session.query(Aulas).filter_by(professor=nome).filter_by(diadasemana='Terça') \
        .filter_by(departamento=depto).all()
    somas = 0
    for aula in aulaster:
        hr, minut, seg = str(dt.strptime(aula.fim, '%H:%M:%S') - dt.strptime(aula.inicio, '%H:%M:%S')).split(':')
        somadia = int(hr) + int(minut) / 60 + int(seg) / 60 * 60
        somadia = round(somadia, 2)
        somas += somadia
    return somas


def somar_aulas_de_quarta(nome: str, depto: str) -> float:
    sessions = sessionmaker(bind=enginefolha)
    session = sessions()
    aulasqua = session.query(Aulas).filter_by(professor=nome).filter_by(diadasemana='Quarta') \
        .filter_by(departamento=depto).all()
    somas = 0
    for aula in aulasqua:
        hr, minut, seg = str(dt.strptime(aula.fim, '%H:%M:%S') - dt.strptime(aula.inicio, '%H:%M:%S')).split(':')
        somadia = int(hr) + int(minut) / 60 + int(seg) / 60 * 60
        somadia = round(somadia, 2)
        somas += somadia
    return somas


def somar_aulas_de_quinta(nome: str, depto: str) -> float:
    sessions = sessionmaker(bind=enginefolha)
    session = sessions()
    aulasqui = session.query(Aulas).filter_by(professor=nome).filter_by(diadasemana='Quinta') \
        .filter_by(departamento=depto).all()
    somas = 0
    for aula in aulasqui:
        hr, minut, seg = str(dt.strptime(aula.fim, '%H:%M:%S') - dt.strptime(aula.inicio, '%H:%M:%S')).split(':')
        somadia = int(hr) + int(minut) / 60 + int(seg) / 60 * 60
        somadia = round(somadia, 2)
        somas += somadia
    return somas


def somar_aulas_de_sexta(nome: str, depto: str) -> float:
    sessions = sessionmaker(bind=enginefolha)
    session = sessions()
    aulassex = session.query(Aulas).filter_by(professor=nome).filter_by(diadasemana='Sexta') \
        .filter_by(departamento=depto).all()
    somas = 0
    for aula in aulassex:
        hr, minut, seg = str(dt.strptime(aula.fim, '%H:%M:%S') - dt.strptime(aula.inicio, '%H:%M:%S')).split(':')
        somadia = int(hr) + int(minut) / 60 + int(seg) / 60 * 60
        somadia = round(somadia, 2)
        somas += somadia
    return somas


def somar_aulas_de_sabado(nome: str, depto: str) -> float:
    sessions = sessionmaker(bind=enginefolha)
    session = sessions()
    aulassab = session.query(Aulas).filter_by(professor=nome).filter_by(diadasemana='Sábado') \
        .filter_by(departamento=depto).all()
    somas = 0
    for aula in aulassab:
        hr, minut, seg = str(dt.strptime(aula.fim, '%H:%M:%S') - dt.strptime(aula.inicio, '%H:%M:%S')).split(':')
        somadia = int(hr) + int(minut) / 60 + int(seg) / 60 * 60
        somadia = round(somadia, 2)
        somas += somadia
    return somas


def somar_aulas_de_domingo(nome: str, depto: str) -> float:
    sessions = sessionmaker(bind=enginefolha)
    session = sessions()
    aulasdom = session.query(Aulas).filter_by(professor=nome).filter_by(diadasemana='Domingo') \
        .filter_by(departamento=depto).all()
    somas = 0
    for aula in aulasdom:
        hr, minut, seg = str(dt.strptime(aula.fim, '%H:%M:%S') - dt.strptime(aula.inicio, '%H:%M:%S')).split(':')
        somadia = int(hr) + int(minut) / 60 + int(seg) / 60 * 60
        somadia = round(somadia, 2)
        somas += somadia
    return somas


def salvar_planilha_soma_final(compet: int):
    sessions = sessionmaker(bind=enginefolha)
    session = sessions()
    folhadehoje = Folha(compet, list(listar_aulas_ativas()), listar_departamentos_ativos())
    somaaulas = {}
    for i in listar_professores_ativos():
        somaaulas[i] = {}
        for d in listar_departamentos_ativos():
            somaaulas[i][d] = {}
    for aulas in listar_aulas_ativas():
        somaaulas[aulas.professor][aulas.departamento][aulas.nome + f' ({aulas.valor})'] = round(
            somar_horas_professor(folhadehoje, aulas.professor, aulas.departamento, aulas.nome, compet), 2)
        dictchav = list(somaaulas.keys())
        dictchav.sort()
        somafinal = {i: somaaulas[i] for i in dictchav}
    arquivo = os.path.relpath(rf'C:\Users\{os.getlogin()}\PycharmProjects\AutomacaoCia\src\models\static\files\Somafinal.xlsx')
    plan = l_w(arquivo, read_only=False)
    folha = plan['Planilha1']
    folha['A1'].value = 'Matrícula'
    folha['B1'].value = 'Nome'
    folha['C1'].value = 'Aula e Valor'
    folha['D1'].value = 'Horas'
    x = 2
    for i in somafinal:
        matr = session.query(Aulas).filter_by(professor=str(i)).first()
        folha[f'A{x}'].value = int(matr.matrprof)
        folha[f'B{x}'].value = str(i)
        for sub in somafinal[i]:
            if somafinal[i][sub] == {}:
                pass
            else:
                for sub2 in somafinal[i][sub]:
                    folha[f'C{x}'].value = str(sub2)
                    folha[f'D{x}'].value = float(str(somafinal[i][sub][sub2]))
                    x += 1
    # folha['F1'].value = 'Total Bruto - Professores'
    # folha['G1'].value = locale.currency(totaldafolha(folhadehoje), grouping=True)
    plan.save(f'Somafinal mes {compet}.xlsx')
    salvar_planilha_grade_horaria(somafinal, compet)
    substitutos = {}
    complementares = {}
    feriasl = {}
    desligadosl = {}
    planilha = l_w(f'Grade {compet}-2023.xlsx')
    aba = planilha['Planilha1']
    for row in aba.iter_cols(min_row=3, min_col=3, max_row=115, max_col=35):
        for cell in row:
            if cell.fill == ferias:
                for i in range(1, 150):
                    if aba.cell(column=1, row=cell.row - i).value is not None:
                        depart = aba.cell(column=1, row=cell.row - i).value
                        break
                for r in aba.iter_cols(min_row=3, min_col=3, max_row=3, max_col=38):
                    for c in r:
                        if c.value == 'Total':
                            tt = c.column
                hrs = 0
                for m in range(3, tt):
                    hrs += aba.cell(column=m, row=cell.row).value
                feriasl[aba.cell(column=2, row=cell.row).value] = {depart: round(hrs, 2)}
            if cell.fill == deslig:
                for i in range(1, 150):
                    if aba.cell(column=1, row=cell.row - i).value is not None:
                        depart = aba.cell(column=1, row=cell.row - i).value
                        break
                for r in aba.iter_cols(min_row=3, min_col=3, max_row=3, max_col=38):
                    for c in r:
                        if c.value == 'Total':
                            tt = c.column
                hrs = 0
                for m in range(3, tt):
                    hrs += aba.cell(column=m, row=cell.row).value
                desligadosl[aba.cell(column=2, row=cell.row).value] = {depart: round(hrs, 2)}
            if cell.fill == subst:
                for i in range(1, 150):
                    if aba.cell(column=1, row=cell.row - i).value is not None:
                        depart = aba.cell(column=1, row=cell.row - i).value
                        break
                for r in aba.iter_cols(min_row=3, min_col=3, max_row=3, max_col=38):
                    for c in r:
                        if c.value == 'Total':
                            tt = c.column
                hrs = 0
                for m in range(3, tt):
                    hrs += aba.cell(column=m, row=cell.row).value
                substitutos[aba.cell(column=2, row=cell.row).value] = {depart: round(hrs, 2)}
            if cell.fill == comple:
                for i in range(1, 150):
                    if aba.cell(column=1, row=cell.row - i).value is not None:
                        depart = aba.cell(column=1, row=cell.row - i).value
                        break
                for r in aba.iter_cols(min_row=3, min_col=3, max_row=3, max_col=38):
                    for c in r:
                        if c.value == 'Total':
                            tt = c.column
                hrs = 0
                for m in range(3, tt):
                    hrs += aba.cell(column=m, row=cell.row).value
                complementares[aba.cell(column=2, row=cell.row).value] = {depart: round(hrs, 2)}

    planilha2 = l_w(f'Somafinal mes {compet}.xlsx', read_only=False)
    aba2 = planilha2['Planilha1']
    for pessoa in feriasl:
        for depart in feriasl[pessoa]:
            for cll in aba2['B']:
                if cll.value is not None:
                    if cll.value == pessoa:
                        aba2.cell(column=4, row=cll.row, value=float(feriasl[pessoa][depart]))
    for pessoa in desligadosl:
        for depart in desligadosl[pessoa]:
            for cell in aba2['B']:
                if cell.value is not None:
                    if cell.value == pessoa:
                        aba2.cell(column=4, row=cell.row).value = float(desligadosl[pessoa][depart])
    for pessoa in substitutos:
        for depart in substitutos[pessoa]:
            for cell in aba2['B']:
                if cell.value is not None:
                    if cell.value == pessoa:
                        aba2.cell(column=4, row=cell.row).value = float(substitutos[pessoa][depart])
    for pessoa in complementares:
        for depart in complementares[pessoa]:
            for cell in aba2['B']:
                if cell.value is not None:
                    if cell.value == pessoa:
                        aba2.cell(column=4, row=cell.row).value = float(complementares[pessoa][depart])
    for i, coluna in enumerate(aba2.columns):
        max_length = 0
        column = coluna[0].column_letter
        for cell in coluna:
            try:
                if len(str(cell.value)) > max_length:
                    max_length = len(str(cell.value))
            except Exception:
                pass
        adjusted_width = max_length + 1
        aba2.column_dimensions[column].width = adjusted_width
    planilha2.save(f'Somafinal mes {compet}.xlsx')
    tkinter.messagebox.showinfo(
        title='Grade ok!',
        message=f'Grade do mês {compet} salva com sucesso!'
    )
    planilha3 = l_w(f'Grade {compet}-2023.xlsx', read_only=False)
    aba3 = planilha3['Planilha1']
    for row in aba3.iter_cols(min_row=3, min_col=3, max_row=120, max_col=35):
        for cell in row:
            if cell.value == 0:
                cell.value = ''

    planilha3.save(f'Grade {compet}-2023.xlsx')
    print('Férias \n', feriasl)
    print('Desligados \n', desligadosl)
    print('Substitutos \n', substitutos)
    print('Hrs Complementares \n', complementares)


def cadastro_funcionario(caminho='', editar=0, ondestou=0, nome='', matricula='', admissao='',
                         horario='', salario='', cargo='', depto='', tipo_contr='',
                         hrsem='', hrmens='', agencia='', conta='', digito=''):
    sessions = sessionmaker(bind=engine)
    session = sessions()
    if caminho == '' or nome == '' or matricula == '' or admissao == '' or horario == '' or salario == '' or \
            cargo == '' or depto == '' or tipo_contr == '' or hrsem == '' or hrmens == '':
        tkinter.messagebox.showinfo(
            title='Erro de preenchimento',
            message='Preencha todos os campos antes de cadastrar o funcionário!'
        )
    else:
        wb = l_w(caminho, read_only=False)
        sh = wb['Respostas ao formulário 1']
        num, name = nome.strip().split(' - ')
        linha = int(num)

        # search for the highest compatibility between the city filled in the form and the cities in the lists to
        # define codmunnas value
        est = str(sh[f'AJ{linha}'].value)
        cidade = str(sh[f'L{linha}'].value).title()
        lista = []
        dicion = {}
        for cid in municipios[est]:
            dicion[SequenceMatcher(None, cidade, cid).ratio()] = cid
            lista.append(SequenceMatcher(None, cidade, cid).ratio())
        codmunnas = municipios[str(sh[f'AJ{linha}'].value).upper().strip()][dicion[max(lista)]]

        # search for the highest compatibility between the city filled in the form and the cities in the lists to
        # define codmunend value
        est = str(sh[f'T{linha}'].value)
        cidade = str(sh[f'S{linha}'].value).title()
        listaend = []
        dicionend = {}
        for cid in municipios[est]:
            dicionend[SequenceMatcher(None, cidade, cid).ratio()] = cid
            listaend.append(SequenceMatcher(None, cidade, cid).ratio())
        codmunend = municipios[str(sh[f'T{linha}'].value).upper().strip()][dicionend[max(listaend)]]

        lotacao = {'UNIDADE PARK SUL - QUALQUER DEPARTAMENTO': '0013', 'KIDS': '0010', 'MUSCULAÇÃO': '0007',
                   'ESPORTES E LUTAS': '0008', 'CROSSFIT': '0012', 'GINÁSTICA': '0006', 'GESTANTES': '0008',
                   'RECEPÇÃO': '0003',
                   'FINANCEIRO': '0001', 'TI': '0001', 'MARKETING': '0001', 'MANUTENÇÃO': '0004'}
        if editar == 0:
            if ondestou == 0:
                # Cadastro iniciado na Cia
                if linha:
                    pess = Colaborador(matricula=matricula, nome=name.upper(), admiss=admissao,
                                       nascimento=str(sh[f'D{linha}'].value),
                                       pis=str(int(sh[f'Y{linha}'].value)).zfill(11),
                                       cpf=str(int(sh[f'V{linha}'].value)).zfill(11),
                                       rg=str(int(sh[f'W{linha}'].value)),
                                       emissor=str(sh[f'X{linha}'].value), email=str(sh[f'B{linha}'].value),
                                       genero=str(sh[f'E{linha}'].value),
                                       estado_civil=str(sh[f'F{linha}'].value), cor=str(sh[f'G{linha}'].value),
                                       instru=str(sh[f'J{linha}'].value),
                                       nacional=str(sh[f'K{linha}'].value),
                                       cod_municipionas=codmunnas,
                                       cid_nas=str(sh[f'L{linha}'].value), uf_nas=str(sh[f'AJ{linha}'].value),
                                       pai=str(sh[f'M{linha}'].value).upper(),
                                       mae=str(sh[f'N{linha}'].value).upper(), endereco=str(sh[f'O{linha}'].value),
                                       num=str(int(sh[f'P{linha}'].value)),
                                       bairro=str(sh[f'Q{linha}'].value), cep=str(int(sh[f'R{linha}'].value)),
                                       cidade=str(sh[f'S{linha}'].value),
                                       uf=str(sh[f'T{linha}'].value),
                                       cod_municipioend=codmunend,
                                       tel=str(int(sh[f'U{linha}'].value)),
                                       tit_eleit=str(sh[f'Z{linha}'].value), zona_eleit=str(sh[f'AA{linha}'].value),
                                       sec_eleit=str(sh[f'AB{linha}'].value),
                                       ctps=str(int(sh[f'AC{linha}'].value)), serie_ctps=str(sh[f'AD{linha}'].value),
                                       uf_ctps=str(sh[f'AE{linha}'].value),
                                       emiss_ctps=str(sh[f'AF{linha}'].value), depto=depto,
                                       cargo=cargo,
                                       horario=horario, salario=salario, tipo_contr=tipo_contr, hr_sem=hrsem,
                                       hr_mens=hrmens,
                                       ag=agencia, conta=conta, cdigito=digito
                                       )
                    session.add(pess)
                    session.commit()
                    pessoa = session.query(Colaborador).filter_by(matricula=matricula).first()
                    pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Dexion.png')))
                    pa.press('alt'), pa.press('a'), pa.press('t'), t.sleep(2), pa.press(
                        'i'), t.sleep(5), pa.write(str(pessoa.matricula)), pa.press('enter'), t.sleep(20)
                    pp.copy(pessoa.nome), pa.hotkey('ctrl', 'v'), pa.press('tab'), pa.write(pessoa.cpf)
                    t.sleep(1), pa.press('tab', 3), pa.write(pessoa.genero), pa.press('tab'), pa.write(pessoa.cor)
                    t.sleep(1), pa.press('tab', 2), pa.write(pessoa.instru)
                    t.sleep(1), pa.press('tab'), pa.write(pessoa.estado_civil)
                    if str(pessoa.estado_civil) == '2 - Casado(a)':
                        pa.press('tab', 6)
                    else:
                        pa.press('tab', 5)
                    pa.write(dt.strftime(dt.strptime(pessoa.nascimento, '%Y-%m-%d %H:%M:%S'), '%d%m%Y'))
                    t.sleep(1), pa.press('tab'), pp.copy(pessoa.cid_nas), pa.hotkey('ctrl', 'v'), pa.press('tab')
                    t.sleep(1), pa.write(pessoa.uf_nas), pa.press('tab'), pa.write(pessoa.cod_municipionas), pa.press(
                        'tab')
                    t.sleep(1), pp.copy(pessoa.pai), pa.hotkey('ctrl', 'v'), pa.press('tab'), pa.write('105'), pa.press(
                        'tab')
                    t.sleep(1), pp.copy(pessoa.mae), pa.hotkey('ctrl', 'v'), pa.press('tab'), pa.write('105')
                    # #clique em documentos
                    pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Documentos.png')))
                    pa.press('tab'), pa.write(str(pessoa.rg)), pa.press('tab'), pa.write(
                        pessoa.emissor), pa.press('tab', 3), pa.write(pessoa.cod_municipioend),
                    pa.press('tab'), pa.write(pessoa.pis), pa.press('enter')
                    pa.press('tab'), pa.write(pessoa.tit_eleit), pa.press('tab'), pa.write(pessoa.zona_eleit), pa.press(
                        'tab'), pa.write(
                        pessoa.sec_eleit), pa.press('tab'), pa.write(pessoa.ctps)
                    pa.press('tab'), pa.write(pessoa.serie_ctps), pa.press('tab'), pa.write(pessoa.uf_ctps), pa.press(
                        'tab'), pa.write(dt.strftime(dt.strptime(pessoa.emiss_ctps, '%Y-%m-%d %H:%M:%S'),
                                                           '%d%m%Y'))
                    pa.press('tab', 5), pa.write('3'), pa.press('tab'), pa.write('341'), pa.press('tab')
                    pa.write(pessoa.ag), pa.press('tab'), pa.write(f'{pessoa.conta}-{pessoa.cdigito}'), pa.press('tab')

                    # #clique em endereço
                    pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Endereco.png')))
                    pa.press('tab', 2), pp.copy(pessoa.endereco), pa.hotkey('ctrl', 'v'), pa.press(
                        'tab'), pa.write(pessoa.num), pa.press('tab', 2)
                    pp.copy(pessoa.bairro), pa.hotkey('ctrl', 'v'), pa.press('tab'), pp.copy(pessoa.cidade), pa.hotkey(
                        'ctrl', 'v'), pa.press('tab'), pa.write(pessoa.uf)
                    pa.press('tab'), pa.write(pessoa.cep), pa.press('tab'), pa.write(pessoa.cod_municipioend), pa.press(
                        'tab'), pa.write(str(pessoa.tel)), pa.press('tab', 2), pa.write(pessoa.email)
                    # #clique em dados contratuais
                    pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Contratuais.png')))
                    pa.press('tab'), pa.write(str(pessoa.admiss).replace('/', '')), t.sleep(1)
                    pa.press('tab', 10), pa.write('2')
                    # #clique em Contrato de Experiência
                    try:
                        pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Experiencia.png')))
                    except pa.ImageNotFoundException:
                        t.sleep(5)
                        pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Experiencia.png')))
                    pa.press('tab'), pa.write('45'), pa.press('tab'), pa.write('45'), pa.press(
                        'tab'), pa.press(
                        'space'), pa.press('tab', 2), pa.write('003')
                    pa.press('tab'), pa.write(str(pessoa.matricula))
                    # #clique em Outros
                    try:
                        pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Outros.png')))
                    except pa.ImageNotFoundException:
                        t.sleep(5)
                        pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Outros.png')))
                    t.sleep(2), pa.write('CARGO GERAL')
                    # #clique em lupa de descrição de cargos
                    pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Lupa.png')))
                    t.sleep(2), pp.copy(pessoa.cargo), pa.hotkey('ctrl', 'v'), t.sleep(1.5), pa.press('enter', 2)
                    t.sleep(1), pa.press('tab'), pa.write(pessoa.salario), pa.press('tab')
                    if str(pessoa.tipo_contr) == 'Horista':
                        pa.press('1')
                    else:
                        pa.press('5')
                    pa.press('tab', 2), t.sleep(1), pa.write(str(pessoa.hr_sem)), pa.press('tab'), t.sleep(1)
                    pa.write(str(pessoa.hr_mens))
                    pa.press('tab', 5), pa.write('00395419000190'), pa.press('tab', 2), pa.write('5')
                    # #clique em eventos trabalhistas
                    pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/EVTrab.png')))
                    t.sleep(1)
                    # #clique em lotação
                    pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Lotacoes.png')))
                    pa.press('tab'), pa.press('tab'), pa.write('i'), pa.write(str(pessoa.admiss).replace('/', ''))
                    t.sleep(1), pa.press('enter'), t.sleep(1)
                    pp.copy(lotacao[f'{pessoa.depto}']), pa.hotkey('ctrl', 'v'), pa.press('enter'), pa.write('f')
                    pa.press('tab'), pa.write('4')
                    pa.press('tab', 6), pa.write('i'), t.sleep(2), pa.press('tab'), pa.write(pessoa.horario)
                    t.sleep(3), pa.press('tab', 3), pa.press('enter'), t.sleep(3)
                    # #clique em cancelar novo registro de horario
                    pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Cancelarhor.png'))), t.sleep(2.5)
                    # #clique em salvar lotação
                    pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Salvarlot.png'))), t.sleep(1)
                    # #clique em fechar lotação
                    try:
                        pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Fecharlot.png'))), t.sleep(1)
                    except pa.ImageNotFoundException:
                        t.sleep(4)
                        pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Fecharlot.png')))
                    # #clique em Compatibilidade
                    pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Compatibilidade.png'))), t.sleep(1)
                    # #clique em Compatibilidade de novo
                    pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Compatibilidade2.png'))), t.sleep(1)
                    # #clique em CAGED
                    pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/CAGED.png')))
                    pa.press('tab', 2), pa.write('20'), pa.press('tab'), pa.write('1'), t.sleep(0.5)
                    # #clique em RAIS
                    pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/RAIS.png')))
                    pp.copy(pessoa.nome), pa.hotkey('ctrl', 'v'), pa.press('tab', 2), pa.write('2')
                    pa.press('tab'), pa.write('10')
                    # #clique em Salvar
                    pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Salvarcadastro.png'))), t.sleep(10)
                    # #clique em fechar novo cadastro
                    pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Fecharnovo1.png'))), t.sleep(2)
                    # #clique em fechar trabalhadores
                    pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Fechartrab1.png'))), t.sleep(0.5)
                    try:
                        os.makedirs(
                            r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e '
                            r'Férias\000 - Pastas Funcionais\00 - ATIVOS\{}'.format(pessoa.nome))
                        os.makedirs(
                            r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e '
                            r'Férias\000 - Pastas Funcionais\00 - ATIVOS\{}\Atestados'.format(pessoa.nome))
                        os.makedirs(
                            r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e '
                            r'Férias\000 - Pastas Funcionais\00 - ATIVOS\{}\Diversos'.format(pessoa.nome))
                        os.makedirs(
                            r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e '
                            r'Férias\000 - Pastas Funcionais\00 - ATIVOS\{}\Contratuais'.format(pessoa.nome))
                        os.makedirs(
                            r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e '
                            r'Férias\000 - Pastas Funcionais\00 - ATIVOS\{}\Férias'.format(pessoa.nome))
                        os.makedirs(
                            r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e '
                            r'Férias\000 - Pastas Funcionais\00 - ATIVOS\{}\Pontos'.format(pessoa.nome))
                        os.makedirs(
                            r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e '
                            r'Férias\000 - Pastas Funcionais\00 - ATIVOS\{}\Recibos'.format(pessoa.nome))
                        os.makedirs(
                            r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e '
                            r'Férias\000 - Pastas Funcionais\00 - ATIVOS\{}\Rescisão'.format(pessoa.nome))
                    except FileExistsError:
                        pass
                    tkinter.messagebox.showinfo(
                        title='Cadastro ok!',
                        message='Cadastro realizado com sucesso!'
                    )
            else:
                # Cadastro iniciado em casa
                wb = l_w(caminho, read_only=False)
                sh = wb['Respostas ao formulário 1']
                num, name = nome.strip().split(' - ')
                linha = int(num)
                if linha:
                    pess = Colaborador(matricula=matricula, nome=name.upper(), admiss=admissao,
                                       nascimento=str(sh[f'D{linha}'].value),
                                       pis=str(int(sh[f'Y{linha}'].value)).zfill(11),
                                       cpf=str(int(sh[f'V{linha}'].value)).zfill(11),
                                       rg=str(int(sh[f'W{linha}'].value)),
                                       emissor=str(sh[f'X{linha}'].value), email=str(sh[f'B{linha}'].value),
                                       genero=str(sh[f'E{linha}'].value),
                                       estado_civil=str(sh[f'F{linha}'].value), cor=str(sh[f'G{linha}'].value),
                                       instru=str(sh[f'J{linha}'].value),
                                       nacional=str(sh[f'K{linha}'].value),
                                       cod_municipionas=codmunnas,
                                       cid_nas=str(sh[f'L{linha}'].value), uf_nas=str(sh[f'AJ{linha}'].value),
                                       pai=str(sh[f'M{linha}'].value).upper(),
                                       mae=str(sh[f'N{linha}'].value).upper(), endereco=str(sh[f'O{linha}'].value),
                                       num=str(int(sh[f'P{linha}'].value)),
                                       bairro=str(sh[f'Q{linha}'].value), cep=str(int(sh[f'R{linha}'].value)),
                                       cidade=str(sh[f'S{linha}'].value),
                                       uf=str(sh[f'T{linha}'].value),
                                       cod_municipioend=codmunend,
                                       tel=str(int(sh[f'U{linha}'].value)),
                                       tit_eleit=str(sh[f'Z{linha}'].value), zona_eleit=str(sh[f'AA{linha}'].value),
                                       sec_eleit=str(sh[f'AB{linha}'].value),
                                       ctps=str(int(sh[f'AC{linha}'].value)), serie_ctps=str(sh[f'AD{linha}'].value),
                                       uf_ctps=str(sh[f'AE{linha}'].value),
                                       emiss_ctps=str(sh[f'AF{linha}'].value), depto=depto,
                                       cargo=cargo,
                                       horario=horario, salario=salario, tipo_contr=tipo_contr, hr_sem=hrsem,
                                       hr_mens=hrmens,
                                       ag=agencia, conta=conta, cdigito=digito
                                       )
                    session.add(pess)
                    session.commit()
                    pessoa = session.query(Colaborador).filter_by(matricula=matricula).first()
                    pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Dexion.png')))
                    pa.press('alt'), pa.press('a'), pa.press('t'), t.sleep(2), pa.press(
                        'i'), t.sleep(5), pa.write(str(pessoa.matricula)), pa.press('enter'), t.sleep(60)
                    pp.copy(pessoa.nome), pa.hotkey('ctrl', 'v'), pa.press('tab'), pa.write(pessoa.cpf)
                    t.sleep(1), pa.press('tab', 3), pa.write(pessoa.genero), pa.press('tab'), pa.write(pessoa.cor)
                    t.sleep(1), pa.press('tab', 2), pa.write(pessoa.instru)
                    t.sleep(1), pa.press('tab'), pa.write(pessoa.estado_civil)
                    if str(pessoa.estado_civil) == '2 - Casado(a)':
                        pa.press('tab', 6)
                    else:
                        pa.press('tab', 5)
                    pa.write(dt.strftime(dt.strptime(pessoa.nascimento, '%Y-%m-%d %H:%M:%S'), '%d%m%Y'))
                    t.sleep(1), pa.press('tab'), pp.copy(pessoa.cid_nas), pa.hotkey('ctrl', 'v'), pa.press('tab')
                    t.sleep(1), pa.write(pessoa.uf_nas), pa.press('tab'), pa.write(pessoa.cod_municipionas), pa.press(
                        'tab')
                    t.sleep(1), pp.copy(pessoa.pai), pa.hotkey('ctrl', 'v'), pa.press('tab'), pa.write('105'), pa.press(
                        'tab')
                    t.sleep(1), pp.copy(pessoa.mae), pa.hotkey('ctrl', 'v'), pa.press('tab'), pa.write('105')
                    # #clique em documentos
                    pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Documentos.png')))
                    pa.press('tab'), pa.write(str(pessoa.rg)), pa.press('tab'), pa.write(
                        pessoa.emissor), pa.press('tab', 3), pa.write(pessoa.cod_municipioend),
                    pa.press('tab'), pa.write(pessoa.pis), pa.press('enter')
                    pa.press('tab'), pa.write(pessoa.tit_eleit), pa.press('tab'), pa.write(pessoa.zona_eleit), pa.press(
                        'tab'), pa.write(
                        pessoa.sec_eleit), pa.press('tab'), pa.write(pessoa.ctps)
                    pa.press('tab'), pa.write(pessoa.serie_ctps), pa.press('tab'), pa.write(pessoa.uf_ctps), pa.press(
                        'tab'), pa.write(dt.strftime(dt.strptime(pessoa.emiss_ctps, '%Y-%m-%d %H:%M:%S'),
                                                           '%d%m%Y'))
                    pa.press('tab', 5), pa.write('3'), pa.press('tab'), pa.write('341'), pa.press('tab')
                    pa.write(pessoa.ag), pa.press('tab'), pa.write(f'{pessoa.conta}-{pessoa.cdigito}'), pa.press('tab')
                    # #clique em endereço
                    pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Endereco.png')))
                    pa.press('tab', 2), pp.copy(pessoa.endereco), pa.hotkey('ctrl', 'v'), pa.press(
                        'tab'), pa.write(pessoa.num), pa.press('tab', 2)
                    pp.copy(pessoa.bairro), pa.hotkey('ctrl', 'v'), pa.press('tab'), pp.copy(pessoa.cidade), pa.hotkey(
                        'ctrl', 'v'), pa.press('tab'), pa.write(pessoa.uf)
                    pa.press('tab'), pa.write(pessoa.cep), pa.press('tab'), pa.write(pessoa.cod_municipioend), pa.press(
                        'tab'), pa.write(str(pessoa.tel)), pa.press('tab', 2), pa.write(pessoa.email)
                    # #clique em dados contratuais
                    pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Contratuais.png')))
                    pa.press('tab'), pa.write(str(pessoa.admiss).replace('/', '')), t.sleep(1)
                    pa.press('tab', 10), pa.write('2')
                    # #clique em Contrato de Experiência
                    try:
                        pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Experiencia.png')))
                    except pa.ImageNotFoundException:
                        t.sleep(5)
                        pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Experiencia.png')))
                    pa.press('tab'), pa.write('45'), pa.press('tab'), pa.write('45'), pa.press(
                        'tab'), pa.press(
                        'space'), pa.press('tab', 2), pa.write('003')
                    pa.press('tab'), pa.write(str(pessoa.matricula))
                    # #clique em Outros
                    try:
                        pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Outros.png')))
                    except pa.ImageNotFoundException:
                        t.sleep(5)
                        pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Outros.png')))
                    t.sleep(2), pa.write('CARGO GERAL')
                    # #clique em lupa de descrição de cargos
                    pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Lupa.png')))
                    t.sleep(2), pp.copy(pessoa.cargo), pa.hotkey('ctrl', 'v'), t.sleep(1.5), pa.press('enter', 2)
                    t.sleep(1), pa.press('tab'), pa.write(pessoa.salario), pa.press('tab')
                    if str(pessoa.tipo_contr) == 'Horista':
                        pa.press('1')
                    else:
                        pa.press('5')
                    pa.press('tab', 2), t.sleep(1), pa.write(str(pessoa.hr_sem)), pa.press('tab'), t.sleep(1)
                    pa.write(str(pessoa.hr_mens))
                    pa.press('tab', 5), pa.write('00395419000190'), pa.press('tab', 2), pa.write('5')
                    # #clique em eventos trabalhistas
                    pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/EVTrab.png')))
                    t.sleep(1)
                    # #clique em lotação
                    pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Lotacoes.png')))
                    pa.press('tab'), pa.press('tab'), pa.write('i'), pa.write(str(pessoa.admiss).replace('/', ''))
                    t.sleep(1), pa.press('enter'), t.sleep(1)
                    pp.copy(lotacao[f'{pessoa.depto}']), pa.hotkey('ctrl', 'v'), pa.press('enter'), pa.write('f')
                    pa.press('tab'), pa.write('4')
                    pa.press('tab', 6), pa.write('i'), t.sleep(2), pa.press('tab'), pa.write(pessoa.horario)
                    t.sleep(3), pa.press('tab', 3), pa.press('enter'), t.sleep(3)
                    # #clique em cancelar novo registro de horario
                    pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Cancelarhor.png'))), t.sleep(2.5)
                    # #clique em salvar lotação
                    pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Salvarlot.png'))), t.sleep(1)
                    # #clique em fechar lotação
                    try:
                        pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Fecharlot.png'))), t.sleep(1)
                    except pa.ImageNotFoundException:
                        t.sleep(4)
                        pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Fecharlot.png')))
                    # #clique em Compatibilidade
                    pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Compatibilidade.png'))), t.sleep(1)
                    # #clique em Compatibilidade de novo
                    pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Compatibilidade2.png'))), t.sleep(1)
                    # #clique em CAGED
                    pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/CAGED.png')))
                    pa.press('tab', 2), pa.write('20'), pa.press('tab'), pa.write('1'), t.sleep(0.5)
                    # #clique em RAIS
                    pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/RAIS.png')))
                    pp.copy(pessoa.nome), pa.hotkey('ctrl', 'v'), pa.press('tab', 2), pa.write('2')
                    pa.press('tab'), pa.write('10')
                    # #clique em Salvar
                    pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Salvarcadastro.png'))), t.sleep(10)
                    # #clique em fechar novo cadastro
                    pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Fecharnovo1.png'))), t.sleep(2)
                    # #clique em fechar trabalhadores
                    pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Fechartrab1.png'))), t.sleep(0.5)

                    os.makedirs(
                        r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e '
                        r'Férias\000 - Pastas Funcionais\00 - ATIVOS\{}'.format(pessoa.nome))
                    os.makedirs(
                        r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e '
                        r'Férias\000 - Pastas Funcionais\00 - ATIVOS\{}\Atestados'.format(pessoa.nome))
                    os.makedirs(
                        r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e '
                        r'Férias\000 - Pastas Funcionais\00 - ATIVOS\{}\Diversos'.format(pessoa.nome))
                    os.makedirs(
                        r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e '
                        r'Férias\000 - Pastas Funcionais\00 - ATIVOS\{}\Contratuais'.format(pessoa.nome))
                    os.makedirs(
                        r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e '
                        r'Férias\000 - Pastas Funcionais\00 - ATIVOS\{}\Férias'.format(pessoa.nome))
                    os.makedirs(
                        r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e '
                        r'Férias\000 - Pastas Funcionais\00 - ATIVOS\{}\Pontos'.format(pessoa.nome))
                    os.makedirs(
                        r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e '
                        r'Férias\000 - Pastas Funcionais\00 - ATIVOS\{}\Recibos'.format(pessoa.nome))
                    os.makedirs(
                        r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e '
                        r'Férias\000 - Pastas Funcionais\00 - ATIVOS\{}\Rescisão'.format(pessoa.nome))
                    tkinter.messagebox.showinfo(
                        title='Cadastro ok!',
                        message='Cadastro realizado com sucesso!'
                    )

        else:
            if ondestou == 0:
                # Cadastro EDITADO na Cia
                num, name = nome.strip().split(' - ')
                linha = int(num)
                if linha:
                    pessoa = session.query(Colaborador).filter_by(matricula=matricula).first()
                    pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Dexion.png')))
                    pa.press('alt'), pa.press('a'), pa.press('t'), t.sleep(2), pa.press(
                        'a'), t.sleep(5), pa.write(str(pessoa.matricula)), pa.press('enter'), t.sleep(15)
                    pp.copy(pessoa.nome), pa.hotkey('ctrl', 'v'), pa.press('tab'), pa.write(pessoa.cpf)
                    pa.press('tab', 3), pa.write(pessoa.genero), pa.press('tab'), pa.write(pessoa.cor)
                    pa.press('tab', 2), pa.write(pessoa.instru)
                    pa.press('tab'), pa.write(pessoa.estado_civil)
                    if str(pessoa.estado_civil) == '2 - Casado(a)':
                        pa.press('tab', 6)
                    else:
                        pa.press('tab', 5)
                    pa.write(dt.strftime(dt.strptime(pessoa.nascimento, '%Y-%m-%d %H:%M:%S'), '%d%m%Y'))
                    pa.press('tab'), pp.copy(pessoa.cid_nas), pa.hotkey('ctrl', 'v'), pa.press('tab')
                    pa.write(pessoa.uf_nas), pa.press('tab'), pa.write(pessoa.cod_municipionas), pa.press('tab')
                    pp.copy(pessoa.pai), pa.hotkey('ctrl', 'v'), pa.press('tab'), pa.write('105'), pa.press('tab')
                    pp.copy(pessoa.mae), pa.hotkey('ctrl', 'v'), pa.press('tab'), pa.write('105')
                    # #clique em documentos
                    pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Documentos.png')))
                    pa.press('tab'), pa.write(str(pessoa.rg)), pa.press('tab'), pa.write(
                        pessoa.emissor), pa.press('tab', 3), pa.write(pessoa.cod_municipioend),
                    pa.press('tab'), pa.write(pessoa.pis), pa.press('enter')
                    pa.press('tab'), pa.write(pessoa.tit_eleit), pa.press('tab'), pa.write(pessoa.zona_eleit), pa.press(
                        'tab'), pa.write(
                        pessoa.sec_eleit), pa.press('tab'), pa.write(pessoa.ctps)
                    pa.press('tab'), pa.write(pessoa.serie_ctps), pa.press('tab'), pa.write(pessoa.uf_ctps), pa.press(
                        'tab'), pa.write(dt.strftime(dt.strptime(pessoa.emiss_ctps, '%Y-%m-%d %H:%M:%S'),
                                                           '%d%m%Y'))
                    pa.press('tab', 5), pa.write('3'), pa.press('tab'), pa.write('341'), pa.press('tab')
                    pa.write(pessoa.ag), pa.press('tab'), pa.write(f'{pessoa.conta}-{pessoa.cdigito}'), pa.press('tab')

                    # #clique em endereço
                    pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Endereco.png')))
                    pa.press('tab', 2), pp.copy(pessoa.endereco), pa.hotkey('ctrl', 'v'), pa.press(
                        'tab'), pa.write(pessoa.num), pa.press('tab', 2)
                    pp.copy(pessoa.bairro), pa.hotkey('ctrl', 'v'), pa.press('tab'), pp.copy(pessoa.cidade), pa.hotkey(
                        'ctrl', 'v'), pa.press('tab'), pa.write(pessoa.uf)
                    pa.press('tab'), pa.write(pessoa.cep), pa.press('tab'), pa.write(pessoa.cod_municipioend), pa.press(
                        'tab'), pa.write(str(pessoa.tel)), pa.press('tab', 2), pa.write(pessoa.email)
                    # #clique em dados contratuais
                    pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Contratuais.png')))
                    pa.press('tab'), pa.write(str(pessoa.admiss).replace('/', '')), t.sleep(1)
                    pa.press('tab', 10), pa.write('2')
                    # #clique em Contrato de Experiência
                    try:
                        pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Experiencia.png')))
                    except pa.ImageNotFoundException:
                        t.sleep(5)
                        pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Experiencia.png')))
                    pa.press('tab'), pa.write('45'), pa.press('tab'), pa.write('45'), pa.press(
                        'tab'), pa.press(
                        'space'), pa.press('tab', 2), pa.write('003')
                    pa.press('tab'), pa.write(str(pessoa.matricula))
                    # #clique em Outros
                    try:
                        pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Outros.png')))
                    except pa.ImageNotFoundException:
                        t.sleep(5)
                        pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Outros.png')))
                    t.sleep(2), pa.write('CARGO GERAL')
                    # #clique em lupa de descrição de cargos
                    pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Lupa.png')))
                    t.sleep(2), pp.copy(pessoa.cargo), pa.hotkey('ctrl', 'v'), t.sleep(1.5), pa.press('enter', 2)
                    t.sleep(1), pa.press('tab'), pa.write(pessoa.salario), pa.press('tab')
                    if str(pessoa.tipo_contr) == 'Horista':
                        pa.press('1')
                    else:
                        pa.press('5')
                    pa.press('tab', 2), t.sleep(1), pa.write(str(pessoa.hr_sem)), pa.press('tab'), t.sleep(1)
                    pa.write(str(pessoa.hr_mens))
                    pa.press('tab', 5), pa.write('00395419000190'), pa.press('tab', 2), pa.write('5')
                    # #clique em Compatibilidade
                    pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Compatibilidade1.png'))), t.sleep(1)
                    # #clique em Compatibilidade de novo
                    pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Compatibilidade2.png'))), t.sleep(1)
                    # #clique em CAGED
                    pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/CAGED.png')))
                    pa.press('tab', 2), pa.write('20'), pa.press('tab'), pa.write('1'), t.sleep(0.5)
                    # #clique em RAIS
                    pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/RAIS.png')))
                    pp.copy(pessoa.nome), pa.hotkey('ctrl', 'v'), pa.press('tab', 2), pa.write('2')
                    pa.press('tab'), pa.write('10')
                    # #clique em Salvar
                    pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Salvarcadastro.png'))), t.sleep(10)
                    # #clique em fechar novo cadastro
                    pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Fecharnovo1.png'))), t.sleep(2)
                    # #clique em fechar trabalhadores
                    pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Fechartrab1.png'))), t.sleep(0.5)

                    os.makedirs(
                        r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e '
                        r'Férias\000 - Pastas Funcionais\00 - ATIVOS\{}'.format(pessoa.nome))
                    os.makedirs(
                        r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e '
                        r'Férias\000 - Pastas Funcionais\00 - ATIVOS\{}\Atestados'.format(pessoa.nome))
                    os.makedirs(
                        r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e '
                        r'Férias\000 - Pastas Funcionais\00 - ATIVOS\{}\Diversos'.format(pessoa.nome))
                    os.makedirs(
                        r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e '
                        r'Férias\000 - Pastas Funcionais\00 - ATIVOS\{}\Contratuais'.format(pessoa.nome))
                    os.makedirs(
                        r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e '
                        r'Férias\000 - Pastas Funcionais\00 - ATIVOS\{}\Férias'.format(pessoa.nome))
                    os.makedirs(
                        r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e '
                        r'Férias\000 - Pastas Funcionais\00 - ATIVOS\{}\Pontos'.format(pessoa.nome))
                    os.makedirs(
                        r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e '
                        r'Férias\000 - Pastas Funcionais\00 - ATIVOS\{}\Recibos'.format(pessoa.nome))
                    os.makedirs(
                        r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e '
                        r'Férias\000 - Pastas Funcionais\00 - ATIVOS\{}\Rescisão'.format(pessoa.nome))
                    tkinter.messagebox.showinfo(title='Cadastro ok!',
                                                message='Cadastro editado com sucesso!')
                else:
                    # Cadastro EDITADO em casa
                    num, name = nome.strip().split(' - ')
                    linha = int(num)
                    if linha:
                        pessoa = session.query(Colaborador).filter_by(matricula=matricula).first()
                        pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Dexion.png')))
                        pa.press('alt'), pa.press('a'), pa.press('t'), t.sleep(2), pa.press(
                            'a'), t.sleep(5), pa.write(str(pessoa.matricula)), pa.press('enter'), t.sleep(20)
                        pp.copy(pessoa.nome), pa.hotkey('ctrl', 'v'), pa.press('tab'), pa.write(pessoa.cpf)
                        t.sleep(1), pa.press('tab', 3), pa.write(pessoa.genero), pa.press('tab'), pa.write(pessoa.cor)
                        t.sleep(1), pa.press('tab', 2), pa.write(pessoa.instru)
                        t.sleep(1), pa.press('tab'), pa.write(pessoa.estado_civil)
                        if str(pessoa.estado_civil) == '2 - Casado(a)':
                            pa.press('tab', 6)
                        else:
                            pa.press('tab', 5)
                        pa.write(dt.strftime(dt.strptime(pessoa.nascimento, '%Y-%m-%d %H:%M:%S'), '%d%m%Y'))
                        t.sleep(1), pa.press('tab'), pp.copy(pessoa.cid_nas), pa.hotkey('ctrl', 'v'), pa.press('tab')
                        t.sleep(1), pa.write(pessoa.uf_nas), pa.press('tab'), pa.write(
                            pessoa.cod_municipionas), pa.press('tab')
                        t.sleep(1), pp.copy(pessoa.pai), pa.hotkey('ctrl', 'v'), pa.press('tab'), pa.write(
                            '105'), pa.press('tab')
                        t.sleep(1), pp.copy(pessoa.mae), pa.hotkey('ctrl', 'v'), pa.press('tab'), pa.write('105')
                        # #clique em documentos
                        pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Documentos.png')))
                        pa.press('tab'), pa.write(str(pessoa.rg)), pa.press('tab'), pa.write(
                            pessoa.emissor), pa.press('tab', 3), pa.write(pessoa.cod_municipioend),
                        pa.press('tab'), pa.write(pessoa.pis), pa.press('enter')
                        pa.press('tab'), pa.write(pessoa.tit_eleit), pa.press('tab'), pa.write(
                            pessoa.zona_eleit), pa.press('tab'), pa.write(
                            pessoa.sec_eleit), pa.press('tab'), pa.write(pessoa.ctps)
                        pa.press('tab'), pa.write(pessoa.serie_ctps), pa.press('tab'), pa.write(
                            pessoa.uf_ctps), pa.press(
                            'tab'), pa.write(
                            dt.strftime(dt.strptime(pessoa.emiss_ctps, '%Y-%m-%d %H:%M:%S'),
                                              '%d%m%Y'))
                        pa.press('tab', 5), pa.write('3'), pa.press('tab'), pa.write('341'), pa.press('tab')
                        pa.write(pessoa.ag), pa.press('tab'), pa.write(f'{pessoa.conta}-{pessoa.cdigito}'), pa.press(
                            'tab')
                        # #clique em endereço
                        pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Endereco.png')))
                        pa.press('tab', 2), pp.copy(pessoa.endereco), pa.hotkey('ctrl', 'v'), pa.press(
                            'tab'), pa.write(pessoa.num), pa.press('tab', 2)
                        pp.copy(pessoa.bairro), pa.hotkey('ctrl', 'v'), pa.press('tab'), pp.copy(
                            pessoa.cidade), pa.hotkey(
                            'ctrl', 'v'), pa.press('tab'), pa.write(pessoa.uf)
                        pa.press('tab'), pa.write(pessoa.cep), pa.press('tab'), pa.write(
                            pessoa.cod_municipioend), pa.press(
                            'tab'), pa.write(str(pessoa.tel)), pa.press('tab', 2), pa.write(pessoa.email)
                        # #clique em dados contratuais
                        pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Contratuais.png')))
                        pa.press('tab'), pa.write(str(pessoa.admiss).replace('/', '')), t.sleep(1)
                        pa.press('tab', 10), pa.write('2')
                        # #clique em Contrato de Experiência
                        try:
                            pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Experiencia.png')))
                        except pa.ImageNotFoundException:
                            t.sleep(5)
                            pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Experiencia.png')))
                        pa.press('tab'), pa.write('45'), pa.press('tab'), pa.write('45'), pa.press(
                            'tab'), pa.press(
                            'space'), pa.press('tab', 2), pa.write('003')
                        pa.press('tab'), pa.write(str(pessoa.matricula))
                        # #clique em Outros
                        try:
                            pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Outros.png')))
                        except pa.ImageNotFoundException:
                            t.sleep(5)
                            pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Outros.png')))
                        t.sleep(2), pa.write('CARGO GERAL')
                        # #clique em lupa de descrição de cargos
                        pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Lupa.png')))
                        t.sleep(2), pp.copy(pessoa.cargo), pa.hotkey('ctrl', 'v'), t.sleep(1.5), pa.press('enter', 2)
                        t.sleep(1), pa.press('tab'), pa.write(pessoa.salario), pa.press('tab')
                        if str(pessoa.tipo_contr) == 'Horista':
                            pa.press('1')
                        else:
                            pa.press('5')
                        pa.press('tab', 2), t.sleep(1), pa.write(str(pessoa.hr_sem)), pa.press('tab'), t.sleep(1)
                        pa.write(str(pessoa.hr_mens))
                        pa.press('tab', 5), pa.write('00395419000190'), pa.press('tab', 2), pa.write('5')
                        # #clique em Compatibilidade
                        pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Compatibilidade.png'))), t.sleep(1)
                        # #clique em Compatibilidade de novo
                        pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Compatibilidade2.png'))), t.sleep(1)
                        # #clique em CAGED
                        pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/CAGED.png')))
                        pa.press('tab', 2), pa.write('20'), pa.press('tab'), pa.write('1'), t.sleep(0.5)
                        # #clique em RAIS
                        pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/RAIS.png')))
                        pp.copy(pessoa.nome), pa.hotkey('ctrl', 'v'), pa.press('tab', 2), pa.write('2')
                        pa.press('tab'), pa.write('10')
                        # #clique em Salvar
                        pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Salvarcadastro.png'))), t.sleep(10)
                        # #clique em fechar novo cadastro
                        pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Fecharnovo1.png'))), t.sleep(2)
                        # #clique em fechar trabalhadores
                        pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Fechartrab1.png'))), t.sleep(0.5)

                        os.makedirs(
                            r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\
                            02 - Funcionários, Departamentos e '
                            r'Férias\000 - Pastas Funcionais\00 - ATIVOS\{}'.format(pessoa.nome))
                        os.makedirs(
                            r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\
                            02 - Funcionários, Departamentos e '
                            r'Férias\000 - Pastas Funcionais\00 - ATIVOS\{}\Atestados'.format(pessoa.nome))
                        os.makedirs(
                            r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\
                            02 - Funcionários, Departamentos e '
                            r'Férias\000 - Pastas Funcionais\00 - ATIVOS\{}\Diversos'.format(pessoa.nome))
                        os.makedirs(
                            r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\
                            02 - Funcionários, Departamentos e '
                            r'Férias\000 - Pastas Funcionais\00 - ATIVOS\{}\Contratuais'.format(pessoa.nome))
                        os.makedirs(
                            r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\
                            02 - Funcionários, Departamentos e '
                            r'Férias\000 - Pastas Funcionais\00 - ATIVOS\{}\Férias'.format(pessoa.nome))
                        os.makedirs(
                            r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\
                            02 - Funcionários, Departamentos e '
                            r'Férias\000 - Pastas Funcionais\00 - ATIVOS\{}\Pontos'.format(pessoa.nome))
                        os.makedirs(
                            r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\
                            02 - Funcionários, Departamentos e '
                            r'Férias\000 - Pastas Funcionais\00 - ATIVOS\{}\Recibos'.format(pessoa.nome))
                        os.makedirs(
                            r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\
                            02 - Funcionários, Departamentos e '
                            r'Férias\000 - Pastas Funcionais\00 - ATIVOS\{}\Rescisão'.format(pessoa.nome))
                        tkinter.messagebox.showinfo(
                            title='Cadastro ok!',
                            message='Cadastro editado com sucesso!'
                        )


def salvar_docs_funcionarios(matricula):
    sessions = sessionmaker(bind=engine)
    session = sessions()
    pessoa = session.query(Colaborador).filter_by(matricula=matricula).first()
    if pessoa is None:
        tkinter.messagebox.showinfo(
            title='Erro de preenchimento',
            message='Não existe funcionário cadastrado com essa matrícula!'
        )
    else:
        pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Dexion.png')))
        p_pessoa = r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e Férias' \
                   r'\000 - Pastas Funcionais\00 - ATIVOS\1 - Ainda nao iniciaram\{}'.format(pessoa.nome)
        p_contr = r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e Férias' \
                  r'\000 - Pastas Funcionais\00 - ATIVOS\1 - Ainda nao iniciaram\{}\Contratuais'.format(pessoa.nome)
        p_recibos = r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e Férias' \
                    r'\000 - Pastas Funcionais\00 - ATIVOS\Modelo\Recibo Crachá e Uniformes MODELO.docx'
        ps_acordo = r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e Férias' \
                    r'\000 - Pastas Funcionais\00 - ATIVOS\1 - Ainda nao iniciaram\{}\Contratuais\Acordo Banco de Horas.pdf'.format(pessoa.nome)
        ps_recctps = r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e Férias' \
                     r'\000 - Pastas Funcionais\00 - ATIVOS\1 - Ainda nao iniciaram\{}\Contratuais\Recibo de Entrega e Dev CTPS.pdf'.format(
                        pessoa.nome)
        ps_anotctps = r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e Férias' \
                      r'\000 - Pastas Funcionais\00 - ATIVOS\1 - Ainda nao iniciaram\{}\Contratuais\Anotacoes CTPS.pdf'.format(pessoa.nome)
        ps_termovt = r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e Férias' \
                     r'\000 - Pastas Funcionais\00 - ATIVOS\1 - Ainda nao iniciaram\{}\Contratuais\Termo Opcao VT.pdf'.format(pessoa.nome)
        ps_contrato = r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e Férias' \
                      r'\000 - Pastas Funcionais\00 - ATIVOS\1 - Ainda nao iniciaram\{}\Contratuais\Contrato de Trabalho.pdf'.format(pessoa.nome)
        ps_ficha = r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e Férias' \
                   r'\000 - Pastas Funcionais\00 - ATIVOS\1 - Ainda nao iniciaram\{}\Contratuais\Folha de Registro.pdf'.format(pessoa.nome)

        recibos = docx.Document(p_recibos)

        # # imprimir recibo entrega e devolução de ctps
        pa.press('alt'), pa.press('r'), pa.press('a'), pa.press('r'), pa.press('e'), pa.press('tab'), pa.write(str(
            pessoa.matricula))
        pa.press('tab', 3), pa.write(str(pessoa.admiss).replace('/', '')), pa.press('tab'), t.sleep(0.5), pa.press('space')
        t.sleep(0.5), pa.press('tab'), pa.write(str(pessoa.admiss).replace('/', '')), pa.press('tab', 2)
        t.sleep(1), pa.press('enter'), t.sleep(2)

        # # clique no endereço de salvamento do recibo
        pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/salvar.png'))), t.sleep(1)
        pp.copy(ps_recctps), pa.hotkey('ctrl', 'v'), t.sleep(0.5)
        pa.press('tab', 2), t.sleep(0.5), pa.press('enter')
        t.sleep(5)
        # # clique para fechar recibo ctps
        pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/fechar_janela.png'))), t.sleep(0.5)
        pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/fechar_janela.png')))

        # # Imprimir Acordo de Banco de horas
        pa.press('alt'), pa.press('r'), pa.press('z'), pa.press('d'), pa.press('d')
        pa.write("(matricula = '00{}')".format(str(pessoa.matricula))), t.sleep(1), pa.press('tab'), pa.write('2')
        pa.press('tab'), pa.press('enter'), t.sleep(10)
        pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/salvar.png')))
        t.sleep(1), pp.copy(ps_acordo)
        pa.hotkey('ctrl', 'v'), t.sleep(1), pa.press('enter'), t.sleep(15)
        # # clique para fechar acordo
        pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/fechar_janela.png'))), t.sleep(0.5)
        pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/fechar_janela.png')))

        # # imprimir Anotações em CTPS
        pa.press('alt'), pa.press('r'), pa.press('a'), pa.press('c'), pa.press('e'), pa.press('tab')
        pa.write(str(pessoa.matricula))
        pa.press('tab', 2), pa.write(str(pessoa.admiss).replace('/', '')), pa.press('tab')
        pa.write(str(pessoa.admiss).replace('/', '')), pa.press('tab', 4), pa.press('space')
        pa.press('tab'), pa.press('enter'), t.sleep(1.5)
        pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/salvar.png'))), t.sleep(1)
        pp.copy(ps_anotctps), pa.hotkey('ctrl', 'v'), t.sleep(0.5), pa.press('enter')
        t.sleep(2), pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/fechar_janela.png'))), t.sleep(0.5)
        pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/fechar_janela.png')))

        # # imprimir Termo VT
        pa.press('alt'), pa.press('r'), pa.press('a'), pa.press('v'), pa.press('e'), pa.press('tab')
        pa.write(str(pessoa.matricula)), pa.press('tab', 2), pa.write(str(pessoa.admiss).replace('/', ''))
        pa.press('tab'), pa.write('d'), pa.press('tab', 4), pa.press('space')
        pa.press('tab', 6), pa.press('enter'), t.sleep(1.5)
        pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/salvar.png')))
        pp.copy(ps_termovt), pa.hotkey('ctrl', 'v'), t.sleep(0.5), pa.press('enter')
        t.sleep(2), pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/fechar_janela.png'))), t.sleep(0.5)
        pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/fechar_janela.png')))

        # # Imprimir Contrato
        pa.press('alt'), pa.press('r'), pa.press('z'), pa.press('d')
        if pessoa.tipo_contr == 'Horista':
            pa.press('c')
        else:
            pa.press('o')

        pa.write("(matricula = '00{}')".format(str(pessoa.matricula))), t.sleep(1), pa.press('tab'), pa.write('2')
        pa.press('tab'), pa.press('enter'), t.sleep(5)
        pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/salvar.png')))
        pp.copy(ps_contrato), pa.hotkey('ctrl', 'v'), t.sleep(0.5), pa.press('enter')
        t.sleep(10), pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/fechar_janela.png'))), t.sleep(0.5)
        pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/fechar_janela.png')))

        # # Imprimir Folha de rosto de Cadastro
        pa.press('alt'), pa.press('r'), pa.press('i'), pa.press('o'), pa.press('r'), pa.press('e'), pa.press('tab')
        pa.write(str(pessoa.matricula)), pa.press('tab', 2), pa.write(str(pessoa.admiss).replace('/', '')), pa.press('tab',
                                                                                                                     2)
        pa.press('enter'), t.sleep(3)
        pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/salvar.png')))
        pp.copy(ps_ficha), pa.hotkey('ctrl', 'v'), t.sleep(0.5), pa.press('enter')
        t.sleep(3), pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/fechar_janela.png'))), t.sleep(0.5)
        pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/fechar_janela.png')))

        # # Alterar Recibos e salvar na pasta
        recibos.paragraphs[4].text = str(recibos.paragraphs[4].text).replace('#nome_completo', pessoa.nome)
        recibos.paragraphs[12].text = str(recibos.paragraphs[12].text).replace('#nome_completo', pessoa.nome)
        recibos.paragraphs[19].text = str(recibos.paragraphs[19].text).replace('#nome_completo', pessoa.nome)
        recibos.paragraphs[27].text = str(recibos.paragraphs[27].text).replace('#nome_completo', pessoa.nome)
        recibos.paragraphs[40].text = str(recibos.paragraphs[40].text).replace('#nome_completo', pessoa.nome)
        recibos.paragraphs[48].text = str(recibos.paragraphs[48].text).replace('#nome_completo', pessoa.nome)
        recibos.save(p_contr + '\\Recibos.docx')
        docx2pdf.convert(p_contr + '\\Recibos.docx', p_contr + '\\Recibos.pdf')
        os.remove(p_contr + '\\Recibos.docx')

        os.rename(p_pessoa, p_pessoa.replace(r'\1 - Ainda nao iniciaram', ''))
        tkinter.messagebox.showinfo(
            title='Documentos ok!',
            message='Documentos salvos com sucesso!'
        )


def enviar_emails_funcionario(matricula):
    """
    Send e-mails to employee about his/her admission.

    :param matricula: Employee registration number
    """
    sessions = sessionmaker(bind=engine)
    session = sessions()
    pessoa = session.query(Colaborador).filter_by(matricula=matricula).first()
    if pessoa is None:
        tkinter.messagebox.showinfo(
            title='Erro de preenchimento',
            message='Não existe funcionário cadastrado com essa matrícula!'
        )
    else:
        p_contr = r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e Férias' \
                  r'\000 - Pastas Funcionais\00 - ATIVOS\{}\Contratuais'.format(pessoa.nome)
        email_remetente = em_rem
        senha = k1
        # set up smtp connection
        s = smtplib.SMTP(host=host, port=port)
        s.starttls()
        s.login(email_remetente, senha)

        # send e-mail to employee with a pdf file so he/she can go to bank to open an account
        msg = MIMEMultipart('alternative')
        msg['From'] = email_remetente
        msg['To'] = pessoa.email
        msg['Subject'] = "Documentos Contratuais"
        arquivo = p_contr + '\\Cod Etica.pdf'
        arquivo2 = p_contr + '\\Contrato de Trabalho.pdf'
        if pessoa.genero == 'Masculino':
            if pessoa.tipo_contr == 'Horista':
                text = MIMEText(f'''Olá, {str(pessoa.nome).title().split(" ")[0]}!<br><br>
                Seja muito bem vindo a Companhia Athletica.<br>
                Seu contrato está pronto para ser assinado.<br>
                Seu tipo de contrato é de horista.<br>
                Nesse tipo de contrato você será remunerado de acordo com a soma total de HORAS trabalhadas.<br>
                É importante lembrar que o tempo dos intervalos entre aulas não é remunerado.<br><br>
                Seguem alguns lembretes do nosso código de ética já enviado a você:<br>
                Você pode treinar todas as modalidades da academia gratuitamente, sempre respeitando a prioridade aos alunos.<br>
                Você tem direito a uma folga no dia do seu aniversário (essa folga não pode ser transferida para outro dia).<br>
                A academia oferece descontos especiais para parentes de funcionários.<br>
                Nosso período de folha é do dia 21 do mês ao dia 20 do mês seguinte. Então, no seu primeiro salário você receberá o valor proporcional do dia que entrou até o próximo dia 20.<br>
                Quando completar um ano de contrato terá direito a férias. Para pedir férias você deve solicitar ao seu superior direto com antecedência de 2 meses da data que pretende sair.<br><br>
                As demais regras você encontra no nosso código de ética. Ok?<br>
                Qualquer dúvida, estou à disposição.<br><br>
                Atenciosamente,<br>
                <img src="cid:image1">''', 'html')
            if pessoa.tipo_contr == 'Mensalista':
                text = MIMEText(f'''Olá, {str(pessoa.nome).title().split(" ")[0]}!<br><br>
                Seja muito bem vindo a Companhia Athletica.<br>
                Seu contrato está pronto para ser assinado.<br><br>
                Seguem alguns lembretes do nosso código de ética já enviado a você:<br>
                Você pode treinar todas as modalidades da academia gratuitamente, sempre respeitando a prioridade aos alunos.<br>
                Você tem direito a uma folga no dia do seu aniversário (essa folga não pode ser transferida para outro dia).<br>
                A academia oferece descontos especiais para parentes de funcionários.<br>
                Nosso período de folha é do dia 21 do mês ao dia 20 do mês seguinte. Então, no seu primeiro salário você receberá o valor proporcional do dia que entrou até o próximo dia 20.<br>
                Quando completar um ano de contrato terá direito a férias. Para pedir férias você deve solicitar ao seu superior direto com antecedência de 2 meses da data que pretende sair.<br><br>
                As demais regras você encontra no nosso código de ética. Ok?<br>
                Qualquer dúvida, estou à disposição.<br><br>
                Atenciosamente,<br>
                <img src="cid:image1">''', 'html')

        if pessoa.genero == 'Feminino':
            if pessoa.tipo_contr == 'Horista':
                text = MIMEText(f'''Olá, {str(pessoa.nome).title().split(" ")[0]}!<br><br>
                Seja muito bem vinda a Companhia Athletica.<br>
                Seu contrato está pronto para ser assinado.<br>
                Seu tipo de contrato é de horista.<br>
                Nesse tipo de contrato você será remunerado de acordo com a soma total de HORAS trabalhadas.<br>
                É importante lembrar que o tempo dos intervalos entre aulas não é remunerado.<br><br>
                Seguem alguns lembretes do nosso código de ética já enviado a você:<br>
                Você pode treinar todas as modalidades da academia gratuitamente, sempre respeitando a prioridade aos alunos.<br>
                Você tem direito a uma folga no dia do seu aniversário (essa folga não pode ser transferida para outro dia).<br>
                A academia oferece descontos especiais para parentes de funcionários.<br>
                Nosso período de folha é do dia 21 do mês ao dia 20 do mês seguinte. Então, no seu primeiro salário você receberá o valor proporcional do dia que entrou até o próximo dia 20.<br>
                Quando completar um ano de contrato terá direito a férias. Para pedir férias você deve solicitar ao seu superior direto com antecedência de 2 meses da data que pretende sair.<br><br>
                As demais regras você encontra no nosso código de ética. Ok?<br>
                Qualquer dúvida, estou à disposição.<br><br>
                Atenciosamente,<br>
                <img src="cid:image1">''', 'html')
            if pessoa.tipo_contr == 'Mensalista':
                text = MIMEText(f'''Olá, {str(pessoa.nome).title().split(" ")[0]}!<br><br>
                Seja muito bem vinda a Companhia Athletica.<br>
                Seu contrato está pronto para ser assinado.<br><br>
                Seguem alguns lembretes do nosso código de ética já enviado a você:<br>
                Você pode treinar todas as modalidades da academia gratuitamente, sempre respeitando a prioridade aos alunos.<br>
                Você tem direito a uma folga no dia do seu aniversário (essa folga não pode ser transferida para outro dia).<br>
                A academia oferece descontos especiais para parentes de funcionários.<br>
                Nosso período de folha é do dia 21 do mês ao dia 20 do mês seguinte. Então, no seu primeiro salário você receberá o valor proporcional do dia que entrou até o próximo dia 20.<br>
                Quando completar um ano de contrato terá direito a férias. Para pedir férias você deve solicitar ao seu superior direto com antecedência de 2 meses da data que pretende sair.<br><br>
                As demais regras você encontra no nosso código de ética. Ok?<br>
                Qualquer dúvida, estou à disposição.<br><br>
                Atenciosamente,<br>
                <img src="cid:image1">''', 'html')

        # set up the parameters of the message
        msg.attach(text)
        image = MIMEImage(
            open(rf'C:\Users\{os.getlogin()}\PycharmProjects\AutomacaoCia\src\models\static\imgs\assinatura.png', 'rb').read())
        image.add_header('Content-ID', '<image1>')
        msg.attach(image)

        # attach pdf file
        part = MIMEBase('application', "octet-stream")
        part.set_payload(open(arquivo, "rb").read())
        encoders.encode_base64(part)
        part.add_header('Content-Disposition', 'attachment',
                        filename=f'Código de Ética.pdf')
        msg.attach(part)

        part2 = MIMEBase('application', "octet-stream")
        part2.set_payload(open(arquivo2, "rb").read())
        encoders.encode_base64(part2)
        part2.add_header('Content-Disposition', 'attachment',
                        filename=f'Contrato.pdf')
        msg.attach(part2)

        s.sendmail(email_remetente, pessoa.email, msg.as_string())
        del msg
        s.quit()
        tkinter.messagebox.showinfo(
            title='E-mails ok!',
            message='E-mails enviados com sucesso'
        )


def enviar_emails_contratacao(caminho: str, nome: str, departamento: str, cargo: str, salario: str, admissao: str):
    """
    This function send e-mails to employee

    :param caminho:
    :param nome:
    :param departamento:
    :param cargo:
    :param salario:
    :param admissao:
    :return:
    """
    if nome == '' or departamento == '' or cargo == '' or salario == '' or admissao == '':
        tkinter.messagebox.showinfo(
            title='Erro de preenchimento',
            message='Preencha corretamente os campos Nome, Admissão, Salário, Cargo e Departamento.'
        )
    else:
        linha, nome = nome.upper().strip().split(' - ')
        departamento = departamento.title()
        plcontr = l_w(caminho, read_only=False)
        fol = plcontr['Respostas ao formulário 1']
        p_pessoa = r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e Férias' \
                   r'\000 - Pastas Funcionais\00 - ATIVOS\1 - Ainda nao iniciaram\{}'.format(nome)
        p_atestado = r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e Férias' \
                     r'\000 - Pastas Funcionais\00 - ATIVOS\1 - Ainda nao iniciaram\{}\Atestado'.format(nome)
        p_contr = r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e Férias' \
                  r'\000 - Pastas Funcionais\00 - ATIVOS\1 - Ainda nao iniciaram\{}\Contratuais'.format(nome)
        p_diversos = r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e Férias' \
                     r'\000 - Pastas Funcionais\00 - ATIVOS\1 - Ainda nao iniciaram\{}\Diversos'.format(nome)
        p_ferias = r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e Férias' \
                   r'\000 - Pastas Funcionais\00 - ATIVOS\1 - Ainda nao iniciaram\{}\Férias'.format(nome)
        p_ponto = r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e Férias' \
                  r'\000 - Pastas Funcionais\00 - ATIVOS\1 - Ainda nao iniciaram\{}\Pontos'.format(nome)
        p_rec = r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e Férias' \
                r'\000 - Pastas Funcionais\00 - ATIVOS\1 - Ainda nao iniciaram\{}\Recibos'.format(nome)
        p_rescisao = r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e Férias' \
                     r'\000 - Pastas Funcionais\00 - ATIVOS\1 - Ainda nao iniciaram\{}\Rescisão'.format(nome)
        p_ac = r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e Férias' \
               r'\000 - Pastas Funcionais\00 - ATIVOS\Modelo\AC Modelo.docx'
        p_abconta = r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e Férias' \
                    r'\000 - Pastas Funcionais\00 - ATIVOS\Modelo\Abertura Conta MODELO.docx'
        p_fcadas = r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e Férias' \
                   r'\000 - Pastas Funcionais\00 - ATIVOS\Modelo\Ficha Cadastral MODELO.docx'
        p_codetic = r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e Férias' \
                    r'\000 - Pastas Funcionais\00 - ATIVOS\Modelo\Cod Etica MODELO.docx'
        try:
            os.mkdir(p_pessoa)
            os.mkdir(p_atestado)
            os.mkdir(p_contr)
            os.mkdir(p_diversos)
            os.mkdir(p_ferias)
            os.mkdir(p_ponto)
            os.mkdir(p_rec)
            os.mkdir(p_rescisao)
        except FileExistsError:
            pass

        lotacao = {
            'Unidade Park Sul - Qualquer Departamento': ['0013', 'Thais Feitosa', 'thais.morais@ciaathletica.com.br',
                                                         'Líder Park Sul'],
            'Kids': ['0010', 'Cindy Stefanie', 'cindy.neves@ciaathletica.com.br', 'Líder Kids'],
            'Musculação': ['0007', 'Thaís Feitosa', 'thais.morais@ciaathletica.com.br', 'Líder Musculação'],
            'Esportes e Lutas': ['0008', 'Morgana Rossini', 'morganalourenco@yahoo.com.br', 'Líder Natação'],
            'Crossfit': ['0012', 'Guilherme Salles', 'gmoreirasalles@gmail.com', 'Líder Crossfit'],
            'Ginástica': ['0006', 'Morgana Rossini', 'morganalourenco@yahoo.com.br', 'Líder Ginástica'],
            'Gestantes': ['0006', 'Filipe Feijó', 'filipe.feijo@ciaathletica.com.br', 'Líder Ginástica'],
            'Recepção': ['0003', 'Paulo Renato', 'paulo.simoes@ciaathletica.com.br', 'Gerente Vendas'],
            'Administrativo': ['0001', 'Felipe Rodrigues', 'felipe.rodrigues@ciaathletica.com.br', 'Gerente RH'],
            'Manutenção': ['0004', 'José Aparecido', 'aparecido.grota@ciaathletica.com.br', 'Gerente Manutenção'],
        }

        abert_c = docx.Document(p_abconta)
        ac = docx.Document(p_ac)
        fch_c = docx.Document(p_fcadas)
        codetic = docx.Document(p_codetic)

        # # Alterar AC e Salvar na pasta
        ac.paragraphs[1].text = str(ac.paragraphs[1].text).replace('#gerente', lotacao[departamento][1])
        ac.paragraphs[2].text = str(ac.paragraphs[2].text).replace('#nome_completo', nome)
        ac.paragraphs[3].text = str(ac.paragraphs[3].text).replace('#cargo', cargo)
        ac.paragraphs[11].text = str(ac.paragraphs[11].text).replace('#salario', salario)
        ac.save(p_contr + '\\AC.docx')
        docx2pdf.convert(p_contr + '\\AC.docx', p_contr + '\\AC.pdf')
        os.remove(p_contr + '\\AC.docx')

        # # Alterar Abertura de Conta e salvar na pasta
        abert_c.paragraphs[14].text = str(abert_c.paragraphs[14].text).replace('#nome_completo', nome).replace(
            '#rg', str(fol[f'W{linha}'].value).replace('.0','')).replace(
            '#cpf', str(fol[f'V{linha}'].value)).replace('#endereco', str(fol[f'O{linha}'].value))\
            .replace('#cep', str(fol[f'R{linha}'].value).replace('.0',''))\
            .replace('#bairro', str(fol[f'Q{linha}'].value).strip()).replace('#cargo', cargo).replace('#data', admissao)
        abert_c.save(p_contr + '\\Abertura Conta.docx')
        docx2pdf.convert(p_contr + '\\Abertura Conta.docx', p_contr + '\\Abertura Conta.pdf')
        os.remove(p_contr + '\\Abertura Conta.docx')

        # Alterar Ficha cadastral e salvar na pasta
        fch_c.paragraphs[34].text = str(fch_c.paragraphs[34].text).replace('#gerente#',
                                                                           lotacao[departamento][1])
        fch_c.paragraphs[9].text = str(fch_c.paragraphs[9].text).replace('#nome_completo', nome)
        fch_c.paragraphs[21].text = str(fch_c.paragraphs[21].text).replace('#cargo', cargo)\
            .replace('#depart', departamento)
        fch_c.paragraphs[19].text = str(fch_c.paragraphs[19].text).replace('#end_eletr', str(fol[f'B{linha}'].value))
        fch_c.paragraphs[17].text = str(fch_c.paragraphs[17].text).replace('#mae#', str(fol[f'N{linha}'].value))
        fch_c.paragraphs[16].text = str(fch_c.paragraphs[16].text).replace('#pai#', str(fol[f'M{linha}'].value))
        fch_c.paragraphs[15].text = str(fch_c.paragraphs[15].text).replace('#ident', str(fol[f'W{linha}'].value).replace('.0','')).replace('#cpf#',
                                                                                                        str(fol[f'V{linha}'].value))
        fch_c.paragraphs[13].text = str(fch_c.paragraphs[13].text).replace('#telefone', str(fol[f'U{linha}'].value).replace('.0',''))
        fch_c.paragraphs[12].text = str(fch_c.paragraphs[12].text).replace('#codigo', str(fol[f'R{linha}'].value).replace('.0',''))\
            .replace('#cid', str(fol[f'S{linha}'].value)).replace('#uf', str(fol[f'T{linha}'].value))
        fch_c.paragraphs[11].text = str(fch_c.paragraphs[11].text).replace('#local', str(fol[f'O{linha}'].value))\
            .replace('#qd', str(fol[f'Q{linha}'].value))
        fch_c.paragraphs[10].text = str(fch_c.paragraphs[10].text)\
            .replace('#nasc', dt.strftime(dt.strptime(str(fol[f'D{linha}'].value), '%Y-%m-%d %H:%M:%S'), '%d/%m/%Y'))\
            .replace('#gen', str(fol[f'E{linha}'].value)).replace('#est_civ', str(fol[f'F{linha}'].value).replace('1 - ', '')
                                                    .replace('2 - ', '').replace('3 - ', '').replace('4 - ', ''))
        fch_c.save(p_contr + '\\Ficha Cadastral.docx')
        docx2pdf.convert(p_contr + '\\Ficha Cadastral.docx', p_contr + '\\Ficha Cadastral.pdf')
        os.remove(p_contr + '\\Ficha Cadastral.docx')

        # Alterar Código de Ética e salvar na pasta
        codetic.paragraphs[534].text = str(codetic.paragraphs[534].text).replace('#nome_completo', nome)
        codetic.paragraphs[535].text = str(codetic.paragraphs[535].text).replace('#func', cargo)
        codetic.paragraphs[537].text = str(codetic.paragraphs[537].text).replace('#nome_completo', nome)
        codetic.paragraphs[541].text = str(codetic.paragraphs[541].text).replace('#admiss', admissao)
        codetic.save(p_contr + '\\Cod Etica.docx')
        docx2pdf.convert(p_contr + '\\Cod Etica.docx', p_contr + '\\Cod Etica.pdf')
        os.remove(p_contr + '\\Cod Etica.docx')

        # send e-mails
        email_remetente = em_rem
        senha = k1
        # set up smtp connection
        s = smtplib.SMTP(host=host, port=port)
        s.starttls()
        s.login(email_remetente, senha)

        # enviar e-mail de boas vindas
        msg = MIMEMultipart('alternative')
        msg['From'] = email_remetente
        msg['To'] = str(fol[f'B{linha}'].value).strip()
        msg['Subject'] = "Boas vindas!"
        arquivo = p_contr + '\\Cod Etica.pdf'
        text = MIMEText(f'''Olá, {str(nome).title().split(" ")[0]}!<br><br>
        Estamos felizes que você fará parte da nossa equipe!<br>
        Em anexo segue nosso código de ética e conduta.<br>
        Nesse documento estão todas as regras da Cia e tudo que a Cia espera de seus funcionários.<br>
        Além de regras, também descreve direitos e benefícios.<br>
        É importante que você leia todo o documento pois precisaremos da sua assinatura nele no dia que for assinar o contrato. Ok?<br>
        Qualquer dúvida, estou à disposição.<br><br>
        Atenciosamente,<br>
        <img src="cid:image1">''', 'html')

        # set up the parameters of the message
        msg.attach(text)
        image = MIMEImage(
            open(rf'C:\Users\{os.getlogin()}\PycharmProjects\AutomacaoCia\src\models\static\imgs\assinatura.png', 'rb').read())
        image.add_header('Content-ID', '<image1>')
        msg.attach(image)
        # attach pdf file
        part = MIMEBase('application', "octet-stream")
        part.set_payload(open(arquivo, "rb").read())
        encoders.encode_base64(part)
        part.add_header('Content-Disposition', 'attachment',
                        filename='Código de Ética Cia Athletica.pdf')
        msg.attach(part)
        s.sendmail(email_remetente, str(fol[f'B{linha}'].value).strip(), msg.as_string())
        del msg

        # enviar AC para líder
        msg = MIMEMultipart('alternative')
        msg['From'] = email_remetente
        msg['To'] = lotacao[departamento][2]
        msg['Subject'] = f'AC - {str(nome).title().split(" ")[0]}'
        arquivo = p_contr + '\\AC.pdf'
        text = MIMEText(f'''Olá, {lotacao[departamento][1].split(" ")[0]}!<br><br>
        Segue a AC do(a) {str(nome).title().split(" ")[0]}.<br>
        A AC é o documento oficial de cadastro de um funcionário na Cia.<br>
        Você deverá imprimir a AC, preencher o horário do funcionário em cada dia e solicitar a assinatura da direção.<br>
        Após a assinatura da direção me entregue a AC no RH, por favor.<br><br>
        Atenciosamente,<br>
        <img src="cid:image1">''', 'html')

        # set up the parameters of the message
        msg.attach(text)
        image = MIMEImage(
            open(rf'C:\Users\{os.getlogin()}\PycharmProjects\AutomacaoCia\src\models\static\imgs\assinatura.png', 'rb').read())
        image.add_header('Content-ID', '<image1>')
        msg.attach(image)
        # attach pdf file
        part = MIMEBase('application', "octet-stream")
        part.set_payload(open(arquivo, "rb").read())
        encoders.encode_base64(part)
        part.add_header('Content-Disposition', 'attachment',
                        filename=f'AC - {str(nome).title().split(" ")[0]}.pdf')
        msg.attach(part)
        s.sendmail(email_remetente, lotacao[departamento][2], msg.as_string())
        del msg

        # send e-mail to employee with a pdf file so he/she can go to bank to open an account
        msg = MIMEMultipart('alternative')
        msg['From'] = email_remetente
        msg['To'] = str(fol[f'B{linha}'].value).strip()
        msg['Subject'] = "Carta para Abertura de conta"
        arquivo = p_contr + '\\Abertura Conta.pdf'
        text = MIMEText(f'''Olá, {str(nome).title().split(" ")[0]}!<br><br>
        Segue sua carta para abertura de conta bancária no Itaú.<br>
        Você deve abrir a conta antes de iniciar seu contrato de trabalho. Ok?<br>
        Assim que conseguir abrir a conta me responda esse e-mail com os dados bancários do Itaú. <br>
        Qualquer dúvida, estou à disposição.<br><br>
        Atenciosamente,<br>
        <img src="cid:image1">''', 'html')

        # set up the parameters of the message
        msg.attach(text)
        image = MIMEImage(
            open(rf'C:\Users\{os.getlogin()}\PycharmProjects\AutomacaoCia\src\models\static\imgs\assinatura.png', 'rb').read())
        image.add_header('Content-ID', '<image1>')
        msg.attach(image)
        # attach pdf file
        part = MIMEBase('application', "octet-stream")
        part.set_payload(open(arquivo, "rb").read())
        encoders.encode_base64(part)
        part.add_header('Content-Disposition', 'attachment',
                        filename=f'Carta Banco {str(nome).title().split(" ")[0]}.pdf')
        msg.attach(part)
        s.sendmail(email_remetente, str(fol[f'B{linha}'].value).strip(), msg.as_string())
        del msg

        # send e-mail to coworker asking to register the ner employee
        msg = MIMEMultipart('alternative')
        arquivo = p_contr + '\\Ficha Cadastral.pdf'
        text = MIMEText(
            f'''Oi, Wallace!<br><br>Segue a ficha cadastral do(a) {nome}.<br><br>Abs.,<br><img src="cid:image1">''',
            'html')
        msg.attach(text)
        image = MIMEImage(
            open(rf'C:\Users\{os.getlogin()}\PycharmProjects\AutomacaoCia\src\models\static\imgs\assinatura.png', 'rb').read())
        image.add_header('Content-ID', '<image1>')
        msg.attach(image)
        # set up the parameters of the message
        msg['From'] = email_remetente
        msg['To'] = em_ti
        msg['Subject'] = f"Ficha Cadastral {str(nome).title().split(' ')[0]}"
        # attach pdf file
        part = MIMEBase('application', "octet-stream")
        part.set_payload(open(arquivo, "rb").read())
        encoders.encode_base64(part)
        part.add_header('Content-Disposition', 'attachment',
                        filename=f'Ficha Cadastral {str(nome).title().split(" ")[0]}.pdf')
        msg.attach(part)
        s.sendmail(email_remetente, em_ti, msg.as_string())
        del msg
        s.quit()
        tkinter.messagebox.showinfo(
            title='E-mails ok!',
            message='E-mails enviados com sucesso'
        )


def cadastro_estagiario(solicitar_contr=0, caminho='', editar=0, ondestou=0, nome='', matricula='', admissao='',
                        cargo='', depto='', tipo_contr='Horista',
                        hrsem='25', hrmens='100', agencia='', conta='', digito=''):
    sessions = sessionmaker(bind=engine)
    session = sessions()
    pa.FAILSAFE = False
    salario = 5.10
    if nome == '':
        tkinter.messagebox.showinfo(
            title='Erro de preenchimento',
            message='Preencha todos os campos antes de cadastrar o estagiário!'
        )
    else:
        if solicitar_contr == 1:
            hoje = dt.today()
            wb = l_w(caminho)
            sh = wb['Respostas ao formulário 1']
            num, name = nome.strip().split(' - ')
            linha = int(num)
            lotacao = {
                'Unidade Park Sul - qualquer departamento': ['0013', 'Thais Feitosa', 'thais.morais@ciaathletica.com.br',
                                                             'Líder Park Sul'],
                'Kids': ['0010', 'Cindy Stefanie', 'cindy.neves@ciaathletica.com.br', 'Líder Kids'],
                'Musculação': ['0007', 'Thais Feitosa', 'thais.morais@ciaathletica.com.br', 'Líder Musculação'],
                'Esportes e Lutas': ['0008', 'Morgana Rossini', 'morganalourenco@yahoo.com.br', 'Líder Natação'],
                'Crossfit': ['0012', 'Guilherme Salles', 'gmoreirasalles@gmail.com', 'Líder Crossfit'],
                'Ginástica': ['0006', 'Morgana Rossini', 'morganalourenco@yahoo.com.br', 'Líder Ginástica'],
                'Gestantes': ['0006', 'Filipe Feijó', 'filipe.feijo@ciaathletica.com.br', 'Gerente Técnico'],
                'Recepção': ['0003', 'Paulo Renato', 'paulo.simoes@ciaathletica.com.br', 'Gerente Vendas'],
                'Administrativo': ['0001', 'Felipe Rodrigues', 'felipe.rodrigues@ciaathletica.com.br', 'Gerente RH'],
                'Manutenção': ['0004', 'José Aparecido', 'aparecido.grota@ciaathletica.com.br', 'Gerente Manutenção'],
            }
            cadastro = {'nome': str(sh[f"C{linha}"].value).title().strip(), 'nasc_ed': sh[f"D{linha}"].value,
                        'genero': str(sh[f"E{linha}"].value), 'est_civ': str(sh[f"F{linha}"].value),
                        'pai': str(sh[f"M{linha}"].value), 'mae': str(sh[f"N{linha}"].value),
                        'end': str(sh[f"O{linha}"].value),
                        'num': str(sh[f"P{linha}"].value), 'bairro': str(sh[f"Q{linha}"].value),
                        'cep': str(sh[f"R{linha}"].value).replace('.', '').replace('-', ''),
                        'cid_end': str(sh[f"S{linha}"].value), 'uf_end': str(sh[f"T{linha}"].value),
                        'tel': str(sh[f"U{linha}"].value).replace('(', '').replace(')', '').replace('-', '').replace(' ',
                                                                                                                     ''),
                        'mun_end': str(sh[f"AP{linha}"].value),
                        'cpf': str(sh[f"V{linha}"].value).strip().replace('.', '').replace('-', '').replace(' ', '').zfill(
                            11),
                        'rg': str(sh[f"W{linha}"].value).strip().replace('.', '').replace('-', '').replace(' ', ''),
                        'emissor': str(sh[f"X{linha}"].value),
                        'lotacao': str(lotacao[f'{sh[f"AG{linha}"].value}'][0]).zfill(4),
                        'cargo': str(sh[f"AH{linha}"].value), 'horario': str(sh[f"AI{linha}"].value),
                        'email': str(sh[f"B{linha}"].value).strip(),
                        'admissao_ed': str(sh[f"AL{linha}"].value),
                        'faculdade': str(sh[f"AV{linha}"].value), 'semestre': str(sh[f"AS{linha}"].value),
                        'turno': str(sh[f"AT{linha}"].value), 'conclusao': str(sh[f"AU{linha}"].value),
                        'salario': str(sh[f"AM{linha}"].value),
                        'hrsemanais': str(sh[f"AQ{linha}"].value), 'hrmensais': str(sh[f"AR{linha}"].value)}
            email_remetente = em_rem
            senha = k1
            lot = lotacao[f'{sh[f"AG{linha}"].value}']
            pasta = r'\192.168.0.250'
            modelo = f'\\{pasta}\\rh\\01 - RH\\01 - Administração.Controles\\02 - Funcionários, Departamentos e Férias\\' \
                     f'000 - Pastas Funcionais\\00 - ATIVOS\\0 - Estagiários\\Modelo'
            try:
                os.makedirs(
                    f'\\{pasta}\\rh\\01 - RH\\01 - Administração.Controles\\02 - Funcionários, Departamentos e Férias\\'
                    f'000 - Pastas Funcionais\\00 - ATIVOS\\0 - Estagiários\\'
                    f'0 - Ainda nao iniciaram\\{str(cadastro["nome"]).upper()}')
                os.makedirs(
                    f'\\{pasta}\\rh\\01 - RH\\01 - Administração.Controles\\02 - Funcionários, Departamentos e Férias\\'
                    f'000 - Pastas Funcionais\\00 - ATIVOS\\0 - Estagiários\\'
                    f'0 - Ainda nao iniciaram\\{str(cadastro["nome"]).upper()}\\Atestados')
                os.makedirs(
                    f'\\{pasta}\\rh\\01 - RH\\01 - Administração.Controles\\02 - Funcionários, Departamentos e Férias\\'
                    f'000 - Pastas Funcionais\\00 - ATIVOS\\0 - Estagiários\\'
                    f'0 - Ainda nao iniciaram\\{str(cadastro["nome"]).upper()}\\Diversos')
                os.makedirs(
                    f'\\{pasta}\\rh\\01 - RH\\01 - Administração.Controles\\02 - Funcionários, Departamentos e Férias\\'
                    f'000 - Pastas Funcionais\\00 - ATIVOS\\0 - Estagiários\\'
                    f'0 - Ainda nao iniciaram\\{str(cadastro["nome"]).upper()}\\Contratuais')
                os.makedirs(
                    f'\\{pasta}\\rh\\01 - RH\\01 - Administração.Controles\\02 - Funcionários, Departamentos e Férias\\'
                    f'000 - Pastas Funcionais\\00 - ATIVOS\\0 - Estagiários\\'
                    f'0 - Ainda nao iniciaram\\{str(cadastro["nome"]).upper()}\\Ferias')
                os.makedirs(
                    f'\\{pasta}\\rh\\01 - RH\\01 - Administração.Controles\\02 - Funcionários, Departamentos e Férias\\'
                    f'000 - Pastas Funcionais\\00 - ATIVOS\\0 - Estagiários\\'
                    f'0 - Ainda nao iniciaram\\{str(cadastro["nome"]).upper()}\\Ponto')
                os.makedirs(
                    f'\\{pasta}\\rh\\01 - RH\\01 - Administração.Controles\\02 - Funcionários, Departamentos e Férias\\'
                    f'000 - Pastas Funcionais\\00 - ATIVOS\\0 - Estagiários\\'
                    f'0 - Ainda nao iniciaram\\{str(cadastro["nome"]).upper()}\\Recibo')
                os.makedirs(
                    f'\\{pasta}\\rh\\01 - RH\\01 - Administração.Controles\\02 - Funcionários, Departamentos e Férias\\'
                    f'000 - Pastas Funcionais\\00 - ATIVOS\\0 - Estagiários\\'
                    f'0 - Ainda nao iniciaram\\{str(cadastro["nome"]).upper()}\\Rescisao')
            except FileExistsError:
                pass
            pasta_contratuais = f'\\{pasta}\\rh\\01 - RH\\01 - Administração.Controles\\' \
                                f'02 - Funcionários, Departamentos e Férias\\000 - Pastas Funcionais\\' \
                                f'00 - ATIVOS\\0 - Estagiários\\0 - Ainda nao iniciaram\\' \
                                f'{str(cadastro["nome"]).upper()}\\Contratuais'

            # change tree docx models files with intern data and save pdfs files
            solicitacao = docx.Document(modelo + r'\Solicitacao MODELO - Copia.docx')
            solicitacao.tables[0].rows[4].cells[0].paragraphs[0].text = str(
                solicitacao.tables[0].rows[4].cells[0].paragraphs[0].text).replace('#supervisor_estagio', f'{lot[1]}')
            solicitacao.tables[0].rows[5].cells[1].paragraphs[0].text = str(
                solicitacao.tables[0].rows[5].cells[1].paragraphs[0].text).replace('#cargo', f'{lot[3]}')
            solicitacao.tables[0].rows[6].cells[0].paragraphs[0].text = str(
                solicitacao.tables[0].rows[6].cells[0].paragraphs[0].text).replace('#email_supervisor', f'{lot[2]}')
            solicitacao.tables[0].rows[9].cells[0].paragraphs[0].text = str(
                solicitacao.tables[0].rows[9].cells[0].paragraphs[0].text).replace('#horario', cadastro['horario'])
            solicitacao.tables[0].rows[14].cells[0].paragraphs[0].text = str(
                solicitacao.tables[0].rows[14].cells[0].paragraphs[0].text).replace('#nome_completo', cadastro['nome'])
            solicitacao.tables[0].rows[15].cells[0].paragraphs[0].text = str(
                solicitacao.tables[0].rows[15].cells[0].paragraphs[0].text).replace('#nasc',
                                                                                    dt.strftime(cadastro['nasc_ed'],
                                                                                                      '%d/%m/%Y')
                                                                                    ).replace('#rg',
                                                                                              cadastro['rg']).replace(
                '#cpf', cadastro['cpf'])
            solicitacao.tables[0].rows[16].cells[0].paragraphs[0].text = str(
                solicitacao.tables[0].rows[16].cells[0].paragraphs[0].text).replace('#sexo', cadastro['genero'])
            solicitacao.tables[0].rows[17].cells[0].paragraphs[0].text = str(
                solicitacao.tables[0].rows[17].cells[0].paragraphs[0].text).replace('#endereco', cadastro['end'])
            solicitacao.tables[0].rows[18].cells[0].paragraphs[0].text = str(
                solicitacao.tables[0].rows[18].cells[0].paragraphs[0].text).replace('#cep', cadastro['cep']).replace(
                '#bairro', cadastro['bairro'])
            solicitacao.tables[0].rows[19].cells[0].paragraphs[0].text = str(
                solicitacao.tables[0].rows[19].cells[0].paragraphs[0].text).replace('#telefone', cadastro['tel'])
            solicitacao.tables[0].rows[20].cells[0].paragraphs[0].text = str(
                solicitacao.tables[0].rows[20].cells[0].paragraphs[0].text).replace('#end_eletr', cadastro['email'])
            solicitacao.tables[0].rows[22].cells[0].paragraphs[0].text = str(
                solicitacao.tables[0].rows[22].cells[0].paragraphs[0].text).replace('#semestre', cadastro['semestre'])
            solicitacao.tables[0].rows[23].cells[0].paragraphs[0].text = str(
                solicitacao.tables[0].rows[23].cells[0].paragraphs[0].text).replace('#turno', cadastro['turno']).replace(
                '#ano_concl', cadastro['conclusao'])
            solicitacao.tables[0].rows[24].cells[0].paragraphs[0].text = str(
                solicitacao.tables[0].rows[24].cells[0].paragraphs[0].text).replace('#faculdade', cadastro['faculdade'])
            solicitacao.save(pasta_contratuais + f'\\Pedido TCE {str(cadastro["nome"]).split(" ")[0]}.docx')
            docx2pdf.convert(pasta_contratuais + f'\\Pedido TCE {str(cadastro["nome"]).split(" ")[0]}.docx',
                             pasta_contratuais + f'\\Pedido TCE {str(cadastro["nome"]).split(" ")[0]}.pdf')
            os.remove(pasta_contratuais + f'\\Pedido TCE {str(cadastro["nome"]).split(" ")[0]}.docx')

            ficha_cadastral = docx.Document(modelo + r'\Ficha Cadastral MODELO - Copia.docx')
            ficha_cadastral.tables[1].rows[0].cells[0].paragraphs[0].text = str(
                ficha_cadastral.tables[1].rows[0].cells[0].paragraphs[0].text).replace('#nome_completo', cadastro['nome'])
            ficha_cadastral.tables[1].rows[1].cells[0].paragraphs[0].text = str(
                ficha_cadastral.tables[1].rows[1].cells[0].paragraphs[0].text)\
                .replace('#nasc', dt.strftime(cadastro['nasc_ed'], '%d/%m/%Y'))
            ficha_cadastral.tables[1].rows[1].cells[2].paragraphs[0].text = str(
                ficha_cadastral.tables[1].rows[1].cells[2].paragraphs[0].text).replace('#gen', cadastro['genero'])
            ficha_cadastral.tables[1].rows[1].cells[4].paragraphs[0].text = str(
                ficha_cadastral.tables[1].rows[1].cells[4].paragraphs[0].text).replace('#est_civ', cadastro['est_civ'])
            ficha_cadastral.tables[1].rows[2].cells[0].paragraphs[0].text = str(
                ficha_cadastral.tables[1].rows[2].cells[0].paragraphs[0].text).replace('#local', cadastro['end'])
            ficha_cadastral.tables[1].rows[2].cells[4].paragraphs[0].text = str(
                ficha_cadastral.tables[1].rows[2].cells[4].paragraphs[0].text).replace('#qd', cadastro['bairro'])
            ficha_cadastral.tables[1].rows[2].cells[7].paragraphs[0].text = str(
                ficha_cadastral.tables[1].rows[2].cells[7].paragraphs[0].text).replace('#codigo', cadastro['cep'])
            ficha_cadastral.tables[1].rows[4].cells[1].paragraphs[0].text = str(
                ficha_cadastral.tables[1].rows[4].cells[1].paragraphs[0].text).replace('#telefone', cadastro['tel'])
            ficha_cadastral.tables[1].rows[4].cells[5].paragraphs[0].text = str(
                ficha_cadastral.tables[1].rows[4].cells[5].paragraphs[0].text).replace('#ident', cadastro['rg'])
            ficha_cadastral.tables[1].rows[5].cells[1].paragraphs[0].text = str(
                ficha_cadastral.tables[1].rows[5].cells[1].paragraphs[0].text).replace('#cpf#', cadastro['cpf'])
            ficha_cadastral.tables[1].rows[6].cells[3].paragraphs[0].text = str(
                ficha_cadastral.tables[1].rows[6].cells[3].paragraphs[0].text).replace('#pai#', cadastro['pai'])
            ficha_cadastral.tables[1].rows[7].cells[1].paragraphs[0].text = str(
                ficha_cadastral.tables[1].rows[7].cells[1].paragraphs[0].text).replace('#mae#', cadastro['mae'])
            ficha_cadastral.tables[1].rows[8].cells[0].paragraphs[0].text = str(
                ficha_cadastral.tables[1].rows[8].cells[0].paragraphs[0].text).replace('#end_eletr', cadastro['email'])
            ficha_cadastral.tables[1].rows[8].cells[1].paragraphs[0].text = str(
                ficha_cadastral.tables[1].rows[8].cells[1].paragraphs[0].text).replace('#depart', str(sh["AG3"].value))
            ficha_cadastral.save(pasta_contratuais + f'\\Ficha Cadastral {str(cadastro["nome"]).split(" ")[0]}.docx')
            docx2pdf.convert(pasta_contratuais + f'\\Ficha Cadastral {str(cadastro["nome"]).split(" ")[0]}.docx',
                             pasta_contratuais + f'\\Ficha Cadastral {str(cadastro["nome"]).split(" ")[0]}.pdf')
            os.remove(pasta_contratuais + f'\\Ficha Cadastral {str(cadastro["nome"]).split(" ")[0]}.docx')

            carta_banco = docx.Document(
                modelo + r'\Abertura Conta MODELO.docx')
            carta_banco.paragraphs[14].text = str(carta_banco.paragraphs[14].text).replace('#nome_completo',
                                                                                           cadastro['nome']
                                                                                           ).replace('#rg', cadastro['rg']
                                                                                                     ).replace(
                '#cpf', cadastro['cpf']).replace('#endereço', cadastro['end']).replace('#cep', cadastro['cep']).replace(
                '#bairro', cadastro['bairro']).replace('#desde#', dt.strftime(hoje, '%d/%m/%Y'))
            carta_banco.save(pasta_contratuais + f'\\Carta Banco {str(cadastro["nome"]).split(" ")[0]}.docx')
            docx2pdf.convert(pasta_contratuais + f'\\Carta Banco {str(cadastro["nome"]).split(" ")[0]}.docx',
                             pasta_contratuais + f'\\Carta Banco {str(cadastro["nome"]).split(" ")[0]}.pdf')
            os.remove(pasta_contratuais + f'\\Carta Banco {str(cadastro["nome"]).split(" ")[0]}.docx')

            # set up smtpp connection
            s = smtplib.SMTP(host=host, port=port)
            s.starttls()
            s.login(email_remetente, senha)

            # send e-mail to intern with a pdf file so he/she can go to bank to open an account
            msg = MIMEMultipart('alternative')
            arquivo = pasta_contratuais + f'\\Carta Banco {str(cadastro["nome"]).split(" ")[0]}.pdf'
            text = MIMEText(f'''Olá, {str(cadastro["nome"]).split(" ")[0]}!<br><br>
            Segue sua carta para abertura de conta bancária no Itaú.<br>
            Você deve abrir a conta antes de iniciar os trabalhos no estágio. Ok?<br>
            Você já pode buscar seu contrato no IF. Será necessário levar uma declaração de matrícula do seu curso.<br><br>
            Atenciosamente,<br>
            <img src="cid:image1">''', 'html')
            msg.attach(text)
            image = MIMEImage(
                open(rf'C:\Users\{os.getlogin()}\PycharmProjects\AutomacaoCia\src\models\static\imgs\assinatura.png', 'rb').read())
            image.add_header('Content-ID', '<image1>')
            msg.attach(image)
            # set up the parameters of the message
            msg['From'] = email_remetente
            msg['To'] = cadastro['email']
            msg['Subject'] = "Carta para Abertura de conta"
            # attach pdf file
            part = MIMEBase('application', "octet-stream")
            part.set_payload(open(arquivo, "rb").read())
            encoders.encode_base64(part)
            part.add_header('Content-Disposition', 'attachment',
                            filename=f'Carta Banco {str(cadastro["nome"]).split(" ")[0]}.pdf')
            msg.attach(part)
            s.sendmail(email_remetente, cadastro['email'], msg.as_string())
            del msg

            # send e-mail to coworker asking to register the intern
            msg = MIMEMultipart('alternative')
            arquivo = pasta_contratuais + f'\\Ficha Cadastral {str(cadastro["nome"]).split(" ")[0]}.pdf'
            text = MIMEText(f'''Oi, Wallace!<br><br>
            Segue a ficha cadastral do(a) estagiário(a) {cadastro["nome"]}.<br><br>
            Abs.,<br>
            <img src="cid:image1">''', 'html')
            msg.attach(text)
            image = MIMEImage(
                open(rf'C:\Users\{os.getlogin()}\PycharmProjects\AutomacaoCia\src\models\static\imgs\assinatura.png', 'rb').read())
            # Define the image's ID as referenced in the HTML body above
            image.add_header('Content-ID', '<image1>')
            msg.attach(image)
            # set up the parameters of the message
            msg['From'] = email_remetente
            msg['To'] = em_ti
            msg['Subject'] = f"Ficha Cadastral {str(cadastro['nome']).split(' ')[0]}"
            # attach pdf file
            part = MIMEBase('application', "octet-stream")
            part.set_payload(open(arquivo, "rb").read())
            encoders.encode_base64(part)
            part.add_header('Content-Disposition', 'attachment',
                            filename=f'Ficha Cadastral {str(cadastro["nome"]).split(" ")[0]}.pdf')
            msg.attach(part)
            s.sendmail(email_remetente, em_ti, msg.as_string())
            del msg

            # send document asking for the intern contract
            msg = MIMEMultipart('alternative')
            arquivo = pasta_contratuais + f'\\Pedido TCE {str(cadastro["nome"]).split(" ")[0]}.pdf'
            text = MIMEText(
                f'''Olá!<br><br>nSegue pedido de TCE do(a) estagiário(a) {cadastro["nome"]}.
                <br><br>Atenciosamente,<br><img src="cid:image1">''',
                'html')
            msg.attach(text)
            image = MIMEImage(
                open(rf'C:\Users\{os.getlogin()}\PycharmProjects\AutomacaoCia\src\models\static\imgs\assinatura.png', 'rb').read())
            # Define the image's ID as referenced in the HTML body above
            image.add_header('Content-ID', '<image1>')
            msg.attach(image)
            # set up the parameters of the message
            msg['From'] = email_remetente
            msg['To'] = em_if
            msg['Subject'] = f"Pedido TCE {str(cadastro['nome']).split(' ')[0]}"
            # attach pdf file
            part = MIMEBase('application', "octet-stream")
            part.set_payload(open(arquivo, "rb").read())
            encoders.encode_base64(part)
            part.add_header('Content-Disposition', 'attachment',
                            filename=f'Pedido TCE {str(cadastro["nome"]).split(" ")[0]}.pdf')
            msg.attach(part)
            s.sendmail(email_remetente, em_if, msg.as_string())
            del msg
            s.quit()
            tkinter.messagebox.showinfo(
                title='E-mails ok!',
                message='E-mails enviados com sucesso'
            )

        else:
            if editar == 0:
                if ondestou == 0:
                    # Cadastro iniciado na Cia
                    wb = l_w(caminho)
                    sh = wb['Respostas ao formulário 1']
                    num, name = nome.strip().split(' - ')
                    linha = int(num)
                    lotacao = {'UNIDADE PARK SUL - QUALQUER DEPARTAMENTO': '0013', 'KIDS': '0010', 'MUSCULAÇÃO': '0007',
                               'ESPORTES E LUTAS': '0008', 'CROSSFIT': '0012', 'GINÁSTICA': '0006', 'GESTANTES': '0008',
                               'RECEPÇÃO': '0003',
                               'FINANCEIRO': '0001', 'TI': '0001', 'MARKETING': '0001', 'MANUTENÇÃO': '0004'}
                    if str(sh[f'E{linha}'].value) == 'Masculino':
                        cargo = 'ESTAGIARIO'
                    else:
                        cargo = 'ESTAGIARIA'

                    estag = session.query(Colaborador).filter_by(matricula=matricula).first()
                    if estag:
                        pass
                    else:
                        estag_cadastrado = Colaborador(
                            matricula=matricula, nome=name.upper(), admiss=admissao,
                            nascimento=str(sh[f'D{linha}'].value),
                            cpf=str(sh[f'V{linha}'].value).replace('.', '').replace('-', '').zfill(11),
                            rg=str(int(sh[f'W{linha}'].value)),
                            emissor='SSP/DF', email=str(sh[f'B{linha}'].value),
                            genero=str(sh[f'E{linha}'].value),
                            estado_civil=str(sh[f'F{linha}'].value), cor='9',
                            instru='08 - Educação Superior Incompleta',
                            nacional='Brasileiro(a)',
                            pai=str(sh[f'M{linha}'].value).upper(),
                            mae=str(sh[f'N{linha}'].value).upper(), endereco=str(sh[f'O{linha}'].value),
                            num='1',
                            bairro=str(sh[f'Q{linha}'].value),
                            cep=str(sh[f'R{linha}'].value).replace('.', '').replace('-', ''),
                            cidade='Brasília', cid_nas='Brasília - DF',
                            uf='DF',
                            cod_municipioend=municipios['DF']['Brasília'],
                            tel=str(sh[f'U{linha}'].value).replace('(', '').replace(')', '').replace('.', '').replace('-',
                                                                                                                      ''),
                            depto=depto, cargo=cargo,
                            horario=str(sh[f'AI{linha}'].value), salario=salario, tipo_contr=tipo_contr,
                            hr_sem='25', hr_mens='100',
                            est_semestre=str(sh[f'AS{linha}'].value),
                            est_turno=str(sh[f'AT{linha}'].value),
                            est_prev_conclu=str(sh[f'AU{linha}'].value),
                            est_faculdade=str(sh[f'AV{linha}'].value),
                            est_endfacul='End',
                            est_numendfacul='1',
                            est_bairroendfacul='Bairro',
                            ag=agencia, conta=conta, cdigito=digito

                        )
                        session.add(estag_cadastrado)
                        session.commit()
                    estag = session.query(Colaborador).filter_by(matricula=matricula).first()
                    pasta = r'\192.168.0.250'
                    try:
                        os.makedirs(
                            f'\\{pasta}\\rh\\01 - RH\\01 - Administração.Controles\\'
                            f'02 - Funcionários, Departamentos e Férias\\000 - Pastas Funcionais\\'
                            f'00 - ATIVOS\\0 - Estagiários\\0 - Ainda nao iniciaram\\{estag.nome}')
                        os.makedirs(
                            f'\\{pasta}\\rh\\01 - RH\\01 - Administração.Controles\\'
                            f'02 - Funcionários, Departamentos e Férias\\000 - Pastas Funcionais\\'
                            f'00 - ATIVOS\\0 - Estagiários\\0 - Ainda nao iniciaram\\{estag.nome}\\Atestados')
                        os.makedirs(
                            f'\\{pasta}\\rh\\01 - RH\\01 - Administração.Controles\\'
                            f'02 - Funcionários, Departamentos e Férias\\000 - Pastas Funcionais\\'
                            f'00 - ATIVOS\\0 - Estagiários\\0 - Ainda nao iniciaram\\{estag.nome}\\Diversos')
                        os.makedirs(
                            f'\\{pasta}\\rh\\01 - RH\\01 - Administração.Controles\\'
                            f'02 - Funcionários, Departamentos e Férias\\000 - Pastas Funcionais\\'
                            f'00 - ATIVOS\\0 - Estagiários\\0 - Ainda nao iniciaram\\{estag.nome}\\Contratuais')
                        os.makedirs(
                            f'\\{pasta}\\rh\\01 - RH\\01 - Administração.Controles\\'
                            f'02 - Funcionários, Departamentos e Férias\\000 - Pastas Funcionais\\'
                            f'00 - ATIVOS\\0 - Estagiários\\0 - Ainda nao iniciaram\\{estag.nome}\\Ferias')
                        os.makedirs(
                            f'\\{pasta}\\rh\\01 - RH\\01 - Administração.Controles\\'
                            f'02 - Funcionários, Departamentos e Férias\\000 - Pastas Funcionais\\'
                            f'00 - ATIVOS\\0 - Estagiários\\0 - Ainda nao iniciaram\\{estag.nome}\\Ponto')
                        os.makedirs(
                            f'\\{pasta}\\rh\\01 - RH\\01 - Administração.Controles\\'
                            f'02 - Funcionários, Departamentos e Férias\\000 - Pastas Funcionais\\'
                            f'00 - ATIVOS\\0 - Estagiários\\0 - Ainda nao iniciaram\\{estag.nome}\\Recibo')
                        os.makedirs(
                            f'\\{pasta}\\rh\\01 - RH\\01 - Administração.Controles\\'
                            f'02 - Funcionários, Departamentos e Férias\\000 - Pastas Funcionais\\'
                            f'00 - ATIVOS\\0 - Estagiários\\0 - Ainda nao iniciaram\\{estag.nome}\\Rescisao')
                    except FileExistsError:
                        pass
                    # abrir cadastro no dexion e atualizar informações campo a campo
                    pessoa = session.query(Colaborador).filter_by(matricula=matricula).first()
                    pastapessoa = f'\\{pasta}\\rh\\01 - RH\\01 - Administração.Controles\\' \
                                  f'02 - Funcionários, Departamentos e Férias\\000 - Pastas Funcionais\\' \
                                  f'00 - ATIVOS\\0 - Estagiários\\0 - Ainda nao iniciaram\\{pessoa.nome}'
                    pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Dexion.png')))
                    pa.press('alt'), pa.press('a'), pa.press('t'), t.sleep(2), pa.press(
                        'i'), t.sleep(5), pa.write(str(pessoa.matricula)), pa.press('enter'), t.sleep(20)
                    pp.copy(pessoa.nome), pa.hotkey('ctrl', 'v'), pa.press('tab'), pa.write(pessoa.cpf)
                    t.sleep(1), pa.press('tab', 3), pa.write(pessoa.genero), pa.press('tab'), pa.write(pessoa.cor)
                    t.sleep(1), pa.press('tab', 2), pa.write(pessoa.instru)
                    t.sleep(1), pa.press('tab')
                    if str(pessoa.estado_civil) == 'Casado(a)':
                        pa.write('2')
                        pa.press('tab', 6)
                    else:
                        pa.write('1')
                        pa.press('tab', 5)
                    pa.write(dt.strftime(dt.strptime(pessoa.nascimento, '%Y-%m-%d %H:%M:%S'), '%d%m%Y'))
                    t.sleep(1), pa.press('tab'), pp.copy(pessoa.cid_nas), pa.hotkey('ctrl', 'v'), pa.press('tab')
                    t.sleep(1), pa.write(pessoa.uf), pa.press('tab'), pa.write(pessoa.cod_municipioend), pa.press('tab')
                    t.sleep(1), pp.copy(pessoa.pai), pa.hotkey('ctrl', 'v'), pa.press('tab'), pa.write('105'), pa.press(
                        'tab')
                    t.sleep(1), pp.copy(pessoa.mae), pa.hotkey('ctrl', 'v'), pa.press('tab'), pa.write('105')
                    # # clique em documentos
                    pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Documentos.png')))
                    pa.press('tab'), pa.write(str(pessoa.rg)), pa.press('tab'), pa.write(
                        pessoa.emissor), pa.press('tab', 3), pa.write(pessoa.cod_municipioend)
                    pa.press('tab', 9), pa.write('3'), pa.press('tab'), pa.write('341'), pa.press('tab')
                    pa.write(pessoa.ag), pa.press('tab'), pa.write(f'{pessoa.conta}-{pessoa.cdigito}'), pa.press('tab')
                    # #clique em endereço
                    pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Endereco.png')))
                    pa.press('tab', 2), pp.copy(pessoa.endereco), pa.hotkey('ctrl', 'v'), pa.press(
                        'tab'), pa.write(pessoa.num), pa.press('tab', 2)
                    pp.copy(pessoa.bairro), pa.hotkey('ctrl', 'v'), pa.press('tab'), pp.copy(pessoa.cidade), pa.hotkey(
                        'ctrl', 'v'), pa.press('tab'), pa.write(pessoa.uf)
                    pa.press('tab'), pa.write(pessoa.cep), pa.press('tab'), pa.write(pessoa.cod_municipioend), pa.press(
                        'tab'), pa.write(str(pessoa.tel)), pa.press('tab', 2), pa.write(pessoa.email)
                    # #clique em dados contratuais
                    pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Contratuais.png')))
                    pa.press('tab'), pa.write(str(pessoa.admiss).replace('/', '')), t.sleep(1)
                    pa.press('tab'), pa.write('9')
                    pa.press('tab', 7), pa.write('n'), pa.press('tab'), pa.write('4')
                    pa.press('tab'), pa.write('Ed. Fisica')
                    pa.press('tab', 2), pa.write(str(int(str(pessoa.admiss).replace('/', '')) + 2).zfill(8))
                    # #clique em instituição de ensino
                    try:
                        pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/faculdade.png')))
                    except pa.ImageNotFoundException:
                        t.sleep(3)
                        pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/faculdade.png')))
                    pa.press('tab'), pp.copy(pessoa.est_faculdade), pa.hotkey('ctrl', 'v')
                    pa.press('tab'), pp.copy(pessoa.est_endfacul), pa.hotkey('ctrl', 'v')
                    pa.press('tab'), pp.copy(pessoa.est_numendfacul), pa.hotkey('ctrl', 'v')
                    pa.press('tab'), pp.copy(pessoa.est_bairroendfacul), pa.hotkey('ctrl', 'v')
                    # #clique em Outros
                    try:
                        pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Outros.png')))
                    except pa.ImageNotFoundException:
                        t.sleep(5)
                        pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Outros.png')))
                    t.sleep(2), pa.write('CARGO GERAL')
                    pa.press('tab'), pa.write(pessoa.cargo)
                    t.sleep(1), pa.press('tab'), pa.write(pessoa.salario), pa.press('tab')
                    if str(pessoa.tipo_contr) == 'Horista':
                        pa.press('1')
                    else:
                        pa.press('5')
                    pa.press('tab', 2), t.sleep(1), pa.write(str(pessoa.hr_sem)), pa.press('tab'), t.sleep(1)
                    pa.write(str(pessoa.hr_mens))
                    # #clique em eventos trabalhistas
                    pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/EVTrab.png')))
                    t.sleep(1)
                    # #clique em lotação
                    pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Lotacoes.png')))
                    pa.press('tab'), pa.press('tab'), pa.write('i'), pa.write(str(pessoa.admiss).replace('/', ''))
                    t.sleep(1), pa.press('enter'), t.sleep(1)
                    pp.copy(lotacao[f'{pessoa.depto}']), pa.hotkey('ctrl', 'v'), pa.press('enter'), pa.write('f')
                    pa.press('tab'), pa.write('4')
                    # #clique em salvar lotação
                    pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Salvarbtn.png'))), t.sleep(1)
                    # #clique em fechar lotação
                    try:
                        pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Fecharlot.png'))), t.sleep(1)
                    except pa.ImageNotFoundException:
                        t.sleep(4)
                        pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Fecharlot.png')))
                    # #clique em Compatibilidade
                    pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Compatibilidade3.png'))), t.sleep(1)
                    # #clique em Compatibilidade de novo
                    pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Compatibilidade2.png'))), t.sleep(1)
                    pa.press('tab', 2), pa.write('9')
                    # #clique em Salvar
                    pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Salvarcadastro.png'))), t.sleep(10)
                    # #clique em fechar novo cadastro
                    pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Fecharnovo1.png'))), t.sleep(2)
                    # #clique em fechar trabalhadores
                    pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Fechartrab1.png'))), t.sleep(0.5)
                    os.rename(pastapessoa,
                              f'\\{pasta}\\rh\\01 - RH\\01 - Administração.Controles\\'
                              f'02 - Funcionários, Departamentos e Férias\\000 - Pastas Funcionais\\'
                              f'00 - ATIVOS\\0 - Estagiários\\{pessoa.nome}')
                    tkinter.messagebox.showinfo(
                        title='Cadastro ok!',
                        message='Cadastro realizado com sucesso!'
                    )
                else:
                    # Cadastro iniciado em casa
                    wb = l_w(caminho)
                    sh = wb['Respostas ao formulário 1']
                    num, name = nome.strip().split(' - ')
                    linha = int(num)
                    lotacao = {'UNIDADE PARK SUL - QUALQUER DEPARTAMENTO': '0013', 'KIDS': '0010', 'MUSCULAÇÃO': '0007',
                               'ESPORTES E LUTAS': '0008', 'CROSSFIT': '0012', 'GINÁSTICA': '0006', 'GESTANTES': '0008',
                               'RECEPÇÃO': '0003',
                               'FINANCEIRO': '0001', 'TI': '0001', 'MARKETING': '0001', 'MANUTENÇÃO': '0004'}
                    if str(sh[f'E{linha}'].value) == 'Masculino':
                        cargo = 'ESTAGIARIO'
                    else:
                        cargo = 'ESTAGIARIA'

                    estag = session.query(Colaborador).filter_by(matricula=matricula).first()
                    if estag:
                        pass
                    else:
                        estag_cadastrado = Colaborador(
                            matricula=matricula, nome=name.upper(), admiss=admissao,
                            nascimento=str(sh[f'D{linha}'].value),
                            cpf=str(int(sh[f'V{linha}'].value)).zfill(11),
                            rg=str(int(sh[f'W{linha}'].value)),
                            emissor='SSP/DF', email=str(sh[f'B{linha}'].value),
                            genero=str(sh[f'E{linha}'].value),
                            estado_civil=str(sh[f'F{linha}'].value), cor='9',
                            instru='08 - Educação Superior Incompleta',
                            nacional='Brasileiro(a)',
                            pai=str(sh[f'M{linha}'].value).upper(),
                            mae=str(sh[f'N{linha}'].value).upper(), endereco=str(sh[f'O{linha}'].value),
                            num='1',
                            bairro=str(sh[f'Q{linha}'].value),
                            cep=str(sh[f'R{linha}'].value).replace('.', '').replace('-', ''),
                            cidade='Brasília', cid_nas='Brasília - DF',
                            uf='DF',
                            cod_municipioend=municipios['DF']['Brasília'],
                            tel=str(sh[f'U{linha}'].value).replace('(', '').replace(')', '').replace('.', '').replace('-',
                                                                                                                      ''),
                            depto=depto, cargo=cargo,
                            horario=str(sh[f'AI{linha}'].value), salario=salario, tipo_contr=tipo_contr,
                            hr_sem='25', hr_mens='100',
                            est_semestre=str(sh[f'AS{linha}'].value),
                            est_turno=str(sh[f'AT{linha}'].value),
                            est_prev_conclu=str(sh[f'AU{linha}'].value),
                            est_faculdade=str(sh[f'AV{linha}'].value),
                            est_endfacul='End',
                            est_numendfacul='1',
                            est_bairroendfacul='Bairro',
                            ag=agencia, conta=conta, cdigito=digito
                        )
                        session.add(estag_cadastrado)
                        session.commit()
                    pasta = r'\192.168.0.250'
                    try:
                        os.makedirs(
                            f'\\{pasta}\\rh\\01 - RH\\01 - Administração.Controles\\'
                            f'02 - Funcionários, Departamentos e Férias\\000 - Pastas Funcionais\\'
                            f'00 - ATIVOS\\0 - Estagiários\\0 - Ainda nao iniciaram\\{estag.nome}')
                        os.makedirs(
                            f'\\{pasta}\\rh\\01 - RH\\01 - Administração.Controles\\'
                            f'02 - Funcionários, Departamentos e Férias\\000 - Pastas Funcionais\\'
                            f'00 - ATIVOS\\0 - Estagiários\\0 - Ainda nao iniciaram\\{estag.nome}\\Atestados')
                        os.makedirs(
                            f'\\{pasta}\\rh\\01 - RH\\01 - Administração.Controles\\'
                            f'02 - Funcionários, Departamentos e Férias\\000 - Pastas Funcionais\\'
                            f'00 - ATIVOS\\0 - Estagiários\\0 - Ainda nao iniciaram\\{estag.nome}\\Diversos')
                        os.makedirs(
                            f'\\{pasta}\\rh\\01 - RH\\01 - Administração.Controles\\'
                            f'02 - Funcionários, Departamentos e Férias\\000 - Pastas Funcionais\\'
                            f'00 - ATIVOS\\0 - Estagiários\\0 - Ainda nao iniciaram\\{estag.nome}\\Contratuais')
                        os.makedirs(
                            f'\\{pasta}\\rh\\01 - RH\\01 - Administração.Controles\\'
                            f'02 - Funcionários, Departamentos e Férias\\000 - Pastas Funcionais\\'
                            f'00 - ATIVOS\\0 - Estagiários\\0 - Ainda nao iniciaram\\{estag.nome}\\Ferias')
                        os.makedirs(
                            f'\\{pasta}\\rh\\01 - RH\\01 - Administração.Controles\\'
                            f'02 - Funcionários, Departamentos e Férias\\000 - Pastas Funcionais\\'
                            f'00 - ATIVOS\\0 - Estagiários\\0 - Ainda nao iniciaram\\{estag.nome}\\Ponto')
                        os.makedirs(
                            f'\\{pasta}\\rh\\01 - RH\\01 - Administração.Controles\\'
                            f'02 - Funcionários, Departamentos e Férias\\000 - Pastas Funcionais\\'
                            f'00 - ATIVOS\\0 - Estagiários\\0 - Ainda nao iniciaram\\{estag.nome}\\Recibo')
                        os.makedirs(
                            f'\\{pasta}\\rh\\01 - RH\\01 - Administração.Controles\\'
                            f'02 - Funcionários, Departamentos e Férias\\000 - Pastas Funcionais\\'
                            f'00 - ATIVOS\\0 - Estagiários\\0 - Ainda nao iniciaram\\{estag.nome}\\Rescisao')
                    except FileExistsError:
                        pass
                    # abrir cadastro no dexion e atualizar informações campo a campo
                    pessoa = session.query(Colaborador).filter_by(matricula=matricula).first()
                    pastapessoa = f'\\{pasta}\\rh\\01 - RH\\01 - Administração.Controles\\' \
                                  f'02 - Funcionários, Departamentos e Férias\\000 - Pastas Funcionais\\' \
                                  f'00 - ATIVOS\\0 - Estagiários\\0 - Ainda nao iniciaram\\{pessoa.nome}'
                    pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Dexion.png')))
                    pa.press('alt'), pa.press('a'), pa.press('t'), t.sleep(2), pa.press(
                        'i'), t.sleep(5), pa.write(str(pessoa.matricula)), pa.press('enter'), t.sleep(20)
                    pp.copy(pessoa.nome), pa.hotkey('ctrl', 'v'), pa.press('tab'), pa.write(pessoa.cpf)
                    t.sleep(1), pa.press('tab', 3), pa.write(pessoa.genero), pa.press('tab'), pa.write(pessoa.cor)
                    t.sleep(1), pa.press('tab', 2), pa.write(pessoa.instru)
                    t.sleep(1), pa.press('tab')
                    if str(pessoa.estado_civil) == 'Casado(a)':
                        pa.write('2')
                        pa.press('tab', 6)
                    else:
                        pa.write('1')
                        pa.press('tab', 5)
                    pa.write(dt.strftime(dt.strptime(pessoa.nascimento, '%Y-%m-%d %H:%M:%S'), '%d%m%Y'))
                    t.sleep(1), pa.press('tab'), pp.copy(pessoa.cid_nas), pa.hotkey('ctrl', 'v'), pa.press('tab')
                    t.sleep(1), pa.write(pessoa.uf), pa.press('tab'), pa.write(pessoa.cod_municipioend), pa.press('tab')
                    t.sleep(1), pp.copy(pessoa.pai), pa.hotkey('ctrl', 'v'), pa.press('tab'), pa.write('105'), pa.press(
                        'tab')
                    t.sleep(1), pp.copy(pessoa.mae), pa.hotkey('ctrl', 'v'), pa.press('tab'), pa.write('105')
                    # # clique em documentos
                    pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Documentos.png')))
                    pa.press('tab'), pa.write(str(pessoa.rg)), pa.press('tab'), pa.write(
                        pessoa.emissor), pa.press('tab', 3), pa.write(pessoa.cod_municipioend)
                    pa.press('tab', 9), pa.write('3'), pa.press('tab'), pa.write('341'), pa.press('tab')
                    pa.write(pessoa.ag), pa.press('tab'), pa.write(f'{pessoa.conta}-{pessoa.cdigito}'), pa.press('tab')
                    # #clique em endereço
                    pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Endereco.png')))
                    pa.press('tab', 2), pp.copy(pessoa.endereco), pa.hotkey('ctrl', 'v'), pa.press(
                        'tab'), pa.write(pessoa.num), pa.press('tab', 2)
                    pp.copy(pessoa.bairro), pa.hotkey('ctrl', 'v'), pa.press('tab'), pp.copy(pessoa.cidade), pa.hotkey(
                        'ctrl', 'v'), pa.press('tab'), pa.write(pessoa.uf)
                    pa.press('tab'), pa.write(pessoa.cep), pa.press('tab'), pa.write(pessoa.cod_municipioend), pa.press(
                        'tab'), pa.write(str(pessoa.tel)), pa.press('tab', 2), pa.write(pessoa.email)
                    # #clique em dados contratuais
                    pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Contratuais.png')))
                    pa.press('tab'), pa.write(str(pessoa.admiss).replace('/', '')), t.sleep(1)
                    pa.press('tab'), pa.write('9')
                    pa.press('tab', 7), pa.write('n'), pa.press('tab'), pa.write('4')
                    pa.press('tab'), pa.write('Ed. Fisica')
                    pa.press('tab', 2), pa.write(str(int(str(pessoa.admiss).replace('/', '')) + 2).zfill(8))
                    # #clique em instituição de ensino
                    try:
                        pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/faculdade.png')))
                    except pa.ImageNotFoundException:
                        t.sleep(3)
                        pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/faculdade.png')))
                    pa.press('tab'), pp.copy(pessoa.est_faculdade), pa.hotkey('ctrl', 'v')
                    pa.press('tab'), pp.copy(pessoa.est_endfacul), pa.hotkey('ctrl', 'v')
                    pa.press('tab'), pp.copy(pessoa.est_numendfacul), pa.hotkey('ctrl', 'v')
                    pa.press('tab'), pp.copy(pessoa.est_bairroendfacul), pa.hotkey('ctrl', 'v')
                    # #clique em Outros
                    try:
                        pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Outros.png')))
                    except pa.ImageNotFoundException:
                        t.sleep(5)
                        pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Outros.png')))
                    t.sleep(2), pa.write('CARGO GERAL')
                    pa.press('tab'), pa.write(pessoa.cargo)
                    t.sleep(1), pa.press('tab'), pa.write(pessoa.salario), pa.press('tab')
                    if str(pessoa.tipo_contr) == 'Horista':
                        pa.press('1')
                    else:
                        pa.press('5')
                    pa.press('tab', 2), t.sleep(1), pa.write(str(pessoa.hr_sem)), pa.press('tab'), t.sleep(1)
                    pa.write(str(pessoa.hr_mens))
                    # #clique em eventos trabalhistas
                    pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/EVTrab.png')))
                    t.sleep(1)
                    # #clique em lotação
                    pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Lotacoes.png')))
                    pa.press('tab'), pa.press('tab'), pa.write('i'), pa.write(str(pessoa.admiss).replace('/', ''))
                    t.sleep(1), pa.press('enter'), t.sleep(1)
                    pp.copy(lotacao[f'{pessoa.depto}']), pa.hotkey('ctrl', 'v'), pa.press('enter'), pa.write('f')
                    pa.press('tab'), pa.write('4')
                    # #clique em salvar lotação
                    pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Salvarbtn.png'))), t.sleep(1)
                    # #clique em fechar lotação
                    try:
                        pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Fecharlot.png'))), t.sleep(1)
                    except pa.ImageNotFoundException:
                        t.sleep(4)
                        pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Fecharlot.png')))
                    # #clique em Compatibilidade
                    pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Compatibilidade3.png'))), t.sleep(1)
                    # #clique em Compatibilidade de novo
                    pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Compatibilidade2.png'))), t.sleep(1)
                    pa.press('tab', 2), pa.write('9')
                    # #clique em Salvar
                    pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Salvarcadastro.png'))), t.sleep(10)
                    # #clique em fechar novo cadastro
                    pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Fecharnovo1.png'))), t.sleep(2)
                    # #clique em fechar trabalhadores
                    pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Fechartrab1.png'))), t.sleep(0.5)
                    os.rename(pastapessoa,
                              f'\\{pasta}\\rh\\01 - RH\\01 - Administração.Controles\\'
                              f'02 - Funcionários, Departamentos e Férias\\000 - Pastas Funcionais\\'
                              f'00 - ATIVOS\\0 - Estagiários\\{pessoa.nome}')
                    tkinter.messagebox.showinfo(
                        title='Cadastro ok!',
                        message='Cadastro realizado com sucesso!'
                    )
            else:
                if ondestou == 0:
                    # Editando o cadastro na Cia
                    wb = l_w(caminho)
                    sh = wb['Respostas ao formulário 1']
                    num, name = nome.strip().split(' - ')
                    linha = int(num)
                    lotacao = {
                        'Unidade Park Sul - qualquer departamento': ['0013', 'Thais Feitosa',
                                                                     'thais.morais@ciaathletica.com.br',
                                                                     'Líder Park Sul'],
                        'Kids': ['0010', 'Cindy Stefanie', 'cindy.neves@ciaathletica.com.br', 'Líder Kids'],
                        'Musculação': ['0007', 'Aline Kanyó', 'aline.kanyo@soucia.com.br', 'Líder Musculação'],
                        'Esportes e Lutas': ['0008', 'Morgana Rossini', 'morganalourenco@yahoo.com.br', 'Líder Natação'],
                        'Crossfit': ['0012', 'Guilherme Salles', 'gmoreirasalles@gmail.com', 'Líder Crossfit'],
                        'Ginástica': ['0006', 'Hugo Albuquerque', 'hugo.albuquerque@ciaathletica.com.br',
                                      'Líder Ginástica'],
                        'Gestantes': ['0006', 'Hugo Albuquerque', 'hugo.albuquerque@ciaathletica.com.br',
                                      'Líder Ginástica'],
                        'Recepção': ['0003', 'Paulo Renato', 'paulo.simoes@ciaathletica.com.br', 'Gerente Vendas'],
                        'Administrativo': ['0001', 'Felipe Rodrigues', 'felipe.rodrigues@ciaathletica.com.br',
                                           'Gerente RH'],
                        'Manutenção': ['0004', 'José Aparecido', 'aparecido.grota@ciaathletica.com.br',
                                       'Gerente Manutenção'],
                    }
                else:
                    # Editando o cadastro em Casa
                    wb = l_w(caminho)
                    sh = wb['Respostas ao formulário 1']
                    num, name = nome.strip().split(' - ')
                    linha = int(num)
                    lotacao = {'UNIDADE PARK SUL - QUALQUER DEPARTAMENTO': '0013', 'KIDS': '0010', 'MUSCULAÇÃO': '0007',
                               'ESPORTES E LUTAS': '0008', 'CROSSFIT': '0012', 'GINÁSTICA': '0006', 'GESTANTES': '0008',
                               'RECEPÇÃO': '0003',
                               'FINANCEIRO': '0001', 'TI': '0001', 'MARKETING': '0001', 'MANUTENÇÃO': '0004'}
                    if str(sh[f'E{linha}'].value) == 'Masculino':
                        cargo = 'ESTAGIARIO'
                    else:
                        cargo = 'ESTAGIARIA'

                    estag = session.query(Colaborador).filter_by(matricula=matricula).first()
                    if estag:
                        pass
                    else:
                        estag_cadastrado = Colaborador(
                            matricula=matricula, nome=name.upper(), admiss=admissao,
                            nascimento=str(sh[f'D{linha}'].value),
                            cpf=str(int(sh[f'V{linha}'].value)).zfill(11),
                            rg=str(int(sh[f'W{linha}'].value)),
                            emissor='SSP/DF', email=str(sh[f'B{linha}'].value),
                            genero=str(sh[f'E{linha}'].value),
                            estado_civil=str(sh[f'F{linha}'].value), cor='9',
                            instru='08 - Educação Superior Incompleta',
                            nacional='Brasileiro(a)',
                            pai=str(sh[f'M{linha}'].value).upper(),
                            mae=str(sh[f'N{linha}'].value).upper(), endereco=str(sh[f'O{linha}'].value),
                            num='1',
                            bairro=str(sh[f'Q{linha}'].value),
                            cep=str(sh[f'R{linha}'].value).replace('.', '').replace('-', ''),
                            cidade='Brasília', cid_nas='Brasília - DF',
                            uf='DF',
                            cod_municipioend=municipios['DF']['Brasília'],
                            tel=str(sh[f'U{linha}'].value).replace('(', '').replace(')', '').replace('.', '').replace('-',
                                                                                                                      ''),
                            depto=depto, cargo=cargo,
                            horario=str(sh[f'AI{linha}'].value), salario=salario, tipo_contr=tipo_contr,
                            hr_sem='25', hr_mens='100',
                            est_semestre=str(sh[f'AS{linha}'].value),
                            est_turno=str(sh[f'AT{linha}'].value),
                            est_prev_conclu=str(sh[f'AU{linha}'].value),
                            est_faculdade=str(sh[f'AV{linha}'].value),
                            est_endfacul='End',
                            est_numendfacul='1',
                            est_bairroendfacul='Bairro',
                            ag=agencia, conta=conta, cdigito=digito
                        )
                        session.add(estag_cadastrado)
                        session.commit()
                    pasta = r'\192.168.0.250'
                    try:
                        os.makedirs(
                            f'\\{pasta}\\rh\\01 - RH\\01 - Administração.Controles\\'
                            f'02 - Funcionários, Departamentos e Férias\\000 - Pastas Funcionais\\'
                            f'00 - ATIVOS\\0 - Estagiários\\0 - Ainda nao iniciaram\\{estag.nome}')
                        os.makedirs(
                            f'\\{pasta}\\rh\\01 - RH\\01 - Administração.Controles\\'
                            f'02 - Funcionários, Departamentos e Férias\\000 - Pastas Funcionais\\'
                            f'00 - ATIVOS\\0 - Estagiários\\0 - Ainda nao iniciaram\\{estag.nome}\\Atestados')
                        os.makedirs(
                            f'\\{pasta}\\rh\\01 - RH\\01 - Administração.Controles\\'
                            f'02 - Funcionários, Departamentos e Férias\\000 - Pastas Funcionais\\'
                            f'00 - ATIVOS\\0 - Estagiários\\0 - Ainda nao iniciaram\\{estag.nome}\\Diversos')
                        os.makedirs(
                            f'\\{pasta}\\rh\\01 - RH\\01 - Administração.Controles\\'
                            f'02 - Funcionários, Departamentos e Férias\\000 - Pastas Funcionais\\'
                            f'00 - ATIVOS\\0 - Estagiários\\0 - Ainda nao iniciaram\\{estag.nome}\\Contratuais')
                        os.makedirs(
                            f'\\{pasta}\\rh\\01 - RH\\01 - Administração.Controles\\'
                            f'02 - Funcionários, Departamentos e Férias\\000 - Pastas Funcionais\\'
                            f'00 - ATIVOS\\0 - Estagiários\\0 - Ainda nao iniciaram\\{estag.nome}\\Ferias')
                        os.makedirs(
                            f'\\{pasta}\\rh\\01 - RH\\01 - Administração.Controles\\'
                            f'02 - Funcionários, Departamentos e Férias\\000 - Pastas Funcionais\\'
                            f'00 - ATIVOS\\0 - Estagiários\\0 - Ainda nao iniciaram\\{estag.nome}\\Ponto')
                        os.makedirs(
                            f'\\{pasta}\\rh\\01 - RH\\01 - Administração.Controles\\'
                            f'02 - Funcionários, Departamentos e Férias\\000 - Pastas Funcionais\\'
                            f'00 - ATIVOS\\0 - Estagiários\\0 - Ainda nao iniciaram\\{estag.nome}\\Recibo')
                        os.makedirs(
                            f'\\{pasta}\\rh\\01 - RH\\01 - Administração.Controles\\'
                            f'02 - Funcionários, Departamentos e Férias\\000 - Pastas Funcionais\\'
                            f'00 - ATIVOS\\0 - Estagiários\\0 - Ainda nao iniciaram\\{estag.nome}\\Rescisao')
                    except FileExistsError:
                        pass
                    # abrir cadastro no dexion e atualizar informações campo a campo
                    pessoa = session.query(Colaborador).filter_by(matricula=matricula).first()
                    pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Dexion.png')))
                    pa.press('alt'), pa.press('a'), pa.press('t'), t.sleep(2), pa.press(
                        'a'), t.sleep(5), pa.write(str(pessoa.matricula)), pa.press('enter'), t.sleep(20)
                    pp.copy(pessoa.nome), pa.hotkey('ctrl', 'v'), pa.press('tab'), pa.write(pessoa.cpf)
                    t.sleep(1), pa.press('tab', 3), pa.write(pessoa.genero), pa.press('tab'), pa.write(pessoa.cor)
                    t.sleep(1), pa.press('tab', 2), pa.write(pessoa.instru)
                    t.sleep(1), pa.press('tab')
                    if str(pessoa.estado_civil) == 'Casado(a)':
                        pa.write('2')
                        pa.press('tab', 6)
                    else:
                        pa.write('1')
                        pa.press('tab', 5)
                    pa.write(dt.strftime(dt.strptime(pessoa.nascimento, '%Y-%m-%d %H:%M:%S'), '%d%m%Y'))
                    t.sleep(1), pa.press('tab'), pp.copy(pessoa.cid_nas), pa.hotkey('ctrl', 'v'), pa.press('tab')
                    t.sleep(1), pa.write(pessoa.uf), pa.press('tab'), pa.write(pessoa.cod_municipioend), pa.press('tab')
                    t.sleep(1), pp.copy(pessoa.pai), pa.hotkey('ctrl', 'v'), pa.press('tab'), pa.write('105'), pa.press(
                        'tab')
                    t.sleep(1), pp.copy(pessoa.mae), pa.hotkey('ctrl', 'v'), pa.press('tab'), pa.write('105')
                    # # clique em documentos
                    pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Documentos.png')))
                    pa.press('tab'), pa.write(str(pessoa.rg)), pa.press('tab'), pa.write(
                        pessoa.emissor), pa.press('tab', 3), pa.write(pessoa.cod_municipioend)
                    pa.press('tab', 9), pa.write('3'), pa.press('tab'), pa.write('341'), pa.press('tab')
                    pa.write(pessoa.ag), pa.press('tab'), pa.write(f'{pessoa.conta}-{pessoa.cdigito}'), pa.press('tab')
                    # #clique em endereço
                    pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Endereco.png')))
                    pa.press('tab', 2), pp.copy(pessoa.endereco), pa.hotkey('ctrl', 'v'), pa.press(
                        'tab'), pa.write(pessoa.num), pa.press('tab', 2)
                    pp.copy(pessoa.bairro), pa.hotkey('ctrl', 'v'), pa.press('tab'), pp.copy(pessoa.cidade), pa.hotkey(
                        'ctrl', 'v'), pa.press('tab'), pa.write(pessoa.uf)
                    pa.press('tab'), pa.write(pessoa.cep), pa.press('tab'), pa.write(pessoa.cod_municipioend), pa.press(
                        'tab'), pa.write(str(pessoa.tel)), pa.press('tab', 2), pa.write(pessoa.email)
                    # #clique em dados contratuais
                    pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Contratuais.png')))
                    pa.press('tab'), pa.write(str(pessoa.admiss).replace('/', '')), t.sleep(1)
                    pa.press('tab'), pa.write('9')
                    pa.press('tab', 7), pa.write('n'), pa.press('tab'), pa.write('4')
                    pa.press('tab'), pa.write('Ed. Fisica')
                    pa.press('tab', 2), pa.write(str(int(str(pessoa.admiss).replace('/', '')) + 2).zfill(8))
                    # #clique em instituição de ensino
                    pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/faculdade.png')))
                    pa.press('tab'), pp.copy(pessoa.est_faculdade), pa.hotkey('ctrl', 'v')
                    pa.press('tab'), pp.copy(pessoa.est_endfacul), pa.hotkey('ctrl', 'v')
                    pa.press('tab'), pp.copy(pessoa.est_numendfacul), pa.hotkey('ctrl', 'v')
                    pa.press('tab'), pp.copy(pessoa.est_bairroendfacul), pa.hotkey('ctrl', 'v')
                    # #clique em Outros
                    try:
                        pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Outros.png')))
                    except pa.ImageNotFoundException:
                        t.sleep(5)
                        pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Outros.png')))
                    t.sleep(2), pa.write('CARGO GERAL')
                    pa.press('tab'), pa.write(pessoa.cargo)
                    t.sleep(1), pa.press('tab'), pa.write(pessoa.salario), pa.press('tab')
                    if str(pessoa.tipo_contr) == 'Horista':
                        pa.press('1')
                    else:
                        pa.press('5')
                    pa.press('tab', 2), t.sleep(1), pa.write(str(pessoa.hr_sem)), pa.press('tab'), t.sleep(1)
                    pa.write(str(pessoa.hr_mens))
                    # #clique em Compatibilidade
                    pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Compatibilidade3.png'))), t.sleep(1)
                    # #clique em Compatibilidade de novo
                    pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Compatibilidade2.png'))), t.sleep(1)
                    pa.press('tab', 2), pa.write('9')
                    # #clique em Salvar
                    pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Salvarcadastro.png'))), t.sleep(10)
                    # #clique em fechar novo cadastro
                    pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Fecharnovo1.png'))), t.sleep(2)
                    # #clique em fechar trabalhadores
                    pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Fechartrab1.png'))), t.sleep(0.5)
                    tkinter.messagebox.showinfo(
                        title='Cadastro ok!',
                        message='Cadastro editado com sucesso!'
                    )


def cadastrar_autonomo(caminhoaut, nomeaut, matriculaaut, admissaoaut, cargoaut, deptoaut, ondeaut):
    sessions = sessionmaker(bind=engine)
    session = sessions()
    # Cadastro iniciado em casa
    wb = l_w(caminhoaut)
    sh = wb['Respostas ao formulário 1']
    num, name = nomeaut.strip().split(' - ')
    linha = int(num)
    lotacao = {'UNIDADE PARK SUL - QUALQUER DEPARTAMENTO': '0013', 'KIDS': '0010', 'MUSCULAÇÃO': '0007',
               'ESPORTES E LUTAS': '0008', 'CROSSFIT': '0012', 'GINÁSTICA': '0006', 'GESTANTES': '0008',
               'RECEPÇÃO': '0003',
               'FINANCEIRO': '0001', 'TI': '0001', 'MARKETING': '0001', 'MANUTENÇÃO': '0004'}
    aut = session.query(Colaborador).filter_by(matricula=matriculaaut).first()
    if aut:
        pass
    else:
        aut_cadastrado = Colaborador(
            matricula=matriculaaut, nome=name.upper(), admiss=admissaoaut,
            nascimento=str(sh[f'D{linha}'].value),
            pis=str(sh[f'S{linha}'].value).replace('.', '').replace('-', '').zfill(11),
            cpf=str(int(sh[f'P{linha}'].value)).zfill(11),
            rg=str(int(sh[f'Q{linha}'].value)),
            emissor='SSP/DF', email=str(sh[f'B{linha}'].value),
            genero=str(sh[f'E{linha}'].value), cor='9',
            instru=str(sh[f'F{linha}'].value),
            nacional='Brasileiro(a)', estado_civil='Solteiro(a)',
            endereco=str(sh[f'I{linha}'].value),
            num=str(sh[f'J{linha}'].value),
            bairro=str(sh[f'K{linha}'].value), cep=str(sh[f'L{linha}'].value).replace('.', '').replace('-', ''),
            cidade=str(sh[f'M{linha}'].value), cid_nas='Brasília - DF', uf='DF',
            cod_municipioend=municipios['DF']['Brasília'],
            tel=str(sh[f'O{linha}'].value).replace('(', '').replace(')', '').replace('.', '').replace('-', ''),
            depto=deptoaut, cargo=cargoaut,
        )
        session.add(aut_cadastrado)
        session.commit()
    # abrir cadastro no dexion e atualizar informações campo a campo
    pessoa = session.query(Colaborador).filter_by(matricula=matriculaaut).first()
    pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Dexion.png')))
    pa.press('alt'), pa.press('a'), pa.press('t'), t.sleep(2), pa.press(
        'i'), t.sleep(5), pa.write(str(pessoa.matricula)), pa.press('enter'), t.sleep(20)
    pp.copy(pessoa.nome), pa.hotkey('ctrl', 'v'), pa.press('tab'), pa.write(pessoa.cpf)
    t.sleep(1), pa.press('tab', 3), pa.write(pessoa.genero), pa.press('tab'), pa.write(pessoa.cor)
    t.sleep(1), pa.press('tab', 2), pa.write(pessoa.instru)
    t.sleep(1), pa.press('tab')
    if str(pessoa.estado_civil) == 'Casado(a)':
        pa.write('2')
        pa.press('tab', 6)
    else:
        pa.write('1')
        pa.press('tab', 5)
    pa.write(dt.strftime(dt.strptime(pessoa.nascimento, '%Y-%m-%d %H:%M:%S'), '%d%m%Y'))
    t.sleep(1), pa.press('tab'), pp.copy(pessoa.cid_nas), pa.hotkey('ctrl', 'v')
    # # clique em documentos
    pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Documentos.png')))
    pa.press('tab'), pa.write(str(pessoa.rg)), pa.press('tab'), pa.write(
        pessoa.emissor), pa.press('tab', 3), pa.write(pessoa.cod_municipioend), pa.press('tab')
    pa.write(pessoa.pis)
    # #clique em endereço
    pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Endereco.png')))
    pa.press('tab', 2), pp.copy(pessoa.endereco), pa.hotkey('ctrl', 'v'), pa.press(
        'tab'), pa.write(pessoa.num), pa.press('tab', 2)
    pp.copy(pessoa.bairro), pa.hotkey('ctrl', 'v'), pa.press('tab'), pp.copy(pessoa.cidade), pa.hotkey(
        'ctrl', 'v'), pa.press('tab'), pa.write(pessoa.uf)
    pa.press('tab'), pa.write(pessoa.cep), pa.press('tab'), pa.write(pessoa.cod_municipioend), pa.press(
        'tab'), pa.write(str(pessoa.tel)), pa.press('tab', 2), pa.write(pessoa.email)
    # #clique em dados contratuais
    pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Contratuais.png')))
    pa.press('tab'), pa.write(str(pessoa.admiss).replace('/', '')), t.sleep(1)
    pa.press('tab'), pa.write('7')
    # #clique em Outros
    try:
        pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Outros.png')))
    except pa.ImageNotFoundException:
        t.sleep(5)
        pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Outros.png')))
    t.sleep(2), pa.write('CARGO GERAL')
    pa.press('tab'), pa.write(pessoa.cargo)
    # #clique em eventos trabalhistas
    pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/EVTrab.png')))
    t.sleep(1)
    # #clique em lotação
    pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Lotacoes.png')))
    pa.press('tab'), pa.press('tab'), pa.write('i'), pa.write(str(pessoa.admiss).replace('/', ''))
    t.sleep(1), pa.press('enter'), t.sleep(1)
    pp.copy(lotacao[f'{pessoa.depto}']), pa.hotkey('ctrl', 'v'), pa.press('enter'), pa.write('3')
    # #clique em salvar lotação
    pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Salvarbtn.png'))), t.sleep(1)
    # #clique em fechar lotação
    try:
        pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Fecharlot.png'))), t.sleep(1)
    except pa.ImageNotFoundException:
        t.sleep(4)
        pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Fecharlot.png')))
    # #clique em Compatibilidade
    pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Compatibilidade3.png'))), t.sleep(1)
    # #clique em Compatibilidade de novo
    pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Compatibilidade2.png'))), t.sleep(1)
    pa.press('tab', 2), pa.write('13')
    # #clique em Salvar
    pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Salvarcadastro.png'))), t.sleep(10)
    # #clique em fechar novo cadastro
    pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Fecharnovo1.png'))), t.sleep(2)
    # #clique em fechar trabalhadores
    pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Fechartrab1.png'))), t.sleep(0.5)
    tkinter.messagebox.showinfo(
        title='Cadastro ok!',
        message='Cadastro realizado com sucesso!'
    )


def validar_pis(local, nome):
    wb = l_w(local, read_only=False)
    sh = wb['Respostas ao formulário 1']
    num, name = nome.strip().split(' - ')
    x = int(num)
    pis = str(sh[f'S{x}'].value).replace('-', '').replace('.', '').zfill(11)
    v1 = int(pis[0]) * 3
    v2 = int(pis[1]) * 2
    v3 = int(pis[2]) * 9
    v4 = int(pis[3]) * 8
    v5 = int(pis[4]) * 7
    v6 = int(pis[5]) * 6
    v7 = int(pis[6]) * 5
    v8 = int(pis[7]) * 4
    v9 = int(pis[8]) * 3
    v10 = int(pis[9]) * 2
    d = int(pis[10])
    soma = v1 + v2 + v3 + v4 + v5 + v6 + v7 + v8 + v9 + v10
    divisao = soma % 11
    resultado = 11 - divisao
    if resultado != d:
        if resultado == 10 & d == 0 | resultado == 11 & d == 0:
            pass
        else:
            tkinter.messagebox.showinfo(
                title='Erro!',
                message='PIS inválido!'
            )
    else:
        tkinter.messagebox.showinfo(
            title='Ok!',
            message='PIS ok!'
        )


def send_email(matriculas):
    # code to send e-mails through smtplib
    # set up smtp connection
    s = smtplib.SMTP(host=host, port=port)
    s.starttls()
    s.login(em_rem, k1)
    # send e-mails to a list of employees
    wb = l_w('Nomes e e-mails.xlsx')
    sh = wb['Dados']
    x = 1
    while x <= len(sh['A']):
        msg = MIMEMultipart()
        message = f'''
        Olá, {str(sh[f'A{x}'].value).title().split(sep=' ')[0]}!\n
        \n
        Para repor o encontro com colaboradores novatos cancelado no dia 21/03, abrimos novo horário hoje:\n
        24/03:
        14h às 15h30 - Sala 3.
        \n
        Atenciosamente,\n
        Felipe Rodrigues
        '''
        # parameters of the message
        msg['From'] = em_rem
        msg['To'] = str(sh[f'B{x}'].value).lower()
        msg['Subject'] = "Reposição Encontro Cinthia Guimarães"
        msg.attach(MIMEText(message, 'plain', _charset='utf-8'))
        s.send_message(msg)
        del msg
        x += 1
    s.quit()


def send_wpp():
    # code to send whatsapp messages through browser
    wb = l_w("AV.xlsm")
    sh = wb['Planilha1']

    for x in range(2, len(sh['A'])):
        pessoa = str(sh[f"B{x}"].value).split(' ')[0]
        email = str(sh[f"D{x}"].value)
        numero = str(sh[f"H{x}"].value)
        url = str(sh[f"F{x}"].value)
        mensagem = f'Oi {pessoa}, te enviei por e-mail(no {email}) o resultado da primeira etapa da sua avaliação de ' \
                   f'desempenho e o link da pesquisa sobre a avaliação. Antes da segunda etapa, precisamos que ' \
                   f'responda a pesquisa. Ok? Se puder responder agora, é bem rápido, dura no máximo 5 minutos. ' \
                   f'Segue o link: {url}'
        texto = urllib.parse.quote(mensagem)
        cel = urllib.parse.quote(numero)
        link = f'https://web.whatsapp.com/send?phone={cel}&text={texto}'
        if numero:
            print(link)
            print(pessoa)
            print(email)
            print(numero)
        x += 1


def desligar_pessoa(nome: str, data: str, tipo: int):
    """
    This function does all the procedures for terminating an employee: it issues documents, sends them by e-mail,
    saves them in the respective folders, moves folders and schedules appointments.

    The function works according to the type of dismissal entered at 'tipo' parameter. For each type of dismissal
    there are specific procedures to be performed.

    Through the 'tipo' parameter, the subfunction of the dictionary 'desligamento' is called by an if condition.

    :param nome: Employee's name
    :param data: Dismiss date
    :param tipo: Dismiss type
    :return: Procedures to dismiss employee
    """
    sessions = sessionmaker(engine)
    session = sessions()
    pessoa = session.query(Colaborador).filter_by(nome=nome).first()

    def desligar_estag():
        pasta_rescisao = rede + rf'\02 - Funcionários, Departamentos e Férias\000 - Pastas Funcionais\00 - ATIVOS\0 - Estagiários\{pessoa.nome}\Rescisao'
        # send e-mails to end intern contract
        email_remetente = em_rem
        senha = k1
        # set up smtpp connection
        s = smtplib.SMTP(host=host, port=port)
        s.starttls()
        s.login(email_remetente, senha)

        # send e-mail to intern
        msg = MIMEMultipart('alternative')
        arquivo = pasta_rescisao + f'\\TRCT.pdf'
        text = MIMEText(f'''Olá, {pessoa.nome.split(" ")[0].title()}!<br><br>
        Obrigado por sua dedicação no Programa Novos Talentos da Companhia Athletica de Brasília!<br>
        Seu desligamento do estágio foi efetuado em {data}.<br>
        Para concluirmos essa etapa precisamos que você compareça a Companhia para devolver uniformes, BTS e assinar o termo de rescisão anexo.<br>
        Algumas faculdades exigem o termo de desligamento do IF, vou solicitar a eles que nos envie.<br>
        Qual o melhor dia para você para nos encontrarmos na Cia Athletica e assinarmos o seu termo de rescisão?<br>
        Assim que o termo for assinado agendamos o pagamento do valor final.<br>
        Aguardo você me informar a melhor data para assinarmos o termo.<br><br>
        
        Atenciosamente,<br>
        <img src="cid:image1">''', 'html')
        msg.attach(text)
        image = MIMEImage(
            open(rf'C:\Users\{os.getlogin()}\PycharmProjects\AutomacaoCia\src\models\static\imgs\assinatura.png', 'rb').read())
        image.add_header('Content-ID', '<image1>')
        msg.attach(image)
        # set up the parameters of the message
        msg['From'] = email_remetente
        msg['To'] = pessoa.email
        msg['Subject'] = "Desligamento Estágio Cia Athletica"
        # attach pdf file
        part = MIMEBase('application', "octet-stream")
        part.set_payload(open(arquivo, "rb").read())
        encoders.encode_base64(part)
        part.add_header('Content-Disposition', 'attachment',
                        filename=f'TRCT {pessoa.nome.split(" ")[0].title()}.pdf')
        msg.attach(part)
        s.sendmail(email_remetente, pessoa.email, msg.as_string())
        del msg

        # send e-mail to coworker asking to exclude intern register
        msg = MIMEMultipart('alternative')
        text = MIMEText(f'''Oi, Wallace!<br><br>
        Favor desativar o(a) estagiário(a) {pessoa.nome.title()}.<br><br>
        Abs.,<br>
        <img src="cid:image1">''', 'html')
        msg.attach(text)
        image = MIMEImage(
            open(rf'C:\Users\{os.getlogin()}\PycharmProjects\AutomacaoCia\src\models\static\imgs\assinatura.png', 'rb').read())
        # define the image's ID as referenced in the HTML body above
        image.add_header('Content-ID', '<image1>')
        msg.attach(image)
        # set up the parameters of the message
        msg['From'] = email_remetente
        msg['To'] = em_ti
        msg['Subject'] = f"Desligamento Estágio - {str(pessoa.nome).split(' ')[0].title()}"
        s.sendmail(email_remetente, em_ti, msg.as_string())
        del msg

        # send document asking for terminate intern's contract
        msg = MIMEMultipart('alternative')
        text = MIMEText(
            f'''Olá!<br><br>
            Favor desligar estagiário(a) {pessoa.nome.title()}, CPF: {pessoa.cpf}, na data {data}.<br><br>
            Atenciosamente,<br><img src="cid:image1">''',
            'html')
        msg.attach(text)
        image = MIMEImage(
            open(rf'C:\Users\{os.getlogin()}\PycharmProjects\AutomacaoCia\src\models\static\imgs\assinatura.png', 'rb').read())
        # define the image's ID as referenced in the HTML body above
        image.add_header('Content-ID', '<image1>')
        msg.attach(image)
        # set up the parameters of the message
        msg['From'] = email_remetente
        msg['To'] = em_if
        msg['Subject'] = f"Desligamento Estágio - {str(pessoa.nome).split(' ')[0]}"
        s.sendmail(email_remetente, em_if, msg.as_string())
        del msg
        s.quit()
        os.rename(pasta_rescisao.replace(r'\Rescisao',''), pasta_rescisao.replace(r'\Rescisao','')
                  .replace('00 - ATIVOS', '01 - Inativos'))

    def desligar_func_apedido_com_aviso():
        # e-mail informdando data de crédito na conta e solicitando data para marcar no sindicato e dev uniformes
        print(f'{nome} foi desligado(a) em {data} com rescisão do tipo: Desligametno de Func à pedido com aviso.')

    def desligar_func_apedido_sem_aviso():
        print(f'{nome} foi desligado(a) em {data} com rescisão do tipo: Desligametno de Func à pedido sem aviso.')

    def desligar_func_por_acordo():
        # após gerada a rescisão e guia e-mail informanda dia do crédito em conta, guias de fgts e seguro
        # explicar quanto saca do fgts
        # e-mail marcando data para ir no sindicato, dev uniformes e bts
        # e-mail para TI informando nome e CPF do funcionário/estagiário e solicitando o desligamento
        print(f'{nome} foi desligado(a) em {data} com rescisão do tipo: Desligametno de Func por acordo.')

    def desligar_func_sem_aviso():
        # gerar docs de homologação no dexion: Rescisão 5 cópias, av prév, comprovantes recolhimento inss, carta preposto,
        # folha de registro, carta abono conduta, guia de seguro desemprego(?)
        # e-mails com data do pgto, orientações do passo a passo, guias de orientação do FGTS e Seguro desemprego
        # explicar quanto saca do fgts
        # solicitar data para agendar no sindicato
        print(f'{nome} foi desligado(a) em {data} com rescisão do tipo: Desligametno de Func sem aviso.')

    def desligar_func_com_aviso():
        # após gerada a rescisão e guia e-mail informanda dia do crédito em conta, guias de fgts e seguro
        # explicar quanto saca do fgts
        # e-mail marcando data para ir no sindicato, dev uniformes e bts



        # FOLHA DE ROSTO
        #
        # TRCT(PG 1 E 2) 5x
        #
        # COMPROVANTES PGTO TRCT (pegar no itau manualmente e salvar na pasta com nome padrao -- cógigo pegar na pasta)
        # EXTRATO FGTS
        # CHAVE MOVIMENT CAIXA
        # MULTA FGTS
        # COMPROV PGTO MULTA FGTS
        # __________________________________________
        # AVISO PREVIO
        #
        # alt r s d e
        # matricula
        # tab, 2
        # desligamento
        # tab, 2
        # espaço
        # tab, 4
        # opção
        # 2
        # h(1)
        # ou
        # 7
        # dias(2)
        # tab
        # desligamento
        # tab, 2
        # 1100
        # tab, 4
        # clique
        # visualizar
        # salvar
        # pdf
        # na
        # pasta
        # do
        # func,
        # cliques
        # para
        # fechar
        #
        # ____________________________________________
        #
        # GUIA SEGURO DESEMPREGO
        # RELAÇÃO SALÁRIOS DE CONTRIBUIÇ
        # DISCRIMINAÇÃO DAS PARCELAS DE CONTRIB
        # CARTA PREPOSTO
        # ATESTADO DEMISSIONAL
        print(f'{nome} foi desligado(a) em {data} com rescisão do tipo: Desligametno de Func com aviso.')

    desligamento = {
        1: desligar_estag,
        2: desligar_func_apedido_com_aviso,
        3: desligar_func_apedido_sem_aviso,
        4: desligar_func_por_acordo,
        5: desligar_func_sem_aviso,
        6: desligar_func_com_aviso
    }
    if tipo in desligamento:
        desligamento[tipo]()

    pessoa.desligamento = data
    session.commit()
    tkinter.messagebox.showinfo(
        title='Desligamento ok!',
        message='Desligamento registrado com sucesso!'
    )


def emitir_contracheque():
# code to automate the process of creation of documents
    wb = l_w('Contracheque.xlsx')
    sh = wb['Planilha1']
    x = 2
    while x <= len(sh['A']):
        competencia = str(sh[f'A{x}'].value).replace('/', '')
        pagamento = str(sh[f'B{x}'].value).replace('/', '')
        de = str(sh[f'C{x}'].value)
        ate = str(sh[f'D{x}'].value)
        caminho = f'C:\\Users\\RH\\PycharmProjects\\AutomacaoCia\\Emissao de contracheques\\Contracheque {competencia}.pdf'
        pa.click(-816, 515), t.sleep(0.5), pa.write(competencia), pa.press('tab'), pa.write(pagamento), pa.press('tab')
        pa.write(de), pa.press('tab'), pa.write(ate), pa.click(-787, 731), t.sleep(4)
        pa.hotkey('ctrl', 's'), pa.write(caminho), t.sleep(0.5), pa.press('enter'), t.sleep(0.5), pa.click(-33, 132)
        x += 1


def gerar_excel_from_ponto_secullum() -> None:
    """
    This function generates excel files from data extracted from external aplication.
    Its generetaes one excel file per employee with schedules information.
    :return: one .xlsx file per employee
    """
    locale.setlocale(locale.LC_ALL, 'pt_pt.UTF-8')
    data_inicial = dt.strptime('01/10/2022', '%d/%m/%Y')
    data_final = dt.strptime('13/11/2022', '%d/%m/%Y')

    # Ler planilha geral
    geral = pd.read_excel('../Ponto/xls/zzPonto Geral.xls')
    geral = geral.rename(
        columns={'CARTÃO PONTO': 'Dia', 'Unnamed: 1': 0, 'Unnamed: 2': 1, 'Unnamed: 3': 2, 'Unnamed: 4': 3,
                 'Unnamed: 5': 4,
                 'Unnamed: 6': 5})
    geral = geral.drop(['Unnamed: 7', 'Unnamed: 8'], axis=1)
    geral = geral[geral.Dia.notnull()]

    # Pegar index onde aparece 'Nome'
    linhasNomes = geral.index[geral['Dia'].str.contains('Nome')]

    # salvar plan com nome do funcionário pasta ponto (dentro da pasta automação)
    for linha in linhasNomes:
        geral = geral.rename(
            columns={'CARTÃO PONTO': 'Dia', 'Unnamed: 1': 0, 'Unnamed: 2': 1, 'Unnamed: 3': 2, 'Unnamed: 4': 3,
                     'Unnamed: 5': 4, 'Unnamed: 6': 5})
        geral = geral[geral.Dia.notnull()]
        geral = geral[geral['Dia'].str.contains(' - ') | geral['Dia'].str.contains('Nome')]
        geral2 = geral.loc[linha:(linha + (linhasNomes[1] - linhasNomes[0] - 1))] \
            .to_excel(f'../Ponto/xls/{geral[0][linha]}.xlsx')

        # Verifica a hora certa na planilha zzBase.xlsx
        wb = l_w(f'../Ponto/xls/zzBase.xlsx')
        ws = wb.active
        for row in ws.rows:
            for cell in row:
                if cell.value == f'{geral[0][linha]}':
                    if ws.cell(row=cell.row, column=5).value is None:
                        ent1 = dt.strptime('00:00:00', '%H:%M:%S')
                    else:
                        ent1 = dt.strptime(str(ws.cell(row=cell.row, column=5).value), '%H:%M:%S')
                    if ws.cell(row=cell.row, column=6).value is None:
                        sai1 = dt.strptime('00:00:00', '%H:%M:%S')
                    else:
                        sai1 = dt.strptime(str(ws.cell(row=cell.row, column=6).value), '%H:%M:%S')
                    if ws.cell(row=cell.row, column=7).value is None:
                        ent2 = dt.strptime('00:00:00', '%H:%M:%S')
                    else:
                        ent2 = dt.strptime(str(ws.cell(row=cell.row, column=7).value), '%H:%M:%S')
                    if ws.cell(row=cell.row, column=8).value is None:
                        sai2 = dt.strptime('00:00:00', '%H:%M:%S')
                    else:
                        sai2 = dt.strptime(str(ws.cell(row=cell.row, column=8).value), '%H:%M:%S')
                    if ws.cell(row=cell.row, column=9).value is None:
                        ent3 = dt.strptime('00:00:00', '%H:%M:%S')
                    else:
                        ent3 = dt.strptime(str(ws.cell(row=cell.row, column=9).value), '%H:%M:%S')
                    if ws.cell(row=cell.row, column=10).value is None:
                        sai3 = dt.strptime('00:00:00', '%H:%M:%S')
                    else:
                        sai3 = dt.strptime(str(ws.cell(row=cell.row, column=10).value), '%H:%M:%S')
        # relacionar dia da semana
        entradacerta = td(hours=ent1.hour, minutes=ent1.minute, seconds=ent1.second)
        saidacerta = td(hours=sai1.hour, minutes=sai1.minute, seconds=sai1.second)
        entradacerta2 = td(hours=ent2.hour, minutes=ent2.minute, seconds=ent2.second)
        saidacerta2 = td(hours=sai2.hour, minutes=sai2.minute, seconds=sai2.second)
        entradacerta3 = td(hours=ent3.hour, minutes=ent3.minute, seconds=ent3.second)
        saidacerta3 = td(hours=sai3.hour, minutes=sai3.minute, seconds=sai3.second)
        regra = td(hours=0, minutes=10, seconds=0)

        # Salvar planilha adicionando colunas de diferenças
        plan = pd.read_excel(f'../Ponto/xls/{geral[0][linha]}.xlsx')
        plan = plan[plan['Dia'].str.contains(' - ')]
        plan = plan.drop(['Unnamed: 0'], axis=1)
        plan = plan.rename(columns={'Dia': 'Data'})
        plan[0] = pd.to_timedelta(plan[0].astype(str))
        plan[1] = pd.to_timedelta(plan[1].astype(str))
        plan[2] = pd.to_timedelta(plan[2].astype(str))
        plan[3] = pd.to_timedelta(plan[3].astype(str))
        plan[4] = pd.to_timedelta(plan[4].astype(str))
        plan[5] = pd.to_timedelta(plan[5].astype(str))
        plan['Entradacerta'] = entradacerta
        plan['Saídacerta'] = saidacerta
        plan['DifEntr'] = abs(plan[0] - plan['Entradacerta'])
        plan['DifSaida'] = abs(plan[1] - plan['Saídacerta'])
        dif = plan
        dif = dif[['Data', 'DifEntr', 'DifSaida']]
        dif = dif.loc[(dif['DifEntr'] >= regra) | (dif['DifEntr'] + dif['DifSaida'] >= regra)]
        dif = dif.astype(str)
        dif['DifEntr'] = dif['DifEntr'].map(
            lambda x: dt.strftime(dt.strptime(str(x).replace('0 days ', '').replace('NaT', '00:00:00'), '%H:%M:%S'),
                      '%H hora(s) e %M minutos')
        )
        dif['DifSaida'] = dif['DifSaida'].map(
            lambda x: dt.strftime(dt.strptime(str(x).replace('0 days ', '').replace('NaT', '00:00:00'),
                                                          '%H:%M:%S'), '%H hora(s) e %M minutos')
        )
        dif['DifEntr'] = dif['DifEntr'].map(
            lambda x: str(x).replace('00 hora(s) e 00 minutos', '-').replace('00 hora(s) e ', '')
        )
        dif['DifSaida'] = dif['DifSaida'].map(
            lambda x: str(x).replace('00 hora(s) e 00 minutos', '-').replace('00 hora(s) e ', '')
        )
        plan = plan[['Data', 0, 1, 2, 3]]
        plan = plan.merge(dif, on='Data', how='outer')
        plan = plan.astype(str)
        # plan = plan.fillna('-')
        # plan['Data'] = plan['Data'].map(lambda x: x.rstrip('- qua qui sex sab ter seg sá dom'))
        plan[0] = plan[0].map(lambda x: x.lstrip('0 days 00:'))
        plan[1] = plan[1].map(lambda x: x.lstrip('0 days 00:'))
        plan[2] = plan[2].map(lambda x: x.lstrip('0 days 00:'))
        plan[3] = plan[3].map(lambda x: x.lstrip('0 days 00:'))
        plan = plan.replace('NaT', '-').replace('nan', '-')
        plan = plan.rename(columns={'Data': 'Dia', 0: 'Entrada1', 1: 'Saída 1', 2: 'Entrada 2', 3: 'Saída 2', 'DifEntr':
            'Diferença Entrada', 'DifSaida': 'Diferença Saída 1'})
        plan = plan[plan.Entrada1 != '-']
        plan = plan.rename(columns={'Entrada1': 'Entrada 1'})
        plan = plan.astype(str)
        plan['Dia'] = plan['Dia'].map(lambda x: dt.strptime(x, '%d/%m/%y - %a'))
        plan = plan[plan.Dia >= data_inicial]
        plan = plan[plan.Dia <= data_final]
        plan['Dia'] = plan['Dia'].map(lambda x: dt.strftime(x, '%d/%m/%y - %a'))
        plan = plan.rename(columns={'Dia': 'Data'})
        plan = plan.set_index(['Data'])
        plan = plan.to_excel(f'../Ponto/xls/{geral[0][linha]}.xlsx')
        func = l_w(f'../Ponto/xls/{geral[0][linha]}.xlsx', read_only=False)
        sh = func['Sheet1']
        sh.column_dimensions['A'].width = 15
        sh.column_dimensions['F'].width = 22
        sh.column_dimensions['G'].width = 22
        sh.column_dimensions['H'].width = 22
        sh.column_dimensions['I'].width = 22
        sh.column_dimensions['J'].width = 22
        sh.column_dimensions['K'].width = 22
        func.save(f'../Ponto/xls/{geral[0][linha]}.xlsx')


def gerar_relatorios_de_atrasos_estagiarios():
    # Ler arquivo txt dos registror rejeitados
    geral = pd.read_csv('Rejeitados.txt', sep=' ', header=None, encoding='iso8859-1')
    geral = geral.rename(columns={0: 'matricula', 16: 'data', 17: 'dia', 18: 'hora'})
    geral = geral[geral.dia != 'Batida']
    geral = geral[geral.matricula < 9999]
    mat = []
    mat_unicas = []
    for matricula in geral['matricula']:
        mat.append(matricula)
        mat_unicas = list(set(mat))
    for matr in mat_unicas:
        geral2 = geral[geral.matricula == matr]
        geral2 = geral2.set_index('matricula')
        geral2 = geral2.dropna(axis='columns')
        geral2 = geral2.drop(geral2.iloc[:, [2, 3, 4, 5]], axis=1)
        geral2['dia'] = pd.to_datetime(geral2['dia'], format='%d/%m/%Y')
        geral2['dia'] = geral2['dia'].apply(lambda y: dt.strftime(y, '%d/%m/%Y'))
        geral2['hora'] = geral2['hora'].apply(lambda x: f'{x}:00')
        geral2['hora'] = pd.to_timedelta(geral2['hora'])
        # geral2 = geral2.set_index('dia')
        # geral2 = geral2.groupby('dia')
        geral2.to_excel(f'../Ponto/xls/Ponto Estágio - {matr}.xlsx')
        wb = l_w(f'../Ponto/xls/Ponto Estágio - {matr}.xlsx')
        sh = wb['Sheet1']
        sh['A1'].value = 'Matrícula'
        sh['B1'].value = 'Data'
        sh['C1'].value = 'Entrada 1'
        sh['D1'].value = 'Saída 1'
        sh['E1'].value = 'Entrada 2'
        sh['F1'].value = 'Saída 2'
        sh['G1'].value = 'Entrada 3'
        sh['H1'].value = 'Saída 3'
        x = 2
        for row in sh:
            if sh[f'B{x}'].value == sh[f'B{x - 1}'].value:
                if sh[f'D{x - 1}'].value is None:
                    sh[f'D{x - 1}'].value = sh[f'C{x}'].value
                    sh.delete_rows(x, 1)
                else:
                    if sh[f'E{x - 1}'].value is None:
                        sh[f'E{x - 1}'].value = sh[f'C{x}'].value
                        sh.delete_rows(x, 1)
                    else:
                        if sh[f'F{x - 1}'].value is None:
                            sh[f'F{x - 1}'].value = sh[f'C{x}'].value
                            sh.delete_rows(x, 1)
                        else:
                            if sh[f'G{x - 1}'].value is None:
                                sh[f'G{x - 1}'].value = sh[f'C{x}'].value
                                sh.delete_rows(x, 1)
                            else:
                                if sh[f'H{x - 1}'].value is None:
                                    sh[f'H{x - 1}'].value = sh[f'C{x}'].value
                                    sh.delete_rows(x, 1)
        for row in sh:
            if sh[f'B{x}'].value == sh[f'B{x - 1}'].value:
                if sh[f'D{x - 1}'].value is None:
                    sh[f'D{x - 1}'].value = sh[f'C{x}'].value
                    sh.delete_rows(x, 1)
                else:
                    if sh[f'E{x - 1}'].value is None:
                        sh[f'E{x - 1}'].value = sh[f'C{x}'].value
                        sh.delete_rows(x, 1)
                    else:
                        if sh[f'F{x - 1}'].value is None:
                            sh[f'F{x - 1}'].value = sh[f'C{x}'].value
                            sh.delete_rows(x, 1)
                        else:
                            if sh[f'G{x - 1}'].value is None:
                                sh[f'G{x - 1}'].value = sh[f'C{x}'].value
                                sh.delete_rows(x, 1)
                            else:
                                if sh[f'H{x - 1}'].value is None:
                                    sh[f'H{x - 1}'].value = sh[f'C{x}'].value
                                    sh.delete_rows(x, 1)
            for row in sh:
                if sh[f'B{x}'].value == sh[f'B{x - 1}'].value:
                    if sh[f'D{x - 1}'].value is None:
                        sh[f'D{x - 1}'].value = sh[f'C{x}'].value
                        sh.delete_rows(x, 1)
                    else:
                        if sh[f'E{x - 1}'].value is None:
                            sh[f'E{x - 1}'].value = sh[f'C{x}'].value
                            sh.delete_rows(x, 1)
                        else:
                            if sh[f'F{x - 1}'].value is None:
                                sh[f'F{x - 1}'].value = sh[f'C{x}'].value
                                sh.delete_rows(x, 1)
                            else:
                                if sh[f'G{x - 1}'].value is None:
                                    sh[f'G{x - 1}'].value = sh[f'C{x}'].value
                                    sh.delete_rows(x, 1)
                                else:
                                    if sh[f'H{x - 1}'].value is None:
                                        sh[f'H{x - 1}'].value = sh[f'C{x}'].value
                                        sh.delete_rows(x, 1)
            x += 1

        estilo_data = NamedStyle(name='data', number_format='DD/MM/YYYY')
        estilo_hora = NamedStyle(name='hora', number_format='HH:MM:SS')
        for cell in sh['B']:
            sh[f'B{int(cell.row)}'].style = estilo_data
        for item in sh['C']:
            sh[f'C{int(item.row)}'].style = estilo_hora
        for item in sh['D']:
            sh[f'D{int(item.row)}'].style = estilo_hora
        for item in sh['E']:
            sh[f'E{int(item.row)}'].style = estilo_hora
        for item in sh['F']:
            sh[f'F{int(item.row)}'].style = estilo_hora
        for item in sh['G']:
            sh[f'G{int(item.row)}'].style = estilo_hora
        for item in sh['H']:
            sh[f'H{int(item.row)}'].style = estilo_hora
        sh.column_dimensions['A'].width = 11
        sh.column_dimensions['B'].width = 11
        sh.column_dimensions['C'].width = 11
        sh.column_dimensions['D'].width = 11
        sh.column_dimensions['E'].width = 11
        sh.column_dimensions['F'].width = 11
        sh.column_dimensions['G'].width = 11
        sh.column_dimensions['H'].width = 11
        wb2 = l_w(f'../Ponto/xls/zzBase.xlsx')
        ws2 = wb2.active
        for row in ws2.rows:
            for cell in row:
                if cell.value == matr:
                    nome = ws2.cell(row=cell.row, column=3).value
                    wb.save(f'../Ponto/xls/Ponto Estágio - {nome}.xlsx')
                    os.remove(f'../Ponto/xls/Ponto Estágio - {matr}.xlsx')
        plan = l_w(f'../Ponto/xls/Ponto Estágio - {nome}.xlsx')
        splan = plan.active
        # Verifica a hora certa na planilha zzBase.xlsx
        wb = l_w(f'../Ponto/xls/zzBase.xlsx')
        ws = wb.active
        for row in ws.rows:
            for cell in row:
                if cell.value == f'{nome}':
                    for rowsplan in splan.rows:
                        for cellplan in rowsplan:
                            if str(splan.cell(row=cellplan.row, column=2).value) == 'Data':
                                pass
                            # se o dia no ponto for segunda
                            if dt.strptime(str(splan.cell(row=cellplan.row, column=2).value),
                                           '%d/%m/%Y').isoweekday() == 1:
                                pass
                            # se o dia no ponto for terça
                            if dt.strptime(str(splan.cell(row=cellplan.row, column=2).value),
                                           '%d/%m/%Y').isoweekday() == 2:
                                pass
                            # se o dia no ponto for quarta
                            if dt.strptime(str(splan.cell(row=cellplan.row, column=2).value),
                                           '%d/%m/%Y').isoweekday() == 3:
                                pass
                            # se o dia no ponto for quinta
                            if dt.strptime(str(splan.cell(row=cellplan.row, column=2).value),
                                           '%d/%m/%Y').isoweekday() == 4:
                                pass
                            # se o dia no ponto for sexta
                            if dt.strptime(str(splan.cell(row=cellplan.row, column=2).value),
                                           '%d/%m/%Y').isoweekday() == 5:
                                pass
                            # se o dia no ponto for sábado
                            if dt.strptime(str(splan.cell(row=cellplan.row, column=2).value),
                                           '%d/%m/%Y').isoweekday() == 6:
                                pass
                            # se o dia no ponto for domingo
                            if dt.strptime(str(splan.cell(row=cellplan.row, column=2).value),
                                           '%d/%m/%Y').isoweekday() == 7:
                                pass
                    if ws.cell(row=cell.row, column=5).value is None:
                        ent1 = dt.strptime('00:00:00', '%H:%M:%S')
                    else:
                        ent1 = dt.strptime(str(ws.cell(row=cell.row, column=5).value), '%H:%M:%S')
                    if ws.cell(row=cell.row, column=6).value is None:
                        sai1 = dt.strptime('00:00:00', '%H:%M:%S')
                    else:
                        sai1 = dt.strptime(str(ws.cell(row=cell.row, column=6).value), '%H:%M:%S')
                    if ws.cell(row=cell.row, column=7).value is None:
                        ent2 = dt.strptime('00:00:00', '%H:%M:%S')
                    else:
                        ent2 = dt.strptime(str(ws.cell(row=cell.row, column=7).value), '%H:%M:%S')
                    if ws.cell(row=cell.row, column=8).value is None:
                        sai2 = dt.strptime('00:00:00', '%H:%M:%S')
                    else:
                        sai2 = dt.strptime(str(ws.cell(row=cell.row, column=8).value), '%H:%M:%S')
                    if ws.cell(row=cell.row, column=9).value is None:
                        ent3 = dt.strptime('00:00:00', '%H:%M:%S')
                    else:
                        ent3 = dt.strptime(str(ws.cell(row=cell.row, column=9).value), '%H:%M:%S')
                    if ws.cell(row=cell.row, column=10).value is None:
                        sai3 = dt.strptime('00:00:00', '%H:%M:%S')
                    else:
                        sai3 = dt.strptime(str(ws.cell(row=cell.row, column=10).value), '%H:%M:%S')
                # relacionar dia da semana
                entradacerta = td(hours=ent1.hour, minutes=ent1.minute, seconds=ent1.second)
                saidacerta = td(hours=sai1.hour, minutes=sai1.minute, seconds=sai1.second)
                entradacerta2 = td(hours=ent2.hour, minutes=ent2.minute, seconds=ent2.second)
                saidacerta2 = td(hours=sai2.hour, minutes=sai2.minute, seconds=sai2.second)
                entradacerta3 = td(hours=ent3.hour, minutes=ent3.minute, seconds=ent3.second)
                saidacerta3 = td(hours=sai3.hour, minutes=sai3.minute, seconds=sai3.second)
                regra = td(hours=0, minutes=10, seconds=0)
    # # procurar na planilha base.xlsx nome e-mail matricula no ponto e horarios de entrada e saida


def gerar_excel_ponto_estagiarios():
    # Ler arquivo txt dos registror rejeitados
    geral = pd.read_csv(
        r'C:\Users\Felipe Rodrigues\Desktop\Relatorios Ponto\Rej - 16-03-2023.txt', sep=' ', header=None,
        encoding='iso8859-1'
    )
    geral = geral.rename(columns={0: 'matricula', 16: 'data', 17: 'dia', 18: 'hora'})
    geral = geral[geral.dia != 'Batida']
    geral = geral[geral.matricula < 9999]
    mat = []
    mat_unicas = []
    for matricula in geral['matricula']:
        mat.append(matricula)
        mat_unicas = list(set(mat))
    for matr in mat_unicas:
        geral2 = geral[geral.matricula == matr]
        geral2 = geral2.set_index('matricula')
        geral2 = geral2.dropna(axis='columns')
        # print(geral2)
        geral2 = geral2.drop(geral2.iloc[:, [2, 3, 4]], axis=1)
        geral2['dia'] = pd.to_datetime(geral2['dia'], format='%d/%m/%Y')
        geral2['dia'] = geral2['dia'].apply(lambda y: dt.strftime(y, '%d/%m/%Y'))
        geral2['hora'] = geral2['hora'].apply(lambda x: f'{x}:00')
        geral2['hora'] = pd.to_timedelta(geral2['hora'])
        # geral2 = geral2.set_index('dia')
        # geral2 = geral2.groupby('dia')
        geral2.to_excel(f'../Ponto/xls/Ponto Estágio - {matr}.xlsx')
        wb = l_w(f'../Ponto/xls/Ponto Estágio - {matr}.xlsx')
        sh = wb['Sheet1']
        sh['A1'].value = 'Matrícula'
        sh['B1'].value = 'Data'
        sh['C1'].value = 'Entrada 1'
        sh['D1'].value = 'Saída 1'
        sh['E1'].value = 'Entrada 2'
        sh['F1'].value = 'Saída 2'
        sh['G1'].value = 'Entrada 3'
        sh['H1'].value = 'Saída 3'
        x = 2
        for row in sh:
            if sh[f'B{x}'].value == sh[f'B{x - 1}'].value:
                if sh[f'D{x - 1}'].value is None:
                    sh[f'D{x - 1}'].value = sh[f'C{x}'].value
                    sh.delete_rows(x, 1)
                else:
                    if sh[f'E{x - 1}'].value is None:
                        sh[f'E{x - 1}'].value = sh[f'C{x}'].value
                        sh.delete_rows(x, 1)
                    else:
                        if sh[f'F{x - 1}'].value is None:
                            sh[f'F{x - 1}'].value = sh[f'C{x}'].value
                            sh.delete_rows(x, 1)
                        else:
                            if sh[f'G{x - 1}'].value is None:
                                sh[f'G{x - 1}'].value = sh[f'C{x}'].value
                                sh.delete_rows(x, 1)
                            else:
                                if sh[f'H{x - 1}'].value is None:
                                    sh[f'H{x - 1}'].value = sh[f'C{x}'].value
                                    sh.delete_rows(x, 1)
        for row in sh:
            if sh[f'B{x}'].value == sh[f'B{x - 1}'].value:
                if sh[f'D{x - 1}'].value is None:
                    sh[f'D{x - 1}'].value = sh[f'C{x}'].value
                    sh.delete_rows(x, 1)
                else:
                    if sh[f'E{x - 1}'].value is None:
                        sh[f'E{x - 1}'].value = sh[f'C{x}'].value
                        sh.delete_rows(x, 1)
                    else:
                        if sh[f'F{x - 1}'].value is None:
                            sh[f'F{x - 1}'].value = sh[f'C{x}'].value
                            sh.delete_rows(x, 1)
                        else:
                            if sh[f'G{x - 1}'].value is None:
                                sh[f'G{x - 1}'].value = sh[f'C{x}'].value
                                sh.delete_rows(x, 1)
                            else:
                                if sh[f'H{x - 1}'].value is None:
                                    sh[f'H{x - 1}'].value = sh[f'C{x}'].value
                                    sh.delete_rows(x, 1)
            for row in sh:
                if sh[f'B{x}'].value == sh[f'B{x - 1}'].value:
                    if sh[f'D{x - 1}'].value is None:
                        sh[f'D{x - 1}'].value = sh[f'C{x}'].value
                        sh.delete_rows(x, 1)
                    else:
                        if sh[f'E{x - 1}'].value is None:
                            sh[f'E{x - 1}'].value = sh[f'C{x}'].value
                            sh.delete_rows(x, 1)
                        else:
                            if sh[f'F{x - 1}'].value is None:
                                sh[f'F{x - 1}'].value = sh[f'C{x}'].value
                                sh.delete_rows(x, 1)
                            else:
                                if sh[f'G{x - 1}'].value is None:
                                    sh[f'G{x - 1}'].value = sh[f'C{x}'].value
                                    sh.delete_rows(x, 1)
                                else:
                                    if sh[f'H{x - 1}'].value is None:
                                        sh[f'H{x - 1}'].value = sh[f'C{x}'].value
                                        sh.delete_rows(x, 1)
            x += 1

        estilo_data = NamedStyle(name='data', number_format='DD/MM/YYYY')
        estilo_hora = NamedStyle(name='hora', number_format='HH:MM:SS')
        for cell in sh['B']:
            sh[f'B{int(cell.row)}'].style = estilo_data
        for item in sh['C']:
            sh[f'C{int(item.row)}'].style = estilo_hora
        for item in sh['D']:
            sh[f'D{int(item.row)}'].style = estilo_hora
        for item in sh['E']:
            sh[f'E{int(item.row)}'].style = estilo_hora
        for item in sh['F']:
            sh[f'F{int(item.row)}'].style = estilo_hora
        for item in sh['G']:
            sh[f'G{int(item.row)}'].style = estilo_hora
        for item in sh['H']:
            sh[f'H{int(item.row)}'].style = estilo_hora
        sh.column_dimensions['A'].width = 11
        sh.column_dimensions['B'].width = 11
        sh.column_dimensions['C'].width = 11
        sh.column_dimensions['D'].width = 11
        sh.column_dimensions['E'].width = 11
        sh.column_dimensions['F'].width = 11
        sh.column_dimensions['G'].width = 11
        sh.column_dimensions['H'].width = 11
        wb2 = l_w(f'../Ponto/xls/zzBase.xlsx')
        ws2 = wb2.active
        for row in ws2.rows:
            for cell in row:
                if cell.value == matr:
                    nome = ws2.cell(row=cell.row, column=3).value
                    linha = int((len(sh['A']) / 2) + 1)
                    wb.save(f'../Ponto/xls/Ponto Estágio - {nome}.xlsx')
                    os.remove(f'../Ponto/xls/Ponto Estágio - {matr}.xlsx')
                    # t.sleep(2)
                    # inserir código de envio de email aqui
                    excel = client.Dispatch('Excel.Application')
                    plan = excel.Workbooks.Open(
                        r'C:\Users\Felipe Rodrigues\PycharmProjects\AutomacaoCia\Ponto\xls\Ponto Estágio - {}.xlsx'
                        .format(nome)
                    )
                    folha = plan.Sheets['Sheet1']
                    excel.visible = 0
                    copyrange = folha.Range(f'B1:F{linha}')
                    copyrange.CopyPicture(Format=2)
                    ImageGrab.grabclipboard().save(f'Ponto {str(nome).title().split(" ")[0]}.png')
                    excel.Quit()
                    html_body = f'''
                    <p>Olá {str(nome).title().split(" ")[0]},</p>
                    <p> Segue em anexo seu relatório de ponto do mês Fevereiro/2023.</p><br>
                    Atenciosamente,<br>
                    Felipe Rodrigues,<br>
                    '''
                    outlook = client.Dispatch('Outlook.Application')
                    message = outlook.CreateItem(0)
                    message.To = 'felipe.rodrigs09@gmail.com'
                    message.Subject = 'Relatório de Ponto'
                    message.HTMLBody = html_body
                    message.Attachments.Add(
                        f'C:\\Users\\{os.getlogin()}\\PycharmProjects\\AutomacaoCia\\Ponto\\'
                        f'Ponto {str(nome).title().split(" ")[0]}.png'
                    )
                    message.Send()
    # # procurar na planilha base.xlsx nome e-mail matricula no ponto e horarios de entrada e saida


def cadastrar_funcionario_no_secullum():
    wb = l_w('PlanPonto.xlsx')
    sh = wb['Planilha1']
    x = 2
    while x <= len(sh['A']):
        matricula = str(sh[f'A{x}'].value)
        nome = str(sh[f'B{x}'].value)
        pis = str(sh[f'C{x}'].value)
        horario = '1'
        funcao = str(sh[f'D{x}'].value)
        depto = str(sh[f'E{x}'].value)
        admiss = str(sh[f'F{x}'].value)
        # # clicar incluir
        pa.click(-1507, 147), t.sleep(3)
        pa.write(matricula), pa.press('tab')
        pp.copy(nome), pa.hotkey('ctrl', 'v'), pa.press('tab')
        pa.write(pis), pa.press('tab', 4)
        pa.write(horario), pa.press('tab')
        pp.copy(funcao), pa.hotkey('ctrl', 'v'), pa.press('tab')
        pp.copy(depto), pa.hotkey('ctrl', 'v'), pa.press('tab')
        pa.write(admiss)
        # click concluir
        pa.click(-1492, 617)
        t.sleep(2)
        x += 1


def gerar_relatorios_ponto_pdf(arq: str, datai: str, dataf: str):
    """
    Generates working time reports in .pdf trought analysis of file .AFD.
    :param arquivo: path of .AFD file.
    :param datai: Beginning date.
    :param dataf: End date.
    :return: Worker reports.
    """
    # edit arq
    edic1 = open(arq, 'r')
    linhas = edic1.readlines()
    novalinha = [x for x in linhas if len(x) <= 39]
    out = open(arq, 'w')
    out.writelines(novalinha)
    out.close()
    arquivo = rf'C:\Users\{os.getlogin()}\PycharmProjects\AutomacaoCia\src\models\static\files\AFD.txt'

    # define excel plans to work with
    base = os.path.relpath(rf'C:\Users\{os.getlogin()}\PycharmProjects\AutomacaoCia\src\models\static\files\zzBase.xlsx')
    wb = l_w(base)
    sh = wb['Funcionários e e-mail']
    hrsflh = os.path.relpath(rf'C:\Users\{os.getlogin()}\PycharmProjects\AutomacaoCia\src\models\static\files\Hrs Folha.xlsx')
    pl = l_w(hrsflh)
    fl = pl['Planilha1']
    dataipt = datai.replace('/', '.')
    datafpt = dataf.replace('/', '.')
    dia, mes, ano = dataf.split('/')
    datainicio = dt.strptime(datai, '%d/%m/%Y')
    datafim = dt.strptime(dataf, '%d/%m/%Y')
    mesext = {'01': 'JAN', '02': 'FEV', '03': 'MAR', '04': 'ABR', '05': 'MAI', '06': 'JUN',
              '07': 'JUL', '08': 'AGO', '09': 'SET', '10': 'OUT', '11': 'NOV', '12': 'DEZ'}

    def intervalo(inicio, fim):
        for n in range(int((fim - inicio).days) + 1):
            yield [inicio + td(n), np.nan, np.nan, np.nan, np.nan, np.nan, np.nan]

    # creating dicts to store employees datas
    planbase = {}
    planmat = {}
    planfunc = {}
    plandept = {}

    x = 2
    while x <= len(sh['B']):
        planbase.update(
            {str(sh[f'B{x}'].value).replace('.', '').replace('-', ''): str(sh[f'C{x}'].value).title().strip()})
        planmat.update(
            {str(sh[f'B{x}'].value).replace('.', '').replace('-', ''): str(sh[f'A{x}'].value).title().strip()})
        planfunc.update({str(sh[f'B{x}'].value).replace('.', '').replace('-', ''): str(sh[f'D{x}'].value).strip()})
        plandept.update(
            {str(sh[f'B{x}'].value).replace('.', '').replace('-', ''): str(sh[f'E{x}'].value).title().strip()})
        x += 1

    geral = pd.read_csv(arquivo, sep=' ', header=None, encoding='iso8859-1')
    geral = geral[geral[0].str.len() <= 34]
    geral.dropna(axis=1, inplace=True)
    geral = geral.rename(columns={0: 'Dados'})

    # dividir ultimos 11 caracteres em outra col
    geral['Matricula'] = geral['Dados'].str[-11:]
    geral = geral.drop('Dados', axis=1)
    mat = []
    matriculas_unicas = []
    for item in geral['Matricula']:
        mat.append(item)
    matnum = list(map(int, mat))
    matriculas_unicas = list(set(matnum))
    for matricula in matriculas_unicas:
        base = pd.DataFrame(list(intervalo(datainicio, datafim)),
                            columns=['Data', 'Entrada 1', 'Saída 1', 'Entrada 2', 'Saída 2', 'Entrada 3', 'Saída 3'])
        base = base.set_index('Data')
        geral = pd.read_csv(arquivo, sep=' ', header=None, encoding='iso8859-1')
        geral = geral[geral[0].str.len() <= 34]
        geral.dropna(axis=1, inplace=True)
        geral = geral.rename(columns={0: 'Dados'})
        # dividir ultimos 11 caracteres em outra col
        geral['Matricula'] = geral['Dados'].str[-11:]
        # dividir caracteres da data em outra col
        geral['Data'] = geral['Dados'].str[-24:-16]
        geral['Data'] = pd.to_datetime(geral['Data'], format='%d%m%Y')
        # dividir caracteres da hr em outra col
        geral['Hora'] = geral['Dados'].str[-16:-12]
        geral['Hora'] = pd.to_datetime(geral['Hora'], format='%H%M')

        geral1 = pd.DataFrame(geral)

        geral = geral.loc[(geral['Data'] < datafim) & (geral['Data'] > datainicio)]
        geral1 = geral1.loc[(geral1['Data'] < datafim) & (geral1['Data'] > datainicio)]

        geral = geral.loc[geral['Matricula'] == str(matricula).zfill(11)]
        geral1 = geral1.loc[geral1['Matricula'] == str(matricula).zfill(11)]

        geral['Hora'] = geral['Hora'].apply(lambda k: dt.strftime(k, '%H:%M'))
        geral1['Hora'] = geral1['Hora'].apply(lambda l: dt.strftime(l, '%H:%M:%S'))

        geral = geral.drop('Dados', axis=1)
        geral1 = geral1.drop('Dados', axis=1)

        geral = geral.drop('Matricula', axis=1)
        geral1 = geral1.drop('Matricula', axis=1)

        geral = geral.reset_index(drop=True)
        geral = geral.pivot_table(index='Data', columns=geral.groupby('Data').cumcount() + 1, values='Hora',
                                  aggfunc='first')
        geral = geral.reset_index(level=[0])
        geral = geral.rename(
            columns={0: 'Data', 1: 'Entrada 1', 2: 'Saída 1', 3: 'Entrada 2', 4: 'Saída 2', 5: 'Entrada 3',
                     6: 'Saída 3'})

        geral1 = geral1.reset_index(drop=True)
        geral1 = geral1.pivot_table(index='Data', columns=geral1.groupby('Data').cumcount() + 1, values='Hora',
                                    aggfunc='first')
        geral1 = geral1.reset_index(level=[0])
        geral1 = geral1.rename(
            columns={0: 'Data', 1: 'Entrada 1', 2: 'Saída 1', 3: 'Entrada 2', 4: 'Saída 2', 5: 'Entrada 3',
                     6: 'Saída 3'})

        geral1['Entrada 1a'] = geral1['Entrada 1'].apply(lambda z: pd.to_timedelta(str(z)))
        geral1['Saída 1a'] = geral1['Saída 1'].apply(lambda z: pd.to_timedelta(str(z)))
        geral1['Tot1'] = geral1['Saída 1a'] - geral1['Entrada 1a']
        total_horas = round(geral1['Tot1'].sum().total_seconds() / 3600, 2)

        if geral1.shape[1] > 6:
            if geral1.shape[1] < 10:
                geral1['Entrada 2a'] = geral1['Entrada 2'].apply(lambda z: pd.to_timedelta(str(z)))
                geral1['Saída 2a'] = geral1['Saída 2'].apply(lambda z: pd.to_timedelta(str(z)))
                geral1['Tot2'] = geral1['Saída 2a'] - geral1['Entrada 2a']
                geral1['Soma'] = geral1['Tot1'] + geral1['Tot2']
                total_horas = round(geral1['Soma'].sum().total_seconds() / 3600, 2)
                geral1 = geral1.drop('Entrada 2a', axis=1)
                geral1 = geral1.drop('Saída 2a', axis=1)
                geral1 = geral1.drop('Tot2', axis=1)
                geral1 = geral1.drop('Soma', axis=1)

        if geral1.shape[1] > 8:
            geral1['Entrada 2a'] = geral1['Entrada 2'].apply(lambda z: pd.to_timedelta(str(z)))
            geral1['Saída 2a'] = geral1['Saída 2'].apply(lambda z: pd.to_timedelta(str(z)))
            geral1['Entrada 3a'] = geral1['Entrada 3'].apply(lambda z: pd.to_timedelta(str(z)))
            geral1['Saída 3a'] = geral1['Saída 3'].apply(lambda z: pd.to_timedelta(str(z)))
            geral1['Tot1'] = geral1['Saída 1a'] - geral1['Entrada 1a']
            geral1['Tot2'] = geral1['Saída 2a'] - geral1['Entrada 2a']
            geral1['Tot3'] = geral1['Saída 3a'] - geral1['Entrada 3a']
            geral1['Soma'] = geral1['Tot1'] + geral1['Tot2'] + geral1['Tot3']
            total_horas = round(geral1['Soma'].sum().total_seconds() / 3600, 2)
            geral1 = geral1.drop('Entrada 2a', axis=1)
            geral1 = geral1.drop('Saída 2a', axis=1)
            geral1 = geral1.drop('Entrada 3a', axis=1)
            geral1 = geral1.drop('Saída 3a', axis=1)
            geral1 = geral1.drop('Tot2', axis=1)
            geral1 = geral1.drop('Tot3', axis=1)
            geral1 = geral1.drop('Soma', axis=1)

        geral1 = geral1.drop('Entrada 1a', axis=1)
        geral1 = geral1.drop('Saída 1a', axis=1)
        geral1 = geral1.drop('Tot1', axis=1)

        geral = geral.set_index('Data')
        base = base.combine_first(geral)
        base = base.reset_index(level=[0])
        base['Data'] = base['Data'].apply(lambda h: dt.strftime(h, '%d/%m/%Y - %a'))
        base = base[['Data', 'Entrada 1', 'Saída 1', 'Entrada 2', 'Saída 2', 'Entrada 3', 'Saída 3']]
        dias = base['Entrada 1'].count()
        base = base.fillna('-')
        ponto = docx.Document(rf'C:\Users\{os.getlogin()}\PycharmProjects\AutomacaoCia\src\models\static\files\modelo.docx')
        ponto.paragraphs[1].text = str(ponto.paragraphs[1].text).replace('#data1',
                                                                         dt.strftime(datainicio, '%d/%m/%Y')).replace(
            '#data2', dt.strftime(datafim, '%d/%m/%Y'))
        ponto.paragraphs[2].text = str(ponto.paragraphs[2].text).replace('#emissao',
                                                                         dt.strftime(dt.today(), '%d/%m/%Y'))
        ponto.tables[0].rows[3].cells[1].paragraphs[0].text = str(
            ponto.tables[0].rows[3].cells[1].paragraphs[0].text).replace('#nome', planbase[str(matricula)])
        ponto.tables[0].rows[4].cells[1].paragraphs[0].text = str(
            ponto.tables[0].rows[4].cells[1].paragraphs[0].text).replace('#cod', str(matricula))
        ponto.tables[0].rows[5].cells[1].paragraphs[0].text = str(
            ponto.tables[0].rows[5].cells[1].paragraphs[0].text).replace('#mat', planmat[str(matricula)])
        ponto.tables[0].rows[6].cells[1].paragraphs[0].text = str(
            ponto.tables[0].rows[6].cells[1].paragraphs[0].text).replace('#func', planfunc[str(matricula)])
        ponto.tables[0].rows[7].cells[1].paragraphs[0].text = str(
            ponto.tables[0].rows[7].cells[1].paragraphs[0].text).replace('#depto', plandept[str(matricula)])
        # ponto.tables[0].rows[5].cells[4].paragraphs[0].text = 'HH:MM'
        style = ponto.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(10)
        style8 = ponto.styles['Default']
        font = style8.font
        font.name = 'Times New Roman'
        font.size = Pt(8)

        ponto.paragraphs[0].style = ponto.styles['Normal']
        ponto.paragraphs[1].style = ponto.styles['Normal']
        ponto.paragraphs[2].style = ponto.styles['Normal']
        ponto.paragraphs[3].style = ponto.styles['Normal']
        ponto.paragraphs[4].style = ponto.styles['Normal']
        ponto.tables[0].rows[5].cells[4].paragraphs[0].style = ponto.styles['Default']
        t = ponto.add_table(base.shape[0] + 1, base.shape[1])
        t.style = 'Estilo2'
        # add the header rows.
        for j in range(base.shape[-1]):
            t.cell(0, j).text = base.columns[j]
        # add the rest of the data frame
        for i in range(base.shape[0]):
            for j in range(base.shape[-1]):
                t.cell(i + 1, j).text = str(base.values[i, j])
                t.cell(i + 1, j).paragraphs[0].alignment = 1
                t.cell(i + 1, 0).paragraphs[0].alignment = 0

        ponto.tables[1].columns[0].cells[0].width = Cm(4.2)
        ponto.tables[1].columns[1].cells[0].width = Cm(2.6)
        ponto.tables[1].columns[2].cells[0].width = Cm(2.6)
        ponto.tables[1].columns[3].cells[0].width = Cm(2.6)
        ponto.tables[1].columns[4].cells[0].width = Cm(2.6)
        ponto.tables[1].columns[5].cells[0].width = Cm(2.6)
        ponto.tables[1].columns[6].cells[0].width = Cm(2.6)
        ponto.tables[1].rows[0].cells[0].paragraphs[0].alignment = 1
        ponto.tables[1].rows[0].cells[1].paragraphs[0].alignment = 1
        ponto.tables[1].rows[0].cells[2].paragraphs[0].alignment = 1
        ponto.tables[1].rows[0].cells[3].paragraphs[0].alignment = 1
        ponto.tables[1].rows[0].cells[4].paragraphs[0].alignment = 1
        ponto.tables[1].rows[0].cells[5].paragraphs[0].alignment = 1
        ponto.tables[1].rows[0].cells[6].paragraphs[0].alignment = 1

        y = len(fl['A']) + 1
        fl[f'A{y}'].value = planbase[str(matricula)]
        fl[f'B{y}'].value = total_horas
        pl.save(rf'C:\Users\{os.getlogin()}\PycharmProjects\AutomacaoCia\src\models\static\files\Hrs Folha.xlsx')

        ponto.add_paragraph(f'Total de dias trabalhados: {dias}', 'Default').alignment = 2
        ponto.add_paragraph('', 'Normal')
        ponto.add_paragraph('', 'Normal')
        ponto.add_paragraph('______________________________________________', 'Normal').alignment = 1
        ponto.add_paragraph(f'{planbase[str(matricula)]}', 'Normal').alignment = 1
        try:
            ponto.save(rede + rf'\04 - Folha de Pgto\{ano}\{mes} - {mesext[mes]}\Relatórios de Ponto\{dataipt} a {datafpt}\Ponto {planbase[str(matricula)]}.docx')
            docx2pdf.convert(
                rede + rf'\04 - Folha de Pgto\{ano}\{mes} - {mesext[mes]}\Relatórios de Ponto\{dataipt} a {datafpt}\Ponto {planbase[str(matricula)]}.docx',
                rede + rf'\04 - Folha de Pgto\{ano}\{mes} - {mesext[mes]}\Relatórios de Ponto\{dataipt} a {datafpt}\Ponto {planbase[str(matricula)]}.pdf')
            os.remove(
                rede + rf'\04 - Folha de Pgto\{ano}\{mes} - {mesext[mes]}\Relatórios de Ponto\{dataipt} a {datafpt}\Ponto {planbase[str(matricula)]}.docx')
        except FileNotFoundError:
            os.mkdir(rede + rf'\04 - Folha de Pgto\{ano}\{mes} - {mesext[mes]}\Relatórios de Ponto')
            os.mkdir(
                rede + rf'\04 - Folha de Pgto\{ano}\{mes} - {mesext[mes]}\Relatórios de Ponto\{dataipt} a {datafpt}')
            ponto.save( rede + rf'\04 - Folha de Pgto\{ano}\{mes} - {mesext[mes]}\Relatórios de Ponto\{dataipt} a {datafpt}\Ponto {planbase[str(matricula)]}.docx')
            docx2pdf.convert(
                rede + rf'\04 - Folha de Pgto\{ano}\{mes} - {mesext[mes]}\Relatórios de Ponto\{dataipt} a {datafpt}\Ponto {planbase[str(matricula)]}.docx',
                rede + rf'\04 - Folha de Pgto\{ano}\{mes} - {mesext[mes]}\Relatórios de Ponto\{dataipt} a {datafpt}\Ponto {planbase[str(matricula)]}.pdf')
            os.remove(
                rede + rf'\04 - Folha de Pgto\{ano}\{mes} - {mesext[mes]}\Relatórios de Ponto\{dataipt} a {datafpt}\Ponto {planbase[str(matricula)]}.docx')

        # print(base)


def emitir_certificados(nome: str, data: str, horas: int, participantes: list):
    modelo = os.path.relpath(rf'C:\Users\{os.getlogin()}\PycharmProjects\AutomacaoCia\src\models\static\files\certificados\Treinamento.docx')

    def extenso(datacompleta):
        dia, mes, ano = datacompleta.split('/')
        mesext = {'01': 'janeiro', '02': 'fevereiro', '03': 'março', '04': 'abril', '05': 'maio', '06': 'junho',
                  '07': 'julho', '08': 'agosto', '09': 'setembro', '10': 'outubro', '11': 'novembro', '12': 'dezembro'}
        return f'{dia} de {mesext[mes]} de {ano}.'

    def exthoras(hr: int):
        horasext = {'1': 'uma', '2': 'duas', '3': 'três', '4': 'quatro', '5': 'cinco', '6': 'seis', '7': 'sete',
                    '8': 'oito', '9': 'nove', '10': 'dez', '11': 'onze', '12': 'doze', '13': 'treze', '14': 'quatorze',
                    '15': 'quinze'}
        return horasext[str(hr)]

    sessions = sessionmaker(bind=engine)
    session = sessions()
    pesq = session.query(Colaborador).filter_by(desligamento=None).all()
    nomes = []
    dicion = {}
    for p in pesq:
        nomes.append(str(p.nome).upper())
    #  Emitir certificado
        for pessoa in nomes:
            for item in participantes:
                dicion[pessoa] = SequenceMatcher(None, item, pessoa).ratio()
        pess = [i for i in dicion if dicion[i] == max(dicion.values())][0]
        doc = docx.Document(modelo)
        for p in doc.paragraphs:
            if '#nome' in p.text:
                inline = p.runs
                # Loop added to work with runs (strings with same style)
                for i in range(len(inline)):
                    if '#nome' in inline[i].text:
                        text = inline[i].text.replace('#nome', pess.title())
                        inline[i].text = text
            if '#treinamento' in p.text:
                inline = p.runs
                # Loop added to work with runs (strings with same style)
                for i in range(len(inline)):
                    if '#treinamento' in inline[i].text:
                        text = inline[i].text.replace('#treinamento', nome.title())
                        inline[i].text = text
            if '#data' in p.text:
                inline = p.runs
                # Loop added to work with runs (strings with same style)
                for i in range(len(inline)):
                    if '#data' in inline[i].text:
                        text = inline[i].text.replace('#data', data)
                        inline[i].text = text
            if '#duracao' in p.text:
                inline = p.runs
                # Loop added to work with runs (strings with same style)
                for i in range(len(inline)):
                    if '#duracao' in inline[i].text:
                        text = inline[i].text.replace('#duracao', horas)
                        inline[i].text = text
            if '#hrsexten' in p.text:
                inline = p.runs
                # Loop added to work with runs (strings with same style)
                for i in range(len(inline)):
                    if '#hrsexten' in inline[i].text:
                        text = inline[i].text.replace('#hrsexten', exthoras(horas))
                        inline[i].text = text
            if '#extens' in p.text:
                inline = p.runs
                # Loop added to work with runs (strings with same style)
                for i in range(len(inline)):
                    if '#extens' in inline[i].text:
                        text = inline[i].text.replace('#extens', extenso(data))
                        inline[i].text = text
        caminho = os.path.relpath(rf'C:\Users\{os.getlogin()}\PycharmProjects\AutomacaoCia\src\models\static\files\certificados')
        doc.save(caminho + f'\\{pess} - {nome} {data.replace("/",".")}.docx')
        convert(caminho + f'\\{pess} - {nome} {data.replace("/",".")}.docx', caminho + f'\\{pess} - {nome} {data.replace("/",".")}.pdf')
        os.remove(caminho + f'\\{pess} - {nome} {data.replace("/",".")}.docx')


def gerar_recibo_uniformes(local, nome, cargo, cpf, genero, tamanho1, tamanho2=''):
    relatorio = l_w(local, read_only=False)
    estoque = relatorio['Estoque']
    entregues = relatorio['Entregues']
    lista = relatorio['Nomes']
    hoje = dt.today()
    tipo, gen = genero.split(': ')
    num, pess = nome.split(' - ')
    label, cpf_ed = cpf.split(': ')
    pessoa = pess.title()

    if tamanho2 != '':
        recibo = docx.Document('recibo_uniforme.docx')
        recibo.paragraphs[11].text = str(recibo.paragraphs[11].text)\
            .replace('#nome', pessoa).replace('#num_cpf', cpf_ed).replace('#tam', tamanho1+' e '+tamanho2).replace(
            '#genero', str(gen).lower())
        recibo.paragraphs[19].text = str(recibo.paragraphs[19].text)\
            .replace('#data', dt.strftime(hoje, '%d/%m/%Y'))
        recibo.paragraphs[24].text = str(recibo.paragraphs[24].text).replace('#nome', pessoa)
        recibo.paragraphs[25].text = str(recibo.paragraphs[25].text).replace('#cargo', cargo)
        recibo.save(f'Recibo_alterado {pessoa}.docx')
        docx2pdf.convert(f'Recibo_alterado {pessoa}.docx', f'Recibo {pessoa}.pdf')
        # diminuir contagem de estoque
        if gen == 'Masculino':
            if tamanho1 == 'P':
                estoque['C4'].value = estoque['C4'].value - 1
            if tamanho1 == 'M':
                estoque['C5'].value = estoque['C5'].value - 1
            if tamanho1 == 'G':
                estoque['C6'].value = estoque['C6'].value - 1
            if tamanho1 == 'XG':
                estoque['C8'].value = estoque['C8'].value - 1
            if tamanho1 == 'GG':
                estoque['C7'].value = estoque['C7'].value - 1
            if tamanho1 == 'XGG':
                estoque['C9'].value = estoque['C9'].value - 1
            if tamanho2 == 'P':
                estoque['C4'].value = estoque['C4'].value - 1
            if tamanho2 == 'M':
                estoque['C5'].value = estoque['C5'].value - 1
            if tamanho2 == 'G':
                estoque['C6'].value = estoque['C6'].value - 1
            if tamanho2 == 'XG':
                estoque['C8'].value = estoque['C8'].value - 1
            if tamanho2 == 'GG':
                estoque['C7'].value = estoque['C7'].value - 1
            if tamanho2 == 'XGG':
                estoque['C9'].value = estoque['C9'].value - 1
        else:
            if tamanho1 == 'P':
                estoque['E4'].value = estoque['E4'].value - 1
            if tamanho1 == 'M':
                estoque['E5'].value = estoque['E5'].value - 1
            if tamanho1 == 'G':
                estoque['E6'].value = estoque['E6'].value - 1
            if tamanho1 == 'GG':
                estoque['E7'].value = estoque['E7'].value - 1
            if tamanho2 == 'P':
                estoque['E4'].value = estoque['E4'].value - 1
            if tamanho2 == 'M':
                estoque['E5'].value = estoque['E5'].value - 1
            if tamanho2 == 'G':
                estoque['E6'].value = estoque['E6'].value - 1
            if tamanho2 == 'GG':
                estoque['E7'].value = estoque['E7'].value - 1
        lista[f'E{num}'].value = tamanho1
        lista[f'F{num}'].value = 'OK'
        x = len(list(entregues.rows)) + 1
        entregues[f'A{x}'].value = pessoa
        entregues[f'B{x}'].value = 1
        entregues[f'C{x}'].value = tamanho1
        entregues[f'D{x}'].value = gen
        entregues[f'E{x}'].value = 1
        entregues[f'F{x}'].value = tamanho2
        entregues[f'G{x}'].value = gen
        relatorio.save(local)
        tkinter.messagebox.showinfo(title='Recibo ok!', message='Recibo impresso com sucesso!')
    else:
        recibo = docx.Document('recibo_uniforme.docx')
        recibo.paragraphs[11].text = str(recibo.paragraphs[11].text)\
            .replace('#nome', pessoa).replace('#num_cpf', cpf_ed)\
            .replace('#tam', tamanho1).replace('#genero', str(gen).lower())
        recibo.paragraphs[19].text = str(recibo.paragraphs[19].text)\
            .replace('#data', dt.strftime(hoje, '%d/%m/%Y'))
        recibo.paragraphs[24].text = str(recibo.paragraphs[24].text).replace('#nome', pessoa)
        recibo.paragraphs[25].text = str(recibo.paragraphs[25].text).replace('#cargo', cargo)
        recibo.save(f'Recibo_alterado {pessoa}.docx')
        docx2pdf.convert(f'Recibo_alterado {pessoa}.docx', f'Recibo {pessoa}.pdf')
        # diminuir contagem de estoque
        if gen == 'Masculino':
            if tamanho1 == 'P':
                estoque['C4'].value = estoque['C4'].value - 2
            if tamanho1 == 'M':
                estoque['C5'].value = estoque['C5'].value - 2
            if tamanho1 == 'G':
                estoque['C6'].value = estoque['C6'].value - 2
            if tamanho1 == 'XG':
                estoque['C8'].value = estoque['C8'].value - 2
            if tamanho1 == 'GG':
                estoque['C7'].value = estoque['C7'].value - 2
            if tamanho1 == 'XGG':
                estoque['C9'].value = estoque['C9'].value - 2
        else:
            if tamanho1 == 'P':
                estoque['E4'].value = estoque['E4'].value - 2
            if tamanho1 == 'M':
                estoque['E5'].value = estoque['E5'].value - 2
            if tamanho1 == 'G':
                estoque['E6'].value = estoque['E6'].value - 2
            if tamanho1 == 'GG':
                estoque['E7'].value = estoque['E7'].value - 2
        lista[f'E{num}'].value = tamanho1
        lista[f'F{num}'].value = 'OK'
        x = len(list(entregues.rows)) + 1
        entregues[f'A{x}'].value = pessoa
        entregues[f'B{x}'].value = 2
        entregues[f'C{x}'].value = tamanho1
        entregues[f'D{x}'].value = gen
        entregues[f'E{x}'].value = '-'
        entregues[f'F{x}'].value = '-'
        entregues[f'G{x}'].value = '-'
        relatorio.save(local)
    tkinter.messagebox.showinfo(title='Recibo ok!', message='Recibo impresso com sucesso!')


def confirmar_pagamento(valor='10,00', tipo1='adiantamento', tipo2='0', tipo3='0', tipo4='0', tipo5='0', tipo6='0',
                        dia='03/07/2023', competencia='07/2023'):
    msg_box = tkinter.messagebox.askquestion('Confirma pagamento',
                                             'Tem certeza que deseja enviar o pagamento ao financeiro?\n'
                                             f'Valor: R$ {valor}\n'
                                             f'Data: {dia}\n'
                                             f'Tipo: {tipo1}\n'
                                             f'Competência: {competencia}\n',
                                             icon='warning')
    if msg_box == 'yes':
        tkinter.messagebox.showinfo('Pagamento enviado!', 'Pagamento enviado ao financeiro com sucesso!')
    else:
        tkinter.messagebox.showinfo('Editar dados', 'Pagamento não enviado. Edite os dados e tente novamente.')


def escrever_valor_por_extenso(total):
    # transformar algarismos do total em número por extenso e com reais e centavos
    reais, centavos = str(format(total, '.2f')).split('.')
    if int(reais) == 1:
        strReal = 'real'
    else:
        strReal = 'reais'
    if int(centavos) == 1:
        strCentavo = 'centavo'
    else:
        strCentavo = 'centavos'
    if int(reais) == 0:
        extenso = f'{nw.num2words(centavos, lang="pt_BR").capitalize()} {strCentavo}.'
    else:
        if int(centavos) == 0:
            extenso = f'{nw.num2words(reais, lang="pt_BR").capitalize()} {strReal}.'
        else:
            extenso = f'{nw.num2words(reais, lang="pt_BR").capitalize()} {strReal} e {nw.num2words(centavos, lang="pt_BR")} {strCentavo}.'
    return extenso


def gerar_planilha_pgto_itau(nome1, nome2, nome3, nome4, nome5, nome6, nome7, nome8, nome9, nome10, nome11, nome12, nome13, nome14, nome15,
                             tipo1, tipo2, tipo3, tipo4, tipo5, tipo6, tipo7, tipo8, tipo9, tipo10, tipo11, tipo12, tipo13, tipo14, tipo15,
                             val1, val2, val3, val4, val5, val6, val7, val8, val9, val10, val11, val12, val13, val14, val15, data):
    mesext = {'01': 'JAN', '02': 'FEV', '03': 'MAR', '04': 'ABR', '05': 'MAI', '06': 'JUN',
              '07': 'JUL', '08': 'AGO', '09': 'SET', '10': 'OUT', '11': 'NOV', '12': 'DEZ'}
    if val1 != '':
        valor1 = float(val1.replace(',','.'))
    else:
        valor1 = ''

    if val2 != '':
        valor2 = float(val2.replace(',','.'))
    else:
        valor2 = ''

    if val3 != '':
        valor3 = float(val3.replace(',','.'))
    else:
        valor3 = ''

    if val4 != '':
        valor4 = float(val4.replace(',','.'))
    else:
        valor4 = ''

    if val5 != '':
        valor5 = float(val5.replace(',','.'))
    else:
        valor5 = ''

    if val6 != '':
        valor6 = float(val6.replace(',','.'))
    else:
        valor6 = ''

    if val7 != '':
        valor7 = float(val7.replace(',','.'))
    else:
        valor7 = ''

    if val8 != '':
        valor8 = float(val8.replace(',','.'))
    else:
        valor8 = ''

    if val9 != '':
        valor9 = float(val9.replace(',','.'))
    else:
        valor9 = ''

    if val10 != '':
        valor10 = float(val10.replace(',','.'))
    else:
        valor10 = ''

    if val11 != '':
        valor11 = float(val11.replace(',','.'))
    else:
        valor11 = ''

    if val12 != '':
        valor12 = float(val12.replace(',','.'))
    else:
        valor12 = ''

    if val13 != '':
        valor13 = float(val13.replace(',','.'))
    else:
        valor13 = ''

    if val14 != '':
        valor14 = float(val14.replace(',','.'))
    else:
        valor14 = ''

    if val15 != '':
        valor15 = float(val15.replace(',','.'))
    else:
        valor15 = ''

    sessions = sessionmaker(bind=engine)
    session = sessions()
    dia = data.replace('/','.')
    tipos = {'':'','Salário': '1', 'Férias': '2', 'Vale Transporte': '3', 'Vale Alimentação': '4', 'Comissão': '5',
             '13º salário': '6', 'Bolsa Estágio': '7', 'Bônus': '8', 'Adiantamento Salarial': '9',
             'Rescisão': '10', 'Bolsa Auxílio': '11', 'Pensão Alimentícia': '12', 'Pgto em C/C': '13',
             'Remuneração': '14'}
    p = os.path.relpath(rf'C:\Users\{os.getlogin()}\PycharmProjects\AutomacaoCia\src\models\static\files\planilha_itau.xlsx')
    wb = l_w(p, read_only=False)
    sh = wb['Planilha1']
    pessoa1 = session.query(Colaborador).filter_by(nome=nome1).first()
    pessoa2 = session.query(Colaborador).filter_by(nome=nome2).first()
    pessoa3 = session.query(Colaborador).filter_by(nome=nome3).first()
    pessoa4 = session.query(Colaborador).filter_by(nome=nome4).first()
    pessoa5 = session.query(Colaborador).filter_by(nome=nome5).first()
    pessoa6 = session.query(Colaborador).filter_by(nome=nome6).first()
    pessoa7 = session.query(Colaborador).filter_by(nome=nome7).first()
    pessoa8 = session.query(Colaborador).filter_by(nome=nome8).first()
    pessoa9 = session.query(Colaborador).filter_by(nome=nome9).first()
    pessoa10 = session.query(Colaborador).filter_by(nome=nome10).first()
    pessoa11 = session.query(Colaborador).filter_by(nome=nome11).first()
    pessoa12 = session.query(Colaborador).filter_by(nome=nome12).first()
    pessoa13 = session.query(Colaborador).filter_by(nome=nome13).first()
    pessoa14 = session.query(Colaborador).filter_by(nome=nome14).first()
    pessoa15 = session.query(Colaborador).filter_by(nome=nome15).first()

    sh['A1'].value = pessoa1.ag
    sh['A2'].value = pessoa2.ag
    sh['A3'].value = pessoa3.ag
    sh['A4'].value = pessoa4.ag
    sh['A5'].value = pessoa5.ag
    sh['A6'].value = pessoa6.ag
    sh['A7'].value = pessoa7.ag
    sh['A8'].value = pessoa8.ag
    sh['A9'].value = pessoa9.ag
    sh['A10'].value = pessoa10.ag
    sh['A11'].value = pessoa11.ag
    sh['A12'].value = pessoa12.ag
    sh['A13'].value = pessoa13.ag
    sh['A14'].value = pessoa14.ag
    sh['A15'].value = pessoa15.ag

    sh['B1'].value = pessoa1.conta
    sh['B2'].value = pessoa2.conta
    sh['B3'].value = pessoa3.conta
    sh['B4'].value = pessoa4.conta
    sh['B5'].value = pessoa5.conta
    sh['B6'].value = pessoa6.conta
    sh['B7'].value = pessoa7.conta
    sh['B8'].value = pessoa8.conta
    sh['B9'].value = pessoa9.conta
    sh['B10'].value = pessoa10.conta
    sh['B11'].value = pessoa11.conta
    sh['B12'].value = pessoa12.conta
    sh['B13'].value = pessoa13.conta
    sh['B14'].value = pessoa14.conta
    sh['B15'].value = pessoa15.conta

    sh['C1'].value = pessoa1.cdigito
    sh['C2'].value = pessoa2.cdigito
    sh['C3'].value = pessoa3.cdigito
    sh['C4'].value = pessoa4.cdigito
    sh['C5'].value = pessoa5.cdigito
    sh['C6'].value = pessoa6.cdigito
    sh['C7'].value = pessoa7.cdigito
    sh['C8'].value = pessoa8.cdigito
    sh['C9'].value = pessoa9.cdigito
    sh['C10'].value = pessoa10.cdigito
    sh['C11'].value = pessoa11.cdigito
    sh['C12'].value = pessoa12.cdigito
    sh['C13'].value = pessoa13.cdigito
    sh['C14'].value = pessoa14.cdigito
    sh['C15'].value = pessoa15.cdigito

    sh['D1'].value = pessoa1.nome
    sh['D2'].value = pessoa2.nome
    sh['D3'].value = pessoa3.nome
    sh['D4'].value = pessoa4.nome
    sh['D5'].value = pessoa5.nome
    sh['D6'].value = pessoa6.nome
    sh['D7'].value = pessoa7.nome
    sh['D8'].value = pessoa8.nome
    sh['D9'].value = pessoa9.nome
    sh['D10'].value = pessoa10.nome
    sh['D11'].value = pessoa11.nome
    sh['D12'].value = pessoa12.nome
    sh['D13'].value = pessoa13.nome
    sh['D14'].value = pessoa14.nome
    sh['D15'].value = pessoa15.nome

    sh['E1'].value = pessoa1.cpf
    sh['E2'].value = pessoa2.cpf
    sh['E3'].value = pessoa3.cpf
    sh['E4'].value = pessoa4.cpf
    sh['E5'].value = pessoa5.cpf
    sh['E6'].value = pessoa6.cpf
    sh['E7'].value = pessoa7.cpf
    sh['E8'].value = pessoa8.cpf
    sh['E9'].value = pessoa9.cpf
    sh['E10'].value = pessoa10.cpf
    sh['E11'].value = pessoa11.cpf
    sh['E12'].value = pessoa12.cpf
    sh['E13'].value = pessoa13.cpf
    sh['E14'].value = pessoa14.cpf
    sh['E15'].value = pessoa15.cpf

    sh['F1'].value = tipos[tipo1]
    sh['F2'].value = tipos[tipo2]
    sh['F3'].value = tipos[tipo3]
    sh['F4'].value = tipos[tipo4]
    sh['F5'].value = tipos[tipo5]
    sh['F6'].value = tipos[tipo6]
    sh['F7'].value = tipos[tipo7]
    sh['F8'].value = tipos[tipo8]
    sh['F9'].value = tipos[tipo9]
    sh['F10'].value = tipos[tipo10]
    sh['F11'].value = tipos[tipo11]
    sh['F12'].value = tipos[tipo12]
    sh['F13'].value = tipos[tipo13]
    sh['F14'].value = tipos[tipo14]
    sh['F15'].value = tipos[tipo15]

    sh['G1'].value = valor1
    sh['G2'].value = valor2
    sh['G3'].value = valor3
    sh['G4'].value = valor4
    sh['G5'].value = valor5
    sh['G6'].value = valor6
    sh['G7'].value = valor7
    sh['G8'].value = valor8
    sh['G9'].value = valor9
    sh['G10'].value = valor10
    sh['G11'].value = valor11
    sh['G12'].value = valor12
    sh['G13'].value = valor13
    sh['G14'].value = valor14
    sh['G15'].value = valor15

    sh.column_dimensions['A'].width = 6
    sh.column_dimensions['B'].width = 8
    sh.column_dimensions['C'].width = 4
    sh.column_dimensions['D'].width = 45
    sh.column_dimensions['E'].width = 20
    sh.column_dimensions['F'].width = 4
    sh.column_dimensions['G'].width = 16
    d, mes, ano = dia.split('.')
    try:
        wb.save(rf'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\04 - Folha de Pgto\{ano}\{mes} - {mesext[mes]}\Pedidos de pagamento\Pagamento Itau {dia}.xlsx')
    except Exception:
        os.makedirs(rf'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\04 - Folha de Pgto\{ano}\{mes} - {mesext[mes]}\Pedidos de pagamento')
        wb.save(
            rf'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\04 - Folha de Pgto\{ano}\{mes} - {mesext[mes]}\Pedidos de pagamento\Pagamento Itau {dia}.xlsx')

    tps = [tipo1, tipo2, tipo3, tipo4, tipo5, tipo6, tipo7, tipo8, tipo9, tipo10, tipo11, tipo12, tipo13, tipo14, tipo15]
    tp = [x for x in tps if x != '']
    tipos_unicos = sorted(list(set(filter(None, tp))))
    tkinter.messagebox.showinfo(title=f'Planilha salva dia {data}!', message=f'{len(tipos_unicos)} tipo(s) de pgto diferente(s).\n {str(tipos_unicos).replace("[","").replace("]","")}.')

