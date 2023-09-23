from datetime import datetime as dt
import docx
import docx2pdf
from email import encoders
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.image import MIMEImage
from email.mime.base import MIMEBase
from inspect import signature
import locale
import num2words as nw
from num2words import num2words
from openpyxl import load_workbook as l_w
import os
from src.models.models import Colaborador, engine
from sqlalchemy.orm import sessionmaker
import smtplib
from src.models.dados_servd import em_rem, k1, host, port, em_fin
import tkinter.filedialog
from tkinter import messagebox
import tkinter.filedialog

locale.setlocale(locale.LC_MONETARY, 'pt_BR.UTF-8')
locale.setlocale(locale.LC_ALL, 'pt_BR.UTF-8')
locale.setlocale(locale.LC_MONETARY, 'pt_BR.UTF-8')


def confirmar_pagamento(tipo1, tipo2, tipo3, tipo4, tipo5, tipo6, tipo7, tipo8, tipo9, tipo10, tipo11, tipo12, tipo13,
                        tipo14, tipo15, valor1, valor2, valor3, valor4, valor5, valor6, valor7, valor8, valor9, valor10,
                        valor11, valor12, valor13, valor14, valor15, data):
    tps = [tipo1, tipo2, tipo3, tipo4, tipo5, tipo6, tipo7, tipo8,
           tipo9, tipo10, tipo11, tipo12, tipo13, tipo14, tipo15]
    tp = [x for x in tps if x != '']
    tipos_unicos = sorted(list(set(filter(None, tp))))
    listatipos = ''
    for item in tipos_unicos:
        listatipos += f'{item} '
    msg_box = tkinter.messagebox.askquestion('Confirma pagamento',
                                             'Tem certeza que deseja enviar o pagamento ao financeiro?\n'
                                             f'Data: {data}\n'
                                             f'Tipos: {listatipos}\n',
                                             icon='warning')
    if msg_box == 'yes':
        gerar_capa_email(tipo1, tipo2, tipo3, tipo4, tipo5, tipo6, tipo7, tipo8, tipo9, tipo10, tipo11, tipo12, tipo13,
                         tipo14, tipo15, valor1, valor2, valor3, valor4, valor5, valor6, valor7, valor8, valor9,
                         valor10, valor11, valor12, valor13, valor14, valor15, data)
        tkinter.messagebox.showinfo('Pagamento enviado!', 'Pagamento enviado ao financeiro com sucesso!')
    else:
        tkinter.messagebox.showinfo('Editar dados', 'Pagamento não enviado. Edite os dados e tente novamente.')


def escrever_valor_por_extenso(total):
    # transformar algarismos do total em número por extenso e com reais e centavos
    reais, centavos = str(format(total, '.2f')).split('.')
    if int(reais) == 1:
        strReal = 'real'
    else:
        strReal = 'reais'
    if int(centavos) == 1:
        strCentavo = 'centavo'
    else:
        strCentavo = 'centavos'
    if int(reais) == 0:
        extenso = f'{nw.num2words(centavos, lang="pt_BR").capitalize()} {strCentavo}.'
    else:
        if int(centavos) == 0:
            extenso = f'{nw.num2words(reais, lang="pt_BR").capitalize()} {strReal}.'
        else:
            extenso = f'{nw.num2words(reais, lang="pt_BR").capitalize()} {strReal} e {nw.num2words(centavos, lang="pt_BR")} {strCentavo}.'
    return extenso


def gerar_planilha_pgto_itau(nome1, nome2, nome3, nome4, nome5, nome6, nome7, nome8, nome9, nome10, nome11, nome12,
                             nome13, nome14, nome15, tipo1, tipo2, tipo3, tipo4, tipo5, tipo6, tipo7, tipo8, tipo9,
                             tipo10, tipo11, tipo12, tipo13, tipo14, tipo15, val1, val2, val3, val4, val5, val6, val7,
                             val8, val9, val10, val11, val12, val13, val14, val15, data):
    mesext = {'01': 'JAN', '02': 'FEV', '03': 'MAR', '04': 'ABR', '05': 'MAI', '06': 'JUN',
              '07': 'JUL', '08': 'AGO', '09': 'SET', '10': 'OUT', '11': 'NOV', '12': 'DEZ'}
    if val1 != '':
        valor1 = float(val1.replace(',', '.'))
    else:
        valor1 = ''

    if val2 != '':
        valor2 = float(val2.replace(',', '.'))
    else:
        valor2 = ''

    if val3 != '':
        valor3 = float(val3.replace(',', '.'))
    else:
        valor3 = ''

    if val4 != '':
        valor4 = float(val4.replace(',', '.'))
    else:
        valor4 = ''

    if val5 != '':
        valor5 = float(val5.replace(',', '.'))
    else:
        valor5 = ''

    if val6 != '':
        valor6 = float(val6.replace(',', '.'))
    else:
        valor6 = ''

    if val7 != '':
        valor7 = float(val7.replace(',', '.'))
    else:
        valor7 = ''

    if val8 != '':
        valor8 = float(val8.replace(',', '.'))
    else:
        valor8 = ''

    if val9 != '':
        valor9 = float(val9.replace(',', '.'))
    else:
        valor9 = ''

    if val10 != '':
        valor10 = float(val10.replace(',', '.'))
    else:
        valor10 = ''

    if val11 != '':
        valor11 = float(val11.replace(',', '.'))
    else:
        valor11 = ''

    if val12 != '':
        valor12 = float(val12.replace(',', '.'))
    else:
        valor12 = ''

    if val13 != '':
        valor13 = float(val13.replace(',', '.'))
    else:
        valor13 = ''

    if val14 != '':
        valor14 = float(val14.replace(',', '.'))
    else:
        valor14 = ''

    if val15 != '':
        valor15 = float(val15.replace(',', '.'))
    else:
        valor15 = ''

    sessions = sessionmaker(bind=engine)
    session = sessions()
    dia = data.replace('/', '.')
    tipos = {'': '', 'Salário': '1', 'Férias': '2', 'Vale Transporte': '3', 'Vale Alimentação': '4', 'Comissão': '5',
             '13º salário': '6', 'Bolsa Estágio': '7', 'Bônus': '8', 'Adiantamento Salarial': '9',
             'Rescisão': '10', 'Bolsa Auxílio': '11', 'Pensão Alimentícia': '12', 'Pgto em C/C': '13',
             'Remuneração': '14'}
    p = os.path.relpath(
        rf'C:\Users\{os.getlogin()}\PycharmProjects\AutomacaoCia\src\models\static\files\planilha_itau.xlsx')
    wb = l_w(p, read_only=False)
    sh = wb['Planilha1']
    pessoa1 = session.query(Colaborador).filter_by(nome=nome1).first()
    pessoa2 = session.query(Colaborador).filter_by(nome=nome2).first()
    pessoa3 = session.query(Colaborador).filter_by(nome=nome3).first()
    pessoa4 = session.query(Colaborador).filter_by(nome=nome4).first()
    pessoa5 = session.query(Colaborador).filter_by(nome=nome5).first()
    pessoa6 = session.query(Colaborador).filter_by(nome=nome6).first()
    pessoa7 = session.query(Colaborador).filter_by(nome=nome7).first()
    pessoa8 = session.query(Colaborador).filter_by(nome=nome8).first()
    pessoa9 = session.query(Colaborador).filter_by(nome=nome9).first()
    pessoa10 = session.query(Colaborador).filter_by(nome=nome10).first()
    pessoa11 = session.query(Colaborador).filter_by(nome=nome11).first()
    pessoa12 = session.query(Colaborador).filter_by(nome=nome12).first()
    pessoa13 = session.query(Colaborador).filter_by(nome=nome13).first()
    pessoa14 = session.query(Colaborador).filter_by(nome=nome14).first()
    pessoa15 = session.query(Colaborador).filter_by(nome=nome15).first()

    sh['A1'].value = pessoa1.ag
    sh['A2'].value = pessoa2.ag
    sh['A3'].value = pessoa3.ag
    sh['A4'].value = pessoa4.ag
    sh['A5'].value = pessoa5.ag
    sh['A6'].value = pessoa6.ag
    sh['A7'].value = pessoa7.ag
    sh['A8'].value = pessoa8.ag
    sh['A9'].value = pessoa9.ag
    sh['A10'].value = pessoa10.ag
    sh['A11'].value = pessoa11.ag
    sh['A12'].value = pessoa12.ag
    sh['A13'].value = pessoa13.ag
    sh['A14'].value = pessoa14.ag
    sh['A15'].value = pessoa15.ag

    sh['B1'].value = pessoa1.conta
    sh['B2'].value = pessoa2.conta
    sh['B3'].value = pessoa3.conta
    sh['B4'].value = pessoa4.conta
    sh['B5'].value = pessoa5.conta
    sh['B6'].value = pessoa6.conta
    sh['B7'].value = pessoa7.conta
    sh['B8'].value = pessoa8.conta
    sh['B9'].value = pessoa9.conta
    sh['B10'].value = pessoa10.conta
    sh['B11'].value = pessoa11.conta
    sh['B12'].value = pessoa12.conta
    sh['B13'].value = pessoa13.conta
    sh['B14'].value = pessoa14.conta
    sh['B15'].value = pessoa15.conta

    sh['C1'].value = pessoa1.cdigito
    sh['C2'].value = pessoa2.cdigito
    sh['C3'].value = pessoa3.cdigito
    sh['C4'].value = pessoa4.cdigito
    sh['C5'].value = pessoa5.cdigito
    sh['C6'].value = pessoa6.cdigito
    sh['C7'].value = pessoa7.cdigito
    sh['C8'].value = pessoa8.cdigito
    sh['C9'].value = pessoa9.cdigito
    sh['C10'].value = pessoa10.cdigito
    sh['C11'].value = pessoa11.cdigito
    sh['C12'].value = pessoa12.cdigito
    sh['C13'].value = pessoa13.cdigito
    sh['C14'].value = pessoa14.cdigito
    sh['C15'].value = pessoa15.cdigito

    sh['D1'].value = pessoa1.nome
    sh['D2'].value = pessoa2.nome
    sh['D3'].value = pessoa3.nome
    sh['D4'].value = pessoa4.nome
    sh['D5'].value = pessoa5.nome
    sh['D6'].value = pessoa6.nome
    sh['D7'].value = pessoa7.nome
    sh['D8'].value = pessoa8.nome
    sh['D9'].value = pessoa9.nome
    sh['D10'].value = pessoa10.nome
    sh['D11'].value = pessoa11.nome
    sh['D12'].value = pessoa12.nome
    sh['D13'].value = pessoa13.nome
    sh['D14'].value = pessoa14.nome
    sh['D15'].value = pessoa15.nome

    sh['E1'].value = pessoa1.cpf
    sh['E2'].value = pessoa2.cpf
    sh['E3'].value = pessoa3.cpf
    sh['E4'].value = pessoa4.cpf
    sh['E5'].value = pessoa5.cpf
    sh['E6'].value = pessoa6.cpf
    sh['E7'].value = pessoa7.cpf
    sh['E8'].value = pessoa8.cpf
    sh['E9'].value = pessoa9.cpf
    sh['E10'].value = pessoa10.cpf
    sh['E11'].value = pessoa11.cpf
    sh['E12'].value = pessoa12.cpf
    sh['E13'].value = pessoa13.cpf
    sh['E14'].value = pessoa14.cpf
    sh['E15'].value = pessoa15.cpf

    sh['F1'].value = tipos[tipo1]
    sh['F2'].value = tipos[tipo2]
    sh['F3'].value = tipos[tipo3]
    sh['F4'].value = tipos[tipo4]
    sh['F5'].value = tipos[tipo5]
    sh['F6'].value = tipos[tipo6]
    sh['F7'].value = tipos[tipo7]
    sh['F8'].value = tipos[tipo8]
    sh['F9'].value = tipos[tipo9]
    sh['F10'].value = tipos[tipo10]
    sh['F11'].value = tipos[tipo11]
    sh['F12'].value = tipos[tipo12]
    sh['F13'].value = tipos[tipo13]
    sh['F14'].value = tipos[tipo14]
    sh['F15'].value = tipos[tipo15]

    sh['G1'].value = valor1
    sh['G2'].value = valor2
    sh['G3'].value = valor3
    sh['G4'].value = valor4
    sh['G5'].value = valor5
    sh['G6'].value = valor6
    sh['G7'].value = valor7
    sh['G8'].value = valor8
    sh['G9'].value = valor9
    sh['G10'].value = valor10
    sh['G11'].value = valor11
    sh['G12'].value = valor12
    sh['G13'].value = valor13
    sh['G14'].value = valor14
    sh['G15'].value = valor15

    sh.column_dimensions['A'].width = 6
    sh.column_dimensions['B'].width = 8
    sh.column_dimensions['C'].width = 4
    sh.column_dimensions['D'].width = 45
    sh.column_dimensions['E'].width = 20
    sh.column_dimensions['F'].width = 4
    sh.column_dimensions['G'].width = 16
    d, mes, ano = dia.split('.')
    try:
        wb.save(
            rf'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\04 - Folha de Pgto\{ano}\{mes} - {mesext[mes]}\Pedidos de pagamento\Pagamento Itau {dia}.xlsx')
    except Exception:
        os.makedirs(
            rf'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\04 - Folha de Pgto\{ano}\{mes} - {mesext[mes]}\Pedidos de pagamento')
        wb.save(
            rf'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\04 - Folha de Pgto\{ano}\{mes} - {mesext[mes]}\Pedidos de pagamento\Pagamento Itau {dia}.xlsx')

    tps = [tipo1, tipo2, tipo3, tipo4, tipo5, tipo6, tipo7, tipo8, tipo9, tipo10, tipo11, tipo12, tipo13, tipo14,
           tipo15]
    tp = [x for x in tps if x != '']
    tipos_unicos = sorted(list(set(filter(None, tp))))
    tkinter.messagebox.showinfo(title=f'Planilha salva dia {data}!',
                                message=f'{len(tipos_unicos)} tipo(s) de pgto diferente(s).\n {str(tipos_unicos).replace("[", "").replace("]", "")}.')


def valor_por_extenso(number_p):
    if number_p.find(',') != -1:
        number_p = number_p.split(',')
        number_p1 = int(number_p[0].replace('.', ''))
        number_p2 = int(number_p[1])
    else:
        number_p1 = int(number_p.replace('.', ''))
        number_p2 = 0

    if number_p1 == 1:
        aux1 = ' real'
    else:
        aux1 = ' reais'

    if number_p2 == 1:
        aux2 = ' centavo'
    else:
        aux2 = ' centavos'

    text1 = ''
    if number_p1 > 0:
        text1 = num2words(number_p1, lang='pt_BR') + str(aux1)
    else:
        text1 = ''

    if number_p2 > 0:
        text2 = num2words(number_p2, lang='pt_BR') + str(aux2)
    else:
        text2 = ''

    if number_p1 > 0 and number_p2 > 0:
        result = text1 + ' e ' + text2
    else:
        result = text1 + text2
    return result


def gerar_capa_email(tipo1, tipo2, tipo3, tipo4, tipo5, tipo6, tipo7, tipo8, tipo9, tipo10, tipo11, tipo12, tipo13,
                     tipo14, tipo15, val1, val2, val3, val4, val5, val6, val7, val8, val9, val10, val11, val12, val13,
                     val14, val15, data):
    somas = {}
    if val1 != '':
        valor1 = float(val1.replace(',', '.'))
    else:
        valor1 = 0

    if val2 != '':
        valor2 = float(val2.replace(',', '.'))
    else:
        valor2 = 0

    if val3 != '':
        valor3 = float(val3.replace(',', '.'))
    else:
        valor3 = 0

    if val4 != '':
        valor4 = float(val4.replace(',', '.'))
    else:
        valor4 = 0

    if val5 != '':
        valor5 = float(val5.replace(',', '.'))
    else:
        valor5 = 0

    if val6 != '':
        valor6 = float(val6.replace(',', '.'))
    else:
        valor6 = 0

    if val7 != '':
        valor7 = float(val7.replace(',', '.'))
    else:
        valor7 = 0

    if val8 != '':
        valor8 = float(val8.replace(',', '.'))
    else:
        valor8 = 0

    if val9 != '':
        valor9 = float(val9.replace(',', '.'))
    else:
        valor9 = 0

    if val10 != '':
        valor10 = float(val10.replace(',', '.'))
    else:
        valor10 = 0

    if val11 != '':
        valor11 = float(val11.replace(',', '.'))
    else:
        valor11 = 0

    if val12 != '':
        valor12 = float(val12.replace(',', '.'))
    else:
        valor12 = 0

    if val13 != '':
        valor13 = float(val13.replace(',', '.'))
    else:
        valor13 = 0

    if val14 != '':
        valor14 = float(val14.replace(',', '.'))
    else:
        valor14 = 0

    if val15 != '':
        valor15 = float(val15.replace(',', '.'))
    else:
        valor15 = 0

    if tipo1 in somas:
        somas[tipo1] = somas[tipo1] + valor1
    else:
        somas[tipo1] = valor1

    if tipo2 in somas:
        somas[tipo2] = somas[tipo2] + valor2
    else:
        somas[tipo2] = valor2

    if tipo3 in somas:
        somas[tipo3] = somas[tipo3] + valor3
    else:
        somas[tipo3] = valor3

    if tipo4 in somas:
        somas[tipo4] = somas[tipo4] + valor4
    else:
        somas[tipo4] = valor4

    if tipo5 in somas:
        somas[tipo5] = somas[tipo5] + valor5
    else:
        somas[tipo5] = valor5

    if tipo6 in somas:
        somas[tipo6] = somas[tipo6] + valor6
    else:
        somas[tipo6] = valor6

    if tipo7 in somas:
        somas[tipo7] = somas[tipo7] + valor7
    else:
        somas[tipo7] = valor7

    if tipo8 in somas:
        somas[tipo8] = somas[tipo8] + valor8
    else:
        somas[tipo8] = valor8

    if tipo9 in somas:
        somas[tipo9] = somas[tipo9] + valor9
    else:
        somas[tipo9] = valor9

    if tipo10 in somas:
        somas[tipo10] = somas[tipo10] + valor10
    else:
        somas[tipo10] = valor10

    if tipo11 in somas:
        somas[tipo11] = somas[tipo11] + valor11
    else:
        somas[tipo11] = valor11

    if tipo12 in somas:
        somas[tipo12] = somas[tipo12] + valor12
    else:
        somas[tipo12] = valor12

    if tipo13 in somas:
        somas[tipo13] = somas[tipo13] + valor13
    else:
        somas[tipo13] = valor13

    if tipo14 in somas:
        somas[tipo14] = somas[tipo14] + valor14
    else:
        somas[tipo14] = valor14

    if tipo15 in somas:
        somas[tipo15] = somas[tipo15] + valor15
    else:
        somas[tipo15] = valor15
    del somas['']
    quantidade_de_pgtos = len(somas)
    total = 0
    for item in somas:
        total += somas[item]
    somas['Total'] = total
    for item in somas:
        somas[item] = '{0:,.2f}'.format(somas[item])
    qtidades = {
        1: umpgto,
        2: doispgtos,
        3: trespgtos,
        4: quatropgtos,
        5: cincopgtos,
        6: seispgtos,
        7: setepgtos
    }
    qtidades[quantidade_de_pgtos](somas, data)


def email_pgto(arquivo: str, somas: dict, data: str):
    mesext = {'01': 'JAN', '02': 'FEV', '03': 'MAR', '04': 'ABR', '05': 'MAI', '06': 'JUN',
              '07': 'JUL', '08': 'AGO', '09': 'SET', '10': 'OUT', '11': 'NOV', '12': 'DEZ'}
    dia, mes, ano = data.split('/')
    pasta_pgto = rf'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\04 - Folha de Pgto\{ano}\{mes} - {mesext[mes]}\Pedidos de pagamento'
    tam = len(somas)
    k = list(somas)
    frase = ''
    for x in range(tam):
        frase += f'{k[x]} - R$ {str(somas[k[x]]).replace(",", "_").replace(".", ",").replace("_", ".")}.<br>'
    # send e-mails
    email_remetente = em_rem
    senha = k1
    # set up smtp connection
    s = smtplib.SMTP(host=host, port=port)
    s.starttls()
    s.login(email_remetente, senha)

    # enviar e-mail de boas vindas
    msg = MIMEMultipart('alternative')
    msg['From'] = email_remetente
    msg['To'] = em_fin
    msg['Subject'] = f'Pagamento Itaú - {data}'
    arquivo = arquivo
    arquivo2 = pasta_pgto + f'\\Itau {data[:2]}.pdf'
    text = MIMEText(f'''Oi, Marcelo!<br><br>
    Segue pagamento agendado no Itaú para {data}.<br><br>
    {frase}<br>
    Atenciosamente,<br>
    <img src="cid:image1">''', 'html')

    # set up the parameters of the message
    msg.attach(text)
    image = MIMEImage(
        open(rf'C:\Users\{os.getlogin()}\PycharmProjects\AutomacaoCia\src\models\static\imgs\assinatura.png',
             'rb').read())
    image.add_header('Content-ID', '<image1>')
    msg.attach(image)
    # attach pdf file
    part = MIMEBase('application', "octet-stream")
    part.set_payload(open(arquivo, "rb").read())
    encoders.encode_base64(part)
    part.add_header('Content-Disposition', 'attachment',
                    filename=f'Pagamento Itaú {data.replace("/", ".")}.pdf')
    msg.attach(part)
    part2 = MIMEBase('application', "octet-stream")
    part2.set_payload(open(arquivo2, "rb").read())
    encoders.encode_base64(part2)
    part2.add_header('Content-Disposition', 'attachment',
                     filename=f'Agendamento Pgto Itau {data.replace("/", ".")}.pdf')
    msg.attach(part2)
    s.sendmail(email_remetente, em_fin, msg.as_string())
    del msg


def umpgto(somas, data):
    mesext = {'01': 'JAN', '02': 'FEV', '03': 'MAR', '04': 'ABR', '05': 'MAI', '06': 'JUN',
              '07': 'JUL', '08': 'AGO', '09': 'SET', '10': 'OUT', '11': 'NOV', '12': 'DEZ'}
    dia, mes, ano = data.split('/')
    hoje = dt.strftime(dt.today(), '%d/%m/%Y')
    competencia = data[-7:]
    pasta_pgto = rf'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\04 - Folha de Pgto\{ano}\{mes} - {mesext[mes]}\Pedidos de pagamento'
    modelo = os.path.relpath(
        rf'C:\Users\{os.getlogin()}\PycharmProjects\AutomacaoCia\src\models\static\files\solicitpgto_modelocapa.docx')
    pgto1 = docx.Document(modelo)
    pagamento = list(somas)[0]
    # # Alterar Modelo e Salvar na pasta do mes correspondente
    for parag in pgto1.paragraphs:
        if '#tipo_pgto' in parag.text:
            inline = parag.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if '#tipo_pgto' in inline[i].text:
                    text = inline[i].text.replace('#tipo_pgto', pagamento)
                    inline[i].text = text
        if '#valor' in parag.text:
            inline = parag.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if '#valor' in inline[i].text:
                    text = inline[i].text.replace('#valor',
                                                  str(somas[pagamento]).replace(',', '_').replace('.', ',').replace('_',
                                                                                                                    '.'))
                    inline[i].text = text
        if '#extenso' in parag.text:
            inline = parag.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if '#extenso' in inline[i].text:
                    text = inline[i].text.replace('#extenso',
                                                  valor_por_extenso(
                                                      str(somas[pagamento]).replace(',', '_').replace('.', ',').replace(
                                                          '_', '.')))
                    inline[i].text = text
        if '#competencia' in parag.text:
            inline = parag.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if '#competencia' in inline[i].text:
                    text = inline[i].text.replace('#competencia', competencia)
                    inline[i].text = text
        if '#dia_pgto' in parag.text:
            inline = parag.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if '#dia_pgto' in inline[i].text:
                    text = inline[i].text.replace('#dia_pgto', data)
                    inline[i].text = text
        if '#hoje' in parag.text:
            inline = parag.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if '#hoje' in inline[i].text:
                    text = inline[i].text.replace('#hoje', hoje)
                    inline[i].text = text
    n = 0
    if os.path.exists(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")}.pdf'):
        n += 1
        if os.path.exists(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")} - {n}.pdf'):
            n += 1
            if os.path.exists(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")} - {n}.pdf'):
                n += 1
                pgto1.save(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")} - {n}.docx')
                docx2pdf.convert(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")} - {n}.docx',
                                 pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")} - {n}.pdf')
                os.remove(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")} - {n}.docx')
                email_pgto(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")} - {n}.pdf', somas, data)
            else:
                pgto1.save(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")} - {n}.docx')
                docx2pdf.convert(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")} - {n}.docx',
                                 pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")} - {n}.pdf')
                os.remove(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")} - {n}.docx')
                email_pgto(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")} - {n}.pdf', somas, data)
        else:
            pgto1.save(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")} - {n}.docx')
            docx2pdf.convert(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")} - {n}.docx',
                             pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")} - {n}.pdf')
            os.remove(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")} - {n}.docx')
            email_pgto(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")} - {n}.pdf', somas, data)
    else:
        pgto1.save(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")}.docx')
        docx2pdf.convert(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")}.docx',
                         pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")}.pdf')
        os.remove(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")}.docx')
        email_pgto(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")}.pdf', somas, data)


def doispgtos(somas, data):
    mesext = {'01': 'JAN', '02': 'FEV', '03': 'MAR', '04': 'ABR', '05': 'MAI', '06': 'JUN',
              '07': 'JUL', '08': 'AGO', '09': 'SET', '10': 'OUT', '11': 'NOV', '12': 'DEZ'}
    dia, mes, ano = data.split('/')
    hoje = dt.strftime(dt.today(), '%d/%m/%Y')
    competencia = data[-7:]
    pasta_pgto = rf'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\04 - Folha de Pgto\{ano}\{mes} - {mesext[mes]}\Pedidos de pagamento'

    modelo = os.path.relpath(
        rf'C:\Users\{os.getlogin()}\PycharmProjects\AutomacaoCia\src\models\static\files\solicitpgto_modelocapa2.docx')
    pgto2 = docx.Document(modelo)
    pagamento = list(somas)[0]
    pagamento2 = list(somas)[1]
    totalpgto = list(somas)[2]
    # # Alterar Modelo e Salvar na pasta do mes correspondente
    for parag in pgto2.paragraphs:
        if 'tipo1' in parag.text:
            inline = parag.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if 'tipo1' in inline[i].text:
                    text = inline[i].text.replace('tipo1', pagamento)
                    inline[i].text = text
        if '#tipo2_pgto' in parag.text:
            inline = parag.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if '#tipo2_pgto' in inline[i].text:
                    text = inline[i].text.replace('#tipo2_pgto', pagamento2)
                    inline[i].text = text
        if 'valor1' in parag.text:
            inline = parag.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if 'valor1' in inline[i].text:
                    text = inline[i].text.replace('valor1',
                                                  str(somas[pagamento]).replace(',', '_').replace('.', ',').replace('_',
                                                                                                                    '.'))
                    inline[i].text = text
        if '#valor2' in parag.text:
            inline = parag.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if '#valor2' in inline[i].text:
                    text = inline[i].text.replace('#valor2',
                                                  str(somas[pagamento2]).replace(',', '_').replace('.', ',').replace(
                                                      '_', '.'))
                    inline[i].text = text

        if '#valortot' in parag.text:
            inline = parag.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if '#valortot' in inline[i].text:
                    text = inline[i].text.replace('#valortot',
                                                  str(somas[totalpgto]).replace(',', '_').replace('.', ',').replace('_',
                                                                                                                    '.'))
                    inline[i].text = text
        if 'textototal' in parag.text:
            inline = parag.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if 'textototal' in inline[i].text:
                    text = inline[i].text.replace('textototal',
                                                  str(somas[totalpgto]).replace(',', '_').replace('.', ',').replace('_',
                                                                                                                    '.'))
                    inline[i].text = text
        if 'valorum' in parag.text:
            inline = parag.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if 'valorum' in inline[i].text:
                    text = inline[i].text.replace('valorum',
                                                  str(somas[pagamento]).replace(',', '_').replace('.', ',').replace('_',
                                                                                                                    '.'))
                    inline[i].text = text
        if '#extenso' in parag.text:
            inline = parag.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if '#extenso' in inline[i].text:
                    text = inline[i].text.replace('#extenso',
                                                  valor_por_extenso(
                                                      str(somas[totalpgto]).replace(',', '_').replace('.', ',').replace(
                                                          '_', '.')))
                    inline[i].text = text
        if '#competencia' in parag.text:
            inline = parag.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if '#competencia' in inline[i].text:
                    text = inline[i].text.replace('#competencia', competencia)
                    inline[i].text = text
        if '#dia_pgto' in parag.text:
            inline = parag.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if '#dia_pgto' in inline[i].text:
                    text = inline[i].text.replace('#dia_pgto', data)
                    inline[i].text = text
        if '#hoje' in parag.text:
            inline = parag.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if '#hoje' in inline[i].text:
                    text = inline[i].text.replace('#hoje', hoje)
                    inline[i].text = text
        if '#tipo1_pgto' in parag.text:
            inline = parag.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if '#tipo1_pgto' in inline[i].text:
                    text = inline[i].text.replace('#tipo1_pgto', pagamento)
                    inline[i].text = text
        if 'tipo1' in parag.text:
            inline = parag.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if 'tipo1' in inline[i].text:
                    text = inline[i].text.replace('tipo1', pagamento)
                    inline[i].text = text
    n = 0
    if os.path.exists(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")}.pdf'):
        n += 1
        if os.path.exists(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")} - {n}.pdf'):
            n += 1
            if os.path.exists(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")} - {n}.pdf'):
                n += 1
                pgto2.save(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")} - {n}.docx')
                docx2pdf.convert(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")} - {n}.docx',
                                 pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")} - {n}.pdf')
                os.remove(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")} - {n}.docx')
                email_pgto(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")} - {n}.pdf', somas, data)
            else:
                pgto2.save(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")} - {n}.docx')
                docx2pdf.convert(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")} - {n}.docx',
                                 pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")} - {n}.pdf')
                os.remove(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")} - {n}.docx')
                email_pgto(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")} - {n}.pdf', somas, data)
        else:
            pgto2.save(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")} - {n}.docx')
            docx2pdf.convert(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")} - {n}.docx',
                             pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")} - {n}.pdf')
            os.remove(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")} - {n}.docx')
            email_pgto(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")} - {n}.pdf', somas, data)
    else:
        pgto2.save(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")}.docx')
        docx2pdf.convert(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")}.docx',
                         pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")}.pdf')
        os.remove(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")}.docx')
        email_pgto(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")}.pdf', somas, data)


def trespgtos(somas, data):
    mesext = {'01': 'JAN', '02': 'FEV', '03': 'MAR', '04': 'ABR', '05': 'MAI', '06': 'JUN',
              '07': 'JUL', '08': 'AGO', '09': 'SET', '10': 'OUT', '11': 'NOV', '12': 'DEZ'}
    dia, mes, ano = data.split('/')
    hoje = dt.strftime(dt.today(), '%d/%m/%Y')
    competencia = data[-7:]
    pasta_pgto = rf'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\04 - Folha de Pgto\{ano}\{mes} - {mesext[mes]}\Pedidos de pagamento'
    modelo = os.path.relpath(
        rf'C:\Users\{os.getlogin()}\PycharmProjects\AutomacaoCia\src\models\static\files\solicitpgto_modelocapa3.docx')
    pgto3 = docx.Document(modelo)
    pagamento = list(somas)[0]
    pagamento2 = list(somas)[1]
    pagamento3 = list(somas)[2]
    totalpgto = list(somas)[3]

    # # Alterar Modelo e Salvar na pasta do mes correspondente
    for parag in pgto3.paragraphs:
        if 'tipo1' in parag.text:
            inline = parag.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if 'tipo1' in inline[i].text:
                    text = inline[i].text.replace('tipo1', pagamento)
                    inline[i].text = text
        if '#tipo2_pgto' in parag.text:
            inline = parag.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if '#tipo2_pgto' in inline[i].text:
                    text = inline[i].text.replace('#tipo2_pgto', pagamento2)
                    inline[i].text = text
        if '#tipo3_pgto' in parag.text:
            inline = parag.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if '#tipo3_pgto' in inline[i].text:
                    text = inline[i].text.replace('#tipo3_pgto', pagamento3)
                    inline[i].text = text
        if 'valor1' in parag.text:
            inline = parag.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if 'valor1' in inline[i].text:
                    text = inline[i].text.replace('valor1',
                                                  str(somas[pagamento]).replace(',', '_').replace('.', ',').replace('_',
                                                                                                                    '.'))
                    inline[i].text = text
        if '#valor2' in parag.text:
            inline = parag.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if '#valor2' in inline[i].text:
                    text = inline[i].text.replace('#valor2',
                                                  str(somas[pagamento2]).replace(',', '_').replace('.', ',').replace(
                                                      '_', '.'))
                    inline[i].text = text
        if '#valor3' in parag.text:
            inline = parag.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if '#valor3' in inline[i].text:
                    text = inline[i].text.replace('#valor3',
                                                  str(somas[pagamento3]).replace(',', '_').replace('.', ',').replace(
                                                      '_', '.'))
                    inline[i].text = text
        if '#valortot' in parag.text:
            inline = parag.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if '#valortot' in inline[i].text:
                    text = inline[i].text.replace('#valortot',
                                                  str(somas[totalpgto]).replace(',', '_').replace('.', ',').replace('_',
                                                                                                                    '.'))
                    inline[i].text = text
        if 'textototal' in parag.text:
            inline = parag.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if 'textototal' in inline[i].text:
                    text = inline[i].text.replace('textototal',
                                                  str(somas[totalpgto]).replace(',', '_').replace('.', ',').replace('_',
                                                                                                                    '.'))
                    inline[i].text = text
        if 'valorum' in parag.text:
            inline = parag.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if 'valorum' in inline[i].text:
                    text = inline[i].text.replace('valorum',
                                                  str(somas[pagamento]).replace(',', '_').replace('.', ',').replace('_',
                                                                                                                    '.'))
                    inline[i].text = text
        if '#extenso' in parag.text:
            inline = parag.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if '#extenso' in inline[i].text:
                    text = inline[i].text.replace('#extenso',
                                                  valor_por_extenso(
                                                      str(somas[totalpgto]).replace(',', '_').replace('.', ',').replace(
                                                          '_', '.')))
                    inline[i].text = text
        if '#competencia' in parag.text:
            inline = parag.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if '#competencia' in inline[i].text:
                    text = inline[i].text.replace('#competencia', competencia)
                    inline[i].text = text
        if '#dia_pgto' in parag.text:
            inline = parag.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if '#dia_pgto' in inline[i].text:
                    text = inline[i].text.replace('#dia_pgto', data)
                    inline[i].text = text
        if '#hoje' in parag.text:
            inline = parag.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if '#hoje' in inline[i].text:
                    text = inline[i].text.replace('#hoje', hoje)
                    inline[i].text = text
        if '#tipo1_pgto' in parag.text:
            inline = parag.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if '#tipo1_pgto' in inline[i].text:
                    text = inline[i].text.replace('#tipo1_pgto', pagamento)
                    inline[i].text = text
        if 'tipo1' in parag.text:
            inline = parag.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if 'tipo1' in inline[i].text:
                    text = inline[i].text.replace('tipo1', pagamento)
                    inline[i].text = text
    n = 0
    if os.path.exists(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")}.pdf'):
        n += 1
        if os.path.exists(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")} - {n}.pdf'):
            n += 1
            if os.path.exists(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")} - {n}.pdf'):
                n += 1
                pgto3.save(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")} - {n}.docx')
                docx2pdf.convert(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")} - {n}.docx',
                                 pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")} - {n}.pdf')
                os.remove(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")} - {n}.docx')
                email_pgto(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")} - {n}.pdf', somas, data)
            else:
                pgto3.save(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")} - {n}.docx')
                docx2pdf.convert(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")} - {n}.docx',
                                 pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")} - {n}.pdf')
                os.remove(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")} - {n}.docx')
                email_pgto(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")} - {n}.pdf', somas, data)
        else:
            pgto3.save(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")} - {n}.docx')
            docx2pdf.convert(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")} - {n}.docx',
                             pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")} - {n}.pdf')
            os.remove(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")} - {n}.docx')
            email_pgto(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")} - {n}.pdf', somas, data)
    else:
        pgto3.save(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")}.docx')
        docx2pdf.convert(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")}.docx',
                         pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")}.pdf')
        os.remove(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")}.docx')
        email_pgto(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")}.pdf', somas, data)


def quatropgtos(somas, data):
    mesext = {'01': 'JAN', '02': 'FEV', '03': 'MAR', '04': 'ABR', '05': 'MAI', '06': 'JUN',
              '07': 'JUL', '08': 'AGO', '09': 'SET', '10': 'OUT', '11': 'NOV', '12': 'DEZ'}
    dia, mes, ano = data.split('/')
    hoje = dt.strftime(dt.today(), '%d/%m/%Y')
    competencia = data[-7:]
    pasta_pgto = rf'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\04 - Folha de Pgto\{ano}\{mes} - {mesext[mes]}\Pedidos de pagamento'
    modelo = os.path.relpath(
        rf'C:\Users\{os.getlogin()}\PycharmProjects\AutomacaoCia\src\models\static\files\solicitpgto_modelocapa4.docx')
    pgto4 = docx.Document(modelo)
    pagamento = list(somas)[0]
    pagamento2 = list(somas)[1]
    pagamento3 = list(somas)[2]
    pagamento4 = list(somas)[3]
    totalpgto = list(somas)[4]

    # # Alterar Modelo e Salvar na pasta do mes correspondente
    for parag in pgto4.paragraphs:
        if '#tipo1_pgto' in parag.text:
            inline = parag.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if '#tipo1_pgto' in inline[i].text:
                    text = inline[i].text.replace('#tipo1_pgto', pagamento)
                    inline[i].text = text
        if '#valor' in parag.text:
            inline = parag.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if '#valor' in inline[i].text:
                    text = inline[i].text.replace('#valor',
                                                  str(somas[pagamento]).replace(',', '_').replace('.', ',').replace('_',
                                                                                                                    '.'))
                    inline[i].text = text
        if '#extenso' in parag.text:
            inline = parag.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if '#extenso' in inline[i].text:
                    text = inline[i].text.replace('#extenso',
                                                  valor_por_extenso(
                                                      str(somas[pagamento]).replace(',', '_').replace('.', ',').replace(
                                                          '_', '.')))
                    inline[i].text = text
        if '#competencia' in parag.text:
            inline = parag.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if '#competencia' in inline[i].text:
                    text = inline[i].text.replace('#competencia', competencia)
                    inline[i].text = text
        if '#dia_pgto' in parag.text:
            inline = parag.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if '#dia_pgto' in inline[i].text:
                    text = inline[i].text.replace('#dia_pgto', data)
                    inline[i].text = text
        if '#hoje' in parag.text:
            inline = parag.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if '#hoje' in inline[i].text:
                    text = inline[i].text.replace('#hoje', hoje)
                    inline[i].text = text
    n = 0
    if os.path.exists(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")}.pdf'):
        n += 1
        if os.path.exists(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")} - {n}.pdf'):
            n += 1
            if os.path.exists(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")} - {n}.pdf'):
                n += 1
                pgto4.save(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")} - {n}.docx')
                docx2pdf.convert(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")} - {n}.docx',
                                 pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")} - {n}.pdf')
                os.remove(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")} - {n}.docx')
                email_pgto(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")} - {n}.pdf', somas, data)
            else:
                pgto4.save(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")} - {n}.docx')
                docx2pdf.convert(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")} - {n}.docx',
                                 pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")} - {n}.pdf')
                os.remove(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")} - {n}.docx')
                email_pgto(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")} - {n}.pdf', somas, data)
        else:
            pgto4.save(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")} - {n}.docx')
            docx2pdf.convert(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")} - {n}.docx',
                             pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")} - {n}.pdf')
            os.remove(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")} - {n}.docx')
            email_pgto(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")} - {n}.pdf', somas, data)
    else:
        pgto4.save(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")}.docx')
        docx2pdf.convert(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")}.docx',
                         pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")}.pdf')
        os.remove(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")}.docx')
        email_pgto(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")}.pdf', somas, data)


def cincopgtos(somas, data):
    mesext = {'01': 'JAN', '02': 'FEV', '03': 'MAR', '04': 'ABR', '05': 'MAI', '06': 'JUN',
              '07': 'JUL', '08': 'AGO', '09': 'SET', '10': 'OUT', '11': 'NOV', '12': 'DEZ'}
    dia, mes, ano = data.split('/')
    hoje = dt.strftime(dt.today(), '%d/%m/%Y')
    competencia = data[-7:]
    pasta_pgto = rf'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\04 - Folha de Pgto\{ano}\{mes} - {mesext[mes]}\Pedidos de pagamento'
    modelo = os.path.relpath(
        rf'C:\Users\{os.getlogin()}\PycharmProjects\AutomacaoCia\src\models\static\files\solicitpgto_modelocapa5.docx')
    pgto5 = docx.Document(modelo)
    pagamento = list(somas)[0]
    pagamento2 = list(somas)[1]
    pagamento3 = list(somas)[2]
    pagamento4 = list(somas)[3]
    pagamento5 = list(somas)[4]
    totalpgto = list(somas)[5]

    # # Alterar Modelo e Salvar na pasta do mes correspondente
    for parag in pgto5.paragraphs:
        if '#tipo1_pgto' in parag.text:
            inline = parag.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if '#tipo1_pgto' in inline[i].text:
                    text = inline[i].text.replace('#tipo1_pgto', pagamento)
                    inline[i].text = text
        if '#valor' in parag.text:
            inline = parag.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if '#valor' in inline[i].text:
                    text = inline[i].text.replace('#valor',
                                                  str(somas[pagamento]).replace(',', '_').replace('.', ',').replace('_',
                                                                                                                    '.'))
                    inline[i].text = text
        if '#extenso' in parag.text:
            inline = parag.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if '#extenso' in inline[i].text:
                    text = inline[i].text.replace('#extenso',
                                                  valor_por_extenso(
                                                      str(somas[pagamento]).replace(',', '_').replace('.', ',').replace(
                                                          '_', '.')))
                    inline[i].text = text
        if '#competencia' in parag.text:
            inline = parag.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if '#competencia' in inline[i].text:
                    text = inline[i].text.replace('#competencia', competencia)
                    inline[i].text = text
        if '#dia_pgto' in parag.text:
            inline = parag.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if '#dia_pgto' in inline[i].text:
                    text = inline[i].text.replace('#dia_pgto', data)
                    inline[i].text = text
        if '#hoje' in parag.text:
            inline = parag.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if '#hoje' in inline[i].text:
                    text = inline[i].text.replace('#hoje', hoje)
                    inline[i].text = text
    n = 0
    if os.path.exists(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")}.pdf'):
        n += 1
        if os.path.exists(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")} - {n}.pdf'):
            n += 1
            if os.path.exists(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")} - {n}.pdf'):
                n += 1
                pgto5.save(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")} - {n}.docx')
                docx2pdf.convert(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")} - {n}.docx',
                                 pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")} - {n}.pdf')
                os.remove(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")} - {n}.docx')
                email_pgto(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")} - {n}.pdf', somas, data)
            else:
                pgto5.save(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")} - {n}.docx')
                docx2pdf.convert(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")} - {n}.docx',
                                 pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")} - {n}.pdf')
                os.remove(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")} - {n}.docx')
                email_pgto(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")} - {n}.pdf', somas, data)
        else:
            pgto5.save(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")} - {n}.docx')
            docx2pdf.convert(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")} - {n}.docx',
                             pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")} - {n}.pdf')
            os.remove(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")} - {n}.docx')
            email_pgto(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")} - {n}.pdf', somas, data)
    else:
        pgto5.save(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")}.docx')
        docx2pdf.convert(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")}.docx',
                         pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")}.pdf')
        os.remove(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")}.docx')
        email_pgto(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")}.pdf', somas, data)


def seispgtos(somas, data):
    mesext = {'01': 'JAN', '02': 'FEV', '03': 'MAR', '04': 'ABR', '05': 'MAI', '06': 'JUN',
              '07': 'JUL', '08': 'AGO', '09': 'SET', '10': 'OUT', '11': 'NOV', '12': 'DEZ'}
    dia, mes, ano = data.split('/')
    hoje = dt.strftime(dt.today(), '%d/%m/%Y')
    competencia = data[-7:]
    pasta_pgto = rf'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\04 - Folha de Pgto\{ano}\{mes} - {mesext[mes]}\Pedidos de pagamento'
    modelo = os.path.relpath(
        rf'C:\Users\{os.getlogin()}\PycharmProjects\AutomacaoCia\src\models\static\files\solicitpgto_modelocapa6.docx')
    pgto6 = docx.Document(modelo)
    pagamento = list(somas)[0]
    pagamento2 = list(somas)[1]
    pagamento3 = list(somas)[2]
    pagamento4 = list(somas)[3]
    pagamento5 = list(somas)[4]
    pagamento6 = list(somas)[5]
    totalpgto = list(somas)[6]

    # # Alterar Modelo e Salvar na pasta do mes correspondente
    for parag in pgto6.paragraphs:
        if '#tipo1_pgto' in parag.text:
            inline = parag.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if '#tipo1_pgto' in inline[i].text:
                    text = inline[i].text.replace('#tipo1_pgto', pagamento)
                    inline[i].text = text
        if '#valor' in parag.text:
            inline = parag.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if '#valor' in inline[i].text:
                    text = inline[i].text.replace('#valor',
                                                  str(somas[pagamento]).replace(',', '_').replace('.', ',').replace('_',
                                                                                                                    '.'))
                    inline[i].text = text
        if '#extenso' in parag.text:
            inline = parag.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if '#extenso' in inline[i].text:
                    text = inline[i].text.replace('#extenso',
                                                  valor_por_extenso(
                                                      str(somas[pagamento]).replace(',', '_').replace('.', ',').replace(
                                                          '_', '.')))
                    inline[i].text = text
        if '#competencia' in parag.text:
            inline = parag.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if '#competencia' in inline[i].text:
                    text = inline[i].text.replace('#competencia', competencia)
                    inline[i].text = text
        if '#dia_pgto' in parag.text:
            inline = parag.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if '#dia_pgto' in inline[i].text:
                    text = inline[i].text.replace('#dia_pgto', data)
                    inline[i].text = text
        if '#hoje' in parag.text:
            inline = parag.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if '#hoje' in inline[i].text:
                    text = inline[i].text.replace('#hoje', hoje)
                    inline[i].text = text
    n = 0
    if os.path.exists(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")}.pdf'):
        n += 1
        if os.path.exists(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")} - {n}.pdf'):
            n += 1
            if os.path.exists(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")} - {n}.pdf'):
                n += 1
                pgto6.save(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")} - {n}.docx')
                docx2pdf.convert(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")} - {n}.docx',
                                 pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")} - {n}.pdf')
                os.remove(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")} - {n}.docx')
                email_pgto(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")} - {n}.pdf', somas, data)
            else:
                pgto6.save(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")} - {n}.docx')
                docx2pdf.convert(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")} - {n}.docx',
                                 pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")} - {n}.pdf')
                os.remove(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")} - {n}.docx')
                email_pgto(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")} - {n}.pdf', somas, data)
        else:
            pgto6.save(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")} - {n}.docx')
            docx2pdf.convert(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")} - {n}.docx',
                             pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")} - {n}.pdf')
            os.remove(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")} - {n}.docx')
            email_pgto(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")} - {n}.pdf', somas, data)
    else:
        pgto6.save(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")}.docx')
        docx2pdf.convert(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")}.docx',
                         pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")}.pdf')
        os.remove(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")}.docx')
        email_pgto(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")}.pdf', somas, data)


def setepgtos(somas, data):
    mesext = {'01': 'JAN', '02': 'FEV', '03': 'MAR', '04': 'ABR', '05': 'MAI', '06': 'JUN',
              '07': 'JUL', '08': 'AGO', '09': 'SET', '10': 'OUT', '11': 'NOV', '12': 'DEZ'}
    dia, mes, ano = data.split('/')
    hoje = dt.strftime(dt.today(), '%d/%m/%Y')
    competencia = data[-7:]
    pasta_pgto = rf'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\04 - Folha de Pgto\{ano}\{mes} - {mesext[mes]}\Pedidos de pagamento'
    modelo = os.path.relpath(
        rf'C:\Users\{os.getlogin()}\PycharmProjects\AutomacaoCia\src\models\static\files\solicitpgto_modelocapa7.docx')
    pgto7 = docx.Document(modelo)
    pagamento = list(somas)[0]
    pagamento2 = list(somas)[1]
    pagamento3 = list(somas)[2]
    pagamento4 = list(somas)[3]
    pagamento5 = list(somas)[4]
    pagamento6 = list(somas)[5]
    pagamento7 = list(somas)[6]
    totalpgto = list(somas)[7]

    # # Alterar Modelo e Salvar na pasta do mes correspondente
    for parag in pgto7.paragraphs:
        if '#tipo1_pgto' in parag.text:
            inline = parag.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if '#tipo1_pgto' in inline[i].text:
                    text = inline[i].text.replace('#tipo1_pgto', pagamento)
                    inline[i].text = text
        if '#valor' in parag.text:
            inline = parag.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if '#valor' in inline[i].text:
                    text = inline[i].text.replace('#valor',
                                                  str(somas[pagamento]).replace(',', '_').replace('.', ',').replace('_',
                                                                                                                    '.'))
                    inline[i].text = text
        if '#extenso' in parag.text:
            inline = parag.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if '#extenso' in inline[i].text:
                    text = inline[i].text.replace('#extenso',
                                                  valor_por_extenso(
                                                      str(somas[pagamento]).replace(',', '_').replace('.', ',').replace(
                                                          '_', '.')))
                    inline[i].text = text
        if '#competencia' in parag.text:
            inline = parag.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if '#competencia' in inline[i].text:
                    text = inline[i].text.replace('#competencia', competencia)
                    inline[i].text = text
        if '#dia_pgto' in parag.text:
            inline = parag.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if '#dia_pgto' in inline[i].text:
                    text = inline[i].text.replace('#dia_pgto', data)
                    inline[i].text = text
        if '#hoje' in parag.text:
            inline = parag.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if '#hoje' in inline[i].text:
                    text = inline[i].text.replace('#hoje', hoje)
                    inline[i].text = text
    n = 0
    if os.path.exists(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")}.pdf'):
        n += 1
        if os.path.exists(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")} - {n}.pdf'):
            n += 1
            if os.path.exists(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")} - {n}.pdf'):
                n += 1
                pgto7.save(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")} - {n}.docx')
                docx2pdf.convert(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")} - {n}.docx',
                                 pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")} - {n}.pdf')
                os.remove(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")} - {n}.docx')
                email_pgto(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")} - {n}.pdf', somas, data)
            else:
                pgto7.save(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")} - {n}.docx')
                docx2pdf.convert(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")} - {n}.docx',
                                 pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")} - {n}.pdf')
                os.remove(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")} - {n}.docx')
                email_pgto(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")} - {n}.pdf', somas, data)
        else:
            pgto7.save(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")} - {n}.docx')
            docx2pdf.convert(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")} - {n}.docx',
                             pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")} - {n}.pdf')
            os.remove(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")} - {n}.docx')
            email_pgto(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")} - {n}.pdf', somas, data)
    else:
        pgto7.save(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")}.docx')
        docx2pdf.convert(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")}.docx',
                         pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")}.pdf')
        os.remove(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")}.docx')
        email_pgto(pasta_pgto + f'\\Pedido de Pagamento - {str(data).replace("/", ".")}.pdf', somas, data)


def gerar_pedido_pgto_por_arquivo(data: str, caminho_arq1: str, caminho_arq2='', caminho_arq3='', caminho_arq4=''):
    """
    Generate document of payment from .xlsx files.
    :param data: Payment date
    :param caminho_arq1: file path
    :param caminho_arq2: file path
    :param caminho_arq3: file path
    :param caminho_arq4: file path
    """
    somas = {}
    sig = signature(gerar_pedido_pgto_por_arquivo)
    params = sig.parameters

    nomepgto = {
        1: 'Salário',
        2: 'Férias',
        3: 'Vale Transporte',
        4: 'Vale Alimentação',
        5: 'Comissão',
        6: '13º salário',
        7: 'Bolsa Estágio',
        8: 'Bônus',
        9: 'Adiantamento Salarial',
        10: 'Rescisão',
        11: 'Bolsa Auxílio',
        12: 'Pensão Alimentícia',
        13: 'Pgto em C/C',
        14: 'Remuneração'
    }

    quantidade_arquivos = -2
    for par in params:
        if params[par] != '':
            quantidade_arquivos += 1
    if quantidade_arquivos == 1:
        plan1 = l_w(caminho_arq1, read_only=False)
        sh1 = plan1['Planilha1']
        for item in sh1['F']:
            if nomepgto[item.value] in somas:
                somas[nomepgto[item.value]] = somas[nomepgto[item.value]] + sh1[f'G{item.row}'].value
            else:
                somas[nomepgto[item.value]] = sh1[f'G{item.row}'].value

    if quantidade_arquivos == 2:
        plan1 = l_w(caminho_arq1, read_only=False)
        sh1 = plan1['Planilha1']
        plan2 = l_w(caminho_arq2, read_only=False)
        sh2 = plan2['Planilha1']
        for item in sh1['F']:
            if nomepgto[item.value] in somas:
                somas[nomepgto[item.value]] = somas[nomepgto[item.value]] + sh1[f'G{item.row}'].value
            else:
                somas[nomepgto[item.value]] = sh1[f'G{item.row}'].value
        for item in sh2['F']:
            if nomepgto[item.value] in somas:
                somas[nomepgto[item.value]] = somas[nomepgto[item.value]] + sh2[f'G{item.row}'].value
            else:
                somas[nomepgto[item.value]] = sh2[f'G{item.row}'].value
    if quantidade_arquivos == 3:
        plan1 = l_w(caminho_arq1, read_only=False)
        sh1 = plan1['Planilha1']
        plan2 = l_w(caminho_arq2, read_only=False)
        sh2 = plan2['Planilha1']
        plan3 = l_w(caminho_arq3, read_only=False)
        sh3 = plan3['Planilha1']
        for item in sh1['F']:
            if nomepgto[item.value] in somas:
                somas[nomepgto[item.value]] = somas[nomepgto[item.value]] + sh1[f'G{item.row}'].value
            else:
                somas[nomepgto[item.value]] = sh1[f'G{item.row}'].value
        for item in sh2['F']:
            if nomepgto[item.value] in somas:
                somas[nomepgto[item.value]] = somas[nomepgto[item.value]] + sh2[f'G{item.row}'].value
            else:
                somas[nomepgto[item.value]] = sh2[f'G{item.row}'].value
        for item in sh3['F']:
            if nomepgto[item.value] in somas:
                somas[nomepgto[item.value]] = somas[nomepgto[item.value]] + sh3[f'G{item.row}'].value
            else:
                somas[nomepgto[item.value]] = sh3[f'G{item.row}'].value
    if quantidade_arquivos == 4:
        plan1 = l_w(caminho_arq1, read_only=False)
        sh1 = plan1['Planilha1']
        plan2 = l_w(caminho_arq2, read_only=False)
        sh2 = plan2['Planilha1']
        plan3 = l_w(caminho_arq3, read_only=False)
        sh3 = plan3['Planilha1']
        plan4 = l_w(caminho_arq4, read_only=False)
        sh4 = plan4['Planilha1']
        for item in sh1['F']:
            if nomepgto[item.value] in somas:
                somas[nomepgto[item.value]] = somas[nomepgto[item.value]] + sh1[f'G{item.row}'].value
            else:
                somas[nomepgto[item.value]] = sh1[f'G{item.row}'].value
        for item in sh2['F']:
            if nomepgto[item.value] in somas:
                somas[nomepgto[item.value]] = somas[nomepgto[item.value]] + sh2[f'G{item.row}'].value
            else:
                somas[nomepgto[item.value]] = sh2[f'G{item.row}'].value
        for item in sh3['F']:
            if nomepgto[item.value] in somas:
                somas[nomepgto[item.value]] = somas[nomepgto[item.value]] + sh3[f'G{item.row}'].value
            else:
                somas[nomepgto[item.value]] = sh3[f'G{item.row}'].value
        for item in sh1['F']:
            if nomepgto[item.value] in somas:
                somas[nomepgto[item.value]] = somas[nomepgto[item.value]] + sh1[f'G{item.row}'].value
            else:
                somas[nomepgto[item.value]] = sh1[f'G{item.row}'].value
    quantidade_de_pgtos = len(somas)
    total = 0
    for item in somas:
        total += somas[item]
    somas['Total'] = total
    for item in somas:
        somas[item] = '{0:,.2f}'.format(somas[item])
    qtidades = {
        1: umpgto,
        2: doispgtos,
        3: trespgtos,
        4: quatropgtos,
        5: cincopgtos,
        6: seispgtos,
        7: setepgtos
    }
    qtidades[quantidade_de_pgtos](somas, data)
    tkinter.messagebox.showinfo('Pagamento enviado!', 'Pagamento enviado ao financeiro com sucesso!')

