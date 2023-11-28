from datetime import datetime as dt
from difflib import SequenceMatcher
import docx
import docx2pdf
from email import encoders
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.image import MIMEImage
from email.mime.base import MIMEBase
import locale
from openpyxl import load_workbook as l_w
import os
import pyautogui as pa
import pyperclip as pp
import shutil
from src.models.models import Colaborador, engine
from src.models.modelsfolha import Aulas, enginefolha
from sqlalchemy.orm import sessionmaker
from sqlalchemy.exc import IntegrityError
from src.models.listas import municipios
import smtplib
from src.models.dados_servd import em_rem, em_ti, em_if, em_pnt, k1, host, port, rede, elidermusc, eliderkids, elidernat, \
    elidercross, elidergin, egerentevend, egerenterh, egerentemanut, egerentetec, lidermusc, liderkids, lidernat, lidercross,\
    lidergin, gerentevend, gerenterh, gerentemanut, gerentetec, liderpnt
import tkinter.filedialog
from tkinter import messagebox
import tkinter.filedialog
import time as t
import urllib
from urllib import parse

locale.setlocale(locale.LC_MONETARY, 'pt_BR.UTF-8')
locale.setlocale(locale.LC_ALL, 'pt_BR.UTF-8')
locale.setlocale(locale.LC_MONETARY, 'pt_BR.UTF-8')


def apenas_registrar_funcionario(caminho='', editar=0, ondestou=0, nome='', matricula='', admissao='',
                                 horario='', salario='', cargo='', depto='', tipo_contr='',
                                 hrsem='', hrmens='', agencia='', conta='', digito=''):
    sessions = sessionmaker(bind=engine)
    session = sessions()
    if caminho == '' or nome == '' or matricula == '' or admissao == '' or horario == '' or salario == '' or \
            cargo == '' or depto == '' or tipo_contr == '' or hrsem == '' or hrmens == '':
        tkinter.messagebox.showinfo(
            title='Erro de preenchimento',
            message='Preencha todos os campos antes de cadastrar o funcionário!'
        )
    else:
        wb = l_w(caminho, read_only=False)
        sh = wb['Respostas ao formulário 1']
        num, name = nome.strip().split(' - ')
        linha = int(num)

        # search for the highest compatibility between the city filled in the form and the cities in the lists to
        # define codmunnas value
        est = str(sh[f'AJ{linha}'].value).upper().strip()
        cidade = str(sh[f'L{linha}'].value).title().strip()
        lista = []
        dicion = {}
        for cid in municipios[est]:
            dicion[SequenceMatcher(None, cidade, cid).ratio()] = cid
            lista.append(SequenceMatcher(None, cidade, cid).ratio())
        codmunnas = municipios[str(sh[f'AJ{linha}'].value).upper().strip()][dicion[max(lista)]]

        # search for the highest compatibility between the city filled in the form and the cities in the lists to
        # define codmunend value
        est = str(sh[f'T{linha}'].value).upper().strip()
        cidade = str(sh[f'S{linha}'].value).title().strip()
        listaend = []
        dicionend = {}
        for cid in municipios[est]:
            dicionend[SequenceMatcher(None, cidade, cid).ratio()] = cid
            listaend.append(SequenceMatcher(None, cidade, cid).ratio())
        codmunend = municipios[str(sh[f'T{linha}'].value).upper().strip()][dicionend[max(listaend)]]

        lotacao = {'UNIDADE PARK SUL - QUALQUER DEPARTAMENTO': '0013', 'KIDS': '0010', 'MUSCULAÇÃO': '0007',
                   'ESPORTES E LUTAS': '0008', 'CROSSFIT': '0012', 'GINÁSTICA': '0006', 'GESTANTES': '0008',
                   'RECEPÇÃO': '0003',
                   'FINANCEIRO': '0001', 'TI': '0001', 'MARKETING': '0001', 'MANUTENÇÃO': '0004'}
        if linha:
            try:
                pess = Colaborador(matricula=matricula, nome=name.upper(), admiss=admissao,
                                   nascimento=str(sh[f'D{linha}'].value),
                                   pis=str(int(sh[f'Y{linha}'].value)).zfill(11),
                                   cpf=str(int(sh[f'V{linha}'].value)).zfill(11),
                                   rg=str(int(sh[f'W{linha}'].value)),
                                   emissor=str(sh[f'X{linha}'].value), email=str(sh[f'B{linha}'].value),
                                   genero=str(sh[f'E{linha}'].value),
                                   estado_civil=str(sh[f'F{linha}'].value), cor=str(sh[f'G{linha}'].value),
                                   instru=str(sh[f'J{linha}'].value),
                                   nacional=str(sh[f'K{linha}'].value),
                                   cod_municipionas=codmunnas,
                                   cid_nas=str(sh[f'L{linha}'].value), uf_nas=str(sh[f'AJ{linha}'].value),
                                   pai=str(sh[f'M{linha}'].value).upper(),
                                   mae=str(sh[f'N{linha}'].value).upper(), endereco=str(sh[f'O{linha}'].value),
                                   num=str(int(sh[f'P{linha}'].value)),
                                   bairro=str(sh[f'Q{linha}'].value), cep=str(int(sh[f'R{linha}'].value)),
                                   cidade=str(sh[f'S{linha}'].value),
                                   uf=str(sh[f'T{linha}'].value),
                                   cod_municipioend=codmunend,
                                   tel=str(int(sh[f'U{linha}'].value)),
                                   tit_eleit=str(sh[f'Z{linha}'].value), zona_eleit=str(sh[f'AA{linha}'].value),
                                   sec_eleit=str(sh[f'AB{linha}'].value),
                                   ctps=str(int(sh[f'AC{linha}'].value)),
                                   serie_ctps=str(sh[f'AD{linha}'].value),
                                   uf_ctps=str(sh[f'AE{linha}'].value),
                                   emiss_ctps=str(sh[f'AF{linha}'].value), depto=depto,
                                   cargo=cargo,
                                   horario=horario, salario=salario, tipo_contr=tipo_contr, hr_sem=hrsem,
                                   hr_mens=hrmens,
                                   ag=agencia, conta=conta, cdigito=digito
                                   )
                session.add(pess)
                session.commit()
                tkinter.messagebox.showinfo(title='Cadastro ok!', message='Cadastro efetuado com sucesso!')
            except IntegrityError:
                tkinter.messagebox.showinfo(title='Erro', message='Funcionário já cadastrado no DB!')


def cadastro_funcionario(caminho='', editar=0, ondestou=0, nome='', matricula='', admissao='',
                         horario='', salario='', cargo='', depto='', tipo_contr='',
                         hrsem='', hrmens='', agencia='', conta='', digito=''):
    sessions = sessionmaker(bind=engine)
    session = sessions()
    if caminho == '' or nome == '' or matricula == '' or admissao == '' or horario == '' or salario == '' or \
            cargo == '' or depto == '' or tipo_contr == '' or hrsem == '' or hrmens == '':
        tkinter.messagebox.showinfo(
            title='Erro de preenchimento',
            message='Preencha todos os campos antes de cadastrar o funcionário!'
        )
    else:
        wb = l_w(caminho, read_only=False)
        sh = wb['Respostas ao formulário 1']
        num, name = nome.strip().split(' - ')
        linha = int(num)

        # search for the highest compatibility between the city filled in the form and the cities in the lists to
        # define codmunnas value
        est = str(sh[f'AJ{linha}'].value).upper().strip()
        cidade = str(sh[f'L{linha}'].value).title().strip()
        lista = []
        dicion = {}
        for cid in municipios[est]:
            dicion[SequenceMatcher(None, cidade, cid).ratio()] = cid
            lista.append(SequenceMatcher(None, cidade, cid).ratio())
        codmunnas = municipios[str(sh[f'AJ{linha}'].value).upper().strip()][dicion[max(lista)]]

        # search for the highest compatibility between the city filled in the form and the cities in the lists to
        # define codmunend value
        est = str(sh[f'T{linha}'].value).upper().strip()
        cidade = str(sh[f'S{linha}'].value).title().strip()
        listaend = []
        dicionend = {}
        for cid in municipios[est]:
            dicionend[SequenceMatcher(None, cidade, cid).ratio()] = cid
            listaend.append(SequenceMatcher(None, cidade, cid).ratio())
        codmunend = municipios[str(sh[f'T{linha}'].value).upper().strip()][dicionend[max(listaend)]]

        lotacao = {'UNIDADE PARK SUL - QUALQUER DEPARTAMENTO': '0013', 'KIDS': '0010', 'MUSCULAÇÃO': '0007',
                   'ESPORTES E LUTAS': '0008', 'CROSSFIT': '0012', 'GINÁSTICA': '0006', 'GESTANTES': '0008',
                   'RECEPÇÃO': '0003',
                   'FINANCEIRO': '0001', 'TI': '0001', 'MARKETING': '0001', 'MANUTENÇÃO': '0004'}
        if editar == 0:
            if linha:
                pess = Colaborador(matricula=matricula, nome=name.upper(), admiss=admissao,
                                   nascimento=str(sh[f'D{linha}'].value),
                                   pis=str(int(sh[f'Y{linha}'].value)).zfill(11),
                                   cpf=str(int(sh[f'V{linha}'].value)).zfill(11),
                                   rg=str(int(sh[f'W{linha}'].value)),
                                   emissor=str(sh[f'X{linha}'].value), email=str(sh[f'B{linha}'].value),
                                   genero=str(sh[f'E{linha}'].value),
                                   estado_civil=str(sh[f'F{linha}'].value), cor=str(sh[f'G{linha}'].value),
                                   instru=str(sh[f'J{linha}'].value),
                                   nacional=str(sh[f'K{linha}'].value),
                                   cod_municipionas=codmunnas,
                                   cid_nas=str(sh[f'L{linha}'].value), uf_nas=str(sh[f'AJ{linha}'].value),
                                   pai=str(sh[f'M{linha}'].value).upper(),
                                   mae=str(sh[f'N{linha}'].value).upper(), endereco=str(sh[f'O{linha}'].value),
                                   num=str(int(sh[f'P{linha}'].value)),
                                   bairro=str(sh[f'Q{linha}'].value), cep=str(int(sh[f'R{linha}'].value)),
                                   cidade=str(sh[f'S{linha}'].value),
                                   uf=str(sh[f'T{linha}'].value),
                                   cod_municipioend=codmunend,
                                   tel=str(int(sh[f'U{linha}'].value)),
                                   tit_eleit=str(sh[f'Z{linha}'].value), zona_eleit=str(sh[f'AA{linha}'].value),
                                   sec_eleit=str(sh[f'AB{linha}'].value),
                                   ctps=str(int(sh[f'AC{linha}'].value)), serie_ctps=str(sh[f'AD{linha}'].value),
                                   uf_ctps=str(sh[f'AE{linha}'].value),
                                   emiss_ctps=str(sh[f'AF{linha}'].value), depto=depto,
                                   cargo=cargo,
                                   horario=horario, salario=salario, tipo_contr=tipo_contr, hr_sem=hrsem,
                                   hr_mens=hrmens,
                                   ag=agencia, conta=conta, cdigito=digito
                                   )
                session.add(pess)
                session.commit()
                pessoa = session.query(Colaborador).filter_by(matricula=matricula).first()
                while 1:
                    if pa.locateOnScreen('../models/static/imgs/Dexion.png'):
                        pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Dexion.png')))
                        break
                    else:
                        t.sleep(5)
                pa.press('alt'), pa.press('a'), pa.press('t'), t.sleep(2), pa.press(
                    'i'), t.sleep(5), pa.write(str(pessoa.matricula)), pa.press('enter')
                if ondestou == 0:
                    t.sleep(20)
                else:
                    t.sleep(60)
                pp.copy(pessoa.nome), pa.hotkey('ctrl', 'v'), pa.press('tab'), pa.write(pessoa.cpf)
                t.sleep(1), pa.press('tab', 3), pa.write(pessoa.genero[0]), pa.press('tab'), pa.write(pessoa.cor)
                t.sleep(1), pa.press('tab', 2), pa.write(pessoa.instru)
                t.sleep(1), pa.press('tab'), pa.write(pessoa.estado_civil)
                if str(pessoa.estado_civil) == '2 - Casado(a)':
                    pa.press('tab', 6)
                else:
                    pa.press('tab', 5)
                pa.write(dt.strftime(dt.strptime(pessoa.nascimento, '%Y-%m-%d %H:%M:%S'), '%d%m%Y'))
                t.sleep(1), pa.press('tab'), pp.copy(pessoa.cid_nas), pa.hotkey('ctrl', 'v'), pa.press('tab')
                t.sleep(1), pa.write(pessoa.uf_nas), pa.press('tab'), pa.write(pessoa.cod_municipionas), pa.press(
                    'tab')
                t.sleep(1), pp.copy(pessoa.pai), pa.hotkey('ctrl', 'v'), pa.press('tab'), pa.write('105'), pa.press(
                    'tab')
                t.sleep(1), pp.copy(pessoa.mae), pa.hotkey('ctrl', 'v'), pa.press('tab'), pa.write('105')
                # #clique em documentos
                while 1:
                    if pa.locateOnScreen('../models/static/imgs/Documentos.png'):
                        pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Documentos.png')))
                        break
                    else:
                        t.sleep(5)
                pa.press('tab'), pa.write(str(pessoa.rg)), pa.press('tab'), pa.write(
                    pessoa.emissor), pa.press('tab', 3), pa.write(pessoa.cod_municipioend),
                pa.press('tab'), pa.write(pessoa.pis), pa.press('enter')
                pa.press('tab'), pa.write(pessoa.tit_eleit), pa.press('tab'), pa.write(pessoa.zona_eleit), pa.press(
                    'tab'), pa.write(
                    pessoa.sec_eleit), pa.press('tab'), pa.write(pessoa.ctps)
                pa.press('tab'), pa.write(pessoa.serie_ctps), pa.press('tab'), pa.write(pessoa.uf_ctps), pa.press(
                    'tab'), pa.write(dt.strftime(dt.strptime(pessoa.emiss_ctps, '%Y-%m-%d %H:%M:%S'),
                                                 '%d%m%Y'))
                pa.press('tab', 5), pa.write('3'), pa.press('tab'), pa.write('341'), pa.press('tab')
                pa.write(pessoa.ag), pa.press('tab'), pa.write(f'{pessoa.conta}-{pessoa.cdigito}'), pa.press('tab')

                # #clique em endereço
                while 1:
                    if pa.locateOnScreen('../models/static/imgs/Endereco.png'):
                        pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Endereco.png')))
                        break
                    else:
                        t.sleep(5)
                pa.press('tab', 2), pp.copy(pessoa.endereco), pa.hotkey('ctrl', 'v'), pa.press(
                    'tab'), pa.write(pessoa.num), pa.press('tab', 2)
                pp.copy(pessoa.bairro), pa.hotkey('ctrl', 'v'), pa.press('tab'), pp.copy(pessoa.cidade), pa.hotkey(
                    'ctrl', 'v'), pa.press('tab'), pa.write(pessoa.uf)
                pa.press('tab'), pa.write(pessoa.cep), pa.press('tab'), pa.write(pessoa.cod_municipioend), pa.press(
                    'tab'), pa.write(str(pessoa.tel)), pa.press('tab', 2), pa.write(pessoa.email)
                # #clique em dados contratuais
                while 1:
                    if pa.locateOnScreen('../models/static/imgs/Contratuais.png'):
                        pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Contratuais.png')))
                        break
                    else:
                        t.sleep(5)
                pa.press('tab'), pa.write(str(pessoa.admiss).replace('/', '')), t.sleep(1)
                pa.press('tab', 8), pa.write('2')
                # #clique em Contrato de Experiência
                while 1:
                    if pa.locateOnScreen('../models/static/imgs/Experiencia.png'):
                        pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Experiencia.png')))
                        break
                    else:
                        t.sleep(5)
                pa.press('tab'), pa.write('45'), pa.press('tab'), pa.write('45'), pa.press(
                    'tab'), pa.press(
                    'space'), pa.press('tab', 6), pa.write('003')
                pa.press('tab'), pa.write(str(pessoa.matricula))
                # #clique em Outros
                while 1:
                    if pa.locateOnScreen('../models/static/imgs/Outros.png'):
                        pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Outros.png')))
                        break
                    else:
                        t.sleep(5)
                t.sleep(2), pa.write('CARGO GERAL'), t.sleep(1), pa.press('tab')
                t.sleep(2), pp.copy(pessoa.cargo), pa.hotkey('ctrl', 'v')
                t.sleep(1), pa.press('tab'), pa.write(pessoa.salario), pa.press('tab')
                if str(pessoa.tipo_contr) == 'Horista':
                    pa.press('1')
                else:
                    pa.press('5')
                pa.press('tab', 2), t.sleep(1), pa.write(str(pessoa.hr_sem)), pa.press('tab'), t.sleep(1)
                pa.write(str(pessoa.hr_mens))
                pa.press('tab', 5), pa.write('00395419000190'), pa.press('tab', 2), pa.write('5')
                # #clique em eventos trabalhistas
                while 1:
                    if pa.locateOnScreen('../models/static/imgs/EVTrab.png'):
                        pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/EVTrab.png')))
                        break
                    else:
                        t.sleep(5)
                t.sleep(1)
                # #clique em lotação
                while 1:
                    if pa.locateOnScreen('../models/static/imgs/Lotacoes.png'):
                        pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Lotacoes.png')))
                        break
                    else:
                        t.sleep(5)
                pa.press('tab'), pa.press('tab'), pa.write('i'), pa.write(str(pessoa.admiss).replace('/', ''))
                t.sleep(1), pa.press('enter'), t.sleep(1)
                pp.copy(lotacao[f'{pessoa.depto}']), pa.hotkey('ctrl', 'v'), pa.press('enter'), pa.write('f')
                pa.press('tab'), pa.write('4')
                pa.press('tab', 6), pa.write('i'), t.sleep(2), pa.press('tab'), pa.write(pessoa.horario)
                t.sleep(3), pa.press('tab', 3), pa.press('enter'), t.sleep(3)
                # #clique em cancelar novo registro de horario
                while 1:
                    if pa.locateOnScreen('../models/static/imgs/Cancelarhor.png'):
                        pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Cancelarhor.png')))
                        break
                    else:
                        t.sleep(5)
                t.sleep(2.5)
                # #clique em salvar lotação
                while 1:
                    if pa.locateOnScreen('../models/static/imgs/Salvarlot.png'):
                        pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Salvarlot.png')))
                        break
                    else:
                        t.sleep(5)
                t.sleep(1)
                # #clique em fechar lotação
                while 1:
                    if pa.locateOnScreen('../models/static/imgs/Fecharlot.png'):
                        pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Fecharlot.png')))
                        break
                    else:
                        t.sleep(5)
                t.sleep(2)
                # #clique em Compatibilidade
                while 1:
                    if pa.locateOnScreen('../models/static/imgs/Compatibilidade.png'):
                        pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Compatibilidade.png')))
                        break
                    else:
                        t.sleep(5)
                t.sleep(1)
                # #clique em Compatibilidade de novo
                while 1:
                    if pa.locateOnScreen('../models/static/imgs/Compatibilidade2.png'):
                        pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Compatibilidade2.png')))
                        break
                    else:
                        t.sleep(5)
                t.sleep(1)
                # #clique em CAGED
                while 1:
                    if pa.locateOnScreen('../models/static/imgs/CAGED.png'):
                        pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/CAGED.png')))
                        break
                    else:
                        t.sleep(5)
                pa.press('tab', 2), pa.write('20'), pa.press('tab'), pa.write('1'), t.sleep(0.5)
                # #clique em RAIS
                while 1:
                    if pa.locateOnScreen('../models/static/imgs/RAIS.png'):
                        pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/RAIS.png')))
                        break
                    else:
                        t.sleep(5)
                pp.copy(pessoa.nome), pa.hotkey('ctrl', 'v'), pa.press('tab', 2), pa.write('2')
                pa.press('tab'), pa.write('10')
                # #clique em Salvar
                while 1:
                    if pa.locateOnScreen('../models/static/imgs/Salvarcadastro.png'):
                        pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Salvarcadastro.png')))
                        break
                    else:
                        t.sleep(5)
                t.sleep(10)
                # #clique em fechar novo cadastro
                while 1:
                    if pa.locateOnScreen('../models/static/imgs/Fecharnovo1.png'):
                        pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Fecharnovo1.png')))
                        break
                    else:
                        t.sleep(5)
                t.sleep(2)
                # #clique em fechar trabalhadores
                while 1:
                    if pa.locateOnScreen('../models/static/imgs/Fechartrab1.png'):
                        pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Fechartrab1.png')))
                        break
                    else:
                        t.sleep(5)
                t.sleep(0.5)
                try:
                    os.makedirs(
                        r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e '
                        r'Férias\000 - Pastas Funcionais\00 - ATIVOS\1 - Ainda nao iniciaram\{}'.format(pessoa.nome))
                    os.makedirs(
                        r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e '
                        r'Férias\000 - Pastas Funcionais\00 - ATIVOS\1 - Ainda nao iniciaram\{}\Atestados'.format(
                            pessoa.nome))
                    os.makedirs(
                        r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e '
                        r'Férias\000 - Pastas Funcionais\00 - ATIVOS\1 - Ainda nao iniciaram\{}\Diversos'.format(
                            pessoa.nome))
                    os.makedirs(
                        r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e '
                        r'Férias\000 - Pastas Funcionais\00 - ATIVOS\1 - Ainda nao iniciaram\{}\Contratuais'.format(
                            pessoa.nome))
                    os.makedirs(
                        r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e '
                        r'Férias\000 - Pastas Funcionais\00 - ATIVOS\1 - Ainda nao iniciaram\{}\Férias'.format(
                            pessoa.nome))
                    os.makedirs(
                        r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e '
                        r'Férias\000 - Pastas Funcionais\00 - ATIVOS\1 - Ainda nao iniciaram\{}\Pontos'.format(
                            pessoa.nome))
                    os.makedirs(
                        r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e '
                        r'Férias\000 - Pastas Funcionais\00 - ATIVOS\1 - Ainda nao iniciaram\{}\Recibos'.format(
                            pessoa.nome))
                    os.makedirs(
                        r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e '
                        r'Férias\000 - Pastas Funcionais\00 - ATIVOS\1 - Ainda nao iniciaram\{}\Rescisão'.format(
                            pessoa.nome))
                except FileExistsError:
                    pass
                while 1:
                    if pa.locateOnScreen('../models/static/imgs/pyt.png'):
                        pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/pyt.png')))
                        break
                    else:
                        t.sleep(5)
                tkinter.messagebox.showinfo(title='Cadastro ok!', message='Cadastro realizado com sucesso!')
        else:
            num, name = nome.strip().split(' - ')
            linha = int(num)
            if linha:
                pessoa = session.query(Colaborador).filter_by(matricula=matricula).first()
                pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Dexion.png')))
                pa.press('alt'), pa.press('a'), pa.press('t'), t.sleep(2), pa.press(
                    'a'), t.sleep(5), pa.write(str(pessoa.matricula)), pa.press('enter')
                if ondestou == 0:
                    t.sleep(15)
                else:
                    t.sleep(40)
                pp.copy(pessoa.nome), pa.hotkey('ctrl', 'v'), pa.press('tab'), pa.write(pessoa.cpf)
                pa.press('tab', 3), pa.write(pessoa.genero[0]), pa.press('tab'), pa.write(pessoa.cor)
                pa.press('tab', 2), pa.write(pessoa.instru)
                pa.press('tab'), pa.write(pessoa.estado_civil)
                if str(pessoa.estado_civil) == '2 - Casado(a)':
                    pa.press('tab', 6)
                else:
                    pa.press('tab', 5)
                pa.write(dt.strftime(dt.strptime(pessoa.nascimento, '%Y-%m-%d %H:%M:%S'), '%d%m%Y'))
                pa.press('tab'), pp.copy(pessoa.cid_nas), pa.hotkey('ctrl', 'v'), pa.press('tab')
                pa.write(pessoa.uf_nas), pa.press('tab'), pa.write(pessoa.cod_municipionas), pa.press('tab')
                pp.copy(pessoa.pai), pa.hotkey('ctrl', 'v'), pa.press('tab'), pa.write('105'), pa.press('tab')
                pp.copy(pessoa.mae), pa.hotkey('ctrl', 'v'), pa.press('tab'), pa.write('105')
                # #clique em documentos
                pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Documentos.png')))
                pa.press('tab'), pa.write(str(pessoa.rg)), pa.press('tab'), pa.write(
                    pessoa.emissor), pa.press('tab', 3), pa.write(pessoa.cod_municipioend),
                pa.press('tab'), pa.write(pessoa.pis), pa.press('enter')
                pa.press('tab'), pa.write(pessoa.tit_eleit), pa.press('tab'), pa.write(pessoa.zona_eleit), pa.press(
                    'tab'), pa.write(
                    pessoa.sec_eleit), pa.press('tab'), pa.write(pessoa.ctps)
                pa.press('tab'), pa.write(pessoa.serie_ctps), pa.press('tab'), pa.write(pessoa.uf_ctps), pa.press(
                    'tab'), pa.write(dt.strftime(dt.strptime(pessoa.emiss_ctps, '%Y-%m-%d %H:%M:%S'),
                                                 '%d%m%Y'))
                pa.press('tab', 5), pa.write('3'), pa.press('tab'), pa.write('341'), pa.press('tab')
                pa.write(pessoa.ag), pa.press('tab'), pa.write(f'{pessoa.conta}-{pessoa.cdigito}'), pa.press('tab')

                # #clique em endereço
                pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Endereco.png')))
                pa.press('tab', 2), pp.copy(pessoa.endereco), pa.hotkey('ctrl', 'v'), pa.press(
                    'tab'), pa.write(pessoa.num), pa.press('tab', 2)
                pp.copy(pessoa.bairro), pa.hotkey('ctrl', 'v'), pa.press('tab'), pp.copy(pessoa.cidade), pa.hotkey(
                    'ctrl', 'v'), pa.press('tab'), pa.write(pessoa.uf)
                pa.press('tab'), pa.write(pessoa.cep), pa.press('tab'), pa.write(pessoa.cod_municipioend), pa.press(
                    'tab'), pa.write(str(pessoa.tel)), pa.press('tab', 2), pa.write(pessoa.email)
                # #clique em dados contratuais
                pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Contratuais.png')))
                pa.press('tab'), pa.write(str(pessoa.admiss).replace('/', '')), t.sleep(1)
                pa.press('tab', 10), pa.write('2')
                # #clique em Contrato de Experiência
                try:
                    pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Experiencia.png')))
                except pa.ImageNotFoundException:
                    t.sleep(5)
                    pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Experiencia.png')))
                pa.press('tab'), pa.write('45'), pa.press('tab'), pa.write('45'), pa.press(
                    'tab'), pa.press(
                    'space'), pa.press('tab', 2), pa.write('003')
                pa.press('tab'), pa.write(str(pessoa.matricula))
                # #clique em Outros
                try:
                    pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Outros.png')))
                except pa.ImageNotFoundException:
                    t.sleep(5)
                    pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Outros.png')))
                t.sleep(2), pa.write('CARGO GERAL')
                # #clique em lupa de descrição de cargos
                pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Lupa.png')))
                t.sleep(2), pp.copy(pessoa.cargo), pa.hotkey('ctrl', 'v'), t.sleep(1.5), pa.press('enter', 2)
                t.sleep(1), pa.press('tab'), pa.write(pessoa.salario), pa.press('tab')
                if str(pessoa.tipo_contr) == 'Horista':
                    pa.press('1')
                else:
                    pa.press('5')
                pa.press('tab', 2), t.sleep(1), pa.write(str(pessoa.hr_sem)), pa.press('tab'), t.sleep(1)
                pa.write(str(pessoa.hr_mens))
                pa.press('tab', 5), pa.write('00395419000190'), pa.press('tab', 2), pa.write('5')
                # #clique em Compatibilidade
                pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Compatibilidade1.png'))), t.sleep(1)
                # #clique em Compatibilidade de novo
                pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Compatibilidade2.png'))), t.sleep(1)
                # #clique em CAGED
                pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/CAGED.png')))
                pa.press('tab', 2), pa.write('20'), pa.press('tab'), pa.write('1'), t.sleep(0.5)
                # #clique em RAIS
                pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/RAIS.png')))
                pp.copy(pessoa.nome), pa.hotkey('ctrl', 'v'), pa.press('tab', 2), pa.write('2')
                pa.press('tab'), pa.write('10')
                # #clique em Salvar
                pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Salvarcadastro.png'))), t.sleep(10)
                # #clique em fechar novo cadastro
                pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Fecharnovo1.png'))), t.sleep(2)
                # #clique em fechar trabalhadores
                pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Fechartrab1.png'))), t.sleep(0.5)

                os.makedirs(
                    r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e '
                    r'Férias\000 - Pastas Funcionais\00 - ATIVOS\1 - Ainda nao iniciaram\{}'.format(pessoa.nome))
                os.makedirs(
                    r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e '
                    r'Férias\000 - Pastas Funcionais\00 - ATIVOS\1 - Ainda nao iniciaram\{}\Atestados'.format(
                        pessoa.nome))
                os.makedirs(
                    r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e '
                    r'Férias\000 - Pastas Funcionais\00 - ATIVOS\1 - Ainda nao iniciaram\{}\Diversos'.format(
                        pessoa.nome))
                os.makedirs(
                    r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e '
                    r'Férias\000 - Pastas Funcionais\00 - ATIVOS\1 - Ainda nao iniciaram\{}\Contratuais'.format(
                        pessoa.nome))
                os.makedirs(
                    r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e '
                    r'Férias\000 - Pastas Funcionais\00 - ATIVOS\1 - Ainda nao iniciaram\{}\Férias'.format(pessoa.nome))
                os.makedirs(
                    r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e '
                    r'Férias\000 - Pastas Funcionais\00 - ATIVOS\1 - Ainda nao iniciaram\{}\Pontos'.format(pessoa.nome))
                os.makedirs(
                    r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e '
                    r'Férias\000 - Pastas Funcionais\00 - ATIVOS\1 - Ainda nao iniciaram\{}\Recibos'.format(
                        pessoa.nome))
                os.makedirs(
                    r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e '
                    r'Férias\000 - Pastas Funcionais\00 - ATIVOS\1 - Ainda nao iniciaram\{}\Rescisão'.format(
                        pessoa.nome))
                while 1:
                    if pa.locateOnScreen('../models/static/imgs/pyt.png'):
                        pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/pyt.png')))
                        break
                    else:
                        t.sleep(5)
                tkinter.messagebox.showinfo(title='Cadastro ok!', message='Cadastro realizado com sucesso!')


def salvar_docs_funcionarios(matricula):
    sessions = sessionmaker(bind=engine)
    session = sessions()
    pessoa = session.query(Colaborador).filter_by(matricula=matricula).first()
    if pessoa is None:
        tkinter.messagebox.showinfo(
            title='Erro de preenchimento',
            message='Não existe funcionário cadastrado com essa matrícula!'
        )
    else:
        try:
            os.makedirs(
                r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e '
                r'Férias\000 - Pastas Funcionais\00 - ATIVOS\1 - Ainda nao iniciaram\{}'.format(pessoa.nome))
            os.makedirs(
                r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e '
                r'Férias\000 - Pastas Funcionais\00 - ATIVOS\1 - Ainda nao iniciaram\{}\Atestados'.format(
                    pessoa.nome))
            os.makedirs(
                r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e '
                r'Férias\000 - Pastas Funcionais\00 - ATIVOS\1 - Ainda nao iniciaram\{}\Diversos'.format(
                    pessoa.nome))
            os.makedirs(
                r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e '
                r'Férias\000 - Pastas Funcionais\00 - ATIVOS\1 - Ainda nao iniciaram\{}\Contratuais'.format(
                    pessoa.nome))
            os.makedirs(
                r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e '
                r'Férias\000 - Pastas Funcionais\00 - ATIVOS\1 - Ainda nao iniciaram\{}\Férias'.format(
                    pessoa.nome))
            os.makedirs(
                r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e '
                r'Férias\000 - Pastas Funcionais\00 - ATIVOS\1 - Ainda nao iniciaram\{}\Pontos'.format(
                    pessoa.nome))
            os.makedirs(
                r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e '
                r'Férias\000 - Pastas Funcionais\00 - ATIVOS\1 - Ainda nao iniciaram\{}\Recibos'.format(
                    pessoa.nome))
            os.makedirs(
                r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e '
                r'Férias\000 - Pastas Funcionais\00 - ATIVOS\1 - Ainda nao iniciaram\{}\Rescisão'.format(
                    pessoa.nome))
        except FileExistsError:
            pass
        pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Dexion.png')))
        p_pessoa = r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e Férias' \
                   r'\000 - Pastas Funcionais\00 - ATIVOS\1 - Ainda nao iniciaram\{}'.format(pessoa.nome)
        p_contr = r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e Férias' \
                  r'\000 - Pastas Funcionais\00 - ATIVOS\1 - Ainda nao iniciaram\{}\Contratuais'.format(pessoa.nome)
        p_recibos = r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e Férias' \
                    r'\000 - Pastas Funcionais\00 - ATIVOS\Modelo\Recibo Crachá e Uniformes MODELO.docx'
        ps_acordo = r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e Férias' \
                    r'\000 - Pastas Funcionais\00 - ATIVOS\1 - Ainda nao iniciaram\{}\Contratuais\Acordo Banco de Horas.pdf'.format(
            pessoa.nome)
        ps_recctps = r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e Férias' \
                     r'\000 - Pastas Funcionais\00 - ATIVOS\1 - Ainda nao iniciaram\{}\Contratuais\Recibo de Entrega e Dev CTPS.pdf'.format(
            pessoa.nome)
        ps_anotctps = r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e Férias' \
                      r'\000 - Pastas Funcionais\00 - ATIVOS\1 - Ainda nao iniciaram\{}\Contratuais\Anotacoes CTPS.pdf'.format(
            pessoa.nome)
        ps_termovt = r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e Férias' \
                     r'\000 - Pastas Funcionais\00 - ATIVOS\1 - Ainda nao iniciaram\{}\Contratuais\Termo Opcao VT.pdf'.format(
            pessoa.nome)
        ps_contrato = r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e Férias' \
                      r'\000 - Pastas Funcionais\00 - ATIVOS\1 - Ainda nao iniciaram\{}\Contratuais\Contrato de Trabalho.pdf'.format(
            pessoa.nome)
        ps_ficha = r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e Férias' \
                   r'\000 - Pastas Funcionais\00 - ATIVOS\1 - Ainda nao iniciaram\{}\Contratuais\Folha de Registro.pdf'.format(
            pessoa.nome)

        recibos = docx.Document(p_recibos)

        # # imprimir recibo entrega e devolução de ctps
        pa.press('alt'), pa.press('r'), pa.press('a'), pa.press('r'), pa.press('e'), pa.press('tab'), pa.write(str(
            pessoa.matricula))
        pa.press('tab', 3), pa.write(str(pessoa.admiss).replace('/', '')), pa.press('tab'), t.sleep(0.5), pa.press(
            'space')
        t.sleep(0.5), pa.press('tab'), pa.write(str(pessoa.admiss).replace('/', '')), pa.press('tab', 2)
        t.sleep(1), pa.press('enter'), t.sleep(2)

        # # clique no endereço de salvamento do recibo

        while 1:
            if pa.locateOnScreen('../models/static/imgs/salvar.png'):
                pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/salvar.png')))
                break
            else:
                t.sleep(5)

        pp.copy(ps_recctps), pa.hotkey('ctrl', 'v'), t.sleep(0.5)
        pa.press('tab', 2), t.sleep(0.5), pa.press('enter')
        t.sleep(5)
        # # clique para fechar recibo ctps

        for i in range(2):
            while 1:
                if pa.locateOnScreen('../models/static/imgs/fechar_janela.png'):
                    pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/fechar_janela.png')))
                    break
                else:
                    t.sleep(5)

        # # Imprimir Acordo de Banco de horas
        pa.press('alt'), pa.press('r'), pa.press('z'), pa.press('d'), pa.press('d')
        pa.write("(matricula = '00{}')".format(str(pessoa.matricula))), t.sleep(1), pa.press('tab'), pa.write('2')
        pa.press('tab'), pa.press('enter'), t.sleep(10)

        while 1:
            if pa.locateOnScreen('../models/static/imgs/salvar.png'):
                pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/salvar.png')))
                break
            else:
                t.sleep(5)

        t.sleep(1), pp.copy(ps_acordo)
        pa.hotkey('ctrl', 'v'), t.sleep(1), pa.press('enter'), t.sleep(15)
        # # clique para fechar acordo
        for i in range(2):
            while 1:
                if pa.locateOnScreen('../models/static/imgs/fechar_janela.png'):
                    pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/fechar_janela.png')))
                    break
                else:
                    t.sleep(5)

        # # imprimir Anotações em CTPS
        pa.press('alt'), pa.press('r'), pa.press('a'), pa.press('c'), pa.press('e'), pa.press('tab')
        pa.write(str(pessoa.matricula))
        pa.press('tab', 2), pa.write(str(pessoa.admiss).replace('/', '')), pa.press('tab')
        pa.write(str(pessoa.admiss).replace('/', '')), pa.press('tab', 4), pa.press('space')
        pa.press('tab'), pa.press('enter'), t.sleep(1.5)

        while 1:
            if pa.locateOnScreen('../models/static/imgs/salvar.png'):
                pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/salvar.png')))
                break
            else:
                t.sleep(5)

        pp.copy(ps_anotctps), pa.hotkey('ctrl', 'v'), t.sleep(0.5), pa.press('enter')
        t.sleep(2)
        for i in range(2):
            while 1:
                if pa.locateOnScreen('../models/static/imgs/fechar_janela.png'):
                    pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/fechar_janela.png')))
                    break
                else:
                    t.sleep(5)

        # # imprimir Termo VT
        pa.press('alt'), pa.press('r'), pa.press('a'), pa.press('v'), pa.press('e'), pa.press('tab')
        pa.write(str(pessoa.matricula)), pa.press('tab', 2), pa.write(str(pessoa.admiss).replace('/', ''))
        pa.press('tab'), pa.write('d'), pa.press('tab', 4), pa.press('space')
        pa.press('tab', 6), pa.press('enter'), t.sleep(1.5)
        while 1:
            if pa.locateOnScreen('../models/static/imgs/salvar.png'):
                pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/salvar.png')))
                break
            else:
                t.sleep(5)
        pp.copy(ps_termovt), pa.hotkey('ctrl', 'v'), t.sleep(0.5), pa.press('enter')
        t.sleep(2)
        for i in range(2):
            while 1:
                if pa.locateOnScreen('../models/static/imgs/fechar_janela.png'):
                    pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/fechar_janela.png')))
                    break
                else:
                    t.sleep(5)

        # # Imprimir Contrato
        pa.press('alt'), pa.press('r'), pa.press('z'), pa.press('d')
        if pessoa.tipo_contr == 'Horista':
            pa.press('c')
        else:
            pa.press('o')

        pa.write("(matricula = '00{}')".format(str(pessoa.matricula))), t.sleep(1), pa.press('tab'), pa.write('2')
        pa.press('tab'), pa.press('enter'), t.sleep(5)
        while 1:
            if pa.locateOnScreen('../models/static/imgs/salvar.png'):
                pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/salvar.png')))
                break
            else:
                t.sleep(5)
        pp.copy(ps_contrato), pa.hotkey('ctrl', 'v'), t.sleep(0.5), pa.press('enter')
        t.sleep(10)
        for i in range(2):
            while 1:
                if pa.locateOnScreen('../models/static/imgs/fechar_janela.png'):
                    pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/fechar_janela.png')))
                    break
                else:
                    t.sleep(5)

        # # Imprimir Folha de rosto de Cadastro
        pa.press('alt'), pa.press('r'), pa.press('i'), pa.press('o'), pa.press('r'), pa.press('e'), pa.press('tab')
        pa.write(str(pessoa.matricula)), pa.press('tab', 2), pa.write(str(pessoa.admiss).replace('/', '')), pa.press(
            'tab',
            2)
        pa.press('enter'), t.sleep(3)
        while 1:
            if pa.locateOnScreen('../models/static/imgs/salvar.png'):
                pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/salvar.png')))
                break
            else:
                t.sleep(5)
        pp.copy(ps_ficha), pa.hotkey('ctrl', 'v'), t.sleep(0.5), pa.press('enter')
        t.sleep(3)
        for i in range(2):
            while 1:
                if pa.locateOnScreen('../models/static/imgs/fechar_janela.png'):
                    pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/fechar_janela.png')))
                    break
                else:
                    t.sleep(5)

        # # Alterar Recibos e salvar na pasta
        recibos.paragraphs[4].text = str(recibos.paragraphs[4].text).replace('#nome_completo', pessoa.nome)
        recibos.paragraphs[12].text = str(recibos.paragraphs[12].text).replace('#nome_completo', pessoa.nome)
        recibos.paragraphs[19].text = str(recibos.paragraphs[19].text).replace('#nome_completo', pessoa.nome)
        recibos.paragraphs[27].text = str(recibos.paragraphs[27].text).replace('#nome_completo', pessoa.nome)
        recibos.paragraphs[40].text = str(recibos.paragraphs[40].text).replace('#nome_completo', pessoa.nome)
        recibos.paragraphs[48].text = str(recibos.paragraphs[48].text).replace('#nome_completo', pessoa.nome)
        recibos.save(p_contr + '\\Recibos.docx')
        docx2pdf.convert(p_contr + '\\Recibos.docx', p_contr + '\\Recibos.pdf')
        os.remove(p_contr + '\\Recibos.docx')
        os.rename(p_pessoa, p_pessoa.replace(r'\1 - Ainda nao iniciaram', ''))
        while 1:
            if pa.locateOnScreen('../models/static/imgs/pyt.png'):
                pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/pyt.png')))
                break
            else:
                t.sleep(5)
        tkinter.messagebox.showinfo(
            title='Documentos ok!',
            message='Documentos salvos com sucesso!'
        )


def enviar_emails_funcionario(matricula):
    """
    Send e-mails to employee about his/her admission.
    :param matricula: Employee registration number
    """
    sessions = sessionmaker(bind=engine)
    session = sessions()
    pessoa = session.query(Colaborador).filter_by(matricula=matricula).first()
    if pessoa is None:
        tkinter.messagebox.showinfo(
            title='Erro de preenchimento',
            message='Não existe funcionário cadastrado com essa matrícula!'
        )
    else:
        p_contr = r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e Férias' \
                  r'\000 - Pastas Funcionais\00 - ATIVOS\{}\Contratuais'.format(pessoa.nome)
        email_remetente = em_rem
        senha = k1
        # set up smtp connection
        s = smtplib.SMTP(host=host, port=port)
        s.starttls()
        s.login(email_remetente, senha)
        # send e-mail to employee with a pdf file so he/she can go to bank to open an account
        msg = MIMEMultipart('alternative')
        msg['From'] = email_remetente
        msg['To'] = pessoa.email
        msg['Subject'] = "Documentos Contratuais"
        arquivo = p_contr + '\\Cod Etica.pdf'
        arquivo2 = p_contr + '\\Contrato de Trabalho.pdf'
        if pessoa.genero == 'Masculino':
            if pessoa.tipo_contr == 'Horista':
                text = MIMEText(f'''Olá, {str(pessoa.nome).title().split(" ")[0]}!<br><br>
                Seja muito bem vindo a Companhia Athletica.<br>
                Seu contrato está pronto para ser assinado.<br>
                Seu tipo de contrato é de horista.<br>
                Nesse tipo de contrato você será remunerado de acordo com a soma total de HORAS trabalhadas.<br>
                É importante lembrar que o tempo dos intervalos entre aulas não é remunerado.<br><br>
                Seguem alguns lembretes do nosso código de ética já enviado a você:<br>
                Você pode treinar todas as modalidades da academia gratuitamente, sempre respeitando a prioridade aos alunos.<br>
                Você tem direito a uma folga no dia do seu aniversário (essa folga não pode ser transferida para outro dia).<br>
                A academia oferece descontos especiais para parentes de funcionários.<br>
                Nosso período de folha é do dia 21 do mês ao dia 20 do mês seguinte. Então, no seu primeiro salário você receberá o valor proporcional do dia que entrou até o próximo dia 20.<br>
                Quando completar um ano de contrato terá direito a férias. Para pedir férias você deve solicitar ao seu superior direto com antecedência de 2 meses da data que pretende sair.<br><br>
                As demais regras você encontra no nosso código de ética. Ok?<br>
                Qualquer dúvida, estou à disposição.<br><br>
                Atenciosamente,<br>
                <img src="cid:image1">''', 'html')
            else:
                text = MIMEText(f'''Olá, {str(pessoa.nome).title().split(" ")[0]}!<br><br>
                Seja muito bem vindo a Companhia Athletica.<br>
                Seu contrato está pronto para ser assinado.<br><br>
                Seguem alguns lembretes do nosso código de ética já enviado a você:<br>
                Você pode treinar todas as modalidades da academia gratuitamente, sempre respeitando a prioridade aos alunos.<br>
                Você tem direito a uma folga no dia do seu aniversário (essa folga não pode ser transferida para outro dia).<br>
                A academia oferece descontos especiais para parentes de funcionários.<br>
                Nosso período de folha é do dia 21 do mês ao dia 20 do mês seguinte. Então, no seu primeiro salário você receberá o valor proporcional do dia que entrou até o próximo dia 20.<br>
                Quando completar um ano de contrato terá direito a férias. Para pedir férias você deve solicitar ao seu superior direto com antecedência de 2 meses da data que pretende sair.<br><br>
                As demais regras você encontra no nosso código de ética. Ok?<br>
                Qualquer dúvida, estou à disposição.<br><br>
                Atenciosamente,<br>
                <img src="cid:image1">''', 'html')

        else:
            if pessoa.tipo_contr == 'Horista':
                text = MIMEText(f'''Olá, {str(pessoa.nome).title().split(" ")[0]}!<br><br>
                Seja muito bem vinda a Companhia Athletica.<br>
                Seu contrato está pronto para ser assinado.<br>
                Seu tipo de contrato é de horista.<br>
                Nesse tipo de contrato você será remunerado de acordo com a soma total de HORAS trabalhadas.<br>
                É importante lembrar que o tempo dos intervalos entre aulas não é remunerado.<br><br>
                Seguem alguns lembretes do nosso código de ética já enviado a você:<br>
                Você pode treinar todas as modalidades da academia gratuitamente, sempre respeitando a prioridade aos alunos.<br>
                Você tem direito a uma folga no dia do seu aniversário (essa folga não pode ser transferida para outro dia).<br>
                A academia oferece descontos especiais para parentes de funcionários.<br>
                Nosso período de folha é do dia 21 do mês ao dia 20 do mês seguinte. Então, no seu primeiro salário você receberá o valor proporcional do dia que entrou até o próximo dia 20.<br>
                Quando completar um ano de contrato terá direito a férias. Para pedir férias você deve solicitar ao seu superior direto com antecedência de 2 meses da data que pretende sair.<br><br>
                As demais regras você encontra no nosso código de ética. Ok?<br>
                Qualquer dúvida, estou à disposição.<br><br>
                Atenciosamente,<br>
                <img src="cid:image1">''', 'html')
            else:
                text = MIMEText(f'''Olá, {str(pessoa.nome).title().split(" ")[0]}!<br><br>
                Seja muito bem vinda a Companhia Athletica.<br>
                Seu contrato está pronto para ser assinado.<br><br>
                Seguem alguns lembretes do nosso código de ética já enviado a você:<br>
                Você pode treinar todas as modalidades da academia gratuitamente, sempre respeitando a prioridade aos alunos.<br>
                Você tem direito a uma folga no dia do seu aniversário (essa folga não pode ser transferida para outro dia).<br>
                A academia oferece descontos especiais para parentes de funcionários.<br>
                Nosso período de folha é do dia 21 do mês ao dia 20 do mês seguinte. Então, no seu primeiro salário você receberá o valor proporcional do dia que entrou até o próximo dia 20.<br>
                Quando completar um ano de contrato terá direito a férias. Para pedir férias você deve solicitar ao seu superior direto com antecedência de 2 meses da data que pretende sair.<br><br>
                As demais regras você encontra no nosso código de ética. Ok?<br>
                Qualquer dúvida, estou à disposição.<br><br>
                Atenciosamente,<br>
                <img src="cid:image1">''', 'html')

        # set up the parameters of the message
        msg.attach(text)
        image = MIMEImage(
            open(rf'C:\Users\{os.getlogin()}\PycharmProjects\AutomacaoCia\src\models\static\imgs\assinatura.png',
                 'rb').read())
        image.add_header('Content-ID', '<image1>')
        msg.attach(image)

        # attach pdf file
        part = MIMEBase('application', "octet-stream")
        part.set_payload(open(arquivo, "rb").read())
        encoders.encode_base64(part)
        part.add_header('Content-Disposition', 'attachment',
                        filename=f'Código de Ética.pdf')
        msg.attach(part)

        part2 = MIMEBase('application', "octet-stream")
        part2.set_payload(open(arquivo2, "rb").read())
        encoders.encode_base64(part2)
        part2.add_header('Content-Disposition', 'attachment',
                         filename=f'Contrato.pdf')
        msg.attach(part2)

        s.sendmail(email_remetente, pessoa.email, msg.as_string())
        del msg
        s.quit()
        tkinter.messagebox.showinfo(
            title='E-mails ok!',
            message='E-mails enviados com sucesso'
        )


def enviar_emails_contratacao(caminho: str, nome: str, departamento: str, cargo: str, salario: str, admissao: str):
    """
    This function send e-mails to employee

    :param caminho:
    :param nome:
    :param departamento:
    :param cargo:
    :param salario:
    :param admissao:
    :return:
    """
    if nome == '' or departamento == '' or cargo == '' or salario == '' or admissao == '':
        tkinter.messagebox.showinfo(
            title='Erro de preenchimento',
            message='Preencha corretamente os campos Nome, Admissão, Salário, Cargo e Departamento.'
        )
    else:
        linha, nome = nome.upper().strip().split(' - ')
        departamento = departamento.title()
        plcontr = l_w(caminho, read_only=False)
        fol = plcontr['Respostas ao formulário 1']
        p_pessoa = r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e Férias' \
                   r'\000 - Pastas Funcionais\00 - ATIVOS\1 - Ainda nao iniciaram\{}'.format(nome)
        p_atestado = r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e Férias' \
                     r'\000 - Pastas Funcionais\00 - ATIVOS\1 - Ainda nao iniciaram\{}\Atestado'.format(nome)
        p_contr = r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e Férias' \
                  r'\000 - Pastas Funcionais\00 - ATIVOS\1 - Ainda nao iniciaram\{}\Contratuais'.format(nome)
        p_diversos = r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e Férias' \
                     r'\000 - Pastas Funcionais\00 - ATIVOS\1 - Ainda nao iniciaram\{}\Diversos'.format(nome)
        p_ferias = r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e Férias' \
                   r'\000 - Pastas Funcionais\00 - ATIVOS\1 - Ainda nao iniciaram\{}\Férias'.format(nome)
        p_ponto = r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e Férias' \
                  r'\000 - Pastas Funcionais\00 - ATIVOS\1 - Ainda nao iniciaram\{}\Pontos'.format(nome)
        p_rec = r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e Férias' \
                r'\000 - Pastas Funcionais\00 - ATIVOS\1 - Ainda nao iniciaram\{}\Recibos'.format(nome)
        p_rescisao = r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e Férias' \
                     r'\000 - Pastas Funcionais\00 - ATIVOS\1 - Ainda nao iniciaram\{}\Rescisão'.format(nome)
        p_ac = r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e Férias' \
               r'\000 - Pastas Funcionais\00 - ATIVOS\Modelo\AC Modelo.docx'
        p_abconta = r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e Férias' \
                    r'\000 - Pastas Funcionais\00 - ATIVOS\Modelo\Abertura Conta MODELO.docx'
        p_fcadas = r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e Férias' \
                   r'\000 - Pastas Funcionais\00 - ATIVOS\Modelo\Ficha Cadastral MODELO.docx'
        p_codetic = r'\\192.168.0.250\rh\01 - RH\01 - Administração.Controles\02 - Funcionários, Departamentos e Férias' \
                    r'\000 - Pastas Funcionais\00 - ATIVOS\Modelo\Cod Etica MODELO.docx'
        try:
            os.mkdir(p_pessoa)
            os.mkdir(p_atestado)
            os.mkdir(p_contr)
            os.mkdir(p_diversos)
            os.mkdir(p_ferias)
            os.mkdir(p_ponto)
            os.mkdir(p_rec)
            os.mkdir(p_rescisao)
        except FileExistsError:
            pass

        lotacao = {
            'Unidade Park Sul - Qualquer Departamento': ['0013', lidermusc, elidermusc, 'Líder Park Sul'],
            'Kids': ['0010', liderkids, eliderkids, 'Líder Kids'],
            'Musculação': ['0007', lidermusc, elidermusc, 'Líder Musculação'],
            'Esportes E Lutas': ['0008', lidernat, elidernat, 'Líder Natação'],
            'Crossfit': ['0012', lidercross, elidercross, 'Líder Crossfit'],
            'Ginástica': ['0006', lidergin, elidergin, 'Líder Ginástica'],
            'Gestantes': ['0006', gerentetec, egerentetec, 'Gerente Técnico'],
            'Recepção': ['0003', gerentevend, egerentevend, 'Gerente Vendas'],
            'Administrativo': ['0001', gerenterh, egerenterh, 'Gerente RH'],
            'Manutenção': ['0004', gerentemanut, egerentemanut, 'Gerente Manutenção'],
        }

        abert_c = docx.Document(p_abconta)
        ac = docx.Document(p_ac)
        fch_c = docx.Document(p_fcadas)
        codetic = docx.Document(p_codetic)

        # # Alterar AC e Salvar na pasta
        ac.paragraphs[1].text = str(ac.paragraphs[1].text).replace('#gerente', lotacao[departamento][1])
        ac.paragraphs[2].text = str(ac.paragraphs[2].text).replace('#nome_completo', nome)
        ac.paragraphs[3].text = str(ac.paragraphs[3].text).replace('#cargo', cargo)
        ac.paragraphs[11].text = str(ac.paragraphs[11].text).replace('#salario', salario)
        ac.save(p_contr + '\\AC.docx')
        docx2pdf.convert(p_contr + '\\AC.docx', p_contr + '\\AC.pdf')
        os.remove(p_contr + '\\AC.docx')

        # # Alterar Abertura de Conta e salvar na pasta
        abert_c.paragraphs[14].text = str(abert_c.paragraphs[14].text).replace('#nome_completo', nome).replace(
            '#rg', str(fol[f'W{linha}'].value).replace('.0', '')).replace(
            '#cpf', str(fol[f'V{linha}'].value)).replace('#endereco', str(fol[f'O{linha}'].value)) \
            .replace('#cep', str(fol[f'R{linha}'].value).replace('.0', '')) \
            .replace('#bairro', str(fol[f'Q{linha}'].value).strip()).replace('#cargo', cargo).replace('#data', admissao)
        abert_c.save(p_contr + '\\Abertura Conta.docx')
        docx2pdf.convert(p_contr + '\\Abertura Conta.docx', p_contr + '\\Abertura Conta.pdf')
        os.remove(p_contr + '\\Abertura Conta.docx')

        # Alterar Ficha cadastral e salvar na pasta
        fch_c.paragraphs[34].text = str(fch_c.paragraphs[34].text).replace('#gerente#',
                                                                           lotacao[departamento][1])
        fch_c.paragraphs[9].text = str(fch_c.paragraphs[9].text).replace('#nome_completo', nome)
        fch_c.paragraphs[21].text = str(fch_c.paragraphs[21].text).replace('#cargo', cargo) \
            .replace('#depart', departamento)
        fch_c.paragraphs[19].text = str(fch_c.paragraphs[19].text).replace('#end_eletr', str(fol[f'B{linha}'].value))
        fch_c.paragraphs[17].text = str(fch_c.paragraphs[17].text).replace('#mae#', str(fol[f'N{linha}'].value))
        fch_c.paragraphs[16].text = str(fch_c.paragraphs[16].text).replace('#pai#', str(fol[f'M{linha}'].value))
        fch_c.paragraphs[15].text = str(fch_c.paragraphs[15].text).replace('#ident',
                                                                           str(fol[f'W{linha}'].value).replace('.0',
                                                                                                               '')).replace(
            '#cpf#',
            str(fol[f'V{linha}'].value))
        fch_c.paragraphs[13].text = str(fch_c.paragraphs[13].text).replace('#telefone',
                                                                           str(fol[f'U{linha}'].value).replace('.0',
                                                                                                               ''))
        fch_c.paragraphs[12].text = str(fch_c.paragraphs[12].text).replace('#codigo',
                                                                           str(fol[f'R{linha}'].value).replace('.0',
                                                                                                               '')) \
            .replace('#cid', str(fol[f'S{linha}'].value)).replace('#uf', str(fol[f'T{linha}'].value))
        fch_c.paragraphs[11].text = str(fch_c.paragraphs[11].text).replace('#local', str(fol[f'O{linha}'].value)) \
            .replace('#qd', str(fol[f'Q{linha}'].value))
        fch_c.paragraphs[10].text = str(fch_c.paragraphs[10].text) \
            .replace('#nasc', dt.strftime(dt.strptime(str(fol[f'D{linha}'].value), '%Y-%m-%d %H:%M:%S'), '%d/%m/%Y')) \
            .replace('#gen', str(fol[f'E{linha}'].value)).replace('#est_civ',
                                                                  str(fol[f'F{linha}'].value).replace('1 - ', '')
                                                                  .replace('2 - ', '').replace('3 - ', '').replace(
                                                                      '4 - ', ''))
        fch_c.save(p_contr + '\\Ficha Cadastral.docx')
        docx2pdf.convert(p_contr + '\\Ficha Cadastral.docx', p_contr + '\\Ficha Cadastral.pdf')
        os.remove(p_contr + '\\Ficha Cadastral.docx')

        # Alterar Código de Ética e salvar na pasta
        codetic.paragraphs[535].text = str(codetic.paragraphs[535].text).replace('#nome_completo', nome)
        codetic.paragraphs[536].text = str(codetic.paragraphs[536].text).replace('#func', cargo)
        codetic.paragraphs[538].text = str(codetic.paragraphs[538].text).replace('#nome_completo', nome)
        codetic.paragraphs[542].text = str(codetic.paragraphs[542].text).replace('#admiss', admissao)
        codetic.save(p_contr + '\\Cod Etica.docx')
        docx2pdf.convert(p_contr + '\\Cod Etica.docx', p_contr + '\\Cod Etica.pdf')
        os.remove(p_contr + '\\Cod Etica.docx')

        # send e-mails
        email_remetente = em_rem
        senha = k1
        # set up smtp connection
        s = smtplib.SMTP(host=host, port=port)
        s.starttls()
        s.login(email_remetente, senha)

        # enviar e-mail de boas vindas
        msg = MIMEMultipart('alternative')
        msg['From'] = email_remetente
        msg['To'] = str(fol[f'B{linha}'].value).strip()
        msg['Subject'] = "Boas vindas!"
        arquivo = p_contr + '\\Cod Etica.pdf'
        text = MIMEText(f'''Olá, {str(nome).title().split(" ")[0]}!<br><br>
        Estamos felizes que você fará parte da nossa equipe!<br>
        Em anexo segue nosso código de ética e conduta.<br>
        Nesse documento estão todas as regras da Cia e tudo que a Cia espera de seus funcionários.<br>
        Além de regras, também descreve direitos e benefícios.<br>
        É importante que você leia todo o documento pois precisaremos da sua assinatura nele no dia que for assinar o contrato. Ok?<br>
        Qualquer dúvida, estou à disposição.<br><br>
        Atenciosamente,<br>
        <img src="cid:image1">''', 'html')

        # set up the parameters of the message
        msg.attach(text)
        image = MIMEImage(
            open(rf'C:\Users\{os.getlogin()}\PycharmProjects\AutomacaoCia\src\models\static\imgs\assinatura.png',
                 'rb').read())
        image.add_header('Content-ID', '<image1>')
        msg.attach(image)
        # attach pdf file
        part = MIMEBase('application', "octet-stream")
        part.set_payload(open(arquivo, "rb").read())
        encoders.encode_base64(part)
        part.add_header('Content-Disposition', 'attachment',
                        filename='Código de Ética Cia Athletica.pdf')
        msg.attach(part)
        s.sendmail(email_remetente, str(fol[f'B{linha}'].value).strip(), msg.as_string())
        del msg

        # enviar AC para líder
        msg = MIMEMultipart('alternative')
        msg['From'] = email_remetente
        msg['To'] = lotacao[departamento][2]
        msg['Subject'] = f'AC - {str(nome).title().split(" ")[0]}'
        arquivo = p_contr + '\\AC.pdf'
        text = MIMEText(f'''Olá, {lotacao[departamento][1].split(" ")[0]}!<br><br>
        Segue a AC do(a) {str(nome).title().split(" ")[0]}.<br>
        A AC é o documento oficial de cadastro de um funcionário na Cia.<br>
        Você deverá imprimir a AC, preencher o horário do funcionário em cada dia e solicitar a assinatura da direção.<br>
        Após a assinatura da direção me entregue a AC no RH, por favor.<br><br>
        Atenciosamente,<br>
        <img src="cid:image1">''', 'html')

        # set up the parameters of the message
        msg.attach(text)
        image = MIMEImage(
            open(rf'C:\Users\{os.getlogin()}\PycharmProjects\AutomacaoCia\src\models\static\imgs\assinatura.png',
                 'rb').read())
        image.add_header('Content-ID', '<image1>')
        msg.attach(image)
        # attach pdf file
        part = MIMEBase('application', "octet-stream")
        part.set_payload(open(arquivo, "rb").read())
        encoders.encode_base64(part)
        part.add_header('Content-Disposition', 'attachment',
                        filename=f'AC - {str(nome).title().split(" ")[0]}.pdf')
        msg.attach(part)
        s.sendmail(email_remetente, lotacao[departamento][2], msg.as_string())
        del msg

        # send e-mail to employee with a pdf file so he/she can go to bank to open an account
        msg = MIMEMultipart('alternative')
        msg['From'] = email_remetente
        msg['To'] = str(fol[f'B{linha}'].value).strip()
        msg['Subject'] = "Carta para Abertura de conta"
        arquivo = p_contr + '\\Abertura Conta.pdf'
        text = MIMEText(f'''Olá, {str(nome).title().split(" ")[0]}!<br><br>
        Segue sua carta para abertura de conta bancária no Itaú.<br>
        Você deve abrir a conta antes de iniciar seu contrato de trabalho. Ok?<br>
        Assim que conseguir abrir a conta me responda esse e-mail com os dados bancários do Itaú. <br>
        Qualquer dúvida, estou à disposição.<br><br>
        Atenciosamente,<br>
        <img src="cid:image1">''', 'html')

        # set up the parameters of the message
        msg.attach(text)
        image = MIMEImage(
            open(rf'C:\Users\{os.getlogin()}\PycharmProjects\AutomacaoCia\src\models\static\imgs\assinatura.png',
                 'rb').read())
        image.add_header('Content-ID', '<image1>')
        msg.attach(image)
        # attach pdf file
        part = MIMEBase('application', "octet-stream")
        part.set_payload(open(arquivo, "rb").read())
        encoders.encode_base64(part)
        part.add_header('Content-Disposition', 'attachment',
                        filename=f'Carta Banco {str(nome).title().split(" ")[0]}.pdf')
        msg.attach(part)
        s.sendmail(email_remetente, str(fol[f'B{linha}'].value).strip(), msg.as_string())
        del msg

        # send e-mail to coworker asking to register the ner employee
        msg = MIMEMultipart('alternative')
        arquivo = p_contr + '\\Ficha Cadastral.pdf'
        text = MIMEText(
            f'''Oi, Wallace!<br><br>Segue a ficha cadastral do(a) {nome}.<br><br>Abs.,<br><img src="cid:image1">''',
            'html')
        msg.attach(text)
        image = MIMEImage(
            open(rf'C:\Users\{os.getlogin()}\PycharmProjects\AutomacaoCia\src\models\static\imgs\assinatura.png',
                 'rb').read())
        image.add_header('Content-ID', '<image1>')
        msg.attach(image)
        # set up the parameters of the message
        msg['From'] = email_remetente
        msg['To'] = em_ti
        msg['Subject'] = f"Ficha Cadastral {str(nome).title().split(' ')[0]}"
        # attach pdf file
        part = MIMEBase('application', "octet-stream")
        part.set_payload(open(arquivo, "rb").read())
        encoders.encode_base64(part)
        part.add_header('Content-Disposition', 'attachment',
                        filename=f'Ficha Cadastral {str(nome).title().split(" ")[0]}.pdf')
        msg.attach(part)
        s.sendmail(email_remetente, em_ti, msg.as_string())
        del msg
        s.quit()
        tkinter.messagebox.showinfo(
            title='E-mails ok!',
            message='E-mails enviados com sucesso!'
        )


def apenas_registrar_estagiario(solicitar_contr=0, caminho='', editar=0, ondestou=0, nome='', matricula='', admissao='',
                                cargo='', depto='', tipo_contr='Horista',
                                hrsem='25', hrmens='100', agencia='', conta='', digito=''):
    sessions = sessionmaker(bind=engine)
    session = sessions()
    pa.FAILSAFE = False
    salario = 458
    if nome == '':
        tkinter.messagebox.showinfo(
            title='Erro de preenchimento',
            message='Preencha todos os campos antes de cadastrar o estagiário!'
        )
    else:
        wb = l_w(caminho)
        sh = wb['Respostas ao formulário 1']
        num, name = nome.strip().split(' - ')
        linha = int(num)
        if str(sh[f'E{linha}'].value) == 'Masculino':
            cargo = 'ESTAGIARIO'
        else:
            cargo = 'ESTAGIARIA'
        try:
            estag_cadastrado = Colaborador(
                matricula=matricula, nome=name.upper(), admiss=admissao,
                nascimento=str(sh[f'D{linha}'].value),
                cpf=str(sh[f'V{linha}'].value).replace('.', '').replace('-', '').zfill(11),
                rg=str(int(sh[f'W{linha}'].value)),
                emissor='SSP/DF', email=str(sh[f'B{linha}'].value),
                genero=str(sh[f'E{linha}'].value),
                estado_civil=str(sh[f'F{linha}'].value), cor='9',
                instru='08 - Educação Superior Incompleta',
                nacional='Brasileiro(a)',
                pai=str(sh[f'M{linha}'].value).upper(),
                mae=str(sh[f'N{linha}'].value).upper(), endereco=str(sh[f'O{linha}'].value),
                num='1',
                bairro=str(sh[f'Q{linha}'].value),
                cep=str(sh[f'R{linha}'].value).replace('.', '').replace('-', ''),
                cidade='Brasília', cid_nas='Brasília - DF',
                uf='DF',
                cod_municipioend=municipios['DF']['Brasília'],
                tel=str(sh[f'U{linha}'].value).replace('(', '').replace(')', '').replace('.', '').replace('-',
                                                                                                          ''),
                depto=depto, cargo=cargo,
                horario=str(sh[f'AI{linha}'].value), salario=salario, tipo_contr=tipo_contr,
                hr_sem='25', hr_mens='100',
                est_semestre=str(sh[f'AS{linha}'].value),
                est_turno=str(sh[f'AT{linha}'].value),
                est_prev_conclu=str(sh[f'AU{linha}'].value),
                est_faculdade=str(sh[f'AV{linha}'].value),
                est_endfacul='End',
                est_numendfacul='1',
                est_bairroendfacul='Bairro',
                ag=agencia, conta=conta, cdigito=digito

            )
            session.add(estag_cadastrado)
            session.commit()
            tkinter.messagebox.showinfo(title='Cadastro ok!', message='Estagiário cadastrado com sucesso!')
        except IntegrityError:
            tkinter.messagebox.showinfo(title='Erro', message='Estagiário já cadastrado no DB!')


def cadastro_estagiario(solicitar_contr=0, caminho='', editar=0, ondestou=0, nome='', matricula='', admissao='',
                        cargo='', depto='', tipo_contr='Mensalista',
                        hrsem='25', hrmens='100', agencia='', conta='', digito='', tce=''):
    sessions = sessionmaker(bind=engine)
    session = sessions()
    pa.FAILSAFE = False
    salario = 458
    if nome == '':
        tkinter.messagebox.showinfo(
            title='Erro de preenchimento',
            message='Preencha todos os campos antes de cadastrar o estagiário!'
        )
    else:
        if solicitar_contr == 1:
            hoje = dt.today()
            wb = l_w(caminho)
            sh = wb['Respostas ao formulário 1']
            num, name = nome.strip().split(' - ')
            linha = int(num)
            lotacao = {
                'Unidade Park Sul - Qualquer Departamento': ['0013', liderpnt, em_pnt, 'Líder PNT'],
                'Kids': ['0010', liderpnt, em_pnt, 'Líder PNT'],
                'Musculação': ['0007', liderpnt, em_pnt, 'Líder PNT'],
                'Esportes E Lutas': ['0008', liderpnt, em_pnt, 'Líder PNT'],
                'Crossfit': ['0012', liderpnt, em_pnt, 'Líder PNT'],
                'Ginástica': ['0006', liderpnt, em_pnt, 'Líder PNT'],
                'Gestantes': ['0006', liderpnt, em_pnt, 'Líder PNT'],
                'Recepção': ['0003', liderpnt, em_pnt, 'Líder PNT'],
                'Administrativo': ['0001', liderpnt, em_pnt, 'Líder PNT'],
                'Manutenção': ['0004', liderpnt, em_pnt, 'Líder PNT'],
            }
            cadastro = {'nome': str(sh[f"C{linha}"].value).title().strip(), 'nasc_ed': sh[f"D{linha}"].value,
                        'genero': str(sh[f"E{linha}"].value), 'est_civ': str(sh[f"F{linha}"].value),
                        'pai': str(sh[f"M{linha}"].value), 'mae': str(sh[f"N{linha}"].value),
                        'end': str(sh[f"O{linha}"].value),
                        'num': str(sh[f"P{linha}"].value), 'bairro': str(sh[f"Q{linha}"].value),
                        'cep': str(sh[f"R{linha}"].value).replace('.', '').replace('-', ''),
                        'cid_end': str(sh[f"S{linha}"].value), 'uf_end': str(sh[f"T{linha}"].value),
                        'tel': str(sh[f"U{linha}"].value).replace('(', '').replace(')', '').replace('-', '').replace(
                            ' ',
                            ''),
                        'mun_end': str(sh[f"AP{linha}"].value),
                        'cpf': str(sh[f"V{linha}"].value).strip().replace('.', '').replace('-', '').replace(' ',
                                                                                                            '').zfill(
                            11),
                        'rg': str(sh[f"W{linha}"].value).strip().replace('.', '').replace('-', '').replace(' ', ''),
                        'emissor': str(sh[f"X{linha}"].value),
                        'lotacao': str(lotacao[f'{sh[f"AG{linha}"].value}'][0]).zfill(4),
                        'cargo': str(sh[f"AH{linha}"].value), 'horario': str(sh[f"AI{linha}"].value),
                        'email': str(sh[f"B{linha}"].value).strip(),
                        'admissao_ed': str(sh[f"AL{linha}"].value),
                        'faculdade': str(sh[f"AV{linha}"].value), 'semestre': str(sh[f"AS{linha}"].value),
                        'turno': str(sh[f"AT{linha}"].value), 'conclusao': str(sh[f"AU{linha}"].value),
                        'salario': str(sh[f"AM{linha}"].value),
                        'hrsemanais': str(sh[f"AQ{linha}"].value), 'hrmensais': str(sh[f"AR{linha}"].value)}
            email_remetente = em_rem
            senha = k1
            lot = lotacao[f'{sh[f"AG{linha}"].value}']
            pasta = r'\192.168.0.250'
            modelo = f'\\{pasta}\\rh\\01 - RH\\01 - Administração.Controles\\02 - Funcionários, Departamentos e Férias\\' \
                     f'000 - Pastas Funcionais\\00 - ATIVOS\\0 - Estagiários\\Modelo'
            try:
                os.makedirs(
                    f'\\{pasta}\\rh\\01 - RH\\01 - Administração.Controles\\02 - Funcionários, Departamentos e Férias\\'
                    f'000 - Pastas Funcionais\\00 - ATIVOS\\0 - Estagiários\\'
                    f'0 - Ainda nao iniciaram\\{str(cadastro["nome"]).upper()}')
                os.makedirs(
                    f'\\{pasta}\\rh\\01 - RH\\01 - Administração.Controles\\02 - Funcionários, Departamentos e Férias\\'
                    f'000 - Pastas Funcionais\\00 - ATIVOS\\0 - Estagiários\\'
                    f'0 - Ainda nao iniciaram\\{str(cadastro["nome"]).upper()}\\Atestados')
                os.makedirs(
                    f'\\{pasta}\\rh\\01 - RH\\01 - Administração.Controles\\02 - Funcionários, Departamentos e Férias\\'
                    f'000 - Pastas Funcionais\\00 - ATIVOS\\0 - Estagiários\\'
                    f'0 - Ainda nao iniciaram\\{str(cadastro["nome"]).upper()}\\Diversos')
                os.makedirs(
                    f'\\{pasta}\\rh\\01 - RH\\01 - Administração.Controles\\02 - Funcionários, Departamentos e Férias\\'
                    f'000 - Pastas Funcionais\\00 - ATIVOS\\0 - Estagiários\\'
                    f'0 - Ainda nao iniciaram\\{str(cadastro["nome"]).upper()}\\Contratuais')
                os.makedirs(
                    f'\\{pasta}\\rh\\01 - RH\\01 - Administração.Controles\\02 - Funcionários, Departamentos e Férias\\'
                    f'000 - Pastas Funcionais\\00 - ATIVOS\\0 - Estagiários\\'
                    f'0 - Ainda nao iniciaram\\{str(cadastro["nome"]).upper()}\\Ferias')
                os.makedirs(
                    f'\\{pasta}\\rh\\01 - RH\\01 - Administração.Controles\\02 - Funcionários, Departamentos e Férias\\'
                    f'000 - Pastas Funcionais\\00 - ATIVOS\\0 - Estagiários\\'
                    f'0 - Ainda nao iniciaram\\{str(cadastro["nome"]).upper()}\\Ponto')
                os.makedirs(
                    f'\\{pasta}\\rh\\01 - RH\\01 - Administração.Controles\\02 - Funcionários, Departamentos e Férias\\'
                    f'000 - Pastas Funcionais\\00 - ATIVOS\\0 - Estagiários\\'
                    f'0 - Ainda nao iniciaram\\{str(cadastro["nome"]).upper()}\\Recibo')
                os.makedirs(
                    f'\\{pasta}\\rh\\01 - RH\\01 - Administração.Controles\\02 - Funcionários, Departamentos e Férias\\'
                    f'000 - Pastas Funcionais\\00 - ATIVOS\\0 - Estagiários\\'
                    f'0 - Ainda nao iniciaram\\{str(cadastro["nome"]).upper()}\\Rescisao')
            except FileExistsError:
                pass
            pasta_contratuais = f'\\{pasta}\\rh\\01 - RH\\01 - Administração.Controles\\' \
                                f'02 - Funcionários, Departamentos e Férias\\000 - Pastas Funcionais\\' \
                                f'00 - ATIVOS\\0 - Estagiários\\0 - Ainda nao iniciaram\\' \
                                f'{str(cadastro["nome"]).upper()}\\Contratuais'

            # change tree docx models files with intern data and save pdfs files
            solicitacao = docx.Document(modelo + r'\Solicitacao MODELO - Copia.docx')
            solicitacao.tables[0].rows[4].cells[0].paragraphs[0].text = str(
                solicitacao.tables[0].rows[4].cells[0].paragraphs[0].text).replace('#supervisor_estagio', f'{lot[1]}')
            solicitacao.tables[0].rows[5].cells[1].paragraphs[0].text = str(
                solicitacao.tables[0].rows[5].cells[1].paragraphs[0].text).replace('#cargo', f'{lot[3]}')
            solicitacao.tables[0].rows[6].cells[0].paragraphs[0].text = str(
                solicitacao.tables[0].rows[6].cells[0].paragraphs[0].text).replace('#email_supervisor', f'{lot[2]}')
            solicitacao.tables[0].rows[9].cells[0].paragraphs[0].text = str(
                solicitacao.tables[0].rows[9].cells[0].paragraphs[0].text).replace('#horario', cadastro['horario'])
            solicitacao.tables[0].rows[14].cells[0].paragraphs[0].text = str(
                solicitacao.tables[0].rows[14].cells[0].paragraphs[0].text).replace('#nome_completo', cadastro['nome'])
            solicitacao.tables[0].rows[15].cells[0].paragraphs[0].text = str(
                solicitacao.tables[0].rows[15].cells[0].paragraphs[0].text).replace('#nasc',
                                                                                    dt.strftime(cadastro['nasc_ed'],
                                                                                                '%d/%m/%Y')
                                                                                    ).replace('#rg',
                                                                                              cadastro['rg']).replace(
                '#cpf', cadastro['cpf'])
            solicitacao.tables[0].rows[16].cells[0].paragraphs[0].text = str(
                solicitacao.tables[0].rows[16].cells[0].paragraphs[0].text).replace('#sexo', cadastro['genero'])
            solicitacao.tables[0].rows[17].cells[0].paragraphs[0].text = str(
                solicitacao.tables[0].rows[17].cells[0].paragraphs[0].text).replace('#endereco', cadastro['end'])
            solicitacao.tables[0].rows[18].cells[0].paragraphs[0].text = str(
                solicitacao.tables[0].rows[18].cells[0].paragraphs[0].text).replace('#cep', cadastro['cep']).replace(
                '#bairro', cadastro['bairro'])
            solicitacao.tables[0].rows[19].cells[0].paragraphs[0].text = str(
                solicitacao.tables[0].rows[19].cells[0].paragraphs[0].text).replace('#telefone',
                                                                                    str(cadastro['tel']).replace('.0',
                                                                                                                 ''))
            solicitacao.tables[0].rows[20].cells[0].paragraphs[0].text = str(
                solicitacao.tables[0].rows[20].cells[0].paragraphs[0].text).replace('#end_eletr', cadastro['email'])
            solicitacao.tables[0].rows[22].cells[0].paragraphs[0].text = str(
                solicitacao.tables[0].rows[22].cells[0].paragraphs[0].text).replace('#semestre',
                                                                                    str(cadastro['semestre']).replace(
                                                                                        '.0', ''))
            solicitacao.tables[0].rows[23].cells[0].paragraphs[0].text = str(
                solicitacao.tables[0].rows[23].cells[0].paragraphs[0].text).replace('#turno',
                                                                                    cadastro['turno']).replace(
                '#ano_concl', str(cadastro['conclusao']).replace('.0', ''))
            solicitacao.tables[0].rows[24].cells[0].paragraphs[0].text = str(
                solicitacao.tables[0].rows[24].cells[0].paragraphs[0].text).replace('#faculdade', cadastro['faculdade'])
            solicitacao.save(pasta_contratuais + f'\\Pedido TCE {str(cadastro["nome"]).split(" ")[0]}.docx')
            docx2pdf.convert(pasta_contratuais + f'\\Pedido TCE {str(cadastro["nome"]).split(" ")[0]}.docx',
                             pasta_contratuais + f'\\Pedido TCE {str(cadastro["nome"]).split(" ")[0]}.pdf')
            os.remove(pasta_contratuais + f'\\Pedido TCE {str(cadastro["nome"]).split(" ")[0]}.docx')

            ficha_cadastral = docx.Document(modelo + r'\Ficha Cadastral MODELO - Copia.docx')
            ficha_cadastral.tables[1].rows[0].cells[0].paragraphs[0].text = str(
                ficha_cadastral.tables[1].rows[0].cells[0].paragraphs[0].text).replace('#nome_completo',
                                                                                       cadastro['nome'])
            ficha_cadastral.tables[1].rows[1].cells[0].paragraphs[0].text = str(
                ficha_cadastral.tables[1].rows[1].cells[0].paragraphs[0].text) \
                .replace('#nasc', dt.strftime(cadastro['nasc_ed'], '%d/%m/%Y'))
            ficha_cadastral.tables[1].rows[1].cells[2].paragraphs[0].text = str(
                ficha_cadastral.tables[1].rows[1].cells[2].paragraphs[0].text).replace('#gen', cadastro['genero'])
            ficha_cadastral.tables[1].rows[1].cells[4].paragraphs[0].text = str(
                ficha_cadastral.tables[1].rows[1].cells[4].paragraphs[0].text).replace('#est_civ', cadastro['est_civ'])
            ficha_cadastral.tables[1].rows[2].cells[0].paragraphs[0].text = str(
                ficha_cadastral.tables[1].rows[2].cells[0].paragraphs[0].text).replace('#local', cadastro['end'])
            ficha_cadastral.tables[1].rows[2].cells[4].paragraphs[0].text = str(
                ficha_cadastral.tables[1].rows[2].cells[4].paragraphs[0].text).replace('#qd', cadastro['bairro'])
            ficha_cadastral.tables[1].rows[2].cells[7].paragraphs[0].text = str(
                ficha_cadastral.tables[1].rows[2].cells[7].paragraphs[0].text).replace('#codigo', cadastro['cep'])
            ficha_cadastral.tables[1].rows[4].cells[1].paragraphs[0].text = str(
                ficha_cadastral.tables[1].rows[4].cells[1].paragraphs[0].text).replace('#telefone', cadastro['tel'])
            ficha_cadastral.tables[1].rows[4].cells[5].paragraphs[0].text = str(
                ficha_cadastral.tables[1].rows[4].cells[5].paragraphs[0].text).replace('#ident', cadastro['rg'])
            ficha_cadastral.tables[1].rows[5].cells[1].paragraphs[0].text = str(
                ficha_cadastral.tables[1].rows[5].cells[1].paragraphs[0].text).replace('#cpf#', cadastro['cpf'])
            ficha_cadastral.tables[1].rows[6].cells[3].paragraphs[0].text = str(
                ficha_cadastral.tables[1].rows[6].cells[3].paragraphs[0].text).replace('#pai#', cadastro['pai'])
            ficha_cadastral.tables[1].rows[7].cells[1].paragraphs[0].text = str(
                ficha_cadastral.tables[1].rows[7].cells[1].paragraphs[0].text).replace('#mae#', cadastro['mae'])
            ficha_cadastral.tables[1].rows[8].cells[0].paragraphs[0].text = str(
                ficha_cadastral.tables[1].rows[8].cells[0].paragraphs[0].text).replace('#end_eletr', cadastro['email'])
            ficha_cadastral.tables[1].rows[8].cells[1].paragraphs[0].text = str(
                ficha_cadastral.tables[1].rows[8].cells[1].paragraphs[0].text).replace('#depart', str(sh["AG3"].value))
            ficha_cadastral.save(pasta_contratuais + f'\\Ficha Cadastral {str(cadastro["nome"]).split(" ")[0]}.docx')
            docx2pdf.convert(pasta_contratuais + f'\\Ficha Cadastral {str(cadastro["nome"]).split(" ")[0]}.docx',
                             pasta_contratuais + f'\\Ficha Cadastral {str(cadastro["nome"]).split(" ")[0]}.pdf')
            os.remove(pasta_contratuais + f'\\Ficha Cadastral {str(cadastro["nome"]).split(" ")[0]}.docx')

            carta_banco = docx.Document(
                modelo + r'\Abertura Conta MODELO.docx')
            carta_banco.paragraphs[14].text = str(carta_banco.paragraphs[14].text).replace('#nome_completo',
                                                                                           cadastro['nome']
                                                                                           ).replace('#rg',
                                                                                                     cadastro['rg']
                                                                                                     ).replace(
                '#cpf', cadastro['cpf']).replace('#endereço', cadastro['end']).replace('#cep', cadastro['cep']).replace(
                '#bairro', cadastro['bairro']).replace('#desde#', dt.strftime(hoje, '%d/%m/%Y'))
            carta_banco.save(pasta_contratuais + f'\\Carta Banco {str(cadastro["nome"]).split(" ")[0]}.docx')
            docx2pdf.convert(pasta_contratuais + f'\\Carta Banco {str(cadastro["nome"]).split(" ")[0]}.docx',
                             pasta_contratuais + f'\\Carta Banco {str(cadastro["nome"]).split(" ")[0]}.pdf')
            os.remove(pasta_contratuais + f'\\Carta Banco {str(cadastro["nome"]).split(" ")[0]}.docx')
            # set up smtpp connection
            s = smtplib.SMTP(host=host, port=port)
            s.starttls()
            s.login(email_remetente, senha)

            # send e-mail to intern with a pdf file so he/she can go to bank to open an account
            msg = MIMEMultipart('alternative')
            arquivo = pasta_contratuais + f'\\Carta Banco {str(cadastro["nome"]).split(" ")[0]}.pdf'
            text = MIMEText(f'''Olá, {str(cadastro["nome"]).split(" ")[0]}!<br><br>
            Segue sua carta para abertura de conta bancária no Itaú.<br>
            Você deve abrir a conta antes de iniciar os trabalhos no estágio. Ok?<br><br>
            Seu contrato já foi solicitado ao Instituto Fecomércio por e-mail, com você em cópia.<br>
            Verifique os dias e horários de funcionamento do IF e compareça no instituto para retirar seu contrato.<br>
            Será necessário levar uma declaração de matrícula atual do seu curso.<br>
            Ao chegar no IF, caso eles não tenham recebido o pedido do seu contrato, favor encaminhar para eles o e-mail do pedido, que você está em cópia.<br><br>
            Atenciosamente,<br>
            <img src="cid:image1">''', 'html')
            msg.attach(text)
            image = MIMEImage(
                open(rf'C:\Users\{os.getlogin()}\PycharmProjects\AutomacaoCia\src\models\static\imgs\assinatura.png',
                     'rb').read())
            image.add_header('Content-ID', '<image1>')
            msg.attach(image)
            # set up the parameters of the message
            msg['From'] = email_remetente
            msg['To'] = cadastro['email']
            msg['Subject'] = "Carta para Abertura de conta"
            # attach pdf file
            part = MIMEBase('application', "octet-stream")
            part.set_payload(open(arquivo, "rb").read())
            encoders.encode_base64(part)
            part.add_header('Content-Disposition', 'attachment',
                            filename=f'Carta Banco {str(cadastro["nome"]).split(" ")[0]}.pdf')
            msg.attach(part)
            s.sendmail(email_remetente, cadastro['email'], msg.as_string())
            del msg

            # send e-mail to coworker asking to register the intern
            msg = MIMEMultipart('alternative')
            arquivo = pasta_contratuais + f'\\Ficha Cadastral {str(cadastro["nome"]).split(" ")[0]}.pdf'
            text = MIMEText(f'''Oi, Wallace!<br><br>
            Segue a ficha cadastral do(a) estagiário(a) {cadastro["nome"]}.<br><br>
            Abs.,<br>
            <img src="cid:image1">''', 'html')
            msg.attach(text)
            image = MIMEImage(
                open(rf'C:\Users\{os.getlogin()}\PycharmProjects\AutomacaoCia\src\models\static\imgs\assinatura.png',
                     'rb').read())
            # Define the image's ID as referenced in the HTML body above
            image.add_header('Content-ID', '<image1>')
            msg.attach(image)
            # set up the parameters of the message
            msg['From'] = email_remetente
            msg['To'] = em_ti
            msg['Subject'] = f"Ficha Cadastral {str(cadastro['nome']).split(' ')[0]}"
            # attach pdf file
            part = MIMEBase('application', "octet-stream")
            part.set_payload(open(arquivo, "rb").read())
            encoders.encode_base64(part)
            part.add_header('Content-Disposition', 'attachment',
                            filename=f'Ficha Cadastral {str(cadastro["nome"]).split(" ")[0]}.pdf')
            msg.attach(part)
            s.sendmail(email_remetente, em_ti, msg.as_string())
            del msg

            # send document asking for the intern contract
            msg = MIMEMultipart('alternative')
            arquivo = pasta_contratuais + f'\\Pedido TCE {str(cadastro["nome"]).split(" ")[0]}.pdf'
            text = MIMEText(
                f'''Olá!<br><br>Segue pedido de TCE do(a) estagiário(a) {cadastro["nome"]}.<br>
                Caso não recebam o anexo, o estagiário está em cópia e pode encaminhar no momento da retirada do TCE.
                <br><br>Atenciosamente,<br><img src="cid:image1">''',
                'html')
            msg.attach(text)
            image = MIMEImage(
                open(rf'C:\Users\{os.getlogin()}\PycharmProjects\AutomacaoCia\src\models\static\imgs\assinatura.png',
                     'rb').read())
            # Define the image's ID as referenced in the HTML body above
            image.add_header('Content-ID', '<image1>')
            msg.attach(image)
            # set up the parameters of the message
            msg['From'] = email_remetente
            msg['To'] = em_if
            msg['Cc'] = cadastro['email']
            msg['Subject'] = f"Pedido TCE {str(cadastro['nome']).split(' ')[0]}"
            # attach pdf file
            part = MIMEBase('application', "octet-stream")
            part.set_payload(open(arquivo, "rb").read())
            encoders.encode_base64(part)
            part.add_header('Content-Disposition', 'attachment',
                            filename=f'Pedido TCE {str(cadastro["nome"]).split(" ")[0]}.pdf')
            msg.attach(part)
            s.sendmail(email_remetente, [em_if, cadastro['email'], em_pnt], msg.as_string())
            del msg
            s.quit()
            tkinter.messagebox.showinfo(title='E-mails ok!', message='E-mails enviados com sucesso')
        else:
            if tce == '':
                tkinter.messagebox.showinfo(
                    title='Erro de preenchimento',
                    message='O TCE em PDF deve ser anexado!'
                )
            else:
                if editar == 0:
                    wb = l_w(caminho)
                    sh = wb['Respostas ao formulário 1']
                    num, name = nome.strip().split(' - ')
                    linha = int(num)
                    lotacao = {'UNIDADE PARK SUL - QUALQUER DEPARTAMENTO': '0013', 'KIDS': '0010', 'MUSCULAÇÃO': '0007',
                               'ESPORTES E LUTAS': '0008', 'CROSSFIT': '0012', 'GINÁSTICA': '0006', 'GESTANTES': '0008',
                               'RECEPÇÃO': '0003',
                               'FINANCEIRO': '0001', 'TI': '0001', 'MARKETING': '0001', 'MANUTENÇÃO': '0004'}
                    if str(sh[f'E{linha}'].value) == 'Masculino':
                        cargo = 'ESTAGIARIO'
                    else:
                        cargo = 'ESTAGIARIA'

                    estag = session.query(Colaborador).filter_by(matricula=matricula).first()
                    if estag:
                        pass
                    else:
                        estag_cadastrado = Colaborador(
                            matricula=matricula, nome=name.upper(), admiss=admissao,
                            nascimento=str(sh[f'D{linha}'].value),
                            cpf=str(sh[f'V{linha}'].value).replace('.', '').replace('-', '').zfill(11),
                            rg=str(int(sh[f'W{linha}'].value)),
                            emissor='SSP/DF', email=str(sh[f'B{linha}'].value),
                            genero=str(sh[f'E{linha}'].value),
                            estado_civil=str(sh[f'F{linha}'].value), cor='9',
                            instru='08 - Educação Superior Incompleta',
                            nacional='Brasileiro(a)',
                            pai=str(sh[f'M{linha}'].value).upper(),
                            mae=str(sh[f'N{linha}'].value).upper(), endereco=str(sh[f'O{linha}'].value),
                            num='1',
                            bairro=str(sh[f'Q{linha}'].value),
                            cep=str(sh[f'R{linha}'].value).replace('.', '').replace('-', ''),
                            cidade='Brasília', cid_nas='Brasília - DF',
                            uf='DF',
                            cod_municipioend=municipios['DF']['Brasília'],
                            tel=str(sh[f'U{linha}'].value).replace('(', '').replace(')', '').replace('.', '').replace(
                                '-',
                                ''),
                            depto=depto, cargo=cargo,
                            horario=str(sh[f'AI{linha}'].value), salario=salario, tipo_contr=tipo_contr,
                            hr_sem='25', hr_mens='100',
                            est_semestre=str(sh[f'AS{linha}'].value),
                            est_turno=str(sh[f'AT{linha}'].value),
                            est_prev_conclu=str(sh[f'AU{linha}'].value),
                            est_faculdade=str(sh[f'AV{linha}'].value),
                            est_endfacul='End',
                            est_numendfacul='1',
                            est_bairroendfacul='Bairro',
                            ag=agencia, conta=conta, cdigito=digito

                        )
                        session.add(estag_cadastrado)
                        session.commit()
                    estag = session.query(Colaborador).filter_by(matricula=matricula).first()
                    pasta = r'\192.168.0.250'
                    try:
                        os.makedirs(
                            f'\\{pasta}\\rh\\01 - RH\\01 - Administração.Controles\\'
                            f'02 - Funcionários, Departamentos e Férias\\000 - Pastas Funcionais\\'
                            f'00 - ATIVOS\\0 - Estagiários\\0 - Ainda nao iniciaram\\{estag.nome}')
                        os.makedirs(
                            f'\\{pasta}\\rh\\01 - RH\\01 - Administração.Controles\\'
                            f'02 - Funcionários, Departamentos e Férias\\000 - Pastas Funcionais\\'
                            f'00 - ATIVOS\\0 - Estagiários\\0 - Ainda nao iniciaram\\{estag.nome}\\Atestados')
                        os.makedirs(
                            f'\\{pasta}\\rh\\01 - RH\\01 - Administração.Controles\\'
                            f'02 - Funcionários, Departamentos e Férias\\000 - Pastas Funcionais\\'
                            f'00 - ATIVOS\\0 - Estagiários\\0 - Ainda nao iniciaram\\{estag.nome}\\Diversos')
                        os.makedirs(
                            f'\\{pasta}\\rh\\01 - RH\\01 - Administração.Controles\\'
                            f'02 - Funcionários, Departamentos e Férias\\000 - Pastas Funcionais\\'
                            f'00 - ATIVOS\\0 - Estagiários\\0 - Ainda nao iniciaram\\{estag.nome}\\Contratuais')
                        os.makedirs(
                            f'\\{pasta}\\rh\\01 - RH\\01 - Administração.Controles\\'
                            f'02 - Funcionários, Departamentos e Férias\\000 - Pastas Funcionais\\'
                            f'00 - ATIVOS\\0 - Estagiários\\0 - Ainda nao iniciaram\\{estag.nome}\\Ferias')
                        os.makedirs(
                            f'\\{pasta}\\rh\\01 - RH\\01 - Administração.Controles\\'
                            f'02 - Funcionários, Departamentos e Férias\\000 - Pastas Funcionais\\'
                            f'00 - ATIVOS\\0 - Estagiários\\0 - Ainda nao iniciaram\\{estag.nome}\\Ponto')
                        os.makedirs(
                            f'\\{pasta}\\rh\\01 - RH\\01 - Administração.Controles\\'
                            f'02 - Funcionários, Departamentos e Férias\\000 - Pastas Funcionais\\'
                            f'00 - ATIVOS\\0 - Estagiários\\0 - Ainda nao iniciaram\\{estag.nome}\\Recibo')
                        os.makedirs(
                            f'\\{pasta}\\rh\\01 - RH\\01 - Administração.Controles\\'
                            f'02 - Funcionários, Departamentos e Férias\\000 - Pastas Funcionais\\'
                            f'00 - ATIVOS\\0 - Estagiários\\0 - Ainda nao iniciaram\\{estag.nome}\\Rescisao')
                    except FileExistsError:
                        pass
                    # abrir cadastro no dexion e atualizar informações campo a campo
                    pessoa = session.query(Colaborador).filter_by(matricula=matricula).first()
                    pastapessoa = f'\\{pasta}\\rh\\01 - RH\\01 - Administração.Controles\\' \
                                  f'02 - Funcionários, Departamentos e Férias\\000 - Pastas Funcionais\\' \
                                  f'00 - ATIVOS\\0 - Estagiários\\0 - Ainda nao iniciaram\\{pessoa.nome}'
                    pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Dexion.png')))
                    pa.press('alt'), pa.press('a'), pa.press('t'), t.sleep(2), pa.press(
                        'i'), t.sleep(5), pa.write(str(pessoa.matricula)), pa.press('enter')
                    if ondestou == 0:
                        t.sleep(40)
                    else:
                        t.sleep(60)
                    pp.copy(pessoa.nome), pa.hotkey('ctrl', 'v'), pa.press('tab'), pa.write(pessoa.cpf)
                    t.sleep(1), pa.press('tab', 3), pa.write(pessoa.genero), pa.press('tab'), pa.write(pessoa.cor)
                    t.sleep(1), pa.press('tab', 2), pa.write(pessoa.instru)
                    t.sleep(1), pa.press('tab')
                    if str(pessoa.estado_civil) == 'Casado(a)':
                        pa.write('2')
                        pa.press('tab', 6)
                    else:
                        pa.write('1')
                        pa.press('tab', 5)
                    pa.write(dt.strftime(dt.strptime(pessoa.nascimento, '%Y-%m-%d %H:%M:%S'), '%d%m%Y'))
                    t.sleep(1), pa.press('tab'), pp.copy(pessoa.cid_nas), pa.hotkey('ctrl', 'v'), pa.press('tab')
                    t.sleep(1), pa.write(pessoa.uf), pa.press('tab'), pa.write(pessoa.cod_municipioend), pa.press('tab')
                    t.sleep(1), pp.copy(pessoa.pai), pa.hotkey('ctrl', 'v'), pa.press('tab'), pa.write('105'), pa.press(
                        'tab')
                    t.sleep(1), pp.copy(pessoa.mae), pa.hotkey('ctrl', 'v'), pa.press('tab'), pa.write('105')
                    # # clique em documentos
                    pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Documentos.png')))
                    pa.press('tab'), pa.write(str(pessoa.rg)), pa.press('tab'), pa.write(
                        pessoa.emissor), pa.press('tab', 3), pa.write(pessoa.cod_municipioend)
                    pa.press('tab', 9), pa.write('3'), pa.press('tab'), pa.write('341'), pa.press('tab')
                    pa.write(pessoa.ag), pa.press('tab'), pa.write(f'{pessoa.conta}-{pessoa.cdigito}'), pa.press('tab')
                    # #clique em endereço
                    pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Endereco.png')))
                    pa.press('tab', 2), pp.copy(pessoa.endereco), pa.hotkey('ctrl', 'v'), pa.press(
                        'tab'), pa.write(pessoa.num), pa.press('tab', 2)
                    pp.copy(pessoa.bairro), pa.hotkey('ctrl', 'v'), pa.press('tab'), pp.copy(pessoa.cidade), pa.hotkey(
                        'ctrl', 'v'), pa.press('tab'), pa.write(pessoa.uf)
                    pa.press('tab'), pa.write(pessoa.cep), pa.press('tab'), pa.write(pessoa.cod_municipioend), pa.press(
                        'tab'), pa.write(str(pessoa.tel)), pa.press('tab', 2), pa.write(pessoa.email)
                    # #clique em dados contratuais
                    pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Contratuais.png')))
                    pa.press('tab'), pa.write(str(pessoa.admiss).replace('/', '')), t.sleep(1)
                    pa.press('tab'), pa.write('9')
                    pa.press('tab', 7), pa.write('n'), pa.press('tab'), pa.write('4')
                    pa.press('tab'), pa.write('Ed. Fisica')
                    pa.press('tab', 2), pa.write(str(int(str(pessoa.admiss).replace('/', '')) + 2).zfill(8))
                    t.sleep(2)
                    # #clique em instituição de ensino
                    while 1:
                        if pa.locateOnScreen('../models/static/imgs/faculdade.png'):
                            pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/faculdade.png')))
                            break
                        else:
                            t.sleep(5)
                    pa.press('tab'), pp.copy(pessoa.est_faculdade), pa.hotkey('ctrl', 'v')
                    pa.press('tab'), pp.copy(pessoa.est_endfacul), pa.hotkey('ctrl', 'v')
                    pa.press('tab'), pp.copy(pessoa.est_numendfacul), pa.hotkey('ctrl', 'v')
                    pa.press('tab'), pp.copy(pessoa.est_bairroendfacul), pa.hotkey('ctrl', 'v')
                    # #clique em Outros
                    while 1:
                        if pa.locateOnScreen('../models/static/imgs/Outros.png'):
                            pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Outros.png')))
                            break
                        else:
                            t.sleep(5)
                    t.sleep(2), pa.write('CARGO GERAL')
                    pa.press('tab'), pa.write(pessoa.cargo)
                    t.sleep(1), pa.press('tab'), pa.write(pessoa.salario), pa.press('tab')
                    if str(pessoa.tipo_contr) == 'Horista':
                        pa.press('1')
                    else:
                        pa.press('5')
                    pa.press('tab', 2), t.sleep(1), pa.write(str(pessoa.hr_sem)), pa.press('tab'), t.sleep(1)
                    pa.write(str(pessoa.hr_mens))
                    # #clique em eventos trabalhistas
                    pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/EVTrab.png')))
                    t.sleep(1)
                    # #clique em lotação
                    pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Lotacoes.png')))
                    pa.press('tab'), pa.press('tab'), pa.write('i'), pa.write(str(pessoa.admiss).replace('/', ''))
                    t.sleep(1), pa.press('enter'), t.sleep(1)
                    pp.copy(lotacao[f'{pessoa.depto}']), pa.hotkey('ctrl', 'v'), pa.press('enter'), pa.write('f')
                    pa.press('tab'), pa.write('4')
                    # #clique em salvar lotação
                    pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Salvarbtn.png'))), t.sleep(1)
                    # #clique em fechar lotação
                    while 1:
                        if pa.locateOnScreen('../models/static/imgs/Fecharlot.png'):
                            pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Fecharlot.png')))
                            break
                        else:
                            t.sleep(5)
                    # #clique em Compatibilidade
                    pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Compatibilidade3.png'))), t.sleep(1)
                    # #clique em Compatibilidade de novo
                    pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Compatibilidade2.png'))), t.sleep(1)
                    pa.press('tab', 2), pa.write('9')
                    # #clique em Salvar
                    pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Salvarcadastro.png'))), t.sleep(10)
                    # #clique em fechar novo cadastro
                    pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Fecharnovo1.png'))), t.sleep(2)
                    # #clique em fechar trabalhadores
                    pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Fechartrab1.png'))), t.sleep(0.5)
                    os.rename(pastapessoa,
                              f'\\{pasta}\\rh\\01 - RH\\01 - Administração.Controles\\'
                              f'02 - Funcionários, Departamentos e Férias\\000 - Pastas Funcionais\\'
                              f'00 - ATIVOS\\0 - Estagiários\\{pessoa.nome}')
                    shutil.move(tce,
                              f'\\{pasta}\\rh\\01 - RH\\01 - Administração.Controles\\'
                              f'02 - Funcionários, Departamentos e Férias\\000 - Pastas Funcionais\\'
                              f'00 - ATIVOS\\0 - Estagiários\\{pessoa.nome}\\Contratuais\\TCE {pessoa.nome}.pdf')
                    while 1:
                        if pa.locateOnScreen('../models/static/imgs/pyt.png'):
                            pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/pyt.png')))
                            break
                        else:
                            t.sleep(5)
                    tkinter.messagebox.showinfo(
                        title='Cadastro ok!',
                        message='Cadastro realizado com sucesso!'
                    )
                else:
                    if ondestou == 0:
                        # Editando o cadastro na Cia
                        wb = l_w(caminho)
                        sh = wb['Respostas ao formulário 1']
                        num, name = nome.strip().split(' - ')
                        linha = int(num)
                        lotacao = {
                            'Unidade Park Sul - Qualquer Departamento': ['0013', lidermusc, elidermusc, 'Líder Park Sul'],
                            'Kids': ['0010', liderkids, eliderkids, 'Líder Kids'],
                            'Musculação': ['0007', lidermusc, elidermusc, 'Líder Musculação'],
                            'Esportes E Lutas': ['0008', lidernat, elidernat, 'Líder Natação'],
                            'Crossfit': ['0012', lidercross, elidercross, 'Líder Crossfit'],
                            'Ginástica': ['0006', lidergin, elidergin, 'Líder Ginástica'],
                            'Gestantes': ['0006', gerentetec, egerentetec, 'Gerente Técnico'],
                            'Recepção': ['0003', gerentevend, egerentevend, 'Gerente Vendas'],
                            'Administrativo': ['0001', gerenterh, egerenterh, 'Gerente RH'],
                            'Manutenção': ['0004', gerentemanut, egerentemanut, 'Gerente Manutenção'],
                        }
                    else:
                        # Editando o cadastro em Casa
                        wb = l_w(caminho)
                        sh = wb['Respostas ao formulário 1']
                        num, name = nome.strip().split(' - ')
                        linha = int(num)
                        lotacao = {'UNIDADE PARK SUL - QUALQUER DEPARTAMENTO': '0013', 'KIDS': '0010', 'MUSCULAÇÃO': '0007',
                                   'ESPORTES E LUTAS': '0008', 'CROSSFIT': '0012', 'GINÁSTICA': '0006', 'GESTANTES': '0008',
                                   'RECEPÇÃO': '0003',
                                   'FINANCEIRO': '0001', 'TI': '0001', 'MARKETING': '0001', 'MANUTENÇÃO': '0004'}
                        if str(sh[f'E{linha}'].value) == 'Masculino':
                            cargo = 'ESTAGIARIO'
                        else:
                            cargo = 'ESTAGIARIA'

                        estag = session.query(Colaborador).filter_by(matricula=matricula).first()
                        if estag:
                            pass
                        else:
                            estag_cadastrado = Colaborador(
                                matricula=matricula, nome=name.upper(), admiss=admissao,
                                nascimento=str(sh[f'D{linha}'].value),
                                cpf=str(int(sh[f'V{linha}'].value)).zfill(11),
                                rg=str(int(sh[f'W{linha}'].value)),
                                emissor='SSP/DF', email=str(sh[f'B{linha}'].value),
                                genero=str(sh[f'E{linha}'].value),
                                estado_civil=str(sh[f'F{linha}'].value), cor='9',
                                instru='08 - Educação Superior Incompleta',
                                nacional='Brasileiro(a)',
                                pai=str(sh[f'M{linha}'].value).upper(),
                                mae=str(sh[f'N{linha}'].value).upper(), endereco=str(sh[f'O{linha}'].value),
                                num='1',
                                bairro=str(sh[f'Q{linha}'].value),
                                cep=str(sh[f'R{linha}'].value).replace('.', '').replace('-', ''),
                                cidade='Brasília', cid_nas='Brasília - DF',
                                uf='DF',
                                cod_municipioend=municipios['DF']['Brasília'],
                                tel=str(sh[f'U{linha}'].value).replace('(', '').replace(')', '').replace('.', '').replace(
                                    '-',
                                    ''),
                                depto=depto, cargo=cargo,
                                horario=str(sh[f'AI{linha}'].value), salario=salario, tipo_contr=tipo_contr,
                                hr_sem='25', hr_mens='100',
                                est_semestre=str(sh[f'AS{linha}'].value),
                                est_turno=str(sh[f'AT{linha}'].value),
                                est_prev_conclu=str(sh[f'AU{linha}'].value),
                                est_faculdade=str(sh[f'AV{linha}'].value),
                                est_endfacul='End',
                                est_numendfacul='1',
                                est_bairroendfacul='Bairro',
                                ag=agencia, conta=conta, cdigito=digito
                            )
                            session.add(estag_cadastrado)
                            session.commit()
                        pasta = r'\192.168.0.250'
                        try:
                            os.makedirs(
                                f'\\{pasta}\\rh\\01 - RH\\01 - Administração.Controles\\'
                                f'02 - Funcionários, Departamentos e Férias\\000 - Pastas Funcionais\\'
                                f'00 - ATIVOS\\0 - Estagiários\\0 - Ainda nao iniciaram\\{estag.nome}')
                            os.makedirs(
                                f'\\{pasta}\\rh\\01 - RH\\01 - Administração.Controles\\'
                                f'02 - Funcionários, Departamentos e Férias\\000 - Pastas Funcionais\\'
                                f'00 - ATIVOS\\0 - Estagiários\\0 - Ainda nao iniciaram\\{estag.nome}\\Atestados')
                            os.makedirs(
                                f'\\{pasta}\\rh\\01 - RH\\01 - Administração.Controles\\'
                                f'02 - Funcionários, Departamentos e Férias\\000 - Pastas Funcionais\\'
                                f'00 - ATIVOS\\0 - Estagiários\\0 - Ainda nao iniciaram\\{estag.nome}\\Diversos')
                            os.makedirs(
                                f'\\{pasta}\\rh\\01 - RH\\01 - Administração.Controles\\'
                                f'02 - Funcionários, Departamentos e Férias\\000 - Pastas Funcionais\\'
                                f'00 - ATIVOS\\0 - Estagiários\\0 - Ainda nao iniciaram\\{estag.nome}\\Contratuais')
                            os.makedirs(
                                f'\\{pasta}\\rh\\01 - RH\\01 - Administração.Controles\\'
                                f'02 - Funcionários, Departamentos e Férias\\000 - Pastas Funcionais\\'
                                f'00 - ATIVOS\\0 - Estagiários\\0 - Ainda nao iniciaram\\{estag.nome}\\Ferias')
                            os.makedirs(
                                f'\\{pasta}\\rh\\01 - RH\\01 - Administração.Controles\\'
                                f'02 - Funcionários, Departamentos e Férias\\000 - Pastas Funcionais\\'
                                f'00 - ATIVOS\\0 - Estagiários\\0 - Ainda nao iniciaram\\{estag.nome}\\Ponto')
                            os.makedirs(
                                f'\\{pasta}\\rh\\01 - RH\\01 - Administração.Controles\\'
                                f'02 - Funcionários, Departamentos e Férias\\000 - Pastas Funcionais\\'
                                f'00 - ATIVOS\\0 - Estagiários\\0 - Ainda nao iniciaram\\{estag.nome}\\Recibo')
                            os.makedirs(
                                f'\\{pasta}\\rh\\01 - RH\\01 - Administração.Controles\\'
                                f'02 - Funcionários, Departamentos e Férias\\000 - Pastas Funcionais\\'
                                f'00 - ATIVOS\\0 - Estagiários\\0 - Ainda nao iniciaram\\{estag.nome}\\Rescisao')
                        except FileExistsError:
                            pass
                        # abrir cadastro no dexion e atualizar informações campo a campo
                        pessoa = session.query(Colaborador).filter_by(matricula=matricula).first()
                        pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Dexion.png')))
                        pa.press('alt'), pa.press('a'), pa.press('t'), t.sleep(2), pa.press(
                            'a'), t.sleep(5), pa.write(str(pessoa.matricula)), pa.press('enter'), t.sleep(20)
                        pp.copy(pessoa.nome), pa.hotkey('ctrl', 'v'), pa.press('tab'), pa.write(pessoa.cpf)
                        t.sleep(1), pa.press('tab', 3), pa.write(pessoa.genero), pa.press('tab'), pa.write(pessoa.cor)
                        t.sleep(1), pa.press('tab', 2), pa.write(pessoa.instru)
                        t.sleep(1), pa.press('tab')
                        if str(pessoa.estado_civil) == 'Casado(a)':
                            pa.write('2')
                            pa.press('tab', 6)
                        else:
                            pa.write('1')
                            pa.press('tab', 5)
                        pa.write(dt.strftime(dt.strptime(pessoa.nascimento, '%Y-%m-%d %H:%M:%S'), '%d%m%Y'))
                        t.sleep(1), pa.press('tab'), pp.copy(pessoa.cid_nas), pa.hotkey('ctrl', 'v'), pa.press('tab')
                        t.sleep(1), pa.write(pessoa.uf), pa.press('tab'), pa.write(pessoa.cod_municipioend), pa.press('tab')
                        t.sleep(1), pp.copy(pessoa.pai), pa.hotkey('ctrl', 'v'), pa.press('tab'), pa.write('105'), pa.press(
                            'tab')
                        t.sleep(1), pp.copy(pessoa.mae), pa.hotkey('ctrl', 'v'), pa.press('tab'), pa.write('105')
                        # # clique em documentos
                        pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Documentos.png')))
                        pa.press('tab'), pa.write(str(pessoa.rg)), pa.press('tab'), pa.write(
                            pessoa.emissor), pa.press('tab', 3), pa.write(pessoa.cod_municipioend)
                        pa.press('tab', 9), pa.write('3'), pa.press('tab'), pa.write('341'), pa.press('tab')
                        pa.write(pessoa.ag), pa.press('tab'), pa.write(f'{pessoa.conta}-{pessoa.cdigito}'), pa.press('tab')
                        # #clique em endereço
                        pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Endereco.png')))
                        pa.press('tab', 2), pp.copy(pessoa.endereco), pa.hotkey('ctrl', 'v'), pa.press(
                            'tab'), pa.write(pessoa.num), pa.press('tab', 2)
                        pp.copy(pessoa.bairro), pa.hotkey('ctrl', 'v'), pa.press('tab'), pp.copy(pessoa.cidade), pa.hotkey(
                            'ctrl', 'v'), pa.press('tab'), pa.write(pessoa.uf)
                        pa.press('tab'), pa.write(pessoa.cep), pa.press('tab'), pa.write(pessoa.cod_municipioend), pa.press(
                            'tab'), pa.write(str(pessoa.tel)), pa.press('tab', 2), pa.write(pessoa.email)
                        # #clique em dados contratuais
                        pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Contratuais.png')))
                        pa.press('tab'), pa.write(str(pessoa.admiss).replace('/', '')), t.sleep(1)
                        pa.press('tab'), pa.write('9')
                        pa.press('tab', 7), pa.write('n'), pa.press('tab'), pa.write('4')
                        pa.press('tab'), pa.write('Ed. Fisica')
                        pa.press('tab', 2), pa.write(str(int(str(pessoa.admiss).replace('/', '')) + 2).zfill(8))
                        # #clique em instituição de ensino
                        pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/faculdade.png')))
                        pa.press('tab'), pp.copy(pessoa.est_faculdade), pa.hotkey('ctrl', 'v')
                        pa.press('tab'), pp.copy(pessoa.est_endfacul), pa.hotkey('ctrl', 'v')
                        pa.press('tab'), pp.copy(pessoa.est_numendfacul), pa.hotkey('ctrl', 'v')
                        pa.press('tab'), pp.copy(pessoa.est_bairroendfacul), pa.hotkey('ctrl', 'v')
                        # #clique em Outros
                        while 1:
                            if pa.locateOnScreen('../models/static/imgs/Outros.png'):
                                pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Outros.png')))
                                break
                            else:
                                t.sleep(5)
                        t.sleep(2), pa.write('CARGO GERAL')
                        pa.press('tab'), pa.write(pessoa.cargo)
                        t.sleep(1), pa.press('tab'), pa.write(pessoa.salario), pa.press('tab')
                        if str(pessoa.tipo_contr) == 'Horista':
                            pa.press('1')
                        else:
                            pa.press('5')
                        pa.press('tab', 2), t.sleep(1), pa.write(str(pessoa.hr_sem)), pa.press('tab'), t.sleep(1)
                        pa.write(str(pessoa.hr_mens))
                        # #clique em Compatibilidade
                        pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Compatibilidade3.png'))), t.sleep(1)
                        # #clique em Compatibilidade de novo
                        pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Compatibilidade2.png'))), t.sleep(1)
                        pa.press('tab', 2), pa.write('9')
                        # #clique em Salvar
                        pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Salvarcadastro.png'))), t.sleep(10)
                        # #clique em fechar novo cadastro
                        pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Fecharnovo1.png'))), t.sleep(2)
                        # #clique em fechar trabalhadores
                        pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Fechartrab1.png'))), t.sleep(0.5)
                        while 1:
                            if pa.locateOnScreen('../models/static/imgs/pyt.png'):
                                pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/pyt.png')))
                                break
                            else:
                                t.sleep(5)
                        tkinter.messagebox.showinfo(
                            title='Cadastro ok!',
                            message='Cadastro editado com sucesso!'
                        )


def cadastrar_autonomo(caminhoaut, nomeaut, matriculaaut, admissaoaut, cargoaut, deptoaut, ondeaut):
    sessions = sessionmaker(bind=engine)
    session = sessions()
    # Cadastro iniciado em casa
    wb = l_w(caminhoaut)
    sh = wb['Respostas ao formulário 1']
    num, name = nomeaut.strip().split(' - ')
    linha = int(num)
    lotacao = {'UNIDADE PARK SUL - QUALQUER DEPARTAMENTO': '0013', 'KIDS': '0010', 'MUSCULAÇÃO': '0007',
               'ESPORTES E LUTAS': '0008', 'CROSSFIT': '0012', 'GINÁSTICA': '0006', 'GESTANTES': '0008',
               'RECEPÇÃO': '0003',
               'FINANCEIRO': '0001', 'TI': '0001', 'MARKETING': '0001', 'MANUTENÇÃO': '0004'}
    aut = session.query(Colaborador).filter_by(matricula=matriculaaut).first()
    if aut:
        pass
    else:
        aut_cadastrado = Colaborador(
            matricula=matriculaaut, nome=name.upper(), admiss=admissaoaut,
            nascimento=str(sh[f'D{linha}'].value),
            pis=str(sh[f'S{linha}'].value).replace('.', '').replace('-', '').zfill(11),
            cpf=str(int(sh[f'P{linha}'].value)).zfill(11),
            rg=str(int(sh[f'Q{linha}'].value)),
            emissor='SSP/DF', email=str(sh[f'B{linha}'].value),
            genero=str(sh[f'E{linha}'].value), cor='9',
            instru=str(sh[f'F{linha}'].value),
            nacional='Brasileiro(a)', estado_civil='Solteiro(a)',
            endereco=str(sh[f'I{linha}'].value),
            num=str(sh[f'J{linha}'].value),
            bairro=str(sh[f'K{linha}'].value), cep=str(sh[f'L{linha}'].value).replace('.', '').replace('-', ''),
            cidade=str(sh[f'M{linha}'].value), cid_nas='Brasília - DF', uf='DF',
            cod_municipioend=municipios['DF']['Brasília'],
            tel=str(sh[f'O{linha}'].value).replace('(', '').replace(')', '').replace('.', '').replace('-', ''),
            depto=deptoaut, cargo=cargoaut,
        )
        session.add(aut_cadastrado)
        session.commit()
    # abrir cadastro no dexion e atualizar informações campo a campo
    pessoa = session.query(Colaborador).filter_by(matricula=matriculaaut).first()
    pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Dexion.png')))
    pa.press('alt'), pa.press('a'), pa.press('t'), t.sleep(2), pa.press(
        'i'), t.sleep(5), pa.write(str(pessoa.matricula)), pa.press('enter'), t.sleep(20)
    pp.copy(pessoa.nome), pa.hotkey('ctrl', 'v'), pa.press('tab'), pa.write(pessoa.cpf)
    t.sleep(1), pa.press('tab', 3), pa.write(pessoa.genero), pa.press('tab'), pa.write(pessoa.cor)
    t.sleep(1), pa.press('tab', 2), pa.write(pessoa.instru)
    t.sleep(1), pa.press('tab')
    if str(pessoa.estado_civil) == 'Casado(a)':
        pa.write('2')
        pa.press('tab', 6)
    else:
        pa.write('1')
        pa.press('tab', 5)
    pa.write(dt.strftime(dt.strptime(pessoa.nascimento, '%Y-%m-%d %H:%M:%S'), '%d%m%Y'))
    t.sleep(1), pa.press('tab'), pp.copy(pessoa.cid_nas), pa.hotkey('ctrl', 'v')
    # # clique em documentos
    pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Documentos.png')))
    pa.press('tab'), pa.write(str(pessoa.rg)), pa.press('tab'), pa.write(
        pessoa.emissor), pa.press('tab', 3), pa.write(pessoa.cod_municipioend), pa.press('tab')
    pa.write(pessoa.pis)
    # #clique em endereço
    pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Endereco.png')))
    pa.press('tab', 2), pp.copy(pessoa.endereco), pa.hotkey('ctrl', 'v'), pa.press(
        'tab'), pa.write(pessoa.num), pa.press('tab', 2)
    pp.copy(pessoa.bairro), pa.hotkey('ctrl', 'v'), pa.press('tab'), pp.copy(pessoa.cidade), pa.hotkey(
        'ctrl', 'v'), pa.press('tab'), pa.write(pessoa.uf)
    pa.press('tab'), pa.write(pessoa.cep), pa.press('tab'), pa.write(pessoa.cod_municipioend), pa.press(
        'tab'), pa.write(str(pessoa.tel)), pa.press('tab', 2), pa.write(pessoa.email)
    # #clique em dados contratuais
    pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Contratuais.png')))
    pa.press('tab'), pa.write(str(pessoa.admiss).replace('/', '')), t.sleep(1)
    pa.press('tab'), pa.write('7')
    # #clique em Outros
    while 1:
        if pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Outros.png'))):
            pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Outros.png')))
            break
        else:
            t.sleep(5)
    t.sleep(2), pa.write('CARGO GERAL')
    pa.press('tab'), pa.write(pessoa.cargo)
    # #clique em eventos trabalhistas
    pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/EVTrab.png')))
    t.sleep(1)
    # #clique em lotação
    pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Lotacoes.png')))
    pa.press('tab'), pa.press('tab'), pa.write('i'), pa.write(str(pessoa.admiss).replace('/', ''))
    t.sleep(1), pa.press('enter'), t.sleep(1)
    pp.copy(lotacao[f'{pessoa.depto}']), pa.hotkey('ctrl', 'v'), pa.press('enter'), pa.write('3')
    # #clique em salvar lotação
    pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Salvarbtn.png'))), t.sleep(1)
    # #clique em fechar lotação
    while 1:
        if pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Fecharlot.png'))):
            pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Fecharlot.png')))
            break
        else:
            t.sleep(5)
    # #clique em Compatibilidade
    pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Compatibilidade3.png'))), t.sleep(1)
    # #clique em Compatibilidade de novo
    pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Compatibilidade2.png'))), t.sleep(1)
    pa.press('tab', 2), pa.write('13')
    # #clique em Salvar
    pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Salvarcadastro.png'))), t.sleep(10)
    # #clique em fechar novo cadastro
    pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Fecharnovo1.png'))), t.sleep(2)
    # #clique em fechar trabalhadores
    pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/Fechartrab1.png'))), t.sleep(0.5)
    while 1:
        if pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/pyt.png'))):
            pa.click(pa.center(pa.locateOnScreen('../models/static/imgs/pyt.png')))
            break
        else:
            t.sleep(5)
    tkinter.messagebox.showinfo(
        title='Cadastro ok!',
        message='Cadastro realizado com sucesso!'
    )


def validar_pis(local, nome):
    wb = l_w(local, read_only=False)
    sh = wb['Respostas ao formulário 1']
    num, name = nome.strip().split(' - ')
    x = int(num)
    pis = str(sh[f'S{x}'].value).replace('-', '').replace('.', '').zfill(11)
    v1 = int(pis[0]) * 3
    v2 = int(pis[1]) * 2
    v3 = int(pis[2]) * 9
    v4 = int(pis[3]) * 8
    v5 = int(pis[4]) * 7
    v6 = int(pis[5]) * 6
    v7 = int(pis[6]) * 5
    v8 = int(pis[7]) * 4
    v9 = int(pis[8]) * 3
    v10 = int(pis[9]) * 2
    d = int(pis[10])
    soma = v1 + v2 + v3 + v4 + v5 + v6 + v7 + v8 + v9 + v10
    divisao = soma % 11
    resultado = 11 - divisao
    if resultado != d:
        if resultado == 10 & d == 0 | resultado == 11 & d == 0:
            pass
        else:
            tkinter.messagebox.showinfo(
                title='Erro!',
                message='PIS inválido!'
            )
    else:
        tkinter.messagebox.showinfo(
            title='Ok!',
            message='PIS ok!'
        )


def updatedb(pessoa, nome, cargo, departamento, agencia, conta, digito):
    sessions = sessionmaker(bind=engine)
    session = sessions()

    def alteranome():
        pess = session.query(Colaborador).filter_by(nome=pessoa).order_by(Colaborador.matricula.desc()).first()
        pess.nome = nome
        session.commit()
        tkinter.messagebox.showinfo(title='Alteração efetuada!',
                                    message=f'{pessoa} teve seu NOME alterado para: {nome}.')

    def alteracargo():
        pess = session.query(Colaborador).filter_by(nome=pessoa).order_by(Colaborador.matricula.desc()).first()
        pess.cargo = cargo
        session.commit()
        tkinter.messagebox.showinfo(title='Alteração efetuada!',
                                    message=f'{pessoa} teve seu CARGO alterado para: {cargo}.')

    def alteradepto():
        pess = session.query(Colaborador).filter_by(nome=pessoa).order_by(Colaborador.matricula.desc()).first()
        pess.depto = departamento
        session.commit()
        tkinter.messagebox.showinfo(title='Alteração efetuada!',
                                    message=f'{pessoa} teve seu DEPTO alterado para: {departamento}.')

    def alteracc():
        pess = session.query(Colaborador).filter_by(nome=pessoa).order_by(Colaborador.matricula.desc()).first()
        pess.ag = agencia
        pess.conta = conta
        pess.cdigito = digito
        session.commit()
        tkinter.messagebox.showinfo(title='Alteração efetuada!',
                                    message=f'{pessoa} teve seus dados bancários alterados para:\nAg: {agencia}\nConta: {conta}\nDígito: {digito}')

    if pessoa == '':
        tkinter.messagebox.showinfo(title='Alteração não efetuada!', message='Escolha uma pessoa para editar dados!')
    else:
        if nome == '' and cargo == '' and departamento == '' and agencia == '' and conta == '' and digito == '':
            tkinter.messagebox.showinfo(title='Alteração não efetuada!',
                                        message='Escolha pelo menos uma opção para editar dados!')
        else:
            todosargs = locals()
            todosargs = {k: v for k, v in todosargs.items() if v}
            altera = {
                'nome': alteranome,
                'cargo': alteracargo,
                'departamento': alteradepto,
                'agencia': alteracc
            }
            altera[list(todosargs)[5]]()


def send_email(matriculas):
    # code to send e-mails through smtplib
    # set up smtp connection
    s = smtplib.SMTP(host=host, port=port)
    s.starttls()
    s.login(em_rem, k1)
    # send e-mails to a list of employees
    wb = l_w('Nomes e e-mails.xlsx')
    sh = wb['Dados']
    x = 1
    while x <= len(sh['A']):
        msg = MIMEMultipart()
        message = f'''
        Olá, {str(sh[f'A{x}'].value).title().split(sep=' ')[0]}!\n
        \n
        Para repor o encontro com colaboradores novatos cancelado no dia 21/03, abrimos novo horário hoje:\n
        24/03:
        14h às 15h30 - Sala 3.
        \n
        Atenciosamente,\n
        Felipe Rodrigues
        '''
        # parameters of the message
        msg['From'] = em_rem
        msg['To'] = str(sh[f'B{x}'].value).lower()
        msg['Subject'] = "Reposição Encontro Cinthia Guimarães"
        msg.attach(MIMEText(message, 'plain', _charset='utf-8'))
        s.send_message(msg)
        del msg
        x += 1
    s.quit()


def send_wpp():
    # code to send whatsapp messages through browser
    wb = l_w("AV.xlsm")
    sh = wb['Planilha1']

    for x in range(2, len(sh['A'])):
        pessoa = str(sh[f"B{x}"].value).split(' ')[0]
        email = str(sh[f"D{x}"].value)
        numero = str(sh[f"H{x}"].value)
        url = str(sh[f"F{x}"].value)
        mensagem = f'Oi {pessoa}, te enviei por e-mail(no {email}) o resultado da primeira etapa da sua avaliação de ' \
                   f'desempenho e o link da pesquisa sobre a avaliação. Antes da segunda etapa, precisamos que ' \
                   f'responda a pesquisa. Ok? Se puder responder agora, é bem rápido, dura no máximo 5 minutos. ' \
                   f'Segue o link: {url}'
        texto = urllib.parse.quote(mensagem)
        cel = urllib.parse.quote(numero)
        link = f'https://web.whatsapp.com/send?phone={cel}&text={texto}'
        if numero:
            print(link)
            print(pessoa)
            print(email)
            print(numero)
        x += 1


def desligar_pessoa(nome: str, data: str, tipo: int):
    """
    This function does all the procedures for terminating an employee: it issues documents, sends them by e-mail,
    saves them in the respective folders, moves folders and schedules appointments.

    The function works according to the type of dismissal entered at 'tipo' parameter. For each type of dismissal
    there are specific procedures to be performed.

    Through the 'tipo' parameter, the subfunction of the dictionary 'desligamento' is called by an if condition.

    :param nome: Employee's name
    :param data: Dismiss date
    :param tipo: Dismiss type
    :return: Procedures to dismiss employee
    """
    sessions = sessionmaker(engine)
    session = sessions()
    pessoa = session.query(Colaborador).filter_by(nome=nome).order_by(Colaborador.matricula.desc()).first()
    sessionsaula = sessionmaker(enginefolha)
    sessionaula = sessionsaula()
    aulas = sessionaula.query(Aulas).filter_by(professor=pessoa.nome).all()
    for aula in aulas:
        aula.status = 'Inativa'
        aula.fimgrade = data
        sessionaula.commit()
    sessionaula.close()

    # desligamento deve inativar aulas do prof a partir da data de deslig
    # enviar e-mail para a pessoa solicitando uniformes, marcando data para assinar rescisao
    # informar que em 10 dias serão pagos os saldos rescisórios
    # registrar a data de desligamento no cadastro do func
    # emitir os docs de homologação no sindicato

    def desligar_estag():
        pasta_rescisao = rede + rf'\02 - Funcionários, Departamentos e Férias\000 - Pastas Funcionais\00 - ATIVOS\0 - Estagiários\{pessoa.nome}\Rescisao'
        # # send e-mails to end intern contract
        # email_remetente = em_rem
        # senha = k1
        # # set up smtpp connection
        # s = smtplib.SMTP(host=host, port=port)
        # s.starttls()
        # s.login(email_remetente, senha)
        #
        # # send e-mail to intern
        # msg = MIMEMultipart('alternative')
        # arquivo = pasta_rescisao + f'\\TRCT.pdf'
        # text = MIMEText(f'''Olá, {pessoa.nome.split(" ")[0].title()}!<br><br>
        # Obrigado por sua dedicação no Programa Novos Talentos da Companhia Athletica de Brasília!<br>
        # Seu desligamento do estágio foi efetuado em {data}.<br>
        # Para concluirmos essa etapa precisamos que você compareça a Companhia para devolver uniformes, BTS e assinar o termo de rescisão anexo.<br>
        # Algumas faculdades exigem o termo de desligamento do IF, vou solicitar a eles que nos envie.<br>
        # Qual o melhor dia para você para nos encontrarmos na Cia Athletica e assinarmos o seu termo de rescisão?<br>
        # Assim que o termo for assinado agendamos o pagamento do valor final.<br>
        # Aguardo você me informar a melhor data para assinarmos o termo.<br><br>
        #
        # Atenciosamente,<br>
        # <img src="cid:image1">''', 'html')
        # msg.attach(text)
        # image = MIMEImage(
        #     open(rf'C:\Users\{os.getlogin()}\PycharmProjects\AutomacaoCia\src\models\static\imgs\assinatura.png',
        #          'rb').read())
        # image.add_header('Content-ID', '<image1>')
        # msg.attach(image)
        # # set up the parameters of the message
        # msg['From'] = email_remetente
        # msg['To'] = pessoa.email
        # msg['Subject'] = "Desligamento Estágio Cia Athletica"
        # # attach pdf file
        # part = MIMEBase('application', "octet-stream")
        # part.set_payload(open(arquivo, "rb").read())
        # encoders.encode_base64(part)
        # part.add_header('Content-Disposition', 'attachment',
        #                 filename=f'TRCT {pessoa.nome.split(" ")[0].title()}.pdf')
        # msg.attach(part)
        # s.sendmail(email_remetente, pessoa.email, msg.as_string())
        # del msg
        #
        # # send e-mail to coworker asking to exclude intern register
        # msg = MIMEMultipart('alternative')
        # text = MIMEText(f'''Oi, Wallace!<br><br>
        # Favor desativar o(a) estagiário(a) {pessoa.nome.title()}.<br><br>
        # Abs.,<br>
        # <img src="cid:image1">''', 'html')
        # msg.attach(text)
        # image = MIMEImage(
        #     open(rf'C:\Users\{os.getlogin()}\PycharmProjects\AutomacaoCia\src\models\static\imgs\assinatura.png',
        #          'rb').read())
        # # define the image's ID as referenced in the HTML body above
        # image.add_header('Content-ID', '<image1>')
        # msg.attach(image)
        # # set up the parameters of the message
        # msg['From'] = email_remetente
        # msg['To'] = em_ti
        # msg['Subject'] = f"Desligamento Estágio - {str(pessoa.nome).split(' ')[0].title()}"
        # s.sendmail(email_remetente, em_ti, msg.as_string())
        # del msg
        #
        # # send document asking for terminate intern's contract
        # msg = MIMEMultipart('alternative')
        # text = MIMEText(
        #     f'''Olá!<br><br>
        #     Favor desligar estagiário(a) {pessoa.nome.title()}, CPF: {pessoa.cpf}, na data {data}.<br><br>
        #     Atenciosamente,<br><img src="cid:image1">''',
        #     'html')
        # msg.attach(text)
        # image = MIMEImage(
        #     open(rf'C:\Users\{os.getlogin()}\PycharmProjects\AutomacaoCia\src\models\static\imgs\assinatura.png',
        #          'rb').read())
        # # define the image's ID as referenced in the HTML body above
        # image.add_header('Content-ID', '<image1>')
        # msg.attach(image)
        # # set up the parameters of the message
        # msg['From'] = email_remetente
        # msg['To'] = em_if
        # msg['Subject'] = f"Desligamento Estágio - {str(pessoa.nome).split(' ')[0]}"
        # s.sendmail(email_remetente, em_if, msg.as_string())
        # del msg
        # s.quit()
        os.rename(pasta_rescisao.replace(r'\Rescisao', ''), pasta_rescisao.replace(r'\Rescisao', '')
                  .replace('00 - ATIVOS', '01 - Inativos'))

    def desligar_func_apedido_com_aviso():
        # e-mail informdando data de crédito na conta e solicitando data para marcar no sindicato e dev uniformes
        print(f'{nome} foi desligado(a) em {data} com rescisão do tipo: Desligametno de Func à pedido com aviso.')

    def desligar_func_apedido_sem_aviso():
        print(f'{nome} foi desligado(a) em {data} com rescisão do tipo: Desligametno de Func à pedido sem aviso.')

    def desligar_func_por_acordo():
        # após gerada a rescisão e guia e-mail informanda dia do crédito em conta, guias de fgts e seguro
        # explicar quanto saca do fgts
        # e-mail marcando data para ir no sindicato, dev. uniformes e bts
        # e-mail para TI informando nome e CPF do funcionário/estagiário e solicitando o desligamento
        print(f'{nome} foi desligado(a) em {data} com rescisão do tipo: Desligametno de Func por acordo.')

    def desligar_func_sem_aviso():
        # gerar docs de homologação no dexion: Rescisão 5 cópias, av prév, comprovantes recolhimento inss, carta preposto,
        # folha de registro, carta abono conduta, guia de seguro desemprego(?)
        # e-mails com data do pgto., orientações do passo a passo, guias de orientação do FGTS e Seguro desemprego
        # explicar quanto saca do fgts
        # solicitar data para agendar no sindicato
        print(f'{nome} foi desligado(a) em {data} com rescisão do tipo: Desligametno de Func sem aviso.')

    def desligar_func_com_aviso():
        # após gerada a rescisão e guia e-mail informanda dia do crédito em conta, guias de fgts e seguro
        # explicar quanto saca do fgts
        # e-mail marcando data para ir no sindicato, dev. uniformes e bts

        # FOLHA DE ROSTO
        #
        # TRCT(PG 1 E 2) 5x
        #
        # COMPROVANTES PGTO TRCT (pegar no itau manualmente e salvar na pasta com nome padrao - cógigo pegar na pasta)
        # EXTRATO FGTS
        # CHAVE MOVIMENT CAIXA
        # MULTA FGTS
        # COMPROV PGTO MULTA FGTS
        # __________________________________________
        # AVISO PREVIO
        #
        # alt r s d e
        # matricula
        # tab, 2
        # desligamento
        # tab, 2
        # espaço
        # tab, 4
        # opção
        # 2
        # h(1)
        # ou
        # 7
        # dias(2)
        # tab
        # desligamento
        # tab, 2
        # 1100
        # tab, 4
        # clique
        # visualizar
        # salvar
        # pdf
        # na
        # pasta
        # do
        # func,
        # cliques
        # para
        # fechar
        #
        # ____________________________________________
        #
        # GUIA SEGURO DESEMPREGO
        # RELAÇÃO SALÁRIOS DE CONTRIBUIÇ
        # DISCRIMINAÇÃO DAS PARCELAS DE CONTRIB
        # CARTA PREPOSTO
        # ATESTADO DEMISSIONAL
        print(f'{nome} foi desligado(a) em {data} com rescisão do tipo: Desligametno de Func com aviso.')

    desligamento = {
        1: desligar_estag,
        2: desligar_func_apedido_com_aviso,
        3: desligar_func_apedido_sem_aviso,
        4: desligar_func_por_acordo,
        5: desligar_func_sem_aviso,
        6: desligar_func_com_aviso
    }
    if tipo in desligamento:
        desligamento[tipo]()

    pessoa.desligamento = data
    session.commit()
    tkinter.messagebox.showinfo(title='Desligamento ok!', message='Desligamento registrado com sucesso!')
