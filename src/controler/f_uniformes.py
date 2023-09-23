from datetime import datetime as dt
import docx
import docx2pdf
import locale
from openpyxl import load_workbook as l_w
import tkinter.filedialog
from tkinter import messagebox
import tkinter.filedialog

locale.setlocale(locale.LC_MONETARY, 'pt_BR.UTF-8')
locale.setlocale(locale.LC_ALL, 'pt_BR.UTF-8')
locale.setlocale(locale.LC_MONETARY, 'pt_BR.UTF-8')


def gerar_recibo_uniformes(local, nome, cargo, cpf, genero, tamanho1, tamanho2=''):
    relatorio = l_w(local, read_only=False)
    estoque = relatorio['Estoque']
    entregues = relatorio['Entregues']
    lista = relatorio['Nomes']
    hoje = dt.today()
    tipo, gen = genero.split(': ')
    num, pess = nome.split(' - ')
    label, cpf_ed = cpf.split(': ')
    pessoa = pess.title()

    if tamanho2 != '':
        recibo = docx.Document('recibo_uniforme.docx')
        recibo.paragraphs[11].text = str(recibo.paragraphs[11].text)\
            .replace('#nome', pessoa).replace('#num_cpf', cpf_ed).replace('#tam', tamanho1+' e '+tamanho2).replace(
            '#genero', str(gen).lower())
        recibo.paragraphs[19].text = str(recibo.paragraphs[19].text)\
            .replace('#data', dt.strftime(hoje, '%d/%m/%Y'))
        recibo.paragraphs[24].text = str(recibo.paragraphs[24].text).replace('#nome', pessoa)
        recibo.paragraphs[25].text = str(recibo.paragraphs[25].text).replace('#cargo', cargo)
        recibo.save(f'Recibo_alterado {pessoa}.docx')
        docx2pdf.convert(f'Recibo_alterado {pessoa}.docx', f'Recibo {pessoa}.pdf')
        # diminuir contagem de estoque
        if gen == 'Masculino':
            if tamanho1 == 'P':
                estoque['C4'].value = estoque['C4'].value - 1
            if tamanho1 == 'M':
                estoque['C5'].value = estoque['C5'].value - 1
            if tamanho1 == 'G':
                estoque['C6'].value = estoque['C6'].value - 1
            if tamanho1 == 'XG':
                estoque['C8'].value = estoque['C8'].value - 1
            if tamanho1 == 'GG':
                estoque['C7'].value = estoque['C7'].value - 1
            if tamanho1 == 'XGG':
                estoque['C9'].value = estoque['C9'].value - 1
            if tamanho2 == 'P':
                estoque['C4'].value = estoque['C4'].value - 1
            if tamanho2 == 'M':
                estoque['C5'].value = estoque['C5'].value - 1
            if tamanho2 == 'G':
                estoque['C6'].value = estoque['C6'].value - 1
            if tamanho2 == 'XG':
                estoque['C8'].value = estoque['C8'].value - 1
            if tamanho2 == 'GG':
                estoque['C7'].value = estoque['C7'].value - 1
            if tamanho2 == 'XGG':
                estoque['C9'].value = estoque['C9'].value - 1
        else:
            if tamanho1 == 'P':
                estoque['E4'].value = estoque['E4'].value - 1
            if tamanho1 == 'M':
                estoque['E5'].value = estoque['E5'].value - 1
            if tamanho1 == 'G':
                estoque['E6'].value = estoque['E6'].value - 1
            if tamanho1 == 'GG':
                estoque['E7'].value = estoque['E7'].value - 1
            if tamanho2 == 'P':
                estoque['E4'].value = estoque['E4'].value - 1
            if tamanho2 == 'M':
                estoque['E5'].value = estoque['E5'].value - 1
            if tamanho2 == 'G':
                estoque['E6'].value = estoque['E6'].value - 1
            if tamanho2 == 'GG':
                estoque['E7'].value = estoque['E7'].value - 1
        lista[f'E{num}'].value = tamanho1
        lista[f'F{num}'].value = 'OK'
        x = len(list(entregues.rows)) + 1
        entregues[f'A{x}'].value = pessoa
        entregues[f'B{x}'].value = 1
        entregues[f'C{x}'].value = tamanho1
        entregues[f'D{x}'].value = gen
        entregues[f'E{x}'].value = 1
        entregues[f'F{x}'].value = tamanho2
        entregues[f'G{x}'].value = gen
        relatorio.save(local)
        tkinter.messagebox.showinfo(title='Recibo ok!', message='Recibo impresso com sucesso!')
    else:
        recibo = docx.Document('recibo_uniforme.docx')
        recibo.paragraphs[11].text = str(recibo.paragraphs[11].text)\
            .replace('#nome', pessoa).replace('#num_cpf', cpf_ed)\
            .replace('#tam', tamanho1).replace('#genero', str(gen).lower())
        recibo.paragraphs[19].text = str(recibo.paragraphs[19].text)\
            .replace('#data', dt.strftime(hoje, '%d/%m/%Y'))
        recibo.paragraphs[24].text = str(recibo.paragraphs[24].text).replace('#nome', pessoa)
        recibo.paragraphs[25].text = str(recibo.paragraphs[25].text).replace('#cargo', cargo)
        recibo.save(f'Recibo_alterado {pessoa}.docx')
        docx2pdf.convert(f'Recibo_alterado {pessoa}.docx', f'Recibo {pessoa}.pdf')
        # diminuir contagem de estoque
        if gen == 'Masculino':
            if tamanho1 == 'P':
                estoque['C4'].value = estoque['C4'].value - 2
            if tamanho1 == 'M':
                estoque['C5'].value = estoque['C5'].value - 2
            if tamanho1 == 'G':
                estoque['C6'].value = estoque['C6'].value - 2
            if tamanho1 == 'XG':
                estoque['C8'].value = estoque['C8'].value - 2
            if tamanho1 == 'GG':
                estoque['C7'].value = estoque['C7'].value - 2
            if tamanho1 == 'XGG':
                estoque['C9'].value = estoque['C9'].value - 2
        else:
            if tamanho1 == 'P':
                estoque['E4'].value = estoque['E4'].value - 2
            if tamanho1 == 'M':
                estoque['E5'].value = estoque['E5'].value - 2
            if tamanho1 == 'G':
                estoque['E6'].value = estoque['E6'].value - 2
            if tamanho1 == 'GG':
                estoque['E7'].value = estoque['E7'].value - 2
        lista[f'E{num}'].value = tamanho1
        lista[f'F{num}'].value = 'OK'
        x = len(list(entregues.rows)) + 1
        entregues[f'A{x}'].value = pessoa
        entregues[f'B{x}'].value = 2
        entregues[f'C{x}'].value = tamanho1
        entregues[f'D{x}'].value = gen
        entregues[f'E{x}'].value = '-'
        entregues[f'F{x}'].value = '-'
        entregues[f'G{x}'].value = '-'
        relatorio.save(local)
    tkinter.messagebox.showinfo(title='Recibo ok!', message='Recibo impresso com sucesso!')


