import pandas as pd
import numpy as np
from openpyxl import load_workbook as l_w
import docx2pdf
import docx
from docx.shared import Pt, Cm
import os
from datetime import datetime as dt
from datetime import timedelta
import locale

locale.setlocale(locale.LC_ALL, 'pt_pt.UTF-8')
# define excel plans to work with
wb = l_w('zzBase.xlsx')
sh = wb['Funcionários e e-mail']
pl = l_w('Hrs Folha.xlsx')
fl = pl['Planilha1']

# datainicio = dt.strptime(input('Digite a data de início: '), '%d/%m/%Y')
# datafim = dt.strptime(input('Digite a data fim: '), '%d/%m/%Y')
datainicio = dt.strptime('01/06/2023', '%d/%m/%Y')
datafim = dt.strptime('30/06/2023', '%d/%m/%Y')


def intervalo(inicio, fim):
    for n in range(int((fim - inicio).days) + 1):
        yield [inicio + timedelta(n), np.nan, np.nan, np.nan, np.nan, np.nan, np.nan]


# creating dicts to save employees datas
planbase = {}
planmat = {}
planfunc = {}
plandept = {}

x = 2
while x <= len(sh['B']):
    planbase.update({str(sh[f'B{x}'].value).replace('.', '').replace('-', ''): str(sh[f'C{x}'].value).title().strip()})
    planmat.update({str(sh[f'B{x}'].value).replace('.', '').replace('-', ''): str(sh[f'A{x}'].value).title().strip()})
    planfunc.update({str(sh[f'B{x}'].value).replace('.', '').replace('-', ''): str(sh[f'D{x}'].value).strip()})
    plandept.update({str(sh[f'B{x}'].value).replace('.', '').replace('-', ''): str(sh[f'E{x}'].value).title().strip()})
    x += 1

geral = pd.read_csv(r'AFD.txt', sep=' ', header=None, encoding='iso8859-1')
geral[0] = geral[0].astype(str).str[:-4]
geral = geral[geral[0].str.len() <= 34]
geral.dropna(axis=1, inplace=True)
geral = geral.rename(columns={0: 'Dados'})

# dividir ultimos 11 caracteres em outra col
geral['Matricula'] = geral['Dados'].str[-11:]
geral = geral.drop('Dados', axis=1)
mat = []
matriculas_unicas = []
for item in geral['Matricula']:
    mat.append(item)
matnum = list(map(int, mat))
matriculas_unicas = list(set(matnum))
for matricula in matriculas_unicas:
    base = pd.DataFrame(list(intervalo(datainicio, datafim)),
                        columns=['Data', 'Entrada 1', 'Saída 1', 'Entrada 2', 'Saída 2', 'Entrada 3', 'Saída 3'])
    base = base.set_index('Data')
    geral = pd.read_csv(r'AFD.txt', sep=' ', header=None, encoding='iso8859-1')
    geral[0] = geral[0].astype(str).str[:-4]
    geral = geral[geral[0].str.len() <= 34]
    geral.dropna(axis=1, inplace=True)
    geral = geral.rename(columns={0: 'Dados'})
    # dividir ultimos 11 caracteres em outra col
    geral['Matricula'] = geral['Dados'].str[-11:]
    # dividir caracteres da data em outra col
    geral['Data'] = geral['Dados'].str[-24:-16]
    geral['Data'] = pd.to_datetime(geral['Data'], format='%d%m%Y')
    # dividir caracteres da hr em outra col
    geral['Hora'] = geral['Dados'].str[-16:-12]
    geral['Hora'] = pd.to_datetime(geral['Hora'], format='%H%M')

    geral1 = pd.DataFrame(geral)

    geral = geral.loc[(geral['Data'] < datafim) & (geral['Data'] > datainicio)]
    geral1 = geral1.loc[(geral1['Data'] < datafim) & (geral1['Data'] > datainicio)]

    geral = geral.loc[geral['Matricula'] == str(matricula).zfill(11)]
    geral1 = geral1.loc[geral1['Matricula'] == str(matricula).zfill(11)]

    geral['Hora'] = geral['Hora'].apply(lambda k: dt.strftime(k, '%H:%M'))
    geral1['Hora'] = geral1['Hora'].apply(lambda l: dt.strftime(l, '%H:%M:%S'))

    geral = geral.drop('Dados', axis=1)
    geral1 = geral1.drop('Dados', axis=1)

    geral = geral.drop('Matricula', axis=1)
    geral1 = geral1.drop('Matricula', axis=1)

    geral = geral.reset_index(drop=True)
    geral = geral.pivot_table(index='Data', columns=geral.groupby('Data').cumcount() + 1, values='Hora',
                              aggfunc='first')
    geral = geral.reset_index(level=[0])
    geral = geral.rename(
        columns={0: 'Data', 1: 'Entrada 1', 2: 'Saída 1', 3: 'Entrada 2', 4: 'Saída 2', 5: 'Entrada 3', 6: 'Saída 3'})

    geral1 = geral1.reset_index(drop=True)
    geral1 = geral1.pivot_table(index='Data', columns=geral1.groupby('Data').cumcount() + 1, values='Hora',
                                aggfunc='first')
    geral1 = geral1.reset_index(level=[0])
    geral1 = geral1.rename(
        columns={0: 'Data', 1: 'Entrada 1', 2: 'Saída 1', 3: 'Entrada 2', 4: 'Saída 2', 5: 'Entrada 3', 6: 'Saída 3'})

    try:
        geral1['Entrada 1a'] = geral1['Entrada 1'].apply(lambda z: pd.to_timedelta(str(z)))
        geral1['Saída 1a'] = geral1['Saída 1'].apply(lambda z: pd.to_timedelta(str(z)))
        geral1['Tot1'] = geral1['Saída 1a'] - geral1['Entrada 1a']
        total_horas = round(geral1['Tot1'].sum().total_seconds() / 3600, 2)

        if geral1.shape[1] > 7:
            if geral1.shape[1] < 11:
                geral1['Entrada 2a'] = geral1['Entrada 2'].apply(lambda z: pd.to_timedelta(str(z)))
                geral1['Saída 2a'] = geral1['Saída 2'].apply(lambda z: pd.to_timedelta(str(z)))
                geral1['Tot2'] = geral1['Saída 2a'] - geral1['Entrada 2a']
                geral1['Soma'] = geral1['Tot1'] + geral1['Tot2']
                total_horas = round(geral1['Soma'].sum().total_seconds() / 3600, 2)
                geral1 = geral1.drop('Entrada 2a', axis=1)
                geral1 = geral1.drop('Saída 2a', axis=1)
                geral1 = geral1.drop('Tot2', axis=1)
                geral1 = geral1.drop('Soma', axis=1)

        if geral1.shape[1] > 9:
            geral1['Entrada 2a'] = geral1['Entrada 2'].apply(lambda z: pd.to_timedelta(str(z)))
            geral1['Saída 2a'] = geral1['Saída 2'].apply(lambda z: pd.to_timedelta(str(z)))
            geral1['Entrada 3a'] = geral1['Entrada 3'].apply(lambda z: pd.to_timedelta(str(z)))
            geral1['Saída 3a'] = geral1['Saída 3'].apply(lambda z: pd.to_timedelta(str(z)))
            geral1['Tot1'] = geral1['Saída 1a'] - geral1['Entrada 1a']
            geral1['Tot2'] = geral1['Saída 2a'] - geral1['Entrada 2a']
            geral1['Tot3'] = geral1['Saída 3a'] - geral1['Entrada 3a']
            geral1['Soma'] = geral1['Tot1'] + geral1['Tot2'] + geral1['Tot3']
            total_horas = round(geral1['Soma'].sum().total_seconds() / 3600, 2)
            geral1 = geral1.drop('Entrada 2a', axis=1)
            geral1 = geral1.drop('Saída 2a', axis=1)
            geral1 = geral1.drop('Entrada 3a', axis=1)
            geral1 = geral1.drop('Saída 3a', axis=1)
            geral1 = geral1.drop('Tot2', axis=1)
            geral1 = geral1.drop('Tot3', axis=1)
            geral1 = geral1.drop('Soma', axis=1)

        geral1 = geral1.drop('Entrada 1a', axis=1)
        geral1 = geral1.drop('Saída 1a', axis=1)
        geral1 = geral1.drop('Tot1', axis=1)
    except:
        total_horas = 'erro'

    geral = geral.set_index('Data')
    base = base.combine_first(geral)
    base = base.reset_index(level=[0])
    base['Data'] = base['Data'].apply(lambda h: dt.strftime(h, '%d/%m/%Y - %a'))
    base = base[['Data', 'Entrada 1', 'Saída 1', 'Entrada 2', 'Saída 2', 'Entrada 3', 'Saída 3']]
    dias = base['Entrada 1'].count()
    base = base.fillna('-')
    ponto = docx.Document('modelo.docx')
    ponto.paragraphs[1].text = str(ponto.paragraphs[1].text).replace('#data1',
                                                                     dt.strftime(datainicio, '%d/%m/%Y')).replace(
        '#data2', dt.strftime(datafim, '%d/%m/%Y'))
    ponto.paragraphs[2].text = str(ponto.paragraphs[2].text).replace('#emissao', dt.strftime(dt.today(), '%d/%m/%Y'))
    ponto.tables[0].rows[3].cells[1].paragraphs[0].text = str(
        ponto.tables[0].rows[3].cells[1].paragraphs[0].text).replace('#nome', planbase[str(matricula)])
    ponto.tables[0].rows[4].cells[1].paragraphs[0].text = str(
        ponto.tables[0].rows[4].cells[1].paragraphs[0].text).replace('#cod', str(matricula))
    ponto.tables[0].rows[5].cells[1].paragraphs[0].text = str(
        ponto.tables[0].rows[5].cells[1].paragraphs[0].text).replace('#mat', planmat[str(matricula)])
    ponto.tables[0].rows[6].cells[1].paragraphs[0].text = str(
        ponto.tables[0].rows[6].cells[1].paragraphs[0].text).replace('#func', planfunc[str(matricula)])
    ponto.tables[0].rows[7].cells[1].paragraphs[0].text = str(
        ponto.tables[0].rows[7].cells[1].paragraphs[0].text).replace('#depto', plandept[str(matricula)])
    # ponto.tables[0].rows[5].cells[4].paragraphs[0].text = 'HH:MM'
    style = ponto.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(10)
    style8 = ponto.styles['Default']
    font = style8.font
    font.name = 'Times New Roman'
    font.size = Pt(8)

    ponto.paragraphs[0].style = ponto.styles['Normal']
    ponto.paragraphs[1].style = ponto.styles['Normal']
    ponto.paragraphs[2].style = ponto.styles['Normal']
    ponto.paragraphs[3].style = ponto.styles['Normal']
    ponto.paragraphs[4].style = ponto.styles['Normal']
    ponto.tables[0].rows[5].cells[4].paragraphs[0].style = ponto.styles['Default']
    t = ponto.add_table(base.shape[0] + 1, base.shape[1])
    t.style = 'Estilo2'
    # add the header rows.
    for j in range(base.shape[-1]):
        t.cell(0, j).text = base.columns[j]
    # add the rest of the data frame
    for i in range(base.shape[0]):
        for j in range(base.shape[-1]):
            t.cell(i + 1, j).text = str(base.values[i, j])
            t.cell(i + 1, j).paragraphs[0].alignment = 1
            t.cell(i + 1, 0).paragraphs[0].alignment = 0

    ponto.tables[1].columns[0].cells[0].width = Cm(4.2)
    ponto.tables[1].columns[1].cells[0].width = Cm(2.6)
    ponto.tables[1].columns[2].cells[0].width = Cm(2.6)
    ponto.tables[1].columns[3].cells[0].width = Cm(2.6)
    ponto.tables[1].columns[4].cells[0].width = Cm(2.6)
    ponto.tables[1].columns[5].cells[0].width = Cm(2.6)
    ponto.tables[1].columns[6].cells[0].width = Cm(2.6)
    ponto.tables[1].rows[0].cells[0].paragraphs[0].alignment = 1
    ponto.tables[1].rows[0].cells[1].paragraphs[0].alignment = 1
    ponto.tables[1].rows[0].cells[2].paragraphs[0].alignment = 1
    ponto.tables[1].rows[0].cells[3].paragraphs[0].alignment = 1
    ponto.tables[1].rows[0].cells[4].paragraphs[0].alignment = 1
    ponto.tables[1].rows[0].cells[5].paragraphs[0].alignment = 1
    ponto.tables[1].rows[0].cells[6].paragraphs[0].alignment = 1

    y = len(fl['A']) + 1
    fl[f'A{y}'].value = planbase[str(matricula)]
    fl[f'B{y}'].value = total_horas
    pl.save('Hrs Folha.xlsx')

    ponto.add_paragraph(f'Total de dias trabalhados: {dias}', 'Default').alignment = 2
    ponto.add_paragraph('', 'Normal')
    ponto.add_paragraph('', 'Normal')
    ponto.add_paragraph('______________________________________________', 'Normal').alignment = 1
    ponto.add_paragraph(f'{planbase[str(matricula)]}', 'Normal').alignment = 1
    ponto.save(
        rf'C:\Users\Felipe\PycharmProjects\AutomacaoCia\Ponto\Novo codigo\Ponto {planbase[str(matricula)]}.docx')
    docx2pdf.convert(
        rf'C:\Users\Felipe\PycharmProjects\AutomacaoCia\Ponto\Novo codigo\Ponto {planbase[str(matricula)]}.docx',
        rf'C:\Users\Felipe\PycharmProjects\AutomacaoCia\Ponto\Novo codigo\Ponto {planbase[str(matricula)]}.pdf')
    os.remove(
        rf'C:\Users\Felipe\PycharmProjects\AutomacaoCia\Ponto\Novo codigo\Ponto {planbase[str(matricula)]}.docx')
