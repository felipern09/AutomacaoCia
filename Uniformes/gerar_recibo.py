from datetime import datetime
import docx
import docx2pdf
from openpyxl import load_workbook as l_w
import tkinter.filedialog
from tkinter import ttk, messagebox
from tkinter import *
import tkinter.filedialog

root = Tk()
root.title("Recibos uniformes - Cia BSB")
img = PhotoImage(file='../Admissao/static/icone.png')
root.iconphoto(False, img)
root.geometry('480x350')
root.columnconfigure(0, weight=5)
root.rowconfigure(0, weight=5)

my_notebook = ttk.Notebook(root)
my_notebook.pack()

geral = StringVar()
caminho = StringVar()
caminhoest = StringVar()
caminhoaut = StringVar()


def selecionarfunc():
    try:
        caminhoplan = tkinter.filedialog.askopenfilename(title='Planilha Funcionários')
        caminho.set(str(caminhoplan))
    except ValueError:
        pass


funcionario = Frame(my_notebook, width=10, height=20)
ttk.Label(funcionario, width=40, text="Escolher planilha relatório de uniformes").grid(column=1, row=1, padx=25, pady=1,
                                                                                       sticky=W)
ttk.Button(funcionario, text="Escolha a planilha", command=selecionarfunc).grid(column=1, row=1, padx=350, pady=1,
                                                                                sticky=W)


def gerar_recibo(local, nome, cargo, cpf, genero, tamanho1, tamanho2=''):
    relatorio = l_w(local, read_only=False)
    estoque = relatorio['Estoque']
    entregues = relatorio['Entregues']
    lista = relatorio['Nomes']
    hoje = datetime.today()
    tipo, gen = genero.split(': ')
    num, pess = nome.split(' - ')
    label, cpf_ed = cpf.split(': ')
    pessoa = pess.title()

    if tamanho2 != '':
        recibo = docx.Document('Recibo.docx')
        recibo.paragraphs[11].text = str(recibo.paragraphs[11].text)\
            .replace('#nome', pessoa).replace('#num_cpf', cpf_ed).replace('#tam', tamanho1+' e '+tamanho2).replace(
            '#genero', str(gen).lower())
        recibo.paragraphs[19].text = str(recibo.paragraphs[19].text)\
            .replace('#data', datetime.strftime(hoje, '%d/%m/%Y'))
        recibo.paragraphs[24].text = str(recibo.paragraphs[24].text).replace('#nome', pessoa)
        recibo.paragraphs[25].text = str(recibo.paragraphs[25].text).replace('#cargo', cargo)
        recibo.save(f'Recibo_alterado {pessoa}.docx')
        docx2pdf.convert(f'Recibo_alterado {pessoa}.docx', f'Recibo {pessoa}.pdf')
        # diminuir contagem de estoque
        if gen == 'Masculino':
            if tamanho1 == 'P':
                estoque['C4'].value = estoque['C4'].value - 1
            if tamanho1 == 'M':
                estoque['C5'].value = estoque['C5'].value - 1
            if tamanho1 == 'G':
                estoque['C6'].value = estoque['C6'].value - 1
            if tamanho1 == 'XG':
                estoque['C8'].value = estoque['C8'].value - 1
            if tamanho1 == 'GG':
                estoque['C7'].value = estoque['C7'].value - 1
            if tamanho1 == 'XGG':
                estoque['C9'].value = estoque['C9'].value - 1
            if tamanho2 == 'P':
                estoque['C4'].value = estoque['C4'].value - 1
            if tamanho2 == 'M':
                estoque['C5'].value = estoque['C5'].value - 1
            if tamanho2 == 'G':
                estoque['C6'].value = estoque['C6'].value - 1
            if tamanho2 == 'XG':
                estoque['C8'].value = estoque['C8'].value - 1
            if tamanho2 == 'GG':
                estoque['C7'].value = estoque['C7'].value - 1
            if tamanho2 == 'XGG':
                estoque['C9'].value = estoque['C9'].value - 1
        else:
            if tamanho1 == 'P':
                estoque['E4'].value = estoque['E4'].value - 1
            if tamanho1 == 'M':
                estoque['E5'].value = estoque['E5'].value - 1
            if tamanho1 == 'G':
                estoque['E6'].value = estoque['E6'].value - 1
            if tamanho1 == 'GG':
                estoque['E7'].value = estoque['E7'].value - 1
            if tamanho2 == 'P':
                estoque['E4'].value = estoque['E4'].value - 1
            if tamanho2 == 'M':
                estoque['E5'].value = estoque['E5'].value - 1
            if tamanho2 == 'G':
                estoque['E6'].value = estoque['E6'].value - 1
            if tamanho2 == 'GG':
                estoque['E7'].value = estoque['E7'].value - 1
        lista[f'E{num}'].value = tamanho1
        lista[f'F{num}'].value = 'OK'
        x = len(list(entregues.rows)) + 1
        entregues[f'A{x}'].value = pessoa
        entregues[f'B{x}'].value = 1
        entregues[f'C{x}'].value = tamanho1
        entregues[f'D{x}'].value = gen
        entregues[f'E{x}'].value = 1
        entregues[f'F{x}'].value = tamanho2
        entregues[f'G{x}'].value = gen
        relatorio.save(local)
        tkinter.messagebox.showinfo(title='Recibo ok!', message='Recibo impresso com sucesso!')
    else:
        recibo = docx.Document('Recibo.docx')
        recibo.paragraphs[11].text = str(recibo.paragraphs[11].text)\
            .replace('#nome', pessoa).replace('#num_cpf', cpf_ed)\
            .replace('#tam', tamanho1).replace('#genero', str(gen).lower())
        recibo.paragraphs[19].text = str(recibo.paragraphs[19].text)\
            .replace('#data', datetime.strftime(hoje, '%d/%m/%Y'))
        recibo.paragraphs[24].text = str(recibo.paragraphs[24].text).replace('#nome', pessoa)
        recibo.paragraphs[25].text = str(recibo.paragraphs[25].text).replace('#cargo', cargo)
        recibo.save(f'Recibo_alterado {pessoa}.docx')
        docx2pdf.convert(f'Recibo_alterado {pessoa}.docx', f'Recibo {pessoa}.pdf')
        # diminuir contagem de estoque
        if gen == 'Masculino':
            if tamanho1 == 'P':
                estoque['C4'].value = estoque['C4'].value - 2
            if tamanho1 == 'M':
                estoque['C5'].value = estoque['C5'].value - 2
            if tamanho1 == 'G':
                estoque['C6'].value = estoque['C6'].value - 2
            if tamanho1 == 'XG':
                estoque['C8'].value = estoque['C8'].value - 2
            if tamanho1 == 'GG':
                estoque['C7'].value = estoque['C7'].value - 2
            if tamanho1 == 'XGG':
                estoque['C9'].value = estoque['C9'].value - 2
        else:
            if tamanho1 == 'P':
                estoque['E4'].value = estoque['E4'].value - 2
            if tamanho1 == 'M':
                estoque['E5'].value = estoque['E5'].value - 2
            if tamanho1 == 'G':
                estoque['E6'].value = estoque['E6'].value - 2
            if tamanho1 == 'GG':
                estoque['E7'].value = estoque['E7'].value - 2
        lista[f'E{num}'].value = tamanho1
        lista[f'F{num}'].value = 'OK'
        x = len(list(entregues.rows)) + 1
        entregues[f'A{x}'].value = pessoa
        entregues[f'B{x}'].value = 2
        entregues[f'C{x}'].value = tamanho1
        entregues[f'D{x}'].value = gen
        entregues[f'E{x}'].value = '-'
        entregues[f'F{x}'].value = '-'
        entregues[f'G{x}'].value = '-'
        relatorio.save(local)
    tkinter.messagebox.showinfo(title='Recibo ok!', message='Recibo impresso com sucesso!')


nome = StringVar()
cargo = StringVar()
cpf = StringVar()
tamanho = StringVar()
tamanho2 = StringVar()
genero = StringVar()
nomesplan = []
tamanhos = []
# aparecer dropdown com nomes da plan
labelnome = ttk.Label(funcionario, width=20, text="Nome:")
labelnome.grid(column=1, row=10, padx=25, pady=1, sticky=W)
combonome = ttk.Combobox(funcionario, values=nomesplan, textvariable=nome, width=38)
combonome.grid(column=1, row=10, padx=125, pady=1, sticky=W)
labelcargo = ttk.Label(funcionario, width=40, text="Cargo:")
labelcargo.grid(column=1, row=11, padx=25, pady=1, sticky=W)
labelcpf = ttk.Label(funcionario, width=20, text="CPF:")
labelcpf.grid(column=1, row=12, padx=25, pady=1, sticky=W)
labelgenero = ttk.Label(funcionario, width=20, text="Tipo:")
labelgenero.grid(column=1, row=13, padx=25, pady=1, sticky=W)

# aparecer entry para preencher tamanho
labeltam = ttk.Label(funcionario, width=20, text="Tamanho:")
labeltam.grid(column=1, row=14, padx=25, pady=1, sticky=W)
combotam = ttk.Combobox(funcionario, values=tamanhos, textvariable=tamanho, width=10)
combotam.grid(column=1, row=14, padx=125, pady=1, sticky=W)
maisum = IntVar()


def adicionartamanho():
    labeltam.config(text='Tamanho 1:')
    labeltam1 = ttk.Label(funcionario, width=20, text='Tamanho 2:')
    labeltam1.grid(column=1, row=15, padx=25, pady=1, sticky=W)
    combotam1 = ttk.Combobox(funcionario, values=tamanhos, textvariable=tamanho2, width=10)
    combotam1.grid(column=1, row=15, padx=125, pady=1, sticky=W)
    if labelgenero['text'] == 'Tipo: Feminino':
        lista_tamanhos = ['P', 'M', 'G', 'GG']
        combotam1.config(values=lista_tamanhos)
    if labelgenero['text'] == 'Tipo: Masculino':
        lista_tamanhos = ['M', 'G', 'XG', 'GG', 'XGG']
        combotam1.config(values=lista_tamanhos)


onde = ttk.Checkbutton(
    funcionario, text='Mais de um tamanho.', variable=maisum, onvalue=1, offvalue=0, command=adicionartamanho
)
onde.grid(column=1, row=17, padx=25, pady=1, sticky=W)


def mostravalores(event):
    nome = event.widget.get()
    num, name = nome.split(' - ')
    linha = int(num)
    planwb = l_w(caminho.get())
    plansh = planwb['Nomes']
    labelcargo.config(text=f'Cargo: {str(plansh[f"B{linha}"].value)}')
    labelcpf.config(text=f'CPF: {str(plansh[f"C{linha}"].value)}')
    labelgenero.config(text=f'Tipo: {str(plansh[f"D{linha}"].value)}')
    if labelgenero['text'] == 'Tipo: Feminino':
        lista_tamanhos = ['P', 'M', 'G', 'GG']
        combotam.config(values=lista_tamanhos)
    if labelgenero['text'] == 'Tipo: Masculino':
        lista_tamanhos = ['M', 'G', 'XG', 'GG', 'XGG']
        combotam.config(values=lista_tamanhos)


combonome.bind("<<ComboboxSelected>>", mostravalores)


def carregarfunc(local):
    planwb = l_w(local)
    plansh = planwb['Nomes']
    lista = []
    for x, pessoa in enumerate(plansh):
        lista.append(f'{x + 1} - {pessoa[0].value}')
    combonome.config(values=lista)


ttk.Button(funcionario, text="Carregar planilha", command=lambda: [carregarfunc(caminho.get())]).grid(column=1, row=9,
                                                                                                      padx=350, pady=25,
                                                                                                      sticky=W)
ttk.Button(funcionario, width=20, text="Gerar recibo",
           command=lambda: [
               gerar_recibo(caminho.get(), nome.get(), str(labelcargo['text']), str(labelcpf['text']),
                            str(labelgenero['text']), tamanho.get(), tamanho2.get())
           ]).grid(column=1, row=28, padx=325, pady=1, sticky=W)
funcionario.pack(fill='both', expand=0)
my_notebook.add(funcionario, text='Gerar Recibo de uniforme')
root.mainloop()
